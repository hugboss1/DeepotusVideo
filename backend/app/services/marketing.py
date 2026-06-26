"""v1.9 — Marketing plan generation + channel publishing + schedule loop.

Three building blocks:

1. generate_plan(...)   — prompt → structured week plan. Uses Anthropic
   (BYO ANTHROPIC_API_KEY) when available; otherwise a deterministic
   format-rotation generator so the feature works with zero LLM keys.
2. publish_telegram(...) — the reference auto-publish adapter. Free, no
   app review, supports video. Other channels (X, YouTube, Instagram)
   stay 'assisted' until their adapters land.
3. schedule_loop()       — fires due posts every 60 s. mode='auto' posts
   publish to capable channels; mode='assisted' posts flip to 'ready'
   so the user can one-click post from the UI. A failure never kills
   the loop.
"""
import asyncio
import json
import re
from datetime import datetime, timedelta
from pathlib import Path
from uuid import uuid4

import httpx
from loguru import logger
from sqlalchemy import select

from app.config import settings, SSL_VERIFY
from app.services.storage import ScheduledPost, JobRecord, async_session_factory

# ---------------------------------------------------------------- plan gen

# Proven memecoin/community content formats, rotated across the week.
# Each maps to a pipeline the tool can actually produce.
_FORMATS = [
    {"format": "image",       "label": "Hook meme",
     "hook": "One-liner over a strong visual. Stop the scroll.",
     "time": "09:30"},
    {"format": "seedance",    "label": "Lore drop",
     "hook": "Cinematic 10s clip extending the persona's universe.",
     "time": "12:00"},
    {"format": "heygen",      "label": "Avatar prophecy",
     "hook": "The avatar speaks: bold claim or market take, 15-20s.",
     "time": "18:00"},
    {"format": "news",        "label": "News reaction",
     "hook": "React to today's headline in persona voice. Timely = reach.",
     "time": "13:30"},
    {"format": "composition", "label": "Clip + avatar",
     "hook": "B-roll clip with the avatar reading the take. Premium feel.",
     "time": "19:00"},
    {"format": "image",       "label": "Community CTA",
     "hook": "Question or ritual for holders. Optimize replies.",
     "time": "21:00"},
    {"format": "seedance",    "label": "Weekly recap",
     "hook": "The week as the persona saw it. Recap + tease next week.",
     "time": "17:00"},
]


def _deterministic_plan(prompt: str, days: int, posts_per_day: int,
                        channels: list[str], language: str,
                        persona: dict | None) -> list[dict]:
    """LLM-free fallback: rotate proven formats, seed captions from the
    user's prompt + persona. Always succeeds."""
    pname = (persona or {}).get("name", "the prophet")
    ticker = ""
    m = re.search(r"\$[A-Z]{2,12}", prompt.upper())
    if m:
        ticker = m.group(0)
    theme = prompt.strip()[:140] or "the project"
    posts: list[dict] = []
    k = 0
    for d in range(days):
        for s in range(posts_per_day):
            f = _FORMATS[k % len(_FORMATS)]
            k += 1
            title = f"{f['label']} — day {d + 1}"
            if language.upper().startswith("FR"):
                caption = (f"{theme}\n\n{f['hook']}\n"
                           f"{('— ' + pname) if pname else ''} "
                           f"{ticker}").strip()
            else:
                caption = (f"{theme}\n\n{f['hook']}\n"
                           f"{('— ' + pname) if pname else ''} "
                           f"{ticker}").strip()
            posts.append({
                "day_offset": d,
                "time": _FORMATS[(k - 1) % len(_FORMATS)]["time"]
                        if posts_per_day == 1 else
                        ["09:30", "13:30", "18:00", "21:00"][s % 4],
                "title": title,
                "format": f["format"],
                "hook": f["hook"],
                "caption": caption,
                "script_idea": f"{f['label']}: {theme}",
                "image_idea": f"{pname} — {f['label'].lower()}, "
                              f"deep-sea palette, 9:16",
                "channels": channels or ["x"],
            })
    return posts


async def _anthropic_plan(prompt: str, days: int, posts_per_day: int,
                          channels: list[str], language: str,
                          persona: dict | None) -> list[dict] | None:
    """Ask Claude for a structured plan. Returns None on ANY failure so
    the caller falls back to the deterministic generator."""
    if not settings.has_summarizer:
        return None
    pdesc = ""
    if persona:
        pdesc = (f"Persona: {persona.get('name', '')}. "
                 f"Tone: {persona.get('tone', '')}. "
                 f"Audience: {persona.get('audience', '')}. ")
    sys = (
        "You are a social media content strategist for short-form video "
        "accounts (memecoins, creator brands). Produce a posting plan as "
        "STRICT JSON, no prose. Schema: {\"posts\":[{\"day_offset\":int "
        f"(0..{days - 1}),\"time\":\"HH:MM\",\"title\":str,"
        "\"format\":\"image|seedance|heygen|composition|news\","
        "\"hook\":str,\"caption\":str (ready to publish, with line breaks "
        "and at most 2 emojis),\"script_idea\":str,\"image_idea\":str,"
        "\"channels\":[\"x\"|\"telegram\"|\"youtube\"|\"instagram\"]}]}. "
        f"Exactly {days * posts_per_day} posts ({posts_per_day}/day over "
        f"{days} days). Language: {language}. {pdesc}"
        "Formats map to the user's video tool: image=meme still, "
        "seedance=cinematic clip, heygen=talking avatar, composition="
        "clip+avatar, news=news reaction reel. Vary formats and times "
        "(morning/noon/evening). Captions must follow the persona voice."
    )
    try:
        async with httpx.AsyncClient(verify=SSL_VERIFY, timeout=60.0) as client:
            r = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key": settings.ANTHROPIC_API_KEY,
                    "anthropic-version": "2023-06-01",
                    "content-type": "application/json",
                },
                json={
                    "model": settings.ANTHROPIC_MODEL,
                    "max_tokens": 4000,
                    "system": sys,
                    "messages": [{"role": "user", "content": prompt}],
                },
            )
            if r.status_code != 200:
                logger.warning(f"plan: anthropic {r.status_code}: "
                               f"{r.text[:200]}")
                return None
            text = "".join(
                b.get("text", "") for b in r.json().get("content", []))
            # Tolerate fences / prose around the JSON object.
            start = text.find("{")
            end = text.rfind("}")
            if start < 0 or end <= start:
                return None
            data = json.loads(text[start:end + 1])
            posts = data.get("posts")
            if not isinstance(posts, list) or not posts:
                return None
            clean = []
            for p in posts:
                try:
                    clean.append({
                        "day_offset": max(0, min(days - 1,
                                                 int(p.get("day_offset", 0)))),
                        "time": str(p.get("time", "12:00"))[:5],
                        "title": str(p.get("title", ""))[:200],
                        "format": str(p.get("format", "image")),
                        "hook": str(p.get("hook", "")),
                        "caption": str(p.get("caption", "")),
                        "script_idea": str(p.get("script_idea", "")),
                        "image_idea": str(p.get("image_idea", "")),
                        "channels": [c for c in (p.get("channels") or ["x"])
                                     if c in ("x", "telegram", "youtube",
                                              "instagram")] or ["x"],
                    })
                except (TypeError, ValueError):
                    continue
            return clean or None
    except Exception as e:
        logger.warning(f"plan: anthropic call failed: {e}")
        return None


async def _openai_plan(prompt: str, days: int, posts_per_day: int,
                       channels: list[str], language: str,
                       persona: dict | None) -> list[dict] | None:
    from app.services.openai_llm import generate_plan as _gp
    return await _gp(prompt, days, posts_per_day, channels, language, persona)


async def _gemini_plan(prompt: str, days: int, posts_per_day: int,
                       channels: list[str], language: str,
                       persona: dict | None) -> list[dict] | None:
    from app.services.gemini_llm import generate_plan as _gp
    return await _gp(prompt, days, posts_per_day, channels, language, persona)


async def _ollama_plan(prompt: str, days: int, posts_per_day: int,
                       channels: list[str], language: str,
                       persona: dict | None) -> list[dict] | None:
    """Same contract as _anthropic_plan but against a local Ollama server.
    Plans never leave the machine. Returns None on ANY failure."""
    if not settings.has_ollama:
        return None
    pdesc = ""
    if persona:
        pdesc = (f"Persona: {persona.get('name', '')}. "
                 f"Tone: {persona.get('tone', '')}. "
                 f"Audience: {persona.get('audience', '')}. ")
    sys = (
        "You are a social media content strategist. Output STRICT JSON only, "
        "no prose, no markdown fences. Schema: {\"posts\":[{\"day_offset\":"
        f"int (0..{days - 1}),\"time\":\"HH:MM\",\"title\":str,"
        "\"format\":\"image|seedance|heygen|composition|news\",\"hook\":str,"
        "\"caption\":str,\"script_idea\":str,\"image_idea\":str,"
        "\"channels\":[\"x\"|\"telegram\"|\"youtube\"|\"instagram\"]}]}. "
        f"Exactly {days * posts_per_day} posts ({posts_per_day}/day over "
        f"{days} days). Language: {language}. {pdesc}"
        "Vary formats and times. Captions ready to publish."
    )
    try:
        async with httpx.AsyncClient(verify=SSL_VERIFY, timeout=180.0) as client:
            r = await client.post(
                f"{settings.OLLAMA_URL.rstrip('/')}/api/chat",
                json={
                    "model": settings.OLLAMA_MODEL,
                    "stream": False,
                    "format": "json",
                    "options": {"temperature": 0.7},
                    "messages": [
                        {"role": "system", "content": sys},
                        {"role": "user", "content": prompt},
                    ],
                },
            )
            if r.status_code != 200:
                logger.warning(f"plan: ollama {r.status_code}: {r.text[:200]}")
                return None
            text = (r.json().get("message") or {}).get("content", "")
            start, end = text.find("{"), text.rfind("}")
            if start < 0 or end <= start:
                return None
            data = json.loads(text[start:end + 1])
            posts = data.get("posts")
            if not isinstance(posts, list) or not posts:
                return None
            clean = []
            for p in posts:
                try:
                    clean.append({
                        "day_offset": max(0, min(days - 1,
                                                 int(p.get("day_offset", 0)))),
                        "time": str(p.get("time", "12:00"))[:5],
                        "title": str(p.get("title", ""))[:200],
                        "format": str(p.get("format", "image")),
                        "hook": str(p.get("hook", "")),
                        "caption": str(p.get("caption", "")),
                        "script_idea": str(p.get("script_idea", "")),
                        "image_idea": str(p.get("image_idea", "")),
                        "channels": [c for c in (p.get("channels") or ["x"])
                                     if c in ("x", "telegram", "youtube",
                                              "instagram")] or ["x"],
                    })
                except (TypeError, ValueError):
                    continue
            return clean or None
    except Exception as e:
        logger.warning(f"plan: ollama call failed: {e}")
        return None


_PLAN_PROVIDERS = {
    "anthropic": _anthropic_plan,
    "openai": _openai_plan,
    "gemini": _gemini_plan,
    "ollama": _ollama_plan,
}
_PLAN_PRIORITY = ["anthropic", "openai", "gemini", "ollama"]


def _plan_available(name: str) -> bool:
    checks = {
        "anthropic": lambda: bool(settings.ANTHROPIC_API_KEY.strip()),
        "openai": lambda: bool(settings.OPENAI_API_KEY.strip()),
        "gemini": lambda: bool(settings.GEMINI_API_KEY.strip()),
        "ollama": lambda: bool(settings.OLLAMA_MODEL.strip()),
    }
    return checks.get(name, lambda: False)()


async def generate_plan(prompt: str, *, days: int = 7,
                        posts_per_day: int = 1,
                        channels: list[str] | None = None,
                        language: str = "EN",
                        persona: dict | None = None) -> dict:
    """Returns {"posts": [...], "engine": "<provider>"|"deterministic"}."""
    channels = channels or ["x"]
    perf = await performance_context()
    full_prompt = f"{prompt}\n\n{perf}" if perf else prompt
    pref = settings.PLANNER_PROVIDER.strip().lower()
    order = [p for p in _PLAN_PRIORITY if _plan_available(p)]
    if pref and pref in order:
        order = [pref] + [p for p in order if p != pref]
    for engine in order:
        fn = _PLAN_PROVIDERS.get(engine)
        if fn:
            posts = await fn(full_prompt, days, posts_per_day, channels,
                             language, persona)
            if posts is not None:
                return {"posts": posts, "engine": engine}
    posts = _deterministic_plan(prompt, days, posts_per_day, channels,
                                language, persona)
    return {"posts": posts, "engine": "deterministic"}


async def materialize_plan(posts: list[dict], *, start_date: str,
                           tz_offset_minutes: int = 0,
                           mode: str = "assisted") -> list[str]:
    """Insert plan posts as scheduled_posts rows. start_date is the user's
    local YYYY-MM-DD for day_offset 0; tz_offset_minutes is JS
    getTimezoneOffset() (UTC - local), so utc = local + offset."""
    base = datetime.fromisoformat(start_date)
    plan_id = str(uuid4())
    ids: list[str] = []
    async with async_session_factory() as session:
        for p in posts:
            hh, mm = (p.get("time") or "12:00").split(":")[:2]
            local = base + timedelta(days=int(p.get("day_offset", 0)),
                                     hours=int(hh), minutes=int(mm))
            run_at = local + timedelta(minutes=tz_offset_minutes)
            row = ScheduledPost(
                id=str(uuid4()),
                title=p.get("title", "")[:200],
                caption=p.get("caption", ""),
                channels=",".join(p.get("channels") or ["x"]),
                run_at=run_at,
                status="scheduled",
                mode=mode if mode in ("auto", "assisted") else "assisted",
                format=p.get("format"),
                hook=p.get("hook"),
                script_idea=p.get("script_idea"),
                image_idea=p.get("image_idea"),
                source_image=(p.get("source_image") or None),
                plan_id=plan_id,
            )
            session.add(row)
            ids.append(row.id)
        await session.commit()
    logger.info(f"plan {plan_id}: materialized {len(ids)} posts")
    return ids


# ---------------------------------------------------------------- doc import

def extract_document_text(filename: str, data: bytes) -> str:
    """Plain text from an uploaded .md / .txt / .docx / .pdf strategy doc.
    Raises ValueError on unsupported types or unreadable files."""
    name = (filename or "").lower()
    if name.endswith((".md", ".txt", ".markdown")):
        return data.decode("utf-8", errors="replace")
    if name.endswith(".docx"):
        import io
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(p for p in parts if p and p.strip())
    if name.endswith(".pdf"):
        import io
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    raise ValueError(f"Unsupported file type: {filename}. "
                     "Use .md, .txt, .docx or .pdf")


async def plan_from_document(text: str, *, days: int = 30,
                             channels: list[str] | None = None,
                             language: str = "EN",
                             persona: dict | None = None) -> dict:
    """Transcribe an existing strategy document into scheduled-post slices.
    Human-in-the-loop by design: this only RETURNS the structured plan; the
    UI shows the preview and the user materializes explicitly."""
    channels = channels or ["x"]
    doc = text.strip()
    if not doc:
        raise ValueError("Document is empty after extraction")
    # Cap what we send to the LLM; strategy docs can be huge.
    doc = doc[:24000]
    prompt = (
        "Transcribe this existing marketing strategy document into the "
        "posting-plan schema. PRESERVE the document's own dates, themes, "
        "milestones and copy ideas — do not invent a different strategy. "
        f"Spread over up to {days} days (day_offset 0 = the plan's first "
        "day). If the document defines weeks, keep the weekly structure.\n\n"
        f"DOCUMENT:\n{doc}"
    )
    pref = settings.PLANNER_PROVIDER.strip().lower()
    order = [p for p in _PLAN_PRIORITY if _plan_available(p)]
    if pref and pref in order:
        order = [pref] + [p for p in order if p != pref]
    posts = None
    engine = "deterministic"
    for eng in order:
        fn = _PLAN_PROVIDERS.get(eng)
        if fn:
            posts = await fn(prompt, days, 1, channels, language, persona)
            if posts is not None:
                engine = eng
                break
    if posts is None:
        # Deterministic fallback: split the document on day/week headings and
        # turn each chunk into one post seed.
        engine = "deterministic"
        chunks = re.split(
            r"(?im)^#{1,3}\s*(?:semaine|week|jour|day)\b.*$", doc)
        chunks = [c.strip() for c in chunks if c.strip()][:days]
        if not chunks:
            chunks = [doc[i:i + 400] for i in range(0, min(len(doc), 400 * days), 400)]
        posts = []
        per_week = max(1, len(chunks))
        for i, chunk in enumerate(chunks):
            first_line = chunk.splitlines()[0][:120] if chunk.splitlines() else f"Plan item {i+1}"
            posts.append({
                "day_offset": min(days - 1, i * max(1, days // per_week)),
                "time": ["09:30", "18:00", "21:00"][i % 3],
                "title": first_line,
                "format": ["image", "seedance", "heygen"][i % 3],
                "hook": first_line,
                "caption": chunk[:500],
                "script_idea": chunk[:300],
                "image_idea": first_line,
                "channels": channels,
            })
    return {"posts": posts, "engine": engine}


# ---------------------------------------------------------------- X adapter

def _x_client():
    import tweepy
    return tweepy.Client(
        consumer_key=settings.X_API_KEY,
        consumer_secret=settings.X_API_SECRET,
        access_token=settings.X_ACCESS_TOKEN,
        access_token_secret=settings.X_ACCESS_SECRET,
    )


def _x_api_v1():
    """v1.1 API object — still required for chunked media upload."""
    import tweepy
    auth = tweepy.OAuth1UserHandler(
        settings.X_API_KEY, settings.X_API_SECRET,
        settings.X_ACCESS_TOKEN, settings.X_ACCESS_SECRET)
    return tweepy.API(auth)


def _publish_x_sync(caption: str, video_path: str | None,
                    image_path: str | None) -> tuple[bool, str, str | None]:
    """Blocking tweepy calls — run via asyncio.to_thread. Returns
    (ok, detail, tweet_id)."""
    try:
        client = _x_client()
        media_ids = None
        if video_path and Path(video_path).is_file():
            api = _x_api_v1()
            media = api.media_upload(
                filename=str(video_path),
                media_category="tweet_video", chunked=True)
            media_ids = [media.media_id_string]
        elif image_path and Path(image_path).is_file():
            api = _x_api_v1()
            media = api.media_upload(filename=str(image_path))
            media_ids = [media.media_id_string]
        resp = client.create_tweet(
            text=(caption or "")[:280], media_ids=media_ids)
        tid = str((resp.data or {}).get("id", ""))
        return True, f"tweet {tid}", tid or None
    except Exception as e:
        return False, f"x error: {e}", None


async def publish_x(caption: str, *, video_path: str | None = None,
                    image_path: str | None = None,
                    retries: int = 2) -> tuple[bool, str, str | None]:
    """Post to X with media. One retry on transient failure. Never raises."""
    if not settings.has_x:
        return False, "X keys not set (X_API_KEY/SECRET + ACCESS_TOKEN/SECRET)", None
    last = ("", None)
    for attempt in range(retries):
        ok, detail, tid = await asyncio.to_thread(
            _publish_x_sync, caption, video_path, image_path)
        if ok:
            return True, detail, tid
        last = (detail, None)
        # Don't retry on clearly non-transient errors.
        if any(s in detail.lower() for s in ("401", "403", "duplicate",
                                             "unauthorized", "forbidden")):
            break
        await asyncio.sleep(3 * (attempt + 1))
    return False, last[0], None


def _fetch_x_metrics_sync(tweet_ids: list[str]) -> dict[str, dict]:
    """public_metrics for up to 100 tweets. Read budget on the free tier is
    tiny (~100/mo) — callers must ration."""
    try:
        client = _x_client()
        resp = client.get_tweets(ids=tweet_ids,
                                 tweet_fields=["public_metrics"])
        out = {}
        for t in (resp.data or []):
            out[str(t.id)] = dict(t.public_metrics or {})
        return out
    except Exception as e:
        logger.warning(f"x metrics fetch failed: {e}")
        return {}


async def refresh_x_metrics(max_posts: int = 10) -> int:
    """Best-effort daily pass: pull public_metrics for the most recent posts
    that went out via the X adapter. Returns number updated."""
    if not settings.has_x:
        return 0
    async with async_session_factory() as session:
        res = await session.execute(
            select(ScheduledPost)
            .where(ScheduledPost.x_post_id.isnot(None))
            .where(ScheduledPost.status == "posted")
            .order_by(ScheduledPost.posted_at.desc())
            .limit(max_posts))
        posts = list(res.scalars().all())
        if not posts:
            return 0
        ids = [p.x_post_id for p in posts if p.x_post_id]
        metrics = await asyncio.to_thread(_fetch_x_metrics_sync, ids)
        n = 0
        for p in posts:
            m = metrics.get(p.x_post_id or "")
            if m:
                p.metrics = json.dumps(m)
                n += 1
        await session.commit()
        return n


async def performance_context(limit: int = 12) -> str:
    """Compact summary of recently posted content + any metrics, injected
    into the plan generator prompt so new plans learn from what shipped."""
    async with async_session_factory() as session:
        res = await session.execute(
            select(ScheduledPost)
            .where(ScheduledPost.status == "posted")
            .order_by(ScheduledPost.posted_at.desc())
            .limit(limit))
        posts = list(res.scalars().all())
    if not posts:
        return ""
    lines = []
    for p in posts:
        m = ""
        if p.metrics:
            try:
                d = json.loads(p.metrics)
                m = (f" [impressions={d.get('impression_count', '?')}, "
                     f"likes={d.get('like_count', '?')}, "
                     f"reposts={d.get('retweet_count', '?')}]")
            except ValueError:
                pass
        lines.append(f"- {p.format or 'post'} · \"{(p.title or '')[:60]}\" "
                     f"on {p.channels}{m}")
    return ("Recent posted content and performance (favor formats and "
            "angles that performed):\n" + "\n".join(lines))


# ---------------------------------------------------------------- telegram

async def publish_telegram(caption: str, *, video_path: str | None = None,
                           image_path: str | None = None) -> tuple[bool, str]:
    """Send to the configured Telegram chat. Video > photo > text.
    Returns (ok, detail). Never raises."""
    if not settings.has_telegram:
        return False, "TELEGRAM_BOT_TOKEN / TELEGRAM_CHAT_ID not set"
    base = f"https://api.telegram.org/bot{settings.TELEGRAM_BOT_TOKEN}"
    cap = (caption or "")[:1024]  # media caption hard limit
    try:
        async with httpx.AsyncClient(verify=SSL_VERIFY, timeout=120.0) as client:
            if video_path and Path(video_path).is_file():
                with open(video_path, "rb") as f:
                    r = await client.post(
                        f"{base}/sendVideo",
                        data={"chat_id": settings.TELEGRAM_CHAT_ID,
                              "caption": cap, "supports_streaming": "true"},
                        files={"video": (Path(video_path).name, f,
                                         "video/mp4")},
                    )
            elif image_path and Path(image_path).is_file():
                with open(image_path, "rb") as f:
                    r = await client.post(
                        f"{base}/sendPhoto",
                        data={"chat_id": settings.TELEGRAM_CHAT_ID,
                              "caption": cap},
                        files={"photo": (Path(image_path).name, f,
                                         "image/png")},
                    )
            else:
                r = await client.post(
                    f"{base}/sendMessage",
                    json={"chat_id": settings.TELEGRAM_CHAT_ID,
                          "text": (caption or "")[:4096]},
                )
        j = r.json()
        if r.status_code == 200 and j.get("ok"):
            return True, "sent"
        return False, f"telegram {r.status_code}: {str(j)[:200]}"
    except Exception as e:
        return False, f"telegram error: {e}"


async def _job_video_path(job_id: str | None) -> str | None:
    if not job_id:
        return None
    async with async_session_factory() as session:
        res = await session.execute(
            select(JobRecord).where(JobRecord.id == job_id))
        j = res.scalar_one_or_none()
        return j.final_video_path if j else None


async def _resolve_post_image(post) -> str | None:
    """Still image to attach when a post has no video: the explicit
    source_image first, else the attached render job's image_filename.
    Returns an absolute path (str) or None. Mirrors the hero-image
    resolution used by GET /schedule/{id}/preview.png so what the user
    previews is what actually gets published."""
    # 1) user-attached / generated still
    if getattr(post, "source_image", None):
        cand = settings.images_path / post.source_image
        if cand.is_file():
            return str(cand)
    # 2) the attached render's still frame
    if getattr(post, "job_id", None):
        async with async_session_factory() as session:
            res = await session.execute(
                select(JobRecord).where(JobRecord.id == post.job_id))
            job = res.scalar_one_or_none()
            if job and job.image_filename:
                cand = settings.images_path / job.image_filename
                if cand.is_file():
                    return str(cand)
    return None


def auto_channels() -> set[str]:
    """Channels with a working auto-publish adapter given current keys."""
    out = set()
    if settings.has_telegram:
        out.add("telegram")
    if settings.has_x:
        out.add("x")
    return out


async def fire_post(post_id: str) -> dict:
    """Publish a post NOW on its auto-capable channels. Returns a result
    dict; flips status to posted/ready/failed accordingly."""
    async with async_session_factory() as session:
        res = await session.execute(
            select(ScheduledPost).where(ScheduledPost.id == post_id))
        post = res.scalar_one_or_none()
        if not post:
            return {"ok": False, "error": "post not found"}
        channels = [c for c in (post.channels or "").split(",") if c]
        video = await _job_video_path(post.job_id)
        # When there's no video, attach the post's still image (source_image
        # or the attached render's frame) so X/Telegram get media, not just
        # text. The adapters prefer video > image > text.
        image = None if video else await _resolve_post_image(post)
        sent, errors = [], []
        for ch in channels:
            if ch == "telegram" and settings.has_telegram:
                ok, detail = await publish_telegram(
                    post.caption or post.title,
                    video_path=video, image_path=image)
                (sent if ok else errors).append(f"{ch}: {detail}")
            elif ch == "x" and settings.has_x:
                ok, detail, tid = await publish_x(
                    post.caption or post.title,
                    video_path=video, image_path=image)
                if ok and tid:
                    post.x_post_id = tid
                (sent if ok else errors).append(f"{ch}: {detail}")
            else:
                errors.append(f"{ch}: assisted (no auto adapter)")
        if sent:
            post.status = "posted"
            post.posted_at = datetime.utcnow()
            post.error = "; ".join(errors) if errors else None
        else:
            post.status = "ready"
            post.error = "; ".join(errors) if errors else None
        await session.commit()
        return {"ok": bool(sent), "sent": sent, "pending": errors,
                "status": post.status}


# ---------------------------------------------------------------- the loop

async def schedule_loop() -> None:
    """Fire due posts every 60 s. Robust: one bad post can't kill the loop.
    Once a day, also refreshes X public metrics for recent posts (rationed:
    the free X tier has a tiny read budget)."""
    logger.info("schedule loop started (60s tick)")
    last_metrics_day: str | None = None
    while True:
        try:
            now = datetime.utcnow()
            today = now.strftime("%Y-%m-%d")
            if last_metrics_day != today:
                last_metrics_day = today
                try:
                    n = await refresh_x_metrics(max_posts=10)
                    if n:
                        logger.info(f"x metrics refreshed for {n} post(s)")
                except Exception as e:
                    logger.warning(f"x metrics pass failed: {e}")
            async with async_session_factory() as session:
                res = await session.execute(
                    select(ScheduledPost)
                    .where(ScheduledPost.status == "scheduled")
                    .where(ScheduledPost.run_at <= now))
                due = list(res.scalars().all())
            for post in due:
                try:
                    if post.mode == "auto":
                        result = await fire_post(post.id)
                        logger.info(f"schedule: auto-fired {post.id} "
                                    f"-> {result.get('status')}")
                    else:
                        async with async_session_factory() as session:
                            res = await session.execute(
                                select(ScheduledPost)
                                .where(ScheduledPost.id == post.id))
                            p = res.scalar_one_or_none()
                            if p and p.status == "scheduled":
                                p.status = "ready"
                                await session.commit()
                        logger.info(f"schedule: {post.id} due -> ready "
                                    f"(assisted)")
                except Exception as e:
                    logger.error(f"schedule: post {post.id} failed: {e}")
        except asyncio.CancelledError:
            raise
        except Exception as e:
            logger.error(f"schedule loop tick failed: {e}")
        await asyncio.sleep(60)
