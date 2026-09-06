"""FastAPI route definitions — v1.3 (batch multi-seeds)."""
import asyncio
import hashlib
import json
import random
import re
from datetime import datetime
from pathlib import Path
from uuid import uuid4

import httpx

from fastapi import APIRouter, BackgroundTasks, HTTPException, UploadFile, File, Form, Request, Response
from fastapi.responses import FileResponse
from PIL import Image as PILImage
from loguru import logger

from app.config import settings, APP_VERSION, SSL_VERIFY
from app.services import library_index as LI
from app.models.schemas import (
    GenerateRequest,
    GenerateResponse,
    GenerateBatchRequest,
    GenerateBatchResponse,
    GenerateHeyGenRequest,
    GenerateHeyGenImageRequest,
    GenerateHeyGenCinematicRequest,
    AvatarPresetCreate,
    CompositionRequest,
    CompositionResponse,
    JobStatus,
    ImageItem,
    BuildPromptRequest,
    BuildPromptResponse,
    HeyGenAvatar,
    HeyGenVoice,
    PhotoAvatarCreateResponse,
    BuildScriptRequest,
    BuildScriptResponse,
    BuildCompositionRequest,
    BuildCompositionResponse,
    TemplateSaveRequest,
    TemplateSaveResponse,
    TemplateRenderRequest,
    TemplateRenderResponse,
    JobRenameRequest,
    AddNewsSourceRequest,
    NewsSourceToggleRequest,
    NewsScriptRequest,
    NewsScriptResponse,
    NewsEssence,
    NewsIllustrationRequest,
    NewsIllustrationResponse,
)
from app.services.pipeline import Pipeline
from app.services.fs_guard import is_virtualized as fs_is_virtualized
from app.services.heygen_service import HeyGenClient, HeyGenError, invalidate_list_cache
from app.services.template_service import TemplateEngine
from app.services.news_service import news_service


router = APIRouter()
pipeline = Pipeline(persona_id="deepotus")
template_engine = TemplateEngine()


# ---- Remote-URL fetching (SSRF guard) ----

def _is_private_host(h: str | None) -> bool:
    """True for hosts that must never be fetched on the user's behalf."""
    import ipaddress as _ipaddr
    h = (h or "").lower()
    if h in ("localhost", "") or h.endswith(".local"):
        return True
    try:
        ip = _ipaddr.ip_address(h)
        return ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved
    except ValueError:
        return False  # plain hostname; DNS-rebinding out of scope for a desktop tool


async def _block_private_redirect(response: httpx.Response) -> None:
    """Re-check every hop: a public URL can 302 to a loopback address, which
    would walk straight past a check made only on the original host."""
    loc = response.headers.get("location")
    if response.is_redirect and loc and _is_private_host(httpx.URL(loc).host):
        raise HTTPException(400, "Refusing to follow a redirect to a private/loopback address")


def _remote_image_client(**kw) -> httpx.AsyncClient:
    """Client for user-supplied image URLs: TLS per settings, redirects
    followed but re-validated at each hop."""
    kw.setdefault("timeout", 30.0)
    return httpx.AsyncClient(verify=SSL_VERIFY, follow_redirects=True,
                             event_hooks={"response": [_block_private_redirect]}, **kw)


# ---- Templates ----

@router.get("/templates")
async def list_templates():
    return {
        "persona": pipeline.engine.persona["display_name"],
        "templates": [t.model_dump() for t in pipeline.engine.list_templates()],
    }


@router.get("/persona")
async def get_persona():
    return pipeline.engine.persona


# ---- v1.6: Layout Templates (node system) ----
# NOTE: namespaced under /layout-templates to avoid colliding with the
# existing /templates endpoint (Seedance prompt templates).

@router.get("/layout-templates")
async def list_layout_templates():
    """List all layout templates (built-in + user-created)."""
    return {"templates": template_engine.list_templates()}


@router.get("/layout-templates/{template_id}")
async def get_layout_template(template_id: str):
    try:
        return template_engine.get_template(template_id)
    except FileNotFoundError:
        raise HTTPException(404, f"Template not found: {template_id}")


@router.get("/layout-templates/{template_id}/slots")
async def list_layout_template_slots(template_id: str):
    try:
        return {"slots": template_engine.list_slots(template_id)}
    except FileNotFoundError:
        raise HTTPException(404, f"Template not found: {template_id}")


@router.post("/layout-templates", response_model=TemplateSaveResponse)
async def save_layout_template(request: TemplateSaveRequest):
    try:
        tid = template_engine.save_template(request.template)
    except ValueError as e:
        raise HTTPException(400, str(e))
    return TemplateSaveResponse(template_id=tid, message="Saved")


@router.delete("/layout-templates/{template_id}")
async def delete_layout_template(template_id: str):
    result = template_engine.delete_template(template_id)
    if result == "builtin":
        raise HTTPException(400, "Built-in templates cannot be deleted")
    if result == "missing":
        raise HTTPException(404, f"Template not found: {template_id}")
    return {"deleted": template_id}

@router.post("/layout-templates/{template_id}/render",
             response_model=TemplateRenderResponse)
async def render_layout_template(
    template_id: str,
    request: TemplateRenderRequest,
    background_tasks: BackgroundTasks,
):
    """Render a layout template with filled slots.

    Seedance/HeyGen slots are generated in parallel via the existing pipeline,
    upload/file slots used as-is, text slots drawn directly. All slots resolve,
    then ffmpeg composites them. Poll GET /api/jobs/{job_id} for progress.
    """
    # Inline template (unsaved editor edits) renders as-is; otherwise the
    # saved template must exist.
    if request.template is None:
        try:
            template_engine.get_template(template_id)
        except FileNotFoundError:
            raise HTTPException(404, f"Template not found: {template_id}")
    else:
        try:
            template_engine._validate(request.template)
        except ValueError as e:
            raise HTTPException(400, f"Invalid template: {e}")

    kinds = {sv.source_kind for sv in request.slot_values.values()}
    if not request.preview:  # preview uses source stills, no provider keys needed
        if "seedance" in kinds and not settings.FAL_KEY:
            raise HTTPException(400, "FAL_KEY not configured. Add it to backend/.env")
        if "heygen" in kinds and not settings.has_heygen:
            raise HTTPException(400, "HEYGEN_API_KEY not configured. Add it to backend/.env")

    job_id = str(uuid4())

    if request.source_graph:
        try:
            import json as _json
            gdir = settings.outputs_path / "_graphs"
            gdir.mkdir(parents=True, exist_ok=True)
            (gdir / f"{job_id}.json").write_text(
                _json.dumps(request.source_graph, ensure_ascii=False),
                encoding="utf-8")
        except Exception as e:
            logger.warning(f"source_graph save failed for {job_id}: {e}")

    async def _run():
        try:
            await pipeline.render_template(
                template_id=template_id,
                slot_values=request.slot_values,
                voice_mode=request.voice_mode,
                job_id=job_id,
                template=request.template,
                title=request.title,
                source_graph=request.source_graph,
                preview=request.preview,
                voiceover=request.voiceover,
            )
        except Exception as e:
            logger.exception(f"Template render {job_id} failed: {e}")

    background_tasks.add_task(_run)
    return TemplateRenderResponse(
        template_id=template_id,
        job_id=job_id,
        message=f"Template render queued. Poll GET /api/jobs/{job_id}.",
    )


@router.get("/jobs/{job_id}/graph")
async def get_job_graph(job_id: str):
    """The Studio node graph that produced this render (saved at render time),
    for "Reopen in Studio". 404 if the render had no stored graph (older
    renders, or non-Studio producers)."""
    import json as _json
    safe = Path(job_id).name
    p = settings.outputs_path / "_graphs" / f"{safe}.json"
    if not p.is_file():
        raise HTTPException(404, "No source graph for this render")
    try:
        return _json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        raise HTTPException(500, "Graph unreadable")


async def _resolve_base_to_path(src):
    """Resolve the Animation node's `base` source to an absolute clip path.

    Accepts None, an absolute/relative path string, or a slot-source dict
    ({source_kind:"job", job_id} / {"upload"|"file", upload_filename|file_path|
    filename}). Mirrors the layout-render slot resolver. Returns a Path or None
    (None -> the compositor falls back to a solid canvas)."""
    if not src:
        return None
    from app.services.storage import JobRecord, async_session_factory
    if isinstance(src, str):
        p = Path(src)
        if p.is_absolute() and p.exists():
            return p
        for d in (settings.images_path, settings.outputs_path):
            q = d / src
            if q.exists():
                return q
        return None
    if isinstance(src, dict):
        kind = src.get("source_kind")
        jid = src.get("job_id")
        if jid and kind in (None, "job"):
            async with async_session_factory() as session:
                jr = await session.get(JobRecord, jid)
            fp = jr and (jr.final_video_path or jr.video_path)
            if fp and Path(fp).exists():
                return Path(fp)
            return None
        name = src.get("upload_filename") or src.get("filename")
        if name:
            for d in (settings.images_path, settings.outputs_path):
                q = d / name
                if q.exists():
                    return q
        fpth = src.get("file_path")
        if fpth and Path(fpth).exists():
            return Path(fpth)
    return None


@router.post("/animate")
async def animate(body: dict, background_tasks: BackgroundTasks):
    """Render the Animation node: composite animated elements over a base clip.

    Mirrors the layout-render flow — mint a job id, save the Studio source graph
    (for "Reopen in Studio"), run the per-frame Pillow+ffmpeg compositor in the
    background (off the event loop), and register the result as a finished
    JobRecord so the Job Dock, Library and reopen-in-Studio all work unchanged.
    Poll GET /api/jobs/{job_id}."""
    from datetime import datetime as _dtu
    from app.services.storage import JobRecord, async_session_factory
    from app.services.animation_service import render_animation

    job_id = str(uuid4())
    short = job_id[:8]
    aspect = str(body.get("aspect") or "9:16")
    out_name = f"anim_{short}.mp4"

    # Save the Studio node graph (same path layout-render uses) for reopen.
    if body.get("source_graph"):
        try:
            gdir = settings.outputs_path / "_graphs"
            gdir.mkdir(parents=True, exist_ok=True)
            (gdir / f"{job_id}.json").write_text(
                json.dumps(body["source_graph"], ensure_ascii=False), encoding="utf-8")
        except Exception as e:
            logger.warning(f"source_graph save failed for animate {job_id}: {e}")

    # Pre-register a job so GET /api/jobs shows it immediately.
    async with async_session_factory() as session:
        session.add(JobRecord(
            id=job_id, status=JobStatus.GENERATING_VIDEO.value, progress=10,
            title=(body.get("title") or "Animation"), image_filename=out_name,
            aspect_ratio=aspect, provider="animation", current_step="Animating",
        ))
        await session.commit()

    async def _run():
        try:
            base = await _resolve_base_to_path(body.get("base"))
            payload = {**body, "base": (str(base) if base else None), "aspect": aspect}
            out = await asyncio.to_thread(render_animation, payload, short)
            dur = await asyncio.to_thread(_probe_seconds, str(out)) or float(body.get("duration_s") or 0)
            async with async_session_factory() as session:
                jr = await session.get(JobRecord, job_id)
                jr.status = JobStatus.DONE.value
                jr.progress = 100
                jr.final_video_path = str(out)
                jr.video_path = str(out)
                jr.image_filename = out.name
                jr.duration_s = int(round(dur)) if dur else None
                jr.current_step = "Complete"
                jr.completed_at = _dtu.utcnow()
                await session.commit()
        except Exception as e:
            logger.exception(f"animate job {job_id} failed: {e}")
            async with async_session_factory() as session:
                jr = await session.get(JobRecord, job_id)
                if jr is not None:
                    jr.status = JobStatus.FAILED.value
                    jr.error = str(e)
                    jr.current_step = "Failed"
                    await session.commit()

    background_tasks.add_task(_run)
    return {"job_id": job_id, "status": "queued"}


@router.post("/assets/3d")
async def assets_3d(body: dict, background_tasks: BackgroundTasks):
    """Game Assets 3D: image -> (optional multi-view) -> 3D engine -> mesh + shots.
    Mirrors /api/animate: pre-register an asset3d JobRecord, run in the background,
    record the produced files/shots in cost_meta. Poll GET /api/jobs/{job_id}."""
    from datetime import datetime as _dtu
    import json as _json
    from app.services.asset3d_service import generate_asset3d, ENGINES
    from app.services.storage import JobRecord, async_session_factory

    engine = str(body.get("engine") or "tripo").lower()
    if engine == "meshy":
        # Le pipeline Meshy complet (preview→refine→remesh→rig→animations)
        # vit dans le 3D Studio (/studio3d + proxy /api/meshy/*), pas dans ce
        # flux fal une-passe.
        raise HTTPException(501, "Meshy runs in the 3D Studio (/studio3d). "
                                 "Use a fal engine here.")
    if engine not in ENGINES:
        raise HTTPException(400, f"Unknown engine: {engine}")
    if not settings.FAL_KEY:
        raise HTTPException(400, "FAL_KEY not configured. Add it in Settings.")
    # fail fast on bad input instead of accepting a job that dies in background
    fn = Path(str(body.get("image_filename") or "")).name
    if not fn or not (settings.images_path / fn).is_file():
        raise HTTPException(400, f"image_filename not found in Library: {body.get('image_filename')!r}")
    try:
        int(body.get("views") or 3)
    except (TypeError, ValueError):
        raise HTTPException(400, "views must be an integer (1-4)")
    fmts = body.get("formats") or ["glb"]
    if not isinstance(fmts, list) or not all(isinstance(f, str) for f in fmts):
        raise HTTPException(400, "formats must be a list of strings")

    job_id = str(uuid4())
    short = job_id[:8]
    async with async_session_factory() as s:
        s.add(JobRecord(
            id=job_id, status=JobStatus.GENERATING_VIDEO.value, progress=10,
            title=(body.get("title") or f"3D · {engine}"), image_filename=f"asset3d_{short}",
            provider="asset3d", current_step="Generating 3D"))
        await s.commit()

    async def on_step(label, pct):
        async with async_session_factory() as s2:
            jr2 = await s2.get(JobRecord, job_id)
            if jr2 is not None:
                jr2.current_step = label
                jr2.progress = int(pct)
                await s2.commit()

    async def _run():
        try:
            r = await generate_asset3d(body, short, on_step=on_step)
            async with async_session_factory() as s:
                jr = await s.get(JobRecord, job_id)
                if jr is not None:
                    jr.status = JobStatus.DONE.value
                    jr.progress = 100
                    jr.final_video_path = r.get("glb")
                    if r.get("preview"):
                        jr.image_filename = "preview.png"
                    jr.current_step = "Complete"
                    jr.completed_at = _dtu.utcnow()
                    jr.cost_meta = _json.dumps({"engine": r["engine"], "files": r["files"],
                                                "shots": r["shots"], "job": short,
                                                "skipped_formats": r.get("skipped_formats") or []})
                    await s.commit()
        except Exception as e:
            logger.exception(f"asset3d {job_id} failed: {e}")
            async with async_session_factory() as s:
                jr = await s.get(JobRecord, job_id)
                if jr is not None:
                    jr.status = JobStatus.FAILED.value
                    jr.error = str(e)
                    jr.current_step = "Failed"
                    await s.commit()

    background_tasks.add_task(_run)
    return {"job_id": job_id, "status": "queued"}


# model.v2.glb, model.v3.glb… : les VERSIONS écrites par le raffinement de la
# phase D. Le manifeste liste des FORMATS d'export — sans ce filtre, il
# annonçait « v2.glb » comme un format téléchargeable inexistant.
_MODEL_VERSION_RE = re.compile(r"^model\.v\d+\.glb$", re.I)


@router.get("/assets/3d/{job}/manifest")
async def get_asset3d_manifest(job: str):
    """List what a 3D asset job produced (read from disk — ground truth):
    mesh formats, shot indices, and whether a preview image exists.
    Declared before the /{fmt} route so it isn't captured as fmt='manifest'."""
    d = settings.outputs_path / "assets3d" / Path(job).name
    if not d.is_dir():
        raise HTTPException(404, "Not found")
    formats, shots = [], []
    for f in d.iterdir():
        n = f.name
        if n == "model.opt.glb":
            continue                    # le GLB optimisé a sa propre UI (10a)
        if _MODEL_VERSION_RE.match(n):
            continue                    # model.v2.glb = une VERSION (phase D),
                                        # pas un format d'export
        if n.startswith("model.") and f.is_file():
            formats.append(n.split(".", 1)[1].lower())
        elif n.startswith("shot_") and n.endswith(".png"):
            try:
                shots.append(int(n[5:-4]))
            except ValueError:
                pass
    return {"formats": sorted(set(formats)), "shots": sorted(set(shots)),
            "has_preview": (d / "preview.png").is_file()}


@router.get("/assets/3d/{job}/preview")
async def get_asset3d_preview(job: str):
    """Stream the engine-generated preview render (if the engine returned one).
    Declared before /{fmt} so it isn't captured as fmt='preview'."""
    p = settings.outputs_path / "assets3d" / Path(job).name / "preview.png"
    if not p.is_file():
        raise HTTPException(404, "Not found")
    return FileResponse(p)


# ---- Game Assets 3D — Optimize (chantier 10a) ----
# Les routes /optimize et /opt-glb sont déclarées AVANT /{fmt} (même règle
# que /preview) pour ne pas être capturées comme fmt="optimize"/"opt-glb".

@router.post("/assets/3d/{job}/optimize")
async def optimize_asset3d(job: str, body: dict = None):
    """Simplify model.glb to a triangle budget (gltfpack, local & free).
    Body: {preset: micro|small|prop|detailed|game|balanced|high|ultra}
    OR {target_tris: int}. Returns before/after stats (persisted in
    optimize.json next to the model)."""
    from app.services import mesh_optimize as MO
    body = body or {}
    try:
        info = await asyncio.get_running_loop().run_in_executor(
            None, lambda: MO.optimize_glb(
                job, body.get("target_tris"), body.get("preset")))
    except FileNotFoundError as e:
        raise HTTPException(404, str(e))
    except ValueError as e:
        raise HTTPException(400, str(e))
    except RuntimeError as e:
        raise HTTPException(400, str(e))
    return info


@router.get("/assets/3d/{job}/optimize")
async def get_asset3d_optimize(job: str):
    """Stats of the last optimize run for this job (optimize.json)."""
    p = settings.outputs_path / "assets3d" / Path(job).name / "optimize.json"
    if not p.is_file():
        raise HTTPException(404, "Not optimized yet")
    import json as _json
    return _json.loads(p.read_text(encoding="utf-8"))


@router.get("/assets/3d/{job}/opt-glb")
async def download_asset3d_optimized(job: str):
    p = settings.outputs_path / "assets3d" / Path(job).name / "model.opt.glb"
    if not p.is_file():
        raise HTTPException(404, "Not found")
    return FileResponse(p, media_type="model/gltf-binary",
                        filename=f"asset3d_{Path(job).name}_optimized.glb")


# ---- Phase D (spec Magnific §9 + §13) : fiche, porte brouillon→final,
# contrôle qualité et comparaison. MÊME RÈGLE que /preview et /optimize :
# ces sous-routes d'UN segment sont déclarées AVANT /{fmt}, sinon elles
# seraient capturées comme fmt="report", fmt="approve", etc.

@router.get("/assets3d/engines")
async def list_asset3d_engines():
    """Registre des moteurs 3D avec leurs CAPABILITY FLAGS (§8 transposé à la
    3D) + la matrice « besoin d'asset → moteur, et pourquoi ».

    Miroir de GET /api/video-models : chaque entrée porte `available`
    (FAL_KEY présente) et son coût unitaire, pour que l'UI montre-mais-grise
    au lieu de laisser choisir un moteur inutilisable."""
    from app.services.asset3d_service import ENGINES, BESOINS_3D
    from app.services import pricing as _pricing
    dispo = bool(settings.FAL_KEY)
    out = []
    for eid, e in ENGINES.items():
        devis = _pricing.estimate({"kind": "asset3d", "engine": eid,
                                   "textures": True, "multiview": False})
        brouillon = _pricing.estimate({"kind": "asset3d", "engine": eid,
                                       "textures": False, "multiview": False})
        # AUCUNE URL de fournisseur vers le client : le filtre porte sur le
        # PRÉFIXE, pas sur la clé exacte — `endpoint_multiview` (H3.1) était
        # passé à travers un filtre qui ne connaissait que `endpoint`.
        out.append({**{k: v for k, v in e.items()
                       if not k.startswith("endpoint")},
                    "id": eid, "available": dispo,
                    "usd_texture": devis["total_usd"],
                    "usd_brouillon": brouillon["total_usd"]})
    out.sort(key=lambda m: m["id"])
    return {"engines": out, "default": "tripo",
            "besoins": [{"id": k, **v} for k, v in BESOINS_3D.items()]}


@router.get("/assets/3d/{job}/report")
async def get_asset3d_report(job: str):
    """Registre versionné des fiches de maillage (§9.2 étape 6 : checksum,
    faces, taille, textures, version). 404 tant qu'aucune fiche n'existe —
    les jobs antérieurs peuvent en obtenir une par POST .../report."""
    from app.services import mesh_report
    try:
        return mesh_report.read_registry(job)
    except FileNotFoundError as e:
        raise HTTPException(404, str(e))


@router.post("/assets/3d/{job}/report")
async def build_asset3d_report(job: str, body: dict = None):
    """(Re)calcule la fiche d'un maillage du job. Body: {file?, version?,
    silhouettes?}. Tout est local et gratuit."""
    from app.services import mesh_report
    body = body or {}
    f = Path(str(body.get("file") or "model.glb")).name
    version = body.get("version")
    if version is not None:
        try:
            version = int(version)
        except (TypeError, ValueError):
            raise HTTPException(400, "version doit être un entier.")
    try:
        return await asyncio.get_running_loop().run_in_executor(
            None, lambda: mesh_report.write_report(
                job, f, version=version,
                avec_silhouettes=bool(body.get("silhouettes", True))))
    except FileNotFoundError as e:
        raise HTTPException(404, str(e))
    except ValueError as e:
        raise HTTPException(400, str(e))


@router.get("/assets/3d/{job}/approve")
async def get_asset3d_approval(job: str):
    """État de la porte humaine du brouillon."""
    from app.services.asset3d_service import approval
    return approval(job)


@router.post("/assets/3d/{job}/approve")
async def approve_asset3d(job: str, body: dict = None):
    """Porte humaine de l'étape 5 : la géométrie du brouillon est validée
    (ou refusée avec un motif). Body: {approved?: bool, note?: str}.
    Rien de payant ne franchit cette porte tout seul."""
    from app.services.asset3d_service import approve
    body = body or {}
    try:
        return approve(job, bool(body.get("approved", True)),
                       str(body.get("note") or ""))
    except FileNotFoundError as e:
        raise HTTPException(404, str(e))


@router.post("/assets/3d/{job}/refine")
async def refine_asset3d_route(job: str, background_tasks: BackgroundTasks,
                               body: dict = None):
    """Rejoue le même moteur sur les mêmes vues en texture haute qualité, et
    écrit model.v{n}.glb — jamais un écrasement (§2.1). Refusé tant que le
    brouillon n'est pas approuvé. Body: {quality?: "hd"}.
    Rend un job_id : POLLER GET /api/jobs/{job_id}."""
    from datetime import datetime as _dtu
    import json as _json
    from app.services import asset3d_service as A3
    from app.services.storage import JobRecord, async_session_factory

    body = body or {}
    quality = str(body.get("quality") or "hd")
    if not settings.FAL_KEY:
        raise HTTPException(400, "FAL_KEY not configured. Add it in Settings.")

    # TOUT ce qui peut refuser refuse MAINTENANT — y compris les deux refus
    # qui ne vivaient que dans le service, et qui rendaient un 200 « queued »
    # suivi d'un job FAILED sans qu'aucun travail ait été tenté.
    try:
        man = A3.read_manifest(job)
        caps = A3.engine_caps(str(man.get("engine") or ""))
        cible = A3.texture_mode(caps["id"], True, quality)
    except FileNotFoundError as e:
        raise HTTPException(404, str(e))
    except ValueError as e:
        raise HTTPException(400, str(e))
    if not A3.approval(job).get("approved"):
        raise HTTPException(
            409, "Brouillon non approuvé : valide la géométrie "
                 "(POST .../approve) avant la passe texturée.")
    if not caps["detailed"]:
        raise HTTPException(
            400, f"{caps['label']} n'a pas de palier haute qualité dans "
                 f"cet adaptateur — régénère avec un moteur qui l'a.")
    if cible == man.get("texture_mode"):
        raise HTTPException(
            409, f"Le maillage est déjà en texture « {cible} » — repayer "
                 "la même passe ne changerait rien.")
    _d3 = settings.outputs_path / "assets3d" / Path(job).name
    if not [s for s in (man.get("shots") or []) if (_d3 / s).is_file()]:
        raise HTTPException(
            400, "Aucune vue conservée sur le disque : le moteur ne peut pas "
                 "être rejoué à l'identique. Relance une génération complète.")

    # Verrou : sans lui, deux clics rapides lancent DEUX passes HD facturées.
    # Il couvre AUSSI le texturage Meshy, qui écrit dans le même registre de
    # versions — deux finitions concurrentes se disputeraient next_version().
    from app.services.storage import JobRecord as _JR, async_session_factory as _sf
    async with _sf() as s:
        res = await s.execute(
            _select(_JR).where(_JR.provider == "asset3d",
                               _JR.image_filename == f"asset3d_{Path(job).name}",
                               _JR.status.notin_(["done", "failed"])))
        if res.scalars().first() is not None:
            raise HTTPException(
                409, "Une passe de finition de ce maillage est déjà en cours — "
                     "attends qu'elle finisse (file des rendus).")

    job_id = str(uuid4())
    async with async_session_factory() as s:
        s.add(JobRecord(
            id=job_id, status=JobStatus.GENERATING_VIDEO.value, progress=10,
            title=f"3D · {caps['label']} · texture {quality}",
            image_filename=f"asset3d_{Path(job).name}",
            provider="asset3d", current_step="Refine 3D"))
        await s.commit()

    async def on_step(label, pct):
        async with async_session_factory() as s2:
            jr2 = await s2.get(JobRecord, job_id)
            if jr2 is not None:
                jr2.current_step, jr2.progress = label, int(pct)
                await s2.commit()

    async def _run():
        try:
            r = await A3.refine_asset3d(job, quality=quality, on_step=on_step)
            async with async_session_factory() as s:
                jr = await s.get(JobRecord, job_id)
                if jr is not None:
                    jr.status, jr.progress = JobStatus.DONE.value, 100
                    jr.final_video_path = str(
                        settings.outputs_path / "assets3d" / Path(job).name / r["file"])
                    jr.current_step = "Complete"
                    jr.completed_at = _dtu.utcnow()
                    jr.cost_meta = _json.dumps({"engine": r["engine"],
                                                "job": Path(job).name,
                                                "refine": True,
                                                "version": r["version"],
                                                "texture_mode": r["texture_mode"]})
                    await s.commit()
        except Exception as e:
            logger.exception(f"asset3d refine {job_id} failed: {e}")
            async with async_session_factory() as s:
                jr = await s.get(JobRecord, job_id)
                if jr is not None:
                    jr.status, jr.error = JobStatus.FAILED.value, str(e)
                    jr.current_step = "Failed"
                    await s.commit()

    background_tasks.add_task(_run)
    return {"job_id": job_id, "status": "queued", "source_job": Path(job).name}


@router.post("/assets/3d/{job}/texturer")
async def texturer_asset3d_route(job: str, background_tasks: BackgroundTasks,
                                 body: dict = None):
    """Texture chez MESHY un maillage déjà généré — la seconde moitié de la
    chaîne Tripo → Meshy (Tripo reconstruit le volume depuis 4 vues, Meshy
    l'habille).

    Body: {resolution?="2k"|"4k"|"8k", pbr?=true, ai_model?="meshy-7",
    style_prompt?, garder_uv?=false}. Les MÊMES vues qui ont servi à la
    géométrie servent de référence de style ; `style_prompt` n'est requis que
    si le job n'a plus aucune vue sur le disque.

    MÊME PORTE que /refine : refusé tant que la géométrie n'est pas approuvée.
    Écrit model.v{n}.glb — jamais un écrasement. Rend un job_id à poller."""
    from datetime import datetime as _dtu
    import json as _json
    from app.services import asset3d_service as A3
    from app.services import meshy_service as MS
    from app.services.storage import JobRecord, async_session_factory

    body = body or {}
    resolution = str(body.get("resolution") or "2k")
    if resolution not in ("2k", "4k", "8k"):
        raise HTTPException(400, "resolution doit être 2k, 4k ou 8k.")
    if not MS.mock_enabled() and not settings.MESHY_API_KEY.strip():
        raise HTTPException(
            400, "MESHY_API_KEY absente — ajoute-la dans les Réglages : le "
                 "texturage passe par ton compte Meshy, pas par fal.")

    # tout ce qui peut refuser refuse AVANT d'ouvrir un job payant
    try:
        man = A3.read_manifest(job)
    except FileNotFoundError as e:
        raise HTTPException(404, str(e))
    if not A3.approval(job).get("approved"):
        raise HTTPException(
            409, "Géométrie non approuvée : valide le volume "
                 "(POST .../approve) avant de payer un texturage.")
    if str(man.get("texturier") or "") == "meshy" \
            and str(man.get("texture_mode") or "") == f"meshy:{resolution}":
        raise HTTPException(
            409, f"Ce maillage est déjà texturé par Meshy en {resolution} — "
                 "repayer la même passe ne changerait rien.")
    _d3 = settings.outputs_path / "assets3d" / Path(job).name
    a_des_vues = any((_d3 / s).is_file() for s in (man.get("shots") or []))
    if not a_des_vues and not str(body.get("style_prompt") or "").strip():
        raise HTTPException(
            400, "Aucune vue conservée sur le disque : donne un `style_prompt` "
                 "ou relance une génération complète.")
    async with async_session_factory() as s:
        res = await s.execute(
            _select(JobRecord).where(
                JobRecord.provider == "asset3d",
                JobRecord.image_filename == f"asset3d_{Path(job).name}",
                JobRecord.status.notin_(["done", "failed"])))
        if res.scalars().first() is not None:
            raise HTTPException(
                409, "Une passe de finition de ce maillage est déjà en cours — "
                     "attends qu'elle finisse (file des rendus).")

    from app.services import pricing as _pricing
    devis = _pricing.estimate({"kind": "asset3d_texture",
                               "texture_resolution": resolution,
                               "pbr": bool(body.get("pbr", True))})
    job_id = str(uuid4())
    async with async_session_factory() as s:
        s.add(JobRecord(
            id=job_id, status=JobStatus.GENERATING_VIDEO.value, progress=10,
            title=f"3D · texturage Meshy {resolution}",
            image_filename=f"asset3d_{Path(job).name}",
            provider="asset3d", current_step="Texturage Meshy"))
        await s.commit()

    async def on_step(label, pct):
        async with async_session_factory() as s2:
            jr2 = await s2.get(JobRecord, job_id)
            if jr2 is not None:
                jr2.current_step, jr2.progress = label, int(pct)
                await s2.commit()

    async def _run():
        try:
            r = await A3.texturer_asset3d(
                job, resolution=resolution, pbr=bool(body.get("pbr", True)),
                ai_model=str(body.get("ai_model") or "meshy-7"),
                style_prompt=body.get("style_prompt"),
                garder_uv=bool(body.get("garder_uv")), on_step=on_step)
            async with async_session_factory() as s:
                jr = await s.get(JobRecord, job_id)
                if jr is not None:
                    jr.status, jr.progress = JobStatus.DONE.value, 100
                    jr.final_video_path = str(
                        settings.outputs_path / "assets3d"
                        / Path(job).name / r["file"])
                    jr.current_step = "Complete"
                    jr.completed_at = _dtu.utcnow()
                    jr.cost_meta = _json.dumps(
                        {"job": Path(job).name, "texturier": "meshy",
                         "meshy_task": r["meshy_task"], "version": r["version"],
                         "texture_mode": r["texture_mode"]})
                    await s.commit()
        except Exception as e:
            logger.exception(f"asset3d texturer {job_id} failed: {e}")
            async with async_session_factory() as s:
                jr = await s.get(JobRecord, job_id)
                if jr is not None:
                    jr.status, jr.error = JobStatus.FAILED.value, str(e)
                    jr.current_step = "Failed"
                    await s.commit()

    background_tasks.add_task(_run)
    return {"job_id": job_id, "status": "queued", "source_job": Path(job).name,
            "texturier": "meshy", "resolution": resolution,
            "credits": devis.get("credits"), "usd_estime": devis["total_usd"]}


@router.post("/assets/3d/{job}/qc")
async def qc_asset3d(job: str, body: dict = None):
    """Scores 0-100 du maillage contre sa référence maître (§9.2 étapes 3 et
    7) + le verdict de compatibilité runtime (§13 phase D).

    Body: {ref_image?, version?, seuils?, vision?}. La mesure principale est
    LOCALE et GRATUITE (IoU de silhouettes) ; `vision: true` ajoute un score
    d'identité par LLM, best-effort (absent sans clé, jamais bloquant)."""
    from app.services import asset3d_qc as QC
    body = body or {}
    # Entrées validées ICI : sans ça, {"seuils": [70]} ou {"version": []}
    # remontaient en TypeError non rattrapé, donc en 500 pour une faute
    # purement cliente.
    seuils = body.get("seuils")
    if seuils is not None:
        if not isinstance(seuils, dict) or not all(
                isinstance(v, (int, float)) and not isinstance(v, bool)
                for v in seuils.values()):
            raise HTTPException(
                400, "seuils doit être un objet {axe: nombre} — ex. "
                     '{"silhouette": 70}.')
    version = body.get("version")
    if version is not None:
        try:
            version = int(version)
        except (TypeError, ValueError):
            raise HTTPException(400, "version doit être un entier.")
    loop = asyncio.get_running_loop()
    try:
        res = await loop.run_in_executor(None, lambda: QC.controler(
            job, ref_image=body.get("ref_image"), version=version,
            seuils=seuils))
    except FileNotFoundError as e:
        raise HTTPException(404, str(e))
    except ValueError as e:
        raise HTTPException(400, str(e))

    if body.get("vision"):
        d = settings.outputs_path / "assets3d" / Path(job).name
        rendu = next((d / n for n in ("preview.png", "shot_0.png")
                      if (d / n).is_file()), None)
        ref_n = Path(str(body.get("ref_image") or "")).name
        ref = (settings.images_path / ref_n) if ref_n else (d / "shot_0.png")
        if rendu and ref.is_file() and rendu != ref:
            res["identite"] = await loop.run_in_executor(
                None, lambda: QC.identite(rendu, ref))
        else:
            res["identite"] = None
            res["identite_note"] = ("aucun rendu du moteur distinct de la "
                                    "référence — passe vision sans objet")
    return res


@router.get("/assets/3d/{job}/compare")
async def compare_asset3d(job: str, other: str, va: int = None, vb: int = None):
    """Compare deux maillages (§13 phase D : « image unique vs quatre vues »,
    « comparer les moteurs selon l'asset »). Rend les écarts mesurés —
    triangles, poids, textures, arêtes de bord, IoU des silhouettes — et
    jamais un gagnant : c'est l'humain qui tranche, chiffres en main."""
    from app.services import asset3d_qc as QC
    try:
        return await asyncio.get_running_loop().run_in_executor(
            None, lambda: QC.comparer(Path(job).name, Path(other).name,
                                      version_a=va, version_b=vb))
    except FileNotFoundError as e:
        raise HTTPException(404, str(e))


@router.get("/assets/3d/{job}/version/{v}")
async def get_asset3d_version(job: str, v: int):
    """Télécharge une version précise du maillage (v1 = model.glb)."""
    d = settings.outputs_path / "assets3d" / Path(job).name
    p = d / ("model.glb" if int(v) == 1 else f"model.v{int(v)}.glb")
    if not p.is_file():
        raise HTTPException(404, "Not found")
    return FileResponse(p, media_type="model/gltf-binary",
                        filename=f"asset3d_{Path(job).name}_v{int(v)}.glb")


@router.get("/assets/3d/{job}/silhouette/{vue}")
async def get_asset3d_silhouette(job: str, vue: str, v: int = 1):
    """Masque projeté (face|profil|dessus) d'une version — l'aperçu honnête
    de la géométrie, sans moteur de rendu."""
    d = settings.outputs_path / "assets3d" / Path(job).name
    p = d / f"sil_v{int(v)}" / f"silhouette_{Path(vue).name}.png"
    if not p.is_file():
        raise HTTPException(404, "Not found")
    return FileResponse(p, media_type="image/png")


def _etabli_vignette_path(job: str, v: int) -> Path:
    """Le chemin de la vignette d'une version — UN SEUL endroit le compose.

    Écrite par `POST /api/etabli/vignette` (la capture du canevas de l'Établi),
    servie plus bas, préférée par `_etabli_vignette`. Trois endroits qui
    doivent nommer le même fichier : le nom se décide ici, pas trois fois.

    IL COMPOSE, IL NE JUGE PAS. `Path(job).name` réduit au nom, ce qui suffit
    pour un `job` VENU DU DISQUE — `mesh_sources.lister()` ne rend que des noms
    de dossiers réels, et c'est le seul appelant qui ne passe pas par
    `_etabli_vignette_cible`. Toute entrée venue du RÉSEAU passe par celle-ci,
    et pour une raison mesurée que la fonction explique.
    """
    return (settings.outputs_path / "assets3d" / Path(str(job)).name
            / f"vignette_v{int(v)}.png")


def _etabli_vignette_cible(job: str, v: int) -> Path:
    """Le même chemin, pour un `job` VENU DU RÉSEAU — ou un refus parlant.

    DEUX GARDES DE NATURES DIFFÉRENTES, ET C'EST TOUT LE SUJET. Ce fichier a
    déjà payé pour apprendre que deux gardes IDENTIQUES ne valent pas mieux
    qu'une : deux `Path(...).name` posés en parallèle se couvraient si bien
    qu'en retirer un laissait le banc de traversée ENTIÈREMENT VERT, et plus
    rien ne disait laquelle tenait. Deux gardes de natures différentes, elles,
    se prouvent SÉPARÉMENT — un banc chacune, une mutation chacune.

    GARDE 1, LE NOM SE REFUSE — il ne s'aplatit pas. MESURÉ : `Path(...).name`
    normalise le point SIMPLE, JAMAIS le point-point.

        Path("../../evade").name -> 'evade'   (aplati, sans danger)
        Path("..").name          -> '..'      (PASSE TEL QUEL)
        Path("a/..").name        -> '..'      (idem)
        Path(".").name           -> ''

    Un `job=".."` composait donc `outputs/assets3d/../vignette_v1.png` et la
    route répondait 200 en écrivant UN CRAN AU-DESSUS du dossier des jobs.
    Portée exacte, pour ne pas dramatiser : l'évasion est d'un SEUL cran
    (`Path("../..").name` vaut encore `..`), et la seule chose qui la bloquait
    était la garde d'existence du maillage — il fallait un `outputs/model.glb`,
    qu'aucun chemin de code ne crée. Non exploitable en l'état, donc ; mais
    l'invariant qui nous sauvait n'était pas celui que le commentaire
    désignait, et une garde documentée comme tenant alors qu'elle ne tient pas
    est exactement ce qui casse au refactor suivant.

    GARDE 2, LE CONFINEMENT — et elle ne regarde pas le nom, elle RÉSOUT le
    chemin. Elle voit donc ce qu'aucune lecture de nom ne peut voir : une
    jonction ou un lien posés dans le dossier des jobs. `vign_jonction` est un
    nom parfaitement honnête, sans `..` ni séparateur — la garde 1 le laisse
    passer, à juste titre — et le chemin sort pourtant du dossier des jobs.
    Les deux COMPOSENT au lieu de se doubler : retirer l'une rougit un banc,
    retirer l'autre en rougit un DIFFÉRENT.
    """
    return _etabli_cible_sous_jobs(job, lambda j: _etabli_vignette_path(j, v),
                                   "vignette")


def _etabli_cible_sous_jobs(job: str, composer, quoi: str) -> Path:
    """Les deux gardes de `_etabli_vignette_cible`, en UN site — le plan de
    plaque franchit la même porte, et deux copies des mêmes gardes
    divergeraient à la première retouche. `composer(job)` compose le chemin
    à partir du `job` BRUT ; `quoi` préfixe les refus, pour qu'on sache
    laquelle des deux routes a parlé.
    """
    nom = Path(str(job)).name
    if nom in ("", ".", ".."):
        raise HTTPException(400, f"{quoi} : nom de job invalide — "
                                 f"« {job} » ne désigne aucun dossier de job")
    # `job` BRUT, et non `nom` : l'aplatissement qui FAÇONNE le chemin vit
    # dans le composeur, et lui seul. Repasser `nom` ici mettrait deux
    # `Path(...).name` IDENTIQUES en série — l'un masquerait l'autre à la
    # mutation, et ce fichier a déjà payé une fois pour cette leçon-là.
    # Ci-dessus, `.name` ne sert qu'à JUGER ; ici, il ne sert plus du tout.
    p = composer(job)
    racine = (settings.outputs_path / "assets3d").resolve()
    if racine not in p.resolve().parents:
        raise HTTPException(400, f"{quoi} : chemin hors du dossier des jobs")
    return p


# ── le plan de plaque : la DISPOSITION, distincte du maillage ────────────────
# Ce que l'utilisateur compose « sur la plaque » de l'Établi — où chaque pièce
# est posée et de combien elle est tournée — vit dans `plaque.v<N>.json`, à
# côté de `model.v<N>.glb`, et JAMAIS dans le GLB : c'est la séparation que le
# 3MF fait entre maillage et disposition. Ranger des pièces sur la plaque
# n'écrit donc pas de version ; transformer en mode Assemblé, si. Le format,
# tel que le navigateur le compose et que ce fichier l'écrit :
#
#   { "format": "plaque/1", "job": "<nom>", "version": N,
#     "axe": "x" | "y" | "z",      l'axe d'empilement (normale du plateau)
#     "pas": nombre > 0,           le pas du plateau, en unités du modèle
#     "unites": "modele", "repere": "monde",
#     "pieces": [ { "index": i,    index de nœud glTF, entier ≥ 0, unique
#                   "dx": nombre,  déplacement du centre depuis la pose
#                   "dy": nombre,  assemblée, sur les deux axes du plan (u, v :
#                                  l'ordre x, y, z privé de `axe`)
#                   "rot": degrés } ] }   rotation autour de +axe, sens direct
#
# CE QU'IL FAUT SAVOIR POUR L'APPLIQUER, et que le format seul ne dit pas —
# écrit ici parce que c'est ici que l'extraction viendra le lire :
#   - dx/dy sont en COORDONNÉES MONDE de la scène glTF (le repère de la scène
#     chargée), PAS dans le repère local du nœud. Le nœud vit sous
#     l'enveloppe `etabli_correction` que `mesh_edit.reparer` a tournée
#     (_ROT["Z"]) et mise à l'échelle : appliquer dx/dy à la translation du
#     nœud sans les convertir déplacerait toutes les pièces d'une réparation
#     en Z. Il faut A⁻¹·d, A étant le bloc linéaire de la matrice monde du
#     PARENT du nœud — ce que frontend/lib3d/plaque.js fait dans
#     versLocalLineaire.
#   - `rot` tourne autour de +axe (règle de la main droite, dans le MONDE),
#     autour du CENTRE de la boîte englobante monde de la pièce dans sa pose
#     ASSEMBLÉE — ni son pivot glTF, ni l'origine de son repère local. La
#     matrice à composer sur le nœud est M⁻¹·T(c)·R·T(−c)·M, M étant la
#     matrice monde du parent (poserPivot, même fichier) ; sous un parent à
#     échelle non uniforme, ce n'est pas une TRS décomposable.
#   - la troisième composante — le posé AU CONTACT du plateau, −min de la
#     boîte sur `axe` — n'est PAS stockée : elle se déduit de la géométrie,
#     et une rotation autour de la normale ne la change pas.
#   - le plan porte le NUMÉRO DE VERSION : après une extraction ou un couteau,
#     la version N+1 n'en a AUCUN, et c'est voulu — ses index de nœud ne sont
#     plus ceux du plan, un plan hérité poserait les mauvaises pièces.
# La doctrine du module navigateur (frontend/lib3d/plaque.js, en tête) est le
# site canonique du format ; ceci en est le miroir côté écriture.

_ETABLI_AXES = ("x", "y", "z")


def _etabli_plaque_path(job: str, v: int) -> Path:
    """Le chemin du plan d'une version — composé ICI et nulle part ailleurs,
    comme `_etabli_vignette_path` pour la vignette."""
    return (settings.outputs_path / "assets3d" / Path(str(job)).name
            / f"plaque.v{int(v)}.json")


def _etabli_plaque_cible(job: str, v: int) -> Path:
    """Le même chemin pour un `job` venu du RÉSEAU : les deux gardes."""
    return _etabli_cible_sous_jobs(job, lambda j: _etabli_plaque_path(j, v),
                                   "plan de plaque")


def _etabli_nombre(v) -> bool:
    """Un nombre JSON fini — et PAS un booléen, que Python compte parmi les
    entiers : `True` passerait pour 1."""
    return (isinstance(v, (int, float)) and not isinstance(v, bool)
            and v == v and v not in (float("inf"), float("-inf")))


def _etabli_entier(v) -> bool:
    return isinstance(v, int) and not isinstance(v, bool)


@router.get("/etabli/plaque")
async def etabli_plaque_lire(job: str, version: int):
    """Relit le plan de plaque d'une version. 404 FRANC quand il n'y en a
    pas : c'est le cas ordinaire — le plan naît à la première retouche, et
    une version jamais rangée n'en a pas. Le navigateur étale alors par
    défaut. Un fichier présent mais illisible est un 500 qui le dit : pris
    pour un plan absent, il serait écrasé à la retouche suivante."""
    if version < 1:
        raise HTTPException(400, f"plan de plaque : version {version} — les "
                                 "versions sont numérotées à partir de 1")
    p = _etabli_plaque_cible(job, version)
    if not p.is_file():
        raise HTTPException(404, "aucun plan de plaque pour cette version")
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except (ValueError, OSError) as e:
        raise HTTPException(500, f"plan de plaque illisible ({p.name}) : {e}")


@router.post("/etabli/plaque")
async def etabli_plaque_ecrire(body: dict):
    """Écrit le plan de plaque d'une version — le navigateur compose, PYTHON
    ÉCRIT, comme pour la vignette. Ce n'est pas une version : aucun GLB n'est
    touché, aucune fiche n'est ajoutée au registre.

    LES GARDES, dans l'ordre où elles mordent. `version` doit être un entier
    ≥ 1 (un `2.0` de JSON est un flottant, refusé) ; `job` franchit les deux
    gardes de chemin de la vignette (nom dégénéré refusé, chemin résolu
    confiné) ; la version doit EXISTER sur le disque — un plan sans maillage
    ne dit rien, et c'est ce qui empêche de fabriquer un dossier de job à
    volonté ; `axe` ∈ x|y|z ; `pas` un nombre > 0 ; chaque pièce un objet à
    `index` entier ≥ 0, unique dans le plan, et `dx`/`dy`/`rot` des nombres
    finis (absents = 0). Rien n'est écrit tant qu'une garde n'a pas fini de
    mordre.

    ÉCRITURE ATOMIQUE (`.tmp` puis `Path.replace`), pour la raison de la
    vignette : un plan tronqué serait relu comme un JSON invalide, donc un
    500 à l'entrée de la plaque suivante.
    """
    version = body.get("version")
    if not _etabli_entier(version) or version < 1:
        raise HTTPException(400, f"plan de plaque : version « {version} » — "
                                 "un entier à partir de 1")
    job = body.get("job")
    # Une chaîne, sinon `Path(str(None)).name` composerait un dossier « None »
    # et répondrait 404 « introuvable » à ce qui est une requête malformée.
    if not isinstance(job, str):
        raise HTTPException(400, f"plan de plaque : job « {job} » — le nom du "
                                 "dossier de job est attendu")
    p = _etabli_plaque_cible(job, version)
    d = p.parent
    glb = d / ("model.glb" if version <= 1 else f"model.v{version}.glb")
    if not glb.is_file():
        raise HTTPException(404, f"plan de plaque : {d.name}/{glb.name} "
                                 "introuvable — un plan sans maillage ne dit "
                                 "rien")
    axe = body.get("axe")
    if axe not in _ETABLI_AXES:
        raise HTTPException(400, f"plan de plaque : axe « {axe} » — x, y ou z")
    pas = body.get("pas")
    if not _etabli_nombre(pas) or pas <= 0:
        raise HTTPException(400, f"plan de plaque : pas « {pas} » — un nombre "
                                 "> 0, en unités du modèle")
    pieces = body.get("pieces")
    if not isinstance(pieces, list):
        raise HTTPException(400, "plan de plaque : `pieces` doit être une liste")
    propres, vus = [], set()
    for rang, pc in enumerate(pieces):
        if not isinstance(pc, dict):
            raise HTTPException(400, f"plan de plaque : pièce {rang} — un objet "
                                     "{index, dx, dy, rot} est attendu")
        index = pc.get("index")
        if not _etabli_entier(index) or index < 0:
            raise HTTPException(400, f"plan de plaque : pièce {rang} — index "
                                     f"« {index} », un entier ≥ 0 (index de "
                                     "nœud glTF)")
        if index in vus:
            raise HTTPException(400, f"plan de plaque : l'index {index} est "
                                     "posé deux fois")
        vus.add(index)
        propre = {"index": index}
        for cle in ("dx", "dy", "rot"):
            val = pc.get(cle, 0)
            if not _etabli_nombre(val):
                raise HTTPException(400, f"plan de plaque : pièce {index} — "
                                         f"{cle} « {val} », un nombre fini")
            propre[cle] = float(val)
        propres.append(propre)
    doc = {"format": "plaque/1", "job": d.name, "version": version,
           "axe": axe, "pas": float(pas), "unites": "modele",
           "repere": "monde", "pieces": propres}
    tmp = d / f"{p.name}.tmp"
    tmp.write_text(json.dumps(doc, ensure_ascii=False, indent=1),
                   encoding="utf-8")
    tmp.replace(p)
    return {"ok": True, "fichier": p.name, "pieces": len(propres)}


@router.get("/assets/3d/{job}/vignette")
async def get_asset3d_vignette(job: str, v: int = 1):
    """La vignette CAPTURÉE par l'Établi au moment de l'écriture d'une version.

    Déclarée AVANT /{fmt} — même règle que /preview et /silhouette : sans cela
    elle serait avalée comme `fmt="vignette"` et irait chercher un
    `model.vignette` qui n'existe nulle part. Un banc mesure ce 200.

    404 FRANC quand la version n'a pas été capturée. C'est le cas NORMAL, et
    ce n'est pas un défaut : la vignette naît à l'écriture SEULEMENT (décision
    de l'utilisateur), donc toute production antérieure à cette tâche n'en a
    pas. `_etabli_vignette` ne compose d'ailleurs cette URL que si le fichier
    est là ; personne ne sert de lien mort.

    `_etabli_vignette_cible` et non `_etabli_vignette_path` : `job` vient ici
    d'un segment d'URL, donc du réseau, et une LECTURE hors du dossier des
    jobs dirait déjà quels fichiers existent ailleurs. Les deux routes qui
    reçoivent un `job` du dehors franchissent la même porte.
    """
    p = _etabli_vignette_cible(job, v)
    if not p.is_file():
        raise HTTPException(404, "Not found")
    return FileResponse(p, media_type="image/png")


@router.get("/assets/3d/{job}/{fmt}")
async def get_asset3d_file(job: str, fmt: str):
    """Stream a generated mesh file (glb|fbx|obj|stl|usdz)."""
    p = settings.outputs_path / "assets3d" / Path(job).name / f"model.{Path(fmt).name}"
    if not p.is_file():
        raise HTTPException(404, "Not found")
    return FileResponse(p)


@router.get("/assets/3d/{job}/shot/{i}")
async def get_asset3d_shot(job: str, i: int):
    """Stream an individual view shot image (shot_0 = source, shot_1..N = views)."""
    p = settings.outputs_path / "assets3d" / Path(job).name / f"shot_{int(i)}.png"
    if not p.is_file():
        raise HTTPException(404, "Not found")
    return FileResponse(p)


@router.post("/assets/3d/{job}/shot/{i}/save")
async def save_asset3d_shot(job: str, i: int):
    """Copy a shot into the Library images folder so it can be reused as a source."""
    import shutil
    src = settings.outputs_path / "assets3d" / Path(job).name / f"shot_{int(i)}.png"
    if not src.is_file():
        raise HTTPException(404, "Not found")
    dest = settings.images_path / f"shot_{Path(job).name}_{int(i)}.png"
    n = 2
    while dest.exists():
        dest = settings.images_path / f"shot_{Path(job).name}_{int(i)}_{n}.png"
        n += 1
    shutil.copy2(src, dest)
    await LI.noter([dest.name], "assets3d", job_id=Path(job).name)
    return {"filename": dest.name}


# ---- 3D Studio Meshy (v2.1) — proxy sécurisé + journal + bibliothèque ----
# Spec : INTEGRATION-MESHY.md (projet Claude Design « DeepOtus Studio »).
# La clé MESHY_API_KEY ne quitte jamais le serveur : le front (/studio3d +
# frontend/meshy/meshy.client.js) parle à /api/meshy/* qui recopie vers
# api.meshy.ai en injectant le Bearer, chemins allowlistés. Ce flux est
# indépendant de Game Assets 3D (fal) ci-dessus, qui reste inchangé.


@router.api_route("/meshy/{meshy_path:path}", methods=["GET", "POST", "DELETE"])
async def meshy_proxy(meshy_path: str, request: Request):
    from fastapi.responses import StreamingResponse
    from app.services import meshy_service as MS

    parsed = MS.parse_proxy_path(request.method, meshy_path)
    if parsed is None:
        raise HTTPException(403, f"Meshy path not allowed: {request.method} /{meshy_path}")
    mock = MS.mock_enabled()
    if not mock and not settings.has_meshy:
        raise HTTPException(503, "MESHY_API_KEY not configured — add it in "
                                 "Settings (or set MESHY_MOCK=1 for the local simulator)")

    body = await request.body()
    payload = None
    if request.method == "POST":
        try:
            payload = json.loads(body) if body else {}
        except ValueError:
            raise HTTPException(400, "Invalid JSON body")

    # SSE : relais en streaming — sinon le client retombe sur le polling.
    if parsed["stream"]:
        if mock:
            async def _mock_events():
                async for ev in MS.get_mock().stream(parsed["task_id"]):
                    if ev.startswith("data:"):
                        try:
                            task = json.loads(ev[5:].strip())
                            if task.get("status") in MS.TERMINAL:
                                await MS.record_state(task, parsed["base"])
                        except ValueError:
                            pass
                    yield ev
            return StreamingResponse(_mock_events(), media_type="text/event-stream")
        return StreamingResponse(MS.proxy_stream(meshy_path),
                                 media_type="text/event-stream")

    if mock:
        mk = MS.get_mock()
        if request.method == "POST":
            code, data = mk.create(parsed["base"], payload or {})
            if code < 400 and data.get("result"):
                await MS.record_created(str(data["result"]), parsed["base"], payload)
        elif request.method == "DELETE":
            code, data = mk.delete(parsed["task_id"])
        elif parsed["base"] == MS.BALANCE_PATH:
            code, data = mk.balance()
        elif parsed["task_id"]:
            code, data = mk.get(parsed["task_id"])
            if code == 200:
                await MS.record_state(data, parsed["base"])
        else:
            code, data = 200, {"result": []}  # liste paginée non simulée
        return Response(json.dumps(data), status_code=code,
                        media_type="application/json")

    code, content, ctype = await MS.proxy_request(
        request.method, meshy_path, body, dict(request.query_params))
    # Journal (spec §6) : créations + états qui transitent par le proxy.
    try:
        data = json.loads(content) if content else {}
        if (request.method == "POST" and code < 400
                and isinstance(data, dict) and data.get("result")):
            await MS.record_created(str(data["result"]), parsed["base"], payload)
        elif (request.method == "GET" and code == 200
                and parsed["task_id"] and isinstance(data, dict)):
            await MS.record_state(data, parsed["base"])
    except ValueError:
        pass
    return Response(content, status_code=code, media_type=ctype)


@router.post("/meshy3d/estimate")
async def meshy3d_estimate(body: dict):
    """Coût estimé AVANT lancement (règle produit) : total + détail lignes."""
    from app.services import meshy_service as MS
    return MS.estimate_pipeline(body or {})


@router.get("/meshy3d/status")
async def meshy3d_status():
    from app.services import meshy_service as MS
    mock = MS.mock_enabled()
    return {"enabled": settings.has_meshy or mock, "mock": mock,
            "configured": settings.has_meshy,
            "host": "simulateur local" if mock else "api.meshy.ai"}


@router.get("/meshy3d/tasks")
async def meshy3d_tasks(limit: int = 60):
    """Journal + bibliothèque persistée (les URLs Meshy expirent, pas nous)."""
    from app.services import meshy_service as MS
    rows = await MS.list_tasks(limit=limit)
    return {"tasks": rows, "expiring": MS.expiring_soon(rows)}


@router.post("/meshy3d/repatriate/{task_id}")
async def meshy3d_repatriate(task_id: str):
    """Rapatrie les binaires d'une tâche SUCCEEDED (normalement automatique)."""
    from app.services import meshy_service as MS
    try:
        return await MS.repatriate(task_id)
    except ValueError as e:
        raise HTTPException(404, str(e))


@router.get("/meshy3d/files/{task_dir}/{fname}")
async def meshy3d_file(task_dir: str, fname: str):
    """Sert un binaire rapatrié (bibliothèque locale, URLs stables)."""
    from app.services import meshy_service as MS
    p = MS.meshy3d_dir() / Path(task_dir).name / Path(fname).name
    if not p.is_file():
        raise HTTPException(404, "Not found")
    media = "model/gltf-binary" if p.suffix == ".glb" else None
    return FileResponse(p, media_type=media)


@router.get("/meshy3d/mockfile/{task_id}/{fname}")
async def meshy3d_mockfile(task_id: str, fname: str):
    """Binaires du simulateur (MESHY_MOCK=1) : GLB/PNG minimaux valides."""
    from app.services import meshy_service as MS
    if not MS.mock_enabled():
        raise HTTPException(404, "Mock mode disabled")
    data, media = MS.mock_file_bytes(Path(fname).name)
    return Response(data, media_type=media)


# ---- Game Assets 2D — Sprite Lab (chantier 9a) ----

def _sprite_dir(job: str) -> Path:
    return settings.outputs_path / "sprites" / Path(job).name


@router.post("/assets/sprite")
async def assets_sprite(body: dict, background_tasks: BackgroundTasks):
    """Game Assets 2D: video render -> frames -> sprite sheet + pack Unity.
    Mirrors /assets/3d: pre-register a sprite2d JobRecord, run in the
    background, record what was produced in cost_meta. Poll GET /api/jobs/{id}.
    Body: {source: {kind: job|upload|video, ...}, fps_sample, max_frames,
    remove_bg: none|api|local, trim: animation|tight, cell: {size, align},
    columns: "auto"|int, pixel?: {target_px, colors|palette, dither} (9b),
    extract_only?: bool (9c: frames-only probe for the filmstrip),
    keep?: [indices] (9c: filmstrip selection, sampling order), title?}."""
    from datetime import datetime as _dtu
    import json as _json
    from app.services import sprite_service as SS
    from app.services.storage import JobRecord, async_session_factory

    # fail fast on bad input instead of accepting a job that dies in background
    try:
        opts = SS.normalize_opts(body)
        src = await SS.resolve_source(body.get("source") or {})
    except ValueError as e:
        raise HTTPException(400, str(e))
    if opts["remove_bg"] == "api" and not settings.FAL_KEY:
        raise HTTPException(400, "FAL_KEY not configured (Settings) — "
                                 "use remove_bg 'local' or 'none'.")
    if opts["remove_bg"] == "local":
        try:
            import rembg  # noqa: F401
        except ImportError:
            raise HTTPException(
                400, "rembg is not installed in this runtime — use remove_bg "
                     "'api' (fal) or 'none', or install it with: pip install rembg")

    job_id = str(uuid4())
    short = job_id[:8]
    async with async_session_factory() as s:
        s.add(JobRecord(
            id=job_id, status=JobStatus.GENERATING_VIDEO.value, progress=5,
            title=(body.get("title") or f"Sprites · {src.stem}"),
            image_filename=f"sprite_{short}",
            provider="sprite2d", current_step="Extracting frames"))
        await s.commit()

    async def on_step(label, pct):
        async with async_session_factory() as s2:
            jr2 = await s2.get(JobRecord, job_id)
            if jr2 is not None:
                jr2.current_step = label
                jr2.progress = int(pct)
                await s2.commit()

    async def _run():
        try:
            r = await SS.generate_sprites(body, short, on_step=on_step)
            async with async_session_factory() as s:
                jr = await s.get(JobRecord, job_id)
                if jr is not None:
                    jr.status = JobStatus.DONE.value
                    jr.progress = 100
                    jr.final_video_path = r.get("sheet")
                    if r.get("sheet"):   # extract_only probes have no sheet
                        jr.image_filename = "sheet.png"
                    jr.current_step = "Complete"
                    jr.completed_at = _dtu.utcnow()
                    jr.cost_meta = _json.dumps({
                        "job": short, "frames": r.get("frames"),
                        "remove_bg": r.get("remove_bg"),
                        "grid": r.get("grid"),
                        "bg_failed": r.get("bg_failed") or []})
                    await s.commit()
        except Exception as e:
            logger.exception(f"sprite2d {job_id} failed: {e}")
            async with async_session_factory() as s:
                jr = await s.get(JobRecord, job_id)
                if jr is not None:
                    jr.status = JobStatus.FAILED.value
                    jr.error = str(e)
                    jr.current_step = "Failed"
                    await s.commit()

    background_tasks.add_task(_run)
    return {"job_id": job_id, "status": "queued"}


@router.get("/assets/sprite/{job}/manifest")
async def get_sprite_manifest(job: str):
    """The generated manifest (grid, frames, fps, offsets) + which files
    actually exist on disk (ground truth)."""
    d = _sprite_dir(job)
    mf = d / "manifest.json"
    if not mf.is_file():
        raise HTTPException(404, "Not found")
    data = json.loads(mf.read_text(encoding="utf-8"))
    fdir = d / "frames"
    data["files"] = {
        "sheet": (d / "sheet.png").is_file(),
        "preview": (d / "preview.gif").is_file(),
        "unity_json": (d / "sheet.unity.json").is_file(),
        "unity_importer": (d / "SpriteSheetImporter.cs").is_file(),
        "frames": len(list(fdir.glob("*.png"))) if fdir.is_dir() else 0,
    }
    return data


@router.get("/assets/sprite/{job}/sheet")
async def get_sprite_sheet(job: str):
    p = _sprite_dir(job) / "sheet.png"
    if not p.is_file():
        raise HTTPException(404, "Not found")
    return FileResponse(p)


@router.get("/assets/sprite/{job}/preview")
async def get_sprite_preview(job: str):
    p = _sprite_dir(job) / "preview.gif"
    if not p.is_file():
        raise HTTPException(404, "Not found")
    return FileResponse(p)


@router.get("/assets/sprite/{job}/frame/{i}")
async def get_sprite_frame(job: str, i: int):
    p = _sprite_dir(job) / "frames" / f"{int(i):03d}.png"
    if not p.is_file():
        raise HTTPException(404, "Not found")
    return FileResponse(p)


@router.get("/assets/sprite/{job}/zip")
async def get_sprite_zip(job: str):
    """Full pack: sheet + frames + manifests + Unity importer."""
    from app.services.sprite_service import build_zip_bytes
    d = _sprite_dir(job)
    if not (d / "sheet.png").is_file():
        raise HTTPException(404, "Not found")
    data = await asyncio.to_thread(build_zip_bytes, d)
    return Response(
        content=data, media_type="application/zip",
        headers={"Content-Disposition":
                 f'attachment; filename="sprites_{Path(job).name}.zip"'})


@router.post("/assets/sprite/{job}/save")
async def save_sprite_sheet(job: str):
    """Copy sheet.png into the Library images folder so it can be reused as an
    ordinary image (Studio node, Seedance start frame, ...)."""
    import shutil
    src = _sprite_dir(job) / "sheet.png"
    if not src.is_file():
        raise HTTPException(404, "Not found")
    dest = settings.images_path / f"gen_sprite_{Path(job).name}.png"
    n = 2
    while dest.exists():
        dest = settings.images_path / f"gen_sprite_{Path(job).name}_{n}.png"
        n += 1
    shutil.copy2(src, dest)
    await LI.noter([dest.name], "sprites", job_id=Path(job).name)
    return {"filename": dest.name}


# ═══════════════════════════════════════════════════════════════════════════
# Catalogue de démarrage (CC0) + génération locale de particules
#
# Raison d'être : sans clé ElevenLabs ni clé fal, l'écran Son & VFX était vide
# et ne proposait que de dépenser. Ces routes servent un catalogue embarqué
# (80 textures, 5 séquences, 606 bruitages, tous CC0) et un générateur de
# particules 100 % local. Aucune n'exige de clé ni de réseau.
# ═══════════════════════════════════════════════════════════════════════════

@router.get("/starter/catalog")
async def starter_catalog_index(kind: str = "", family: str = "",
                                q: str = "", limit: int = 0):
    """Catalogue de démarrage. Sans `kind`, renvoie l'index complet (familles
    + sources) que l'UI charge une fois ; avec `kind`, la liste filtrée."""
    from app.services import starter_catalog as SC
    cat = SC.load()
    if not kind:
        return {
            "available": cat.get("available", False),
            "sources": cat.get("sources", []),
            "sfx_families": cat.get("sfx_families", []),
            "particle_families": cat.get("particle_families", []),
            "anims": cat.get("anims", []),
            "counts": {"particles": len(cat.get("particles", [])),
                       "sfx": len(cat.get("sfx", [])),
                       "anims": len(cat.get("anims", []))},
        }
    try:
        items = SC.browse(kind, family=family, query=q, limit=limit)
    except SC.StarterError as e:
        raise HTTPException(e.status, e.message)
    return {"kind": kind, "items": items, "total": len(items)}


@router.post("/starter/import")
async def starter_import(body: dict):
    """Copie des éléments du catalogue dans la Bibliothèque de l'utilisateur.

    C'est le point qui évite un cas particulier permanent : une fois copié, un
    son de démarrage est un son de la Bibliothèque comme un autre, et tout
    l'aval (tiroir Sons, Montage, rendu) le traite sans rien savoir de lui.
    Body: {kind: "sfx"|"particle", ids: [...]}"""
    from app.services import starter_catalog as SC
    kind = str((body or {}).get("kind") or "").strip()
    ids = (body or {}).get("ids") or []
    if not isinstance(ids, list) or not ids:
        raise HTTPException(400, "ids: liste non vide attendue.")
    if len(ids) > 100:
        raise HTTPException(400, "100 éléments maximum par import.")
    fn = {"sfx": SC.import_sfx, "particle": SC.import_particle}.get(kind)
    if fn is None:
        raise HTTPException(400, "kind doit valoir 'sfx' ou 'particle'.")
    try:
        items = await asyncio.get_running_loop().run_in_executor(
            None, lambda: fn([str(i) for i in ids]))
    except SC.StarterError as e:
        raise HTTPException(e.status, e.message)
    return {"ok": True, "imported": len(items), "items": items}


@router.get("/particles/presets")
async def particles_presets():
    """Presets de l'écran « VFX particules » : chacun est (texture CC0 +
    réglages d'émetteur), donc directement exécutable."""
    from app.services import particle_service as PS
    from app.services import starter_catalog as SC
    cat = SC.load()
    thumbs = {p["id"]: p.get("thumb") for p in cat.get("particles", [])}
    return {
        "available": cat.get("available", False),
        "presets": [{
            "id": p["id"], "name": p["name"], "type": p["type"],
            "desc": p["desc"], "texture": p["texture"],
            "thumb": (f"/starter/{thumbs[p['texture']]}"
                      if thumbs.get(p["texture"]) else None),
            "frames": p["emitter"]["frames"], "fps": p["emitter"]["fps"],
            "blend": p["emitter"]["blend"],
        } for p in PS.PRESETS],
        "anims": [{"id": a["id"], "name": a["name"], "frames": a["frames"],
                   "thumb": f"/starter/{a['thumb']}"}
                  for a in cat.get("anims", [])],
    }


def _sprite_job(title: str, step: str):
    """Pré-enregistre un JobRecord sprite2d et rend (job_id, short, on_step).

    Les particules et les séquences importées produisent EXACTEMENT le même
    artefact que le Sprite Lab (planche + frames + GIF + pack Unity), donc
    elles réutilisent le type de job `sprite2d` : l'onglet Sprites de la
    Bibliothèque, la visionneuse et l'export ZIP marchent sans une ligne de
    plus.
    """
    from app.services.storage import JobRecord, async_session_factory

    job_id = str(uuid4())
    short = job_id[:8]

    async def register():
        async with async_session_factory() as s:
            s.add(JobRecord(
                id=job_id, status=JobStatus.GENERATING_VIDEO.value, progress=5,
                title=title, image_filename=f"sprite_{short}",
                provider="sprite2d", current_step=step))
            await s.commit()

    async def on_step(label, pct):
        async with async_session_factory() as s2:
            jr = await s2.get(JobRecord, job_id)
            if jr is not None:
                jr.current_step = label
                jr.progress = int(pct)
                await s2.commit()

    return job_id, short, register, on_step


def _sprite_job_finisher(job_id: str, short: str, coro_factory, extra_meta):
    """Exécute le travail et écrit l'issue dans le JobRecord (même contrat
    d'erreur que /assets/sprite : un échec laisse un job FAILED lisible)."""
    from datetime import datetime as _dtu
    import json as _json
    from app.services.storage import JobRecord, async_session_factory

    async def _run():
        try:
            r = await coro_factory()
            async with async_session_factory() as s:
                jr = await s.get(JobRecord, job_id)
                if jr is not None:
                    jr.status = JobStatus.DONE.value
                    jr.progress = 100
                    jr.final_video_path = r.get("sheet")
                    if r.get("sheet"):
                        jr.image_filename = "sheet.png"
                    jr.current_step = "Complete"
                    jr.completed_at = _dtu.utcnow()
                    jr.cost_meta = _json.dumps(dict(
                        {"job": short, "frames": r.get("frames"),
                         "grid": r.get("grid"), "webm": r.get("webm"),
                         "usd": 0.0, "bg_failed": []},
                        **extra_meta(r)))
                    await s.commit()
        except Exception as e:
            logger.exception(f"sprite2d {job_id} failed: {e}")
            async with async_session_factory() as s:
                jr = await s.get(JobRecord, job_id)
                if jr is not None:
                    jr.status = JobStatus.FAILED.value
                    jr.error = str(e)
                    jr.current_step = "Failed"
                    await s.commit()

    return _run


@router.post("/assets/particles")
async def assets_particles(body: dict, background_tasks: BackgroundTasks):
    """Génère un sprite de particules — local, gratuit, sans clé ni réseau.

    Body: {preset?: id, texture?: id, emitter?: {...}, seed?: int,
    webm?: bool, title?: str}. Un preset seul suffit ; `emitter` ne surcharge
    que les réglages fournis. Sondez ensuite GET /api/jobs/{id}."""
    from app.services import particle_service as PS
    from app.services import starter_catalog as SC

    try:
        opts = PS.normalize_opts(body or {})
        SC.asset_path("particle", opts["texture"])   # fail-fast: texture connue
    except ValueError as e:
        raise HTTPException(400, str(e))
    except SC.StarterError as e:
        raise HTTPException(e.status, e.message)

    label = PS.PRESET_BY_ID.get(opts["preset"], {}).get("name") \
        or opts["texture"]
    title = (body or {}).get("title") or f"Particules · {label}"
    job_id, short, register, on_step = _sprite_job(
        title, "Simulation des particules")
    await register()
    background_tasks.add_task(_sprite_job_finisher(
        job_id, short,
        lambda: PS.generate_particles(body or {}, short, on_step=on_step),
        lambda r: {"preset": r.get("preset"), "texture": r.get("texture"),
                   "seed": r.get("seed"), "kind": "particles"}))
    return {"job_id": job_id, "status": "queued"}


@router.post("/assets/starter-anim")
async def assets_starter_anim(body: dict, background_tasks: BackgroundTasks):
    """Assemble une séquence animée CC0 du catalogue en planche sprite2d,
    sans simulation. Body: {anim: id, cell?: 128|256|512, title?: str}."""
    from app.services import particle_service as PS
    from app.services import starter_catalog as SC

    anim = str((body or {}).get("anim") or "").strip()
    cell = (body or {}).get("cell") or 512
    try:
        cell = int(cell)
        item = SC.get("anim", anim)
        SC.anim_frames(anim)
        if cell not in (128, 256, 512):
            raise ValueError("cell doit valoir 128, 256 ou 512")
    except (TypeError, ValueError) as e:
        raise HTTPException(400, str(e))
    except SC.StarterError as e:
        raise HTTPException(e.status, e.message)

    title = (body or {}).get("title") or f"Séquence · {item['name']}"
    job_id, short, register, on_step = _sprite_job(
        title, "Assemblage de la planche")
    await register()
    background_tasks.add_task(_sprite_job_finisher(
        job_id, short,
        lambda: PS.import_anim(anim, short, cell=cell, on_step=on_step),
        lambda r: {"anim": r.get("anim"), "kind": "starter-anim"}))
    return {"job_id": job_id, "status": "queued"}


# ── Studio named-graph store (v1.15.6): save / reload node graphs by name,
# separate from the render-time source_graph dump. Lives in the data dir
# (DATA_ROOT/assets/studio_graphs) so it survives updates/reinstalls.
def _studio_graphs_dir():
    d = settings.outputs_path.parent / "studio_graphs"
    d.mkdir(parents=True, exist_ok=True)
    return d


@router.get("/studio-graphs")
async def list_studio_graphs():
    """Saved Studio graphs (metadata only, newest first)."""
    import json as _json
    out = []
    for f in _studio_graphs_dir().glob("*.json"):
        try:
            d = _json.loads(f.read_text(encoding="utf-8"))
            out.append({"id": d.get("id", f.stem),
                        "name": d.get("name", f.stem),
                        "updated_at": d.get("updated_at")})
        except Exception:
            continue
    out.sort(key=lambda g: g.get("updated_at") or "", reverse=True)
    return {"graphs": out}


@router.get("/studio-graphs/{graph_id}")
async def get_studio_graph(graph_id: str):
    import json as _json
    p = _studio_graphs_dir() / f"{Path(graph_id).name}.json"
    if not p.is_file():
        raise HTTPException(404, "Graph not found")
    try:
        return _json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        raise HTTPException(500, "Graph unreadable")


@router.post("/studio-graphs")
async def save_studio_graph(body: dict, request: Request):
    """Save or overwrite a named Studio graph. Body: {id?, name, graph}."""
    _require_localhost(request)
    import json as _json
    from datetime import datetime as _dtnow
    graph = body.get("graph")
    if not isinstance(graph, dict) or not graph.get("nodes"):
        raise HTTPException(400, "graph (with nodes) is required")
    gid = (str(body.get("id") or "").strip()) or f"g_{uuid4().hex[:8]}"
    gid = Path(gid).name
    name = (str(body.get("name") or graph.get("name") or "Untitled graph").strip())[:120]
    rec = {"id": gid, "name": name, "graph": graph,
           "updated_at": _dtnow.utcnow().isoformat()}
    (_studio_graphs_dir() / f"{gid}.json").write_text(
        _json.dumps(rec, ensure_ascii=False), encoding="utf-8")
    return {"id": gid, "name": name}


@router.delete("/studio-graphs/{graph_id}")
async def delete_studio_graph(graph_id: str):
    p = _studio_graphs_dir() / f"{Path(graph_id).name}.json"
    if not p.is_file():
        raise HTTPException(404, "Graph not found")
    p.unlink()
    return {"deleted": graph_id}


@router.get("/emojis")
async def list_emojis():
    """Curated native emoji set (categories -> [{e: char, f: png basename}])
    for the Studio emoji picker. PNGs are served at /emoji/<f>.png and used by
    the renderer (Twemoji overlay) so picker and video match."""
    import json
    from pathlib import Path as _P
    p = _P(__file__).resolve().parent.parent / "assets" / "emoji" / "manifest.json"
    if p.exists():
        try:
            return json.loads(p.read_text(encoding="utf-8"))
        except Exception:
            pass
    return []


def _emoji_custom_dir() -> Path:
    from app.config import DATA_ROOT
    p = DATA_ROOT / "assets" / "emoji_custom"
    p.mkdir(parents=True, exist_ok=True)
    return p


@router.get("/emojis/custom")
async def list_custom_emojis():
    """User-imported custom emojis (stored in the data dir). Each entry ->
    {name, file, url, code}. The picker inserts `code` (:name:) into ticker /
    text-overlay text; the renderer resolves it to the PNG so video matches."""
    d = _emoji_custom_dir()
    out = [{"name": f.stem, "file": f.name,
            "url": f"/emoji-custom/{f.name}", "code": f":{f.stem}:"}
           for f in sorted(d.glob("*.png"))]
    return {"emojis": out}


@router.post("/emojis/custom")
async def upload_custom_emoji(request: Request,
                             file: UploadFile = File(...),
                             name: str = Form("")):
    """Import a custom emoji image -> RGBA PNG (<=160px) in the data dir under a
    :shortcode: from `name` (or the filename). Survives app reinstall."""
    _require_localhost(request)
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in (".png", ".jpg", ".jpeg", ".webp", ".gif"):
        raise HTTPException(400, "Emoji must be .png, .jpg, .webp or .gif")
    data = await file.read()
    if len(data) > 5 * 1024 * 1024:
        raise HTTPException(400, "Image too large (max 5 MB)")
    base = _slug(name or Path(file.filename or "emoji").stem) or "emoji"
    d = _emoji_custom_dir()
    nm, i = base, 2
    while (d / (nm + ".png")).exists():
        nm = f"{base}-{i}"
        i += 1
    try:
        import io
        from PIL import Image as PILImg
        img = PILImg.open(io.BytesIO(data)).convert("RGBA")
        img.thumbnail((160, 160), PILImg.LANCZOS)
        img.save(d / (nm + ".png"), format="PNG")
    except Exception as e:
        raise HTTPException(400, f"Not a valid image: {e}")
    logger.info(f"custom emoji imported: {nm}")
    return {"name": nm, "file": nm + ".png",
            "url": f"/emoji-custom/{nm}.png", "code": f":{nm}:"}


@router.delete("/emojis/custom/{name}")
async def delete_custom_emoji(name: str, request: Request):
    _require_localhost(request)
    p = _emoji_custom_dir() / (_slug(name) + ".png")
    if p.exists():
        try:
            p.unlink()
        except Exception as e:
            raise HTTPException(500, f"Could not delete: {e}")
    return {"ok": True, "name": _slug(name)}


# ---- v1.7: News / RSS pipeline ----

@router.get("/news/sources")
async def list_news_sources():
    return {"sources": news_service.list_sources()}


@router.post("/news/sources")
async def add_news_source(request: AddNewsSourceRequest):
    try:
        src = news_service.add_source(
            request.url, request.name, request.type)
    except ValueError as e:
        raise HTTPException(400, str(e))
    return {"source": src}


@router.delete("/news/sources/{source_id}")
async def delete_news_source(source_id: str):
    if not news_service.remove_source(source_id):
        raise HTTPException(404, f"Source not found: {source_id}")
    return {"deleted": source_id}


@router.post("/news/sources/{source_id}/toggle")
async def toggle_news_source(source_id: str, request: NewsSourceToggleRequest):
    if not news_service.set_enabled(source_id, request.enabled):
        raise HTTPException(404, f"Source not found: {source_id}")
    return {"source_id": source_id, "enabled": request.enabled}


@router.post("/news/sources/defaults")
async def seed_default_news_sources():
    """Add the curated default feed pack (crypto / geopolitics / economy /
    politics EU·China·USA). Idempotent — skips sources already present."""
    return news_service.seed_defaults()


@router.post("/news/refresh")
async def refresh_news():
    try:
        return await news_service.refresh()
    except Exception as e:
        logger.exception("news refresh failed")
        raise HTTPException(500, f"Refresh failed: {e}")


@router.get("/news/items")
async def list_news_items():
    return news_service.get_items()


@router.post("/news/script", response_model=NewsScriptResponse)
async def generate_news_script_route(request: NewsScriptRequest):
    """Read the selected articles (when read_articles), extract the essence
    + lead images (saved to assets/images for Seedance), then render a
    deepotus 'prophet' (cynical/humorous) script + caption."""
    try:
        items = [i.model_dump() for i in request.items]
        essences: list[NewsEssence] = []
        images: list[str] = []
        if request.read_articles:
            items = await news_service.enrich_items(
                items, summary_words=request.summary_words)
            for it in items:
                if it.get("image"):
                    images.append(it["image"])
                essences.append(NewsEssence(
                    title=it.get("title", ""),
                    essence=it.get("essence", ""),
                    image=it.get("image"),
                    link=it.get("link", ""),
                    status=it.get("scrape_status", ""),
                ))
        base = await asyncio.to_thread(
            pipeline.engine.generate_news_script,
            items,
            voice_mode=request.voice_mode,
            language=request.language,
            max_words=request.max_words,
            angle=request.angle,
        )
        return NewsScriptResponse(
            **base.model_dump(),
            sources_read=len(essences),
            images=images,
            essences=essences,
        )
    except Exception as e:
        logger.exception("news script generation failed")
        raise HTTPException(500, f"Script generation failed: {e}")


@router.post("/news/illustration", response_model=NewsIllustrationResponse)
async def generate_news_illustration_route(
    request: NewsIllustrationRequest,
    background_tasks: BackgroundTasks,
):
    """Render a branded 1080x1920 news-illustration reel from selected items.
    Silent (the avatar carries audio when composed). Poll GET /api/jobs."""
    job_id = str(uuid4())

    async def _run():
        try:
            await pipeline.run_news_illustration(
                [i.model_dump() for i in request.items],
                per_card_s=request.per_card_s,
                engine=request.engine,
                job_id=job_id,
            )
        except Exception as e:
            logger.exception(f"news illustration {job_id} failed: {e}")

    background_tasks.add_task(_run)
    return NewsIllustrationResponse(
        job_id=job_id,
        message=f"News illustration queued. Poll GET /api/jobs/{job_id}.",
    )


# ---- Images ----

@router.get("/images")
async def list_images():
    folder = settings.images_path
    if not folder.exists():
        return {"folder": str(folder), "images": [], "warning": "Folder does not exist"}

    extensions = {".png", ".jpg", ".jpeg", ".webp"}
    items: list[ImageItem] = []
    for p in sorted(folder.iterdir()):
        if p.is_file() and p.suffix.lower() in extensions:
            try:
                with PILImage.open(p) as img:
                    width, height = img.size
            except Exception:
                width = height = None
            items.append(ImageItem(
                filename=p.name,
                path=str(p),
                size_kb=p.stat().st_size // 1024,
                width=width,
                height=height,
                mtime=p.stat().st_mtime,
            ))
    # provenance (28/08) : l'index dit la fonction productrice ; un fichier
    # jamais indexé est classé par son nom, en le disant (heuristique)
    prov = await LI.carte()
    for it in items:
        connu = prov.get(it.filename)
        if connu:
            it.source, it.source_origin = connu[0], connu[1]
        else:
            it.source = LI.heuristique(it.filename)
            it.source_origin = "heuristique"
    return {"folder": str(folder), "images": [i.model_dump() for i in items]}


@router.post("/images/upload")
async def upload_image(file: UploadFile = File(...)):
    folder = settings.images_path
    folder.mkdir(parents=True, exist_ok=True)
    safe = Path(file.filename or "image.png").name
    if not safe or safe in (".", "..") or "/" in safe or "\\" in safe:
        raise HTTPException(400, "Invalid filename")
    dest = folder / safe
    contents = await file.read()
    dest.write_bytes(contents)
    # provenance : les exports du Vectorlab passent par cette route avec
    # leur préfixe ; tout le reste est un import utilisateur
    await LI.noter([safe], "vectorlab" if safe.startswith("vector_")
                   else "import")
    return {"saved": str(dest), "filename": safe, "size_kb": len(contents) // 1024}


@router.get("/images/{filename}")
async def get_image_file(filename: str):
    # Containment via name sanitization only. The previous resolve()+startswith
    # comparison broke under Windows filesystem virtualization (MSIX/sandboxed
    # launchers redirect AppData\Local to a Packages\...\LocalCache view, so
    # the file's resolve() no longer string-matches the base dir) -> every
    # image 404'd. Path(filename).name strips any directory part, which
    # already guarantees the path stays inside images_path.
    safe = Path(filename).name
    if not safe or safe in (".", "..") or safe != filename:
        raise HTTPException(400, "Invalid filename")
    p = settings.images_path / safe
    if not p.is_file():
        raise HTTPException(404, f"Image not found: {filename}")
    return FileResponse(p)


@router.delete("/images/{filename}")
async def delete_image_file(filename: str):
    """Delete a generated/uploaded image from the images folder."""
    # Same virtualization-safe containment check as get_image_file above.
    safe = Path(filename).name
    if not safe or safe in (".", "..") or safe != filename:
        raise HTTPException(400, "Invalid filename")
    p = settings.images_path / safe
    try:
        if not p.is_file():
            raise HTTPException(404, f"Image not found: {filename}")
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(404, f"Image not found: {filename}")
    p.unlink()
    await LI.retirer(safe)
    return {"deleted": safe}


def _safe_rename_image(old_name: str, new_name: str) -> str:
    """Rename an image in images_path. Keeps the source extension, sanitizes the
    stem, strips any path components, auto-suffixes on collision, and refuses to
    escape images_path. Returns the final filename. File-only (no ref rewrite)."""
    src_name = Path(old_name).name
    src = settings.images_path / src_name
    if not src.is_file():
        raise FileNotFoundError(src_name)
    ext = src.suffix
    raw = Path(new_name).name
    raw_stem = raw[:-len(Path(raw).suffix)] if Path(raw).suffix else raw
    stem = re.sub(r"[^A-Za-z0-9._ -]", "_", raw_stem)
    stem = re.sub(r"_+", "_", stem).strip(" _") or "image"
    dest = settings.images_path / f"{stem}{ext}"
    if dest.resolve() == src.resolve():
        return src_name  # no-op
    n = 2
    while dest.exists():
        dest = settings.images_path / f"{stem}_{n}{ext}"
        n += 1
    if not str(dest.resolve()).startswith(str(settings.images_path.resolve())):
        raise ValueError("Invalid destination")
    src.rename(dest)
    return dest.name


@router.post("/images/{filename}/rename")
async def rename_image_file(filename: str, body: dict):
    """Rename an uploaded/generated image (file-only). Body: {new_name}."""
    new = str(body.get("new_name") or "").strip()
    if not new:
        raise HTTPException(400, "new_name required")
    try:
        final = _safe_rename_image(filename, new)
    except FileNotFoundError:
        raise HTTPException(404, f"Image not found: {filename}")
    except Exception as e:
        raise HTTPException(400, str(e))
    await LI.renommer(Path(filename).name, final)
    return {"old": Path(filename).name, "new": final}


_AUDIO_EXTS = {".mp3", ".wav", ".m4a", ".aac", ".ogg", ".flac", ".opus"}


def _audio_dir() -> Path:
    """User audio assets (music / SFX / voice), in the stable data dir."""
    p = settings.images_path.parent / "audio"
    p.mkdir(parents=True, exist_ok=True)
    return p


@router.get("/audio")
async def list_audio():
    """List uploaded audio assets for the Library + Studio audio nodes."""
    d = _audio_dir()
    out = []
    for p in sorted(d.glob("*"), key=lambda f: f.stat().st_mtime, reverse=True):
        if p.is_file() and p.suffix.lower() in _AUDIO_EXTS:
            out.append({"name": p.name, "url": f"/api/audio/{p.name}",
                        "size_kb": p.stat().st_size // 1024})
    return {"audio": out}


@router.post("/audio/upload")
async def upload_audio(file: UploadFile = File(...)):
    folder = _audio_dir()
    safe = Path(file.filename or "audio.mp3").name
    if not safe or safe in (".", "..") or "/" in safe or "\\" in safe:
        raise HTTPException(400, "Invalid filename")
    if Path(safe).suffix.lower() not in _AUDIO_EXTS:
        raise HTTPException(400, "Unsupported audio format")
    contents = await file.read()
    if len(contents) > 50 * 1024 * 1024:
        raise HTTPException(400, "Audio too large (max 50 MB)")
    dest = folder / safe
    dest.write_bytes(contents)
    # R1 — sidecar meta (tags du tiroir Sons) : import ou musique selon le nom.
    from app.services import sfx_service
    await asyncio.get_running_loop().run_in_executor(
        None, lambda: sfx_service.record_meta(safe, {
            "kind": sfx_service.classify_kind(safe),
            "created": datetime.now().isoformat(timespec="seconds")}))
    await LI.noter([safe], "import", kind="audio")
    return {"saved": str(dest), "filename": safe, "size_kb": len(contents) // 1024}


# ── R1 gauntlet SFX — meta, génération ElevenLabs, audition ────────────────
# NB : /audio/meta DOIT être déclaré AVANT /audio/{filename} (ordre FastAPI).

@router.get("/audio/meta")
async def get_audio_meta():
    """Sidecar _sfx_meta.json filtré aux fichiers présents — {meta:{<fn>:
    {prompt?, kind, created, …}}}. Alimente tags + recherche du tiroir Sons."""
    from app.services import sfx_service
    meta = await asyncio.get_running_loop().run_in_executor(
        None, sfx_service.known_meta)
    return {"meta": meta}


@router.post("/audio/sfx")
async def generate_sfx_audio(request: Request):
    """Génération de bruitages ElevenLabs (sound-generation).

    Body: {prompt: str ≤ 450, duration_s: 0.5–22 | null (auto),
    prompt_influence: 0..1 = 0.3, variations: 1..4 = 1}. Variations =
    appels séquentiels ; chaque .mp3 atterrit dans le dossier audio de la
    Bibliothèque (sidecar kind « sfx ») → {ok, items:[{filename, url, name,
    size_kb, dur}], warning?}. Erreurs provider → 4xx {detail:
    « ElevenLabs: … »} (toasts fournisseur)."""
    try:
        payload = await request.json()
    except Exception:
        payload = {}
    prompt = str(payload.get("prompt") or "").strip()
    if not prompt:
        raise HTTPException(400, "Décris le son à générer (prompt vide).")
    if len(prompt) > 450:
        raise HTTPException(400, "Prompt trop long (450 caractères max).")
    raw_dur = payload.get("duration_s")
    duration_s = None
    if raw_dur not in (None, "", 0, "auto"):
        try:
            duration_s = float(raw_dur)
        except (TypeError, ValueError):
            raise HTTPException(400, "duration_s invalide (0.5–22 s ou null).")
    try:
        prompt_influence = float(payload.get("prompt_influence", 0.3))
    except (TypeError, ValueError):
        prompt_influence = 0.3
    try:
        variations = int(payload.get("variations", 1))
    except (TypeError, ValueError):
        variations = 1
    from app.services import sfx_service
    loop = asyncio.get_running_loop()
    try:
        items, warning = await loop.run_in_executor(
            None, lambda: sfx_service.generate_sfx(
                prompt, duration_s=duration_s,
                prompt_influence=prompt_influence, variations=variations))
    except sfx_service.SfxError as e:
        raise HTTPException(e.status, e.message)
    except Exception as e:
        raise HTTPException(502, f"ElevenLabs: {str(e)[:300]}")
    out = {"ok": True, "items": items}
    if warning:
        out["warning"] = warning
    return out


@router.get("/music-models")
async def list_music_models():
    """Modèles de musique fal.ai + ambiances proposées. `enabled` dit si la
    clé fal est configurée — l'UI grise la génération plutôt que de laisser
    l'utilisateur découvrir l'échec après coup."""
    from app.services import music_service
    return music_service.catalog()


@router.post("/audio/music")
async def generate_music_audio(request: Request):
    """Génération de musique via fal.ai (même clé que la vidéo).

    Body: {model?: id, prompt: str, mood?: id, duration_s?: int,
    instrumental?: bool, lyrics?: str, seed?: int}. La piste rejoint le
    dossier audio de la Bibliothèque (sidecar kind « musique »), donc le
    tiroir Sons du Montage et le sélecteur de piste de fond la voient
    immédiatement. `notes` liste ce que le modèle choisi ne sait pas faire."""
    try:
        payload = await request.json()
    except Exception:
        payload = {}
    from app.services import music_service
    try:
        return await music_service.generate_music(payload or {})
    except music_service.MusicError as e:
        raise HTTPException(e.status, e.message)
    except Exception as e:
        raise HTTPException(502, f"fal.ai: {str(e)[:300]}")


@router.post("/audio/audition")
async def audition_audio(request: Request):
    """Aperçu « rendu » d'un extrait audio traité — parité ffmpeg du Rack.

    Body: {filename (dossier audio) | job_id (son d'un plan), src_in: 0,
    len: ≤ 12 s, gain_db: 0, speed: 1, fx: [...] (vocabulaire sfx_service)}.
    → WAV 44.1 k stéréo (FileResponse, fichier temporaire nettoyé après
    envoi). -ss avant -i + -t + aucun décodage vidéo : < ~2 s."""
    try:
        payload = await request.json()
    except Exception:
        payload = {}
    from app.services import sfx_service
    src: Path | None = None
    fn = str(payload.get("filename") or "").strip()
    if fn:
        p = _audio_dir() / Path(fn).name
        if not p.is_file():
            raise HTTPException(404, f"Audio introuvable : {fn}")
        src = p
    elif payload.get("job_id"):
        from app.services.storage import JobRecord, async_session_factory
        async with async_session_factory() as session:
            jr = await session.get(JobRecord, str(payload["job_id"]))
        fp = jr and (jr.final_video_path or jr.video_path)
        if not fp or not Path(fp).exists():
            raise HTTPException(404, "Rendu introuvable pour ce job_id.")
        src = Path(fp)
    else:
        raise HTTPException(400, "filename (ou job_id) requis.")

    def _f(key, dv, lo, hi):
        try:
            v = float(payload.get(key, dv))
        except (TypeError, ValueError):
            return dv
        return dv if v != v else max(lo, min(hi, v))
    src_in = _f("src_in", 0.0, 0.0, 36000.0)
    length = _f("len", 6.0, 0.1, 12.0)
    gain_db = _f("gain_db", 0.0, -24.0, 12.0)
    speed = sfx_service.clamp_speed(payload.get("speed"))
    fx = sfx_service.sanitize_fx(payload.get("fx"), "audition")

    from app.services.montage_service import _has_audio_stream
    loop = asyncio.get_running_loop()
    if not await loop.run_in_executor(None, _has_audio_stream, src):
        raise HTTPException(400, "Cette source n'a pas de piste audio.")
    out = (settings.outputs_path / "audio"
           / f"audition_{uuid4().hex[:10]}.wav")
    cmd = sfx_service.build_audition_command(
        src, out, src_in=src_in, length=length, gain_db=gain_db,
        speed=speed, fx=fx)

    def _run():
        import subprocess
        return subprocess.run(cmd, capture_output=True, text=True, timeout=30)
    try:
        r = await loop.run_in_executor(None, _run)
    except Exception as e:
        raise HTTPException(502, f"Audition impossible : {e}")
    if r.returncode != 0 or not out.is_file() or out.stat().st_size == 0:
        out.unlink(missing_ok=True)
        raise HTTPException(502, "Audition échouée : "
                                 f"{(r.stderr or '')[-300:]}")
    from starlette.background import BackgroundTask
    return FileResponse(out, media_type="audio/wav",
                        filename="audition.wav",
                        background=BackgroundTask(
                            lambda: out.unlink(missing_ok=True)))


@router.get("/audio/{filename}")
async def get_audio_file(filename: str):
    safe = Path(filename).name
    p = _audio_dir() / safe
    try:
        if not str(p.resolve()).startswith(str(_audio_dir().resolve())) \
                or not p.is_file():
            raise HTTPException(404, f"Audio not found: {filename}")
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(404, f"Audio not found: {filename}")
    return FileResponse(p)


@router.delete("/audio/{filename}")
async def delete_audio_file(filename: str):
    safe = Path(filename).name
    p = _audio_dir() / safe
    if not p.is_file():
        raise HTTPException(404, f"Audio not found: {filename}")
    p.unlink()
    return {"deleted": safe}


def _clean_vo_error(e: Exception) -> str:
    """W-b — message actionnable pour les échecs fournisseur fréquents :
    402 (voix library/community sur un plan sans crédit) et 403 (clé
    restreinte). Le détail brut du provider reste joint, tronqué."""
    msg = str(e)
    low = msg.lower()
    if ("payment_required" in low or "paid_plan" in low
            or "status_code: 402" in low):
        return ("Voiceover failed: cette voix exige un plan/crédit ElevenLabs "
                "(402) — choisis une voix premade du catalogue ou recharge le "
                f"compte. Détail : {msg[:200]}")
    if "status_code: 403" in low or "forbidden" in low:
        return ("Voiceover failed: accès refusé par ElevenLabs (403) — vérifie "
                "la clé et ses permissions (Réglages → Clés). "
                f"Détail : {msg[:200]}")
    return f"Voiceover failed: {msg}"


@router.post("/audio/voiceover")
async def create_voiceover(request: Request):
    """Synthesize a voiceover (provider-aware) and save it as a reusable audio asset.

    Used by Quick's « Voice Over » tab, the Studio Voiceover node and the
    Chapitres flow: the script is spoken by the active voice provider and the
    .mp3 lands in the Library audio dir, selectable in audio nodes.
    Body: {script, language?: "en"|"fr", name?, voice_id?, model?, settings?}
    — voice_id omitted = default voice from .env (ELEVENLABS_VOICE_ID_{EN,FR}
    per language) ; model omitted = ELEVENLABS_MODEL (W-b, catalogue
    /api/voice-models) ; settings = {stability, similarity_boost, style,
    speed} clampés/filtrés serveur selon le modèle.
    """
    try:
        payload = await request.json()
    except Exception:
        payload = {}
    script = (payload.get("script") or "").strip()
    if not script:
        raise HTTPException(400, "Empty script")
    from app.services.elevenlabs_service import VoiceoverService
    voice = VoiceoverService()
    loop = asyncio.get_running_loop()
    if not await loop.run_in_executor(None, VoiceoverService.is_enabled):
        raise HTTPException(400, "Aucune voix disponible — configure la clé "
                                 "ElevenLabs ou lance Voicebox (Réglages).")
    voice_id = (payload.get("voice_id") or "").strip() or None
    model = (payload.get("model") or "").strip() or None
    v_settings = payload.get("settings")
    if not isinstance(v_settings, dict):
        v_settings = None
    lang = str(payload.get("language") or "en").lower()
    if lang not in ("en", "fr"):
        lang = "en"
    base = re.sub(r"[^A-Za-z0-9_-]+", "_", str(payload.get("name") or "narration")).strip("_")[:40]
    fn = f"{base or 'narration'}-{random.randint(100000, 999999)}.mp3"
    dest = _audio_dir() / fn
    try:
        loop = asyncio.get_running_loop()
        await loop.run_in_executor(
            None, lambda: voice.generate_long(text=script, output_path=dest,
                                              language=lang, voice_id=voice_id,
                                              model_id=model,
                                              settings_override=v_settings))
    except ValueError as e:                       # modèle inconnu → 400 propre
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(502, _clean_vo_error(e))
    if not dest.is_file():
        raise HTTPException(502, "Voiceover produced no file")
    # R1 — sidecar meta (tiroir Sons) : kind « voix » + début du script.
    from app.services import sfx_service
    await loop.run_in_executor(
        None, lambda: sfx_service.record_meta(fn, {
            "kind": "voix", "prompt": script[:200],
            "created": datetime.now().isoformat(timespec="seconds")}))
    return {"ok": True, "filename": fn, "url": f"/api/audio/{fn}",
            "size_kb": dest.stat().st_size // 1024}


@router.get("/voice-models")
async def list_voice_models():
    """W-b — modèles TTS ElevenLabs pour Quick Voice Over + nœud Voiceover.

    Chaque entrée du catalogue sort avec `available` (clé ElevenLabs
    présente), `max_chars` (borne le textarea et le chunking), la liste
    `settings` des curseurs supportés par le modèle, et `usd_per_char`
    (tarif × multiplicateur, overrides pricing.json honorés) pour le coût
    affiché. Les Chapitres restent au défaut app (`default`)."""
    from app.services.elevenlabs_service import ELEVEN_MODELS, default_model_id
    from app.services import pricing as _pricing
    p = _pricing.load()
    out = []
    for mid, m in ELEVEN_MODELS.items():
        out.append({
            "id": mid,
            "label": m["label"],
            "max_chars": m["max_chars"],
            "settings": list(m["settings"]),
            "mult": _pricing.elevenlabs_mult(mid, p),
            "usd_per_char": _pricing.elevenlabs_rate(mid, p),
            "available": settings.has_voiceover,
        })
    return {"models": out, "default": default_model_id()}


@router.get("/voices")
async def list_voices():
    """Voice picker (Episodes / VO) — catalogue du provider de voix actif
    (ElevenLabs ou Voicebox local, v1.26 étape 3). Même forme qu'avant :
    {voice_id, name, category, language, labels, preview_url}."""
    try:
        provider, voices = await _fetch_casting_voices()
    except HTTPException:                  # aucun provider utilisable
        return {"voices": [], "enabled": False}
    except Exception as e:
        logger.warning(f"voices fetch failed: {e}")
        return {"voices": [], "enabled": True, "error": str(e)}
    out = [{**v, "language": (v.get("labels") or {}).get("language")
            or (v.get("labels") or {}).get("accent")} for v in voices]
    return {"voices": out, "enabled": True, "provider": provider}


def _extract_chapter_text(name: str, data: bytes) -> str:
    """Plain text of a .txt / .docx / .pdf chapter. Blocking (pypdf on a 25 MB
    scan takes seconds) — call it off the event loop."""
    import io as _io
    if name.endswith(".docx"):
        import docx
        doc = docx.Document(_io.BytesIO(data))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(_io.BytesIO(data))
        parts = []
        for pg in reader.pages:
            t = (pg.extract_text() or "").strip()
            if t:
                parts.append(t)
        return "\n\n".join(parts)
    for enc in ("utf-8", "utf-8-sig", "latin-1"):  # .txt or unknown
        try:
            return data.decode(enc)
        except Exception:
            continue
    return ""


@router.post("/episodes/extract-text")
async def extract_chapter_text(file: UploadFile = File(...)):
    """Extract plain text from an uploaded chapter file (.txt / .docx / .pdf)
    for the Episodes narration. Returns {text, words, chars}."""
    name = (file.filename or "").lower()
    data = await file.read()
    if len(data) > 25 * 1024 * 1024:
        raise HTTPException(400, "File too large (max 25 MB)")
    try:
        text = await asyncio.to_thread(_extract_chapter_text, name, data)
    except Exception as e:
        raise HTTPException(422, f"Could not read this file: {e}")
    text = (text or "").strip()
    if not text:
        raise HTTPException(
            422, "No selectable text found (a scanned PDF is an image — export a text PDF).")
    return {"text": text, "words": len(text.split()), "chars": len(text)}


def _scene_prompt_from(text: str, limit: int = 160) -> str:
    """Rough illustration prompt for paragraph mode: the first sentence."""
    m = re.split(r"(?<=[.!?…])\s", text.strip(), maxsplit=1)
    s = (m[0] if m else text).strip()
    return (s[:limit].rstrip() + "…") if len(s) > limit else s


def _paragraph_scenes(script: str) -> list[dict]:
    paras = [p.strip() for p in re.split(r"\n\s*\n", script) if p.strip()]
    if len(paras) <= 1:  # no blank-line paragraphs → fall back to single lines
        paras = [p.strip() for p in script.splitlines() if p.strip()]
    return [{"text": p, "illustration_prompt": _scene_prompt_from(p)}
            for p in paras[:60]]


def _ai_scenes(script: str, lang: str, sujet_seul: bool = False) -> list[dict]:
    """sujet_seul=True (option de style du chantier vitrail 27/08): l'agent
    écrit des prompts SUJET — personnages, décor, action, ambiance — SANS
    vocabulaire de style ni thème de marque; le bloc de style est appliqué
    ensuite, de façon déterministe, par la route."""
    from app.services.summarizer import _chat_dispatch
    langname = "French" if lang.startswith("fr") else "English"
    n = max(3, min(12, len(script.split()) // 80 + 1))
    if sujet_seul:
        system = ("You are a storyboard director. You split a narrated novel "
                  "chapter into visual scenes and write a concise "
                  "subject-focused image prompt for each. Return ONLY valid "
                  "JSON.")
        consigne = (f"\"illustration_prompt\" = a concise subject-focused "
                    f"image prompt in {langname} (characters, setting, "
                    f"action, mood) WITHOUT any style vocabulary — no "
                    f"medium, palette or art-movement words, the visual "
                    f"style is applied separately")
    else:
        system = ("You are a storyboard director for DEEPOTUS, a deep-sea / abyssal "
                  "themed brand. You split a narrated novel chapter into visual scenes "
                  "and write a vivid image prompt for each. Return ONLY valid JSON.")
        consigne = (f"\"illustration_prompt\" = a vivid cinematic image "
                    f"prompt in {langname} (deep-sea, bioluminescent, "
                    f"atmospheric)")
    prompt = (
        f"Split this chapter into about {n} sequential scenes for a narrated video. "
        f"For each scene return: \"text\" = the chapter text for that scene, COPIED "
        f"VERBATIM and in order so concatenating all texts reproduces the chapter; "
        f"and {consigne}. Return ONLY a JSON array "
        f"[{{\"text\":\"...\",\"illustration_prompt\":\"...\"}}].\n\nChapter:\n{script[:12000]}")
    out, _prov = _chat_dispatch(prompt, system, 4000)
    if not out:
        return []
    txt = out.strip()
    if txt.startswith("```"):
        txt = re.sub(r"^```[a-zA-Z]*\n?", "", txt)
        txt = re.sub(r"\n?```$", "", txt).strip()
    i, j = txt.find("["), txt.rfind("]")
    if i >= 0 and j > i:
        txt = txt[i:j + 1]
    try:
        data = json.loads(txt)
    except Exception:
        return []
    scenes = []
    for it in (data if isinstance(data, list) else []):
        if isinstance(it, dict) and str(it.get("text") or "").strip():
            scenes.append({"text": str(it["text"]).strip(),
                           "illustration_prompt": str(it.get("illustration_prompt") or "").strip()})
    return scenes


@router.post("/episodes/scenes")
async def episode_scenes(request: Request):
    """Split a chapter into scenes for the storyboard.
    Body: {script, language?, method:"paragraph"|"ai"}.
    Returns {scenes:[{text, illustration_prompt}], method, count}."""
    try:
        payload = await request.json()
    except Exception:
        payload = {}
    script = (payload.get("script") or "").strip()
    if not script:
        raise HTTPException(400, "Empty script")
    method = (payload.get("method") or "paragraph").lower()
    lang = str(payload.get("language") or "en").lower()
    # option de style (chantier vitrail 27/08): les prompts d'illustration
    # naissent SUJET (l'agent n'écrit pas de style) puis le bloc de la
    # famille est appliqué ici, de façon déterministe — les prompts restent
    # éditables dans les cartes de scène.
    style = (payload.get("style") or "").strip().lower()
    if style:
        from app.services import style_vitrail as SV
        try:
            SV.bloc_style(style)         # valide la famille avant tout travail
        except KeyError as e:
            raise HTTPException(400, f"Style inconnu: {e}")
        except FileNotFoundError as e:
            raise HTTPException(400, f"Style '{style}': {e}")

    def _styliser(scenes: list[dict]) -> list[dict]:
        if not style:
            return scenes
        from app.services import style_vitrail as SV
        for s in scenes:
            p = (s.get("illustration_prompt") or "").strip()
            if not p:
                continue
            try:
                s["illustration_prompt"] = SV.appliquer(SV.epurer_noms(p),
                                                        style)
            except ValueError:           # le prompt n'était QU'un nom épuré
                s["illustration_prompt"] = ""
        return scenes

    if method == "ai":
        from app.services.summarizer import available
        if not available():
            return {"scenes": [], "method": "ai",
                    "error": "Aucun LLM configuré (Réglages → clés API). Utilise le découpage par paragraphe."}
        loop = asyncio.get_running_loop()
        scenes = await loop.run_in_executor(
            None, lambda: _ai_scenes(script, lang, sujet_seul=bool(style)))
        if not scenes:
            return {"scenes": [], "method": "ai",
                    "error": "Le découpage IA a échoué — réessaie, ou utilise les paragraphes."}
        return {"scenes": _styliser(scenes), "method": "ai",
                "count": len(scenes)}
    scenes = _styliser(_paragraph_scenes(script))
    return {"scenes": scenes, "method": "paragraph", "count": len(scenes)}


@router.post("/episodes/render")
async def render_episode(request: Request, background_tasks: BackgroundTasks):
    """Assemble a narrated illustrated episode (per-scene TTS narration + Ken
    Burns / still over each scene's image, concatenated into one 9:16 video).
    Returns a job_id; poll GET /api/jobs/{job_id}.
    Body: {title, voice_id, language, scenes:[{text, image_filename, motion}]}."""
    try:
        payload = await request.json()
    except Exception:
        payload = {}
    scenes = payload.get("scenes") or []
    if not isinstance(scenes, list) or not scenes:
        raise HTTPException(400, "No scenes to render")
    if not any((s.get("text") or "").strip() for s in scenes if isinstance(s, dict)):
        raise HTTPException(400, "Scenes have no narration text")
    from app.services.elevenlabs_service import VoiceoverService
    if not await asyncio.get_running_loop().run_in_executor(
            None, VoiceoverService.is_enabled):
        raise HTTPException(400, "Aucune voix disponible — configure la clé "
                                 "ElevenLabs ou lance Voicebox (Réglages).")
    job_id = str(uuid4())

    async def _run():
        try:
            await pipeline.run_episode(
                job_id=job_id, title=payload.get("title"),
                voice_id=(payload.get("voice_id") or "").strip() or None,
                language=str(payload.get("language") or "en"),
                scenes=scenes)
        except Exception as e:
            logger.exception(f"Episode render {job_id} failed: {e}")

    background_tasks.add_task(_run)
    return {"ok": True, "job_id": job_id,
            "message": f"Episode render queued. Poll GET /api/jobs/{job_id}."}


@router.post("/videos/upload")
async def upload_video(file: UploadFile = File(...)):
    """Upload a user-shot video (UGC — e.g. a phone selfie clip).

    Stored under outputs/uploads and registered as a FINISHED job, so it:
      - shows up immediately in the Library (Renders tab),
      - can be attached to a scheduled post,
      - can be dropped into the Studio as a source clip,
      - and — crucially — its real (ffprobe) duration can drive the MASTER
        duration of a composition: a layout whose audio.master_track points at
        the UGC slot renders to the UGC length, so generated animations
        (Seedance) are calibrated around the real human clip.
    """
    from datetime import datetime as _dtu
    from app.services.storage import JobRecord, async_session_factory

    base = file.filename or "ugc.mp4"
    safe = "".join(c for c in base if c.isalnum() or c in "._- ").strip() or "ugc.mp4"
    if not safe.lower().endswith((".mp4", ".mov", ".webm", ".m4v", ".avi")):
        safe += ".mp4"
    folder = settings.outputs_path / "uploads"
    folder.mkdir(parents=True, exist_ok=True)
    dest = folder / safe
    stem, ext = dest.stem, dest.suffix
    n = 1
    while dest.exists():
        dest = folder / f"{stem}_{n}{ext}"
        n += 1
    contents = await file.read()
    dest.write_bytes(contents)

    dur = await asyncio.to_thread(_probe_seconds, str(dest)) or 0.0
    job_id = str(uuid4())
    async with async_session_factory() as session:
        session.add(JobRecord(
            id=job_id,
            status=JobStatus.DONE.value,
            progress=100,
            title=dest.stem,
            image_filename=dest.name,        # column is non-null; use the file name
            final_video_path=str(dest),
            video_path=str(dest),
            duration_s=int(round(dur)) if dur else None,
            aspect_ratio="9:16",
            provider="ugc",
            current_step="Uploaded",
            completed_at=_dtu.utcnow(),
        ))
        await session.commit()
    return {
        "ok": True, "job_id": job_id, "filename": dest.name,
        "duration_s": round(dur, 2), "final_video_path": str(dest),
    }


# ---- Prompt preview & builder ----

@router.post("/prompt/preview")
async def preview_prompt(request: GenerateRequest):
    try:
        prompt, negative = pipeline.engine.build_prompt(request)
        caption = pipeline.engine.build_caption(request)
        vo = pipeline.engine.build_voiceover_script(request)
        return {
            "prompt": prompt,
            "negative_prompt": negative,
            "caption": caption,
            "voiceover_script": vo,
        }
    except Exception as e:
        raise HTTPException(400, str(e))


@router.post("/prompt/build", response_model=BuildPromptResponse)
async def build_prompt_from_intent(request: BuildPromptRequest):
    """Generate a Seedance prompt from free-text keywords/intent.
    Injects deepotus DNA and structures the output for Seedance 2.0.
    """
    try:
        return pipeline.engine.generate_from_intent(request)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        logger.exception("Prompt builder failed")
        raise HTTPException(500, f"Builder error: {e}")


# ---- Generate ----

@router.post("/generate", response_model=GenerateResponse)
async def generate(request: GenerateRequest, background_tasks: BackgroundTasks):
    # W-a — the required key depends on the selected model's provider
    from app.services.fal_service import VIDEO_MODELS, DEFAULT_VIDEO_MODEL
    _mdl = VIDEO_MODELS.get(request.video_model or DEFAULT_VIDEO_MODEL)
    if _mdl and _mdl["provider"] == "google":
        if not settings.has_gemini:
            raise HTTPException(400, "GEMINI_API_KEY not configured. "
                                     "Add it in Settings -> Keys")
    elif not settings.FAL_KEY:
        raise HTTPException(400, "FAL_KEY not configured. Add it to backend/.env")

    image_path = settings.images_path / request.image_filename
    if not image_path.exists():
        raise HTTPException(404, f"Start image not found: {request.image_filename}")

    if request.image_filename_end:
        end_path = settings.images_path / request.image_filename_end
        if not end_path.exists():
            raise HTTPException(404, f"End image not found: {request.image_filename_end}")

    if not request.template_id and not request.custom_prompt:
        raise HTTPException(400, "Must provide either template_id or custom_prompt")

    async def _run():
        try:
            await pipeline.run(request)
        except Exception as e:
            logger.error(f"Background pipeline error: {e}")

    background_tasks.add_task(_run)

    return GenerateResponse(
        job_id="pending",
        status=JobStatus.QUEUED,
        message="Job queued. Poll GET /jobs to see latest status.",
    )


# v1.3: Batch generate — N variations with offset seeds
@router.post("/generate/batch", response_model=GenerateBatchResponse)
async def generate_batch(request: GenerateBatchRequest, background_tasks: BackgroundTasks):
    """Queue N jobs sharing the same config but with offset seeds.

    Behavior:
    - If request.seed is provided, seeds = [seed, seed+1, ..., seed+N-1].
    - If request.seed is None, a random base seed is generated and used.
    - All N jobs share a batch_id (returned) so the UI can group them.
    """
    if not settings.FAL_KEY:
        raise HTTPException(400, "FAL_KEY not configured. Add it to backend/.env")

    image_path = settings.images_path / request.image_filename
    if not image_path.exists():
        raise HTTPException(404, f"Start image not found: {request.image_filename}")

    if request.image_filename_end:
        end_path = settings.images_path / request.image_filename_end
        if not end_path.exists():
            raise HTTPException(404, f"End image not found: {request.image_filename_end}")

    if not request.template_id and not request.custom_prompt:
        raise HTTPException(400, "Must provide either template_id or custom_prompt")

    if request.variations_count < 1 or request.variations_count > 8:
        raise HTTPException(400, "variations_count must be between 1 and 8")

    # Determine base seed
    base_seed = request.seed if request.seed is not None else random.randint(1, 2_000_000_000)
    seeds = [base_seed + i for i in range(request.variations_count)]
    batch_id = str(uuid4())

    # Build per-variation requests
    base_dict = request.model_dump(exclude={"variations_count", "seed"})

    async def _run_one(variation_seed: int, idx: int):
        try:
            sub_req = GenerateRequest(**base_dict, seed=variation_seed)
            await pipeline.run(
                sub_req,
                batch_id=batch_id,
                batch_index=idx,
                batch_size=request.variations_count,
            )
        except Exception as e:
            logger.error(f"Batch {batch_id} variation {idx} failed: {e}")

    for idx, s in enumerate(seeds):
        background_tasks.add_task(_run_one, s, idx)

    return GenerateBatchResponse(
        batch_id=batch_id,
        job_count=request.variations_count,
        base_seed=base_seed,
        seeds=seeds,
        message=f"Queued {request.variations_count} variations with seeds {seeds[0]}-{seeds[-1]}.",
    )


# ============ HEYGEN ENDPOINTS (v1.4) ============

@router.get("/heygen/health")
async def heygen_health():
    """Check whether HeyGen is configured and reachable.

    Uses the lightweight /v2/user/remaining_quota probe (sub-second) rather
    than listing avatars (/v2/avatars can take 60s+), so the status badge
    stays responsive and a slow avatar catalogue never makes HeyGen look
    'unreachable'.
    """
    if not settings.has_heygen:
        return {"configured": False, "reachable": False,
                "message": "HEYGEN_API_KEY not set in backend/.env"}
    try:
        client = HeyGenClient()
        quota = await client.remaining_quota()
        rem = quota.get("remaining_quota") if isinstance(quota, dict) else None
        msg = "OK -- key valid"
        if rem is not None:
            msg += f", {rem} credits remaining"
        return {"configured": True, "reachable": True,
                "remaining_quota": rem, "message": msg}
    except HeyGenError as e:
        return {"configured": True, "reachable": False, "message": str(e)}
    except Exception as e:
        return {"configured": True, "reachable": False, "message": f"Network error: {e}"}


@router.get("/heygen/avatars")
async def list_heygen_avatars():
    """List avatars available on your HeyGen account.

    First load can take up to ~2 min (HeyGen's /v2/avatars is slow for large
    catalogues); the result is cached so later loads are instant.
    """
    if not settings.has_heygen:
        raise HTTPException(400, "HEYGEN_API_KEY not configured")
    try:
        client = HeyGenClient()
        avatars = await client.list_avatars()
        return {"count": len(avatars), "avatars": avatars}
    except HeyGenError as e:
        raise HTTPException(502, f"HeyGen error: {e}")
    except (httpx.TimeoutException, httpx.HTTPError) as e:
        raise HTTPException(504, f"HeyGen timed out listing avatars: {e}")


@router.get("/heygen/voices")
async def list_heygen_voices():
    """List voices available on your HeyGen account."""
    if not settings.has_heygen:
        raise HTTPException(400, "HEYGEN_API_KEY not configured")
    try:
        client = HeyGenClient()
        voices = await client.list_voices()
        return {"count": len(voices), "voices": voices}
    except HeyGenError as e:
        raise HTTPException(502, f"HeyGen error: {e}")
    except (httpx.TimeoutException, httpx.HTTPError) as e:
        raise HTTPException(504, f"HeyGen timed out listing voices: {e}")


@router.get("/heygen/presets")
async def list_avatar_presets():
    """List saved avatar+voice casting presets (newest first)."""
    from app.services.storage import AvatarPreset, async_session_factory
    from sqlalchemy import select
    async with async_session_factory() as session:
        rows = (await session.execute(
            select(AvatarPreset).order_by(AvatarPreset.created_at.desc())
        )).scalars().all()
    return {"presets": [{
        "id": p.id, "name": p.name,
        "avatar_id": p.avatar_id, "avatar_type": p.avatar_type,
        "avatar_img": p.avatar_img,
        "voice_id": p.voice_id, "voice_name": p.voice_name,
        "voice_prev": p.voice_prev, "voice_lang": p.voice_lang,
        "speed": p.speed, "engine": p.engine or "",
        "created_at": p.created_at.isoformat() if p.created_at else None,
    } for p in rows]}


@router.post("/heygen/presets")
async def create_avatar_preset(body: AvatarPresetCreate):
    """Save an avatar+voice casting preset."""
    from app.services.storage import AvatarPreset, async_session_factory
    from uuid import uuid4
    pid = str(uuid4())
    async with async_session_factory() as session:
        session.add(AvatarPreset(
            id=pid, name=body.name.strip(),
            avatar_id=body.avatar_id, avatar_type=body.avatar_type,
            avatar_img=body.avatar_img,
            voice_id=body.voice_id, voice_name=body.voice_name,
            voice_prev=body.voice_prev, voice_lang=body.voice_lang,
            speed=body.speed, engine=(body.engine or None),
        ))
        await session.commit()
    return {"id": pid, "name": body.name.strip(),
            "avatar_id": body.avatar_id, "avatar_type": body.avatar_type,
            "avatar_img": body.avatar_img,
            "voice_id": body.voice_id, "voice_name": body.voice_name,
            "voice_prev": body.voice_prev, "voice_lang": body.voice_lang,
            "speed": body.speed, "engine": body.engine or ""}


@router.delete("/heygen/presets/{preset_id}")
async def delete_avatar_preset(preset_id: str):
    """Delete a casting preset by id."""
    from app.services.storage import AvatarPreset, async_session_factory
    async with async_session_factory() as session:
        row = await session.get(AvatarPreset, preset_id)
        if not row:
            raise HTTPException(404, "Preset not found")
        await session.delete(row)
        await session.commit()
    return {"ok": True}


@router.post("/generate/heygen")
async def generate_heygen(request: GenerateHeyGenRequest, background_tasks: BackgroundTasks):
    """Queue a HeyGen avatar video generation."""
    if not settings.has_heygen:
        raise HTTPException(400, "HEYGEN_API_KEY not configured. Add it to backend/.env")
    if not request.script.strip():
        raise HTTPException(400, "Script must not be empty")

    async def _run():
        try:
            await pipeline.run_heygen(request)
        except Exception as e:
            logger.error(f"Background HeyGen pipeline error: {e}")

    background_tasks.add_task(_run)
    return GenerateResponse(
        job_id="pending",
        status=JobStatus.QUEUED,
        message="HeyGen job queued. Poll GET /jobs to see status.",
    )


@router.post("/generate/heygen-image")
async def generate_heygen_image(request: GenerateHeyGenImageRequest,
                                background_tasks: BackgroundTasks):
    """v1.16 (D) — animate a Library still into a talking video (HeyGen v3)."""
    if not settings.has_heygen:
        raise HTTPException(400, "HEYGEN_API_KEY not configured. Add it to backend/.env")
    if not request.script.strip():
        raise HTTPException(400, "Script must not be empty")
    img = settings.images_path / request.image_filename
    if not img.exists():
        raise HTTPException(404, f"Image not found in Library: {request.image_filename}")

    async def _run():
        try:
            await pipeline.run_heygen_image(request)
        except Exception as e:
            logger.error(f"Background HeyGen-image pipeline error: {e}")

    background_tasks.add_task(_run)
    return GenerateResponse(job_id="pending", status=JobStatus.QUEUED,
                            message="HeyGen image-animation job queued. Poll GET /jobs.")


@router.post("/generate/heygen-cinematic")
async def generate_heygen_cinematic(request: GenerateHeyGenCinematicRequest,
                                    background_tasks: BackgroundTasks):
    """v1.16 (D) — HeyGen v3 cinematic avatar (prompt-driven, 1–3 looks)."""
    if not settings.has_heygen:
        raise HTTPException(400, "HEYGEN_API_KEY not configured. Add it to backend/.env")
    missing = [f for f in (request.reference_images or [])
               if not (settings.images_path / f).exists()]
    if missing:
        raise HTTPException(404, f"Reference image(s) not in Library: {missing}")

    async def _run():
        try:
            await pipeline.run_heygen_cinematic(request)
        except Exception as e:
            logger.error(f"Background HeyGen-cinematic pipeline error: {e}")

    background_tasks.add_task(_run)
    return GenerateResponse(job_id="pending", status=JobStatus.QUEUED,
                            message="HeyGen cinematic job queued. Poll GET /jobs.")


@router.post("/generate/composition", response_model=CompositionResponse)
async def generate_composition(request: CompositionRequest, background_tasks: BackgroundTasks):
    """Queue a composition job: Seedance clip + HeyGen clip combined.

    Layout: sequential, split_vstack, or split_hstack.
    Both clips are generated in parallel, then composed via ffmpeg.
    """
    if not settings.FAL_KEY:
        raise HTTPException(400, "FAL_KEY not configured")
    if not settings.has_heygen:
        raise HTTPException(400, "HEYGEN_API_KEY not configured")

    # Validate Seedance side
    img_path = settings.images_path / request.seedance.image_filename
    if not img_path.exists():
        raise HTTPException(404, f"Seedance start image not found: {request.seedance.image_filename}")
    if not request.seedance.template_id and not request.seedance.custom_prompt:
        raise HTTPException(400, "Seedance side needs template_id or custom_prompt")

    # Validate HeyGen side
    if not request.heygen.script.strip():
        raise HTTPException(400, "HeyGen script must not be empty")
    if not request.heygen.avatar_id or not request.heygen.voice_id:
        raise HTTPException(400, "HeyGen avatar_id and voice_id are required")

    async def _run():
        try:
            await pipeline.run_composition(request)
        except Exception as e:
            logger.error(f"Background composition pipeline error: {e}")

    background_tasks.add_task(_run)
    return CompositionResponse(
        composition_id="pending",
        job_id="pending",
        message=f"Composition queued ({request.layout.value}). Poll GET /jobs.",
    )


# ============ v1.5: PHOTO AVATAR UPLOAD ============

@router.post("/heygen/photo-avatar/create", response_model=PhotoAvatarCreateResponse)
async def create_photo_avatar_endpoint(
    file: UploadFile = File(...),
    avatar_name: str = "Custom deepotus avatar",
    group_name: str = "",
    do_train: bool = True,
):
    """Upload an image and create a HeyGen photo avatar from it.

    Flow:
      1. Save uploaded file to a temp location
      2. Call HeyGenClient.create_photo_avatar() which:
         - uploads image to HeyGen storage
         - creates an avatar group
         - adds the photo as a look
         - polls until ready (5-30s typical)
         - optionally triggers training (do_train=True)
      3. Returns the photo_avatar_id usable as a talking_photo in video generation.
    """
    if not settings.has_heygen:
        raise HTTPException(400, "HEYGEN_API_KEY not configured")

    allowed = {".png", ".jpg", ".jpeg", ".webp"}
    suffix = Path(file.filename).suffix.lower() if file.filename else ".png"
    if suffix not in allowed:
        raise HTTPException(400, f"Unsupported file type. Use one of: {', '.join(allowed)}")

    # Save to a temp file in images_path/_avatar_uploads (auto-created)
    tmp_dir = settings.images_path / "_avatar_uploads"
    tmp_dir.mkdir(parents=True, exist_ok=True)
    safe_name = "".join(c for c in (file.filename or "avatar.png") if c.isalnum() or c in "._-") or "avatar.png"
    tmp_path = tmp_dir / f"{uuid4().hex[:8]}_{safe_name}"
    data = await file.read()
    tmp_path.write_bytes(data)
    logger.info(f"Photo avatar upload received: {tmp_path} ({len(data)} bytes)")

    try:
        client = HeyGenClient()
        result = await client.create_photo_avatar(
            file_path=tmp_path,
            avatar_name=avatar_name or "Custom deepotus avatar",
            group_name=group_name or None,
            do_train=do_train,
        )
        # Drop the cached avatar list so the new talking photo shows up on the
        # next /heygen/avatars call instead of waiting for the TTL to expire.
        invalidate_list_cache()
        return PhotoAvatarCreateResponse(
            photo_avatar_id=result["photo_avatar_id"],
            group_id=result["group_id"],
            status=result["status"],
            avatar_name=result["avatar_name"],
            asset_url=result.get("asset_url"),
            message=(
                f"Avatar '{result['avatar_name']}' created and ready. "
                f"Use it in HeyGen mode with avatar_type='talking_photo'."
            ),
        )
    except HeyGenError as e:
        raise HTTPException(502, f"HeyGen error: {e}")
    except Exception as e:
        logger.exception("Photo avatar create failed")
        raise HTTPException(500, f"Photo avatar creation failed: {e}")
    finally:
        # Cleanup the temp file
        try:
            if tmp_path.exists():
                tmp_path.unlink()
        except Exception as e:
            logger.warning(f"Could not clean up temp avatar file: {e}")


# ============ v1.5: UNIVERSAL BUILDER ============

@router.post("/prompt/build-script", response_model=BuildScriptResponse)
async def build_script_endpoint(request: BuildScriptRequest):
    """Generate a HeyGen avatar SCRIPT from a free-text intent.

    Returns: spoken script + suggested caption.
    Different from /prompt/build (which generates VISUAL prompts for Seedance).
    """
    if not request.intent.strip():
        raise HTTPException(400, "Intent must not be empty")
    try:
        return await asyncio.to_thread(
            pipeline.engine.generate_script_from_intent,
            intent=request.intent,
            voice_mode=request.voice_mode,
            language=request.voiceover_language,
            max_words=request.max_words,
            inject_persona=request.inject_persona,
        )
    except Exception as e:
        logger.exception("build_script failed")
        raise HTTPException(500, f"Builder error: {e}")


@router.post("/prompt/refine")
async def refine_text(body: dict):
    """AI-refine a Text node's copy for natural spoken (avatar) delivery.

    Controls: tone, humor, avoid[]. Fail-safe — returns the original text
    with ai=false when no LLM key is configured, so the UI never breaks and
    the user spends nothing.
    """
    text = (body.get("text") or "").strip()
    if not text:
        raise HTTPException(400, "text is required")
    tone = str(body.get("tone") or "").strip()
    humor = str(body.get("humor") or "").strip()
    avoid = body.get("avoid") or []
    mode = str(body.get("mode") or "spoken").strip().lower()
    language = str(body.get("language") or "").strip()
    from app.services import summarizer
    _LMAP = {"FR": "French", "EN": "English", "ES": "Spanish", "DE": "German",
             "IT": "Italian", "PT": "Portuguese", "NL": "Dutch", "JA": "Japanese",
             "KO": "Korean", "ZH": "Chinese", "AR": "Arabic", "HI": "Hindi"}
    if not language:
        lang = "English"
    elif len(language) <= 3:
        lang = _LMAP.get(language.upper(), "English")
    else:
        lang = language  # already a full name from the voice (e.g. "French")
    if mode == "visual":
        # Prompt node: expand an idea into a vivid image/video generation prompt.
        system = (
            "You write prompts for an AI image/video generator (Seedance). "
            "Turn the idea into ONE vivid, concrete visual prompt: subject, "
            "setting, action, camera angle/movement, lighting, color, and art "
            "style. Use comma-separated descriptors, not narration. No "
            "markdown, no preamble, no quotation marks. Return ONLY the prompt."
        )
        parts = ["Expand this into a single cinematic image/video generation prompt."]
        if tone:
            parts.append(f"Visual style: {tone}.")
        if humor and humor.lower() != "none":
            parts.append(f"Mood: {humor}.")
        if isinstance(avoid, list) and avoid:
            parts.append("Do NOT include: " + "; ".join(str(a) for a in avoid) + ".")
        parts.append("Keep it under 60 words.")
        prompt = " ".join(parts) + "\n\nIdea:\n" + text[:2000]
    else:
        system = (
            "You rewrite short scripts that an AI avatar will SPEAK ALOUD. "
            "Optimize for natural, human spoken delivery and clean diction: "
            "short sentences, easy-to-pronounce words, no tongue-twisters, no "
            "markdown, no emojis, no stage directions, no quotation marks, no "
            "preamble. Return ONLY the rewritten script."
        )
        parts = [f"Rewrite this script in {lang} for an avatar to read aloud."]
        if tone:
            parts.append(f"Tone: {tone}.")
        if humor and humor.lower() != "none":
            parts.append(f"Humor: {humor}.")
        if isinstance(avoid, list) and avoid:
            parts.append("Avoid: " + "; ".join(str(a) for a in avoid) + ".")
        parts.append("Keep roughly the same length.")
        prompt = " ".join(parts) + "\n\nScript:\n" + text[:4000]
    try:
        out, prov = await asyncio.to_thread(
            summarizer._chat_dispatch, prompt, system, 800)
    except Exception as e:
        logger.warning(f"refine error: {e}")
        out, prov = None, ""
    if out:
        return {"text": out, "provider": prov, "ai": True}
    return {"text": text, "provider": "", "ai": False}


@router.post("/prompt/build-composition", response_model=BuildCompositionResponse)
async def build_composition_endpoint(request: BuildCompositionRequest):
    """Generate BOTH a Seedance prompt AND a HeyGen script from one intent.

    The two outputs are coherent per the layout:
      - Sequential: avatar SETS UP, Seedance PAYS OFF
      - Split: avatar NARRATES, Seedance SHOWS in parallel
    """
    if not request.intent.strip():
        raise HTTPException(400, "Intent must not be empty")
    try:
        return pipeline.engine.generate_composition_from_intent(
            intent=request.intent,
            layout=request.layout,
            style=request.style,
            aspect_ratio=request.aspect_ratio,
            duration_s=request.duration_s,
            voice_mode=request.voice_mode,
            language=request.voiceover_language,
            max_script_words=request.max_script_words,
            inject_persona=request.inject_persona,
        )
    except Exception as e:
        logger.exception("build_composition failed")
        raise HTTPException(500, f"Builder error: {e}")


# ---- Jobs ----

def _job_to_dict(j) -> dict:
    return {
        "job_id": j.id,
        "status": j.status,
        "progress": j.progress,
        "title": getattr(j, "title", None),
        "current_step": j.current_step,
        "image_filename": j.image_filename,
        "image_filename_end": j.image_filename_end,
        "final_prompt": j.final_prompt,
        "negative_prompt": j.negative_prompt,
        "video_path": j.video_path,
        "audio_path": j.audio_path,
        "final_video_path": j.final_video_path,
        "caption_text": j.caption_text,
        "caption_path": j.caption_path,
        "seed": j.seed,
        "duration_s": j.duration_s,
        "aspect_ratio": j.aspect_ratio,
        "style": j.style,
        "template_id": j.template_id,
        "voiceover_language": j.voiceover_language,
        "voice_mode": j.voice_mode,
        "video_model": getattr(j, "video_model", None),
        "provider": j.provider,
        "composition_id": j.composition_id,
        "composition_layout": j.composition_layout,
        "layer_index": j.layer_index,
        "error": j.error,
        "batch_id": j.batch_id,
        "batch_index": j.batch_index,
        "batch_size": j.batch_size,
        "created_at": j.created_at.isoformat() if j.created_at else None,
        "completed_at": j.completed_at.isoformat() if j.completed_at else None,
    }


@router.get("/jobs")
async def list_jobs(limit: int = 50):
    jobs = await Pipeline.list_jobs(limit=limit)
    return [_job_to_dict(j) for j in jobs]


#: (path, mtime_ns) -> duration. A finished render never changes, and the
#: dock re-polls /jobs/{id} until the user closes it.
_PROBE_CACHE: dict[tuple[str, int], float] = {}


def _probe_seconds(path: str | None) -> float | None:
    """Real media duration of a finished render (ffprobe). Used by the
    timeline 'Fit to avatar' so animation clips can be calibrated to the
    avatar's exact length instead of a guessed target."""
    if not path:
        return None
    import subprocess
    try:
        key = (str(path), Path(path).stat().st_mtime_ns)
    except OSError:
        return None
    if key in _PROBE_CACHE:
        return _PROBE_CACHE[key]
    try:
        # timeout: /jobs/{id} is polled continuously by the dock, so a hung
        # ffprobe on a truncated file must not pin a worker thread forever.
        out = subprocess.run(
            ["ffprobe", "-v", "error", "-show_entries", "format=duration",
             "-of", "csv=p=0", str(path)],
            check=False, capture_output=True, text=True, timeout=15).stdout.strip()
        dur = round(float(out), 3)
    except (ValueError, FileNotFoundError, OSError, subprocess.TimeoutExpired):
        return None
    if len(_PROBE_CACHE) > 512:
        _PROBE_CACHE.clear()
    _PROBE_CACHE[key] = dur
    return dur


@router.get("/jobs/{job_id}")
async def get_job(job_id: str):
    j = await Pipeline.get_job(job_id)
    if not j:
        raise HTTPException(404, "Job not found")
    d = _job_to_dict(j)
    d["duration_real_s"] = await asyncio.to_thread(_probe_seconds, j.final_video_path)
    return d


@router.patch("/jobs/{job_id}")
async def rename_job(job_id: str, request: JobRenameRequest):
    """Rename a render so it's identifiable in the queue and the
    'existing' clip / audio pickers."""
    j = await Pipeline.rename_job(job_id, request.title)
    if not j:
        raise HTTPException(404, "Job not found")
    return _job_to_dict(j)


@router.delete("/jobs/{job_id}")
async def delete_job(job_id: str):
    """Delete job DB record + all files (video, audio, caption)."""
    success = await Pipeline.delete_job(job_id)
    if not success:
        raise HTTPException(404, "Job not found")
    return {"deleted": True, "job_id": job_id}


# v1.3: Bulk delete entire batch
@router.delete("/batches/{batch_id}")
async def delete_batch(batch_id: str):
    """Delete all jobs in a batch and their files."""
    count = await Pipeline.delete_batch(batch_id)
    if count == 0:
        raise HTTPException(404, "Batch not found or empty")
    return {"deleted": True, "batch_id": batch_id, "jobs_deleted": count}


@router.get("/jobs/{job_id}/video")
async def download_job_video(job_id: str):
    j = await Pipeline.get_job(job_id)
    if not j or not j.final_video_path:
        raise HTTPException(404, "Final video not available")
    p = Path(j.final_video_path)
    if not p.exists():
        raise HTTPException(404, "Video file missing on disk")
    return FileResponse(p, media_type="video/mp4", filename=p.name)


# ---- Health ----

@router.get("/health")
async def health():
    from app.services.elevenlabs_service import VoiceoverService
    # provider-aware (v1.26 étape 4) : clé 11L OU Voicebox local joignable
    # (détection cachée 5 s dans voice_providers — pas un ping par poll)
    vo_enabled = await asyncio.get_running_loop().run_in_executor(
        None, VoiceoverService.is_enabled)
    return {
        "ok": True,
        "version": APP_VERSION,
        "telegram_enabled": settings.has_telegram,
        "x_enabled": settings.has_x,
        "ollama_enabled": settings.has_ollama,
        "fal_configured": bool(settings.FAL_KEY),
        "voiceover_enabled": vo_enabled,
        "heygen_enabled": settings.has_heygen,
        # v2.1 (3D Studio) : clé réelle OU simulateur local MESHY_MOCK
        "has_meshy": settings.has_meshy,
        "meshy_enabled": settings.has_meshy or bool(settings.MESHY_MOCK),
        "meshy_mock": bool(settings.MESHY_MOCK),
        "summarizer_enabled": settings.has_summarizer,
        "has_summarizer": settings.has_summarizer,
        "openai_enabled": settings.has_openai,
        "gemini_enabled": settings.has_gemini,
        "any_llm": settings.has_any_llm,
        "images_folder": str(settings.images_path),
        "outputs_folder": str(settings.outputs_path),
        # True = backend conteneurisé (MSIX) : ses écritures partent dans un
        # overlay invisible → relancer hors conteneur (voir fs_guard).
        "fs_virtualized": fs_is_virtualized(),
    }


# ============ v1.8: settings / .env editor ============
# LOCAL SINGLE-USER ONLY. The backend already binds 127.0.0.1:8765 and the
# user explicitly approved this surface. Hardening in depth:
#   - strict allowlist of writable keys (no arbitrary env vars)
#   - masked previews on read (raw values never leave the server)
#   - structured upsert (preserves comments, .env layout)
#   - explicit "restart required" signal (pydantic-settings doesn't reload)

_ALLOWED_ENV_KEYS = {
    "FAL_KEY", "HEYGEN_API_KEY", "MESHY_API_KEY", "ELEVENLABS_API_KEY",
    "ELEVENLABS_VOICE_ID_EN", "ELEVENLABS_VOICE_ID_FR",
    "ANTHROPIC_API_KEY", "ANTHROPIC_MODEL",
    "OPENAI_API_KEY", "OPENAI_MODEL",
    "GEMINI_API_KEY", "GEMINI_MODEL",
    "SUMMARIZER_PROVIDER", "PLANNER_PROVIDER",
    "OLLAMA_URL", "OLLAMA_MODEL",
    "ARTICLE_READER_FALLBACK",
    # Connected accounts (Scheduler — UI-only for now, but the keys live here)
    "X_API_KEY", "X_API_SECRET", "X_ACCESS_TOKEN", "X_ACCESS_SECRET",
    "TELEGRAM_BOT_TOKEN", "TELEGRAM_CHAT_ID",
    "YOUTUBE_CLIENT_ID", "YOUTUBE_CLIENT_SECRET",
    "YOUTUBE_REFRESH_TOKEN", "YOUTUBE_CHANNEL_ID",
    "IG_ACCESS_TOKEN", "IG_BUSINESS_ID",
}


def _env_path() -> Path:
    """The per-user .env in the stable data dir (survives reinstalls)."""
    from app.config import ENV_FILE
    return ENV_FILE


def _read_env_file() -> dict[str, str]:
    p = _env_path()
    if not p.exists():
        return {}
    out: dict[str, str] = {}
    for line in p.read_text(encoding="utf-8").splitlines():
        s = line.strip()
        if not s or s.startswith("#") or "=" not in s:
            continue
        k, v = s.split("=", 1)
        out[k.strip()] = v.strip()
    return out


def _mask(v: str | None) -> str:
    if not v:
        return ""
    if len(v) <= 8:
        return "•" * len(v)
    return v[:4] + "•" * max(len(v) - 8, 4) + v[-4:]


def _require_localhost(request: Request) -> None:
    """The settings surface reads/writes API keys — refuse any client that
    isn't loopback, even if HOST was misconfigured to 0.0.0.0."""
    host = (request.client.host if request.client else "") or ""
    if host not in ("127.0.0.1", "::1", "localhost", "testclient"):
        raise HTTPException(403, "Settings are only accessible from localhost")


@router.get("/settings/keys")
async def list_keys(request: Request):
    """Every allowed key with `set` (bool) + masked `preview`. Raw values
    are never returned."""
    _require_localhost(request)
    env = _read_env_file()
    out = []
    for k in sorted(_ALLOWED_ENV_KEYS):
        v = env.get(k, "")
        out.append({"key": k, "set": bool(v), "preview": _mask(v)})
    return {"keys": out, "env_path": str(_env_path())}


@router.post("/settings/keys")
async def set_key(body: dict, request: Request):
    """Upsert one or more keys into backend/.env.
    Accepts { name, value } or { entries: [{name, value}, …] }.
    Empty value clears the key. The backend must be restarted for
    changes to take effect (pydantic-settings doesn't hot-reload .env).
    """
    _require_localhost(request)
    entries = body.get("entries") if isinstance(body, dict) else None
    if entries is None:
        name = (body or {}).get("name")
        value = (body or {}).get("value", "")
        entries = [{"name": name, "value": value}] if name else []
    if not entries:
        raise HTTPException(400, "No entries to write")
    for e in entries:
        n = (e.get("name") or "").strip()
        if n not in _ALLOWED_ENV_KEYS:
            raise HTTPException(400, f"Key not allowed: {n}")

    p = _env_path()
    lines = p.read_text(encoding="utf-8").splitlines() if p.exists() else []
    changes = {(e.get("name") or "").strip(): (e.get("value") or "").strip()
               for e in entries}
    seen: set[str] = set()
    new_lines: list[str] = []
    for line in lines:
        s = line.strip()
        if (not s) or s.startswith("#") or "=" not in s:
            new_lines.append(line)
            continue
        k, _, _v = line.partition("=")
        k = k.strip()
        if k in changes:
            new_lines.append(f"{k}={changes[k]}")
            seen.add(k)
        else:
            new_lines.append(line)
    for k, v in changes.items():
        if k not in seen:
            new_lines.append(f"{k}={v}")
    p.write_text("\n".join(new_lines) + "\n", encoding="utf-8")
    logger.info(f"Wrote {len(changes)} key(s) to {p}")
    return {
        "ok": True,
        "written": list(changes.keys()),
        "restart_required": True,
        "message": "Saved. Restart the backend for changes to apply.",
    }


# ============ v1.15: provider defaults ============

@router.get("/settings/provider-defaults")
async def get_provider_defaults(request: Request):
    _require_localhost(request)
    from app.services.summarizer import active_provider as sum_active, _available_providers as sum_avail
    from app.services.marketing import _plan_available, _PLAN_PRIORITY
    plan_avail = [p for p in _PLAN_PRIORITY if _plan_available(p)]
    pref_plan = settings.PLANNER_PROVIDER.strip().lower()
    active_plan = pref_plan if (pref_plan and pref_plan in plan_avail) else (plan_avail[0] if plan_avail else "")
    return {
        "roles": {
            "summarizer": {
                "available": sum_avail(),
                "active": sum_active(),
                "preference": settings.SUMMARIZER_PROVIDER,
            },
            "planner": {
                "available": plan_avail,
                "active": active_plan,
                "preference": settings.PLANNER_PROVIDER,
            },
        },
    }


@router.post("/settings/provider-defaults")
async def set_provider_defaults(body: dict, request: Request):
    _require_localhost(request)
    allowed_roles = {"summarizer": "SUMMARIZER_PROVIDER", "planner": "PLANNER_PROVIDER"}
    changes: dict[str, str] = {}
    for role, value in (body or {}).items():
        env_key = allowed_roles.get(role)
        if not env_key:
            continue
        changes[env_key] = str(value or "").strip().lower()
    if not changes:
        raise HTTPException(400, "No valid role/provider pairs")
    p = _env_path()
    lines = p.read_text(encoding="utf-8").splitlines() if p.exists() else []
    seen: set[str] = set()
    new_lines: list[str] = []
    for line in lines:
        s = line.strip()
        if (not s) or s.startswith("#") or "=" not in s:
            new_lines.append(line)
            continue
        k, _, _v = line.partition("=")
        k = k.strip()
        if k in changes:
            new_lines.append(f"{k}={changes[k]}")
            seen.add(k)
        else:
            new_lines.append(line)
    for k, v in changes.items():
        if k not in seen:
            new_lines.append(f"{k}={v}")
    p.write_text("\n".join(new_lines) + "\n", encoding="utf-8")
    return {
        "ok": True,
        "written": changes,
        "restart_required": True,
        "message": "Saved. Restart the backend for changes to apply.",
    }


# ============ v1.11: WHITE-LABEL BRANDING ============
# The shipped product boots as "deepotus"; everything user-facing is
# rebrandable. Config lives in assets/branding/branding.json + logo.png —
# under assets/, so upgrades never touch it.

BRAND_DEFAULTS = {
    "app_name": "DEEPOTUS",
    "app_sub": "VIDEO",
    "tagline_1": "From the deep,",
    "tagline_2": "for the deep.",
    "brand_color": "#ef4444",
    "accent_color": "#00e5ff",
}
_BRAND_COLOR_RE = r"^#[0-9a-fA-F]{6}$"


def _branding_dir() -> Path:
    from app.config import DATA_ROOT
    p = DATA_ROOT / "assets" / "branding"
    p.mkdir(parents=True, exist_ok=True)
    return p


def _read_branding() -> dict:
    f = _branding_dir() / "branding.json"
    data = dict(BRAND_DEFAULTS)
    if f.is_file():
        try:
            user = json.loads(f.read_text(encoding="utf-8"))
            for k in BRAND_DEFAULTS:
                if isinstance(user.get(k), str) and user[k].strip():
                    data[k] = user[k].strip()
        except (ValueError, OSError) as e:
            logger.warning(f"branding.json unreadable, using defaults: {e}")
    data["has_custom_logo"] = (_branding_dir() / "logo.png").is_file()
    data["is_default"] = not (_branding_dir() / "branding.json").is_file() \
        and not data["has_custom_logo"]
    return data


@router.get("/branding")
async def get_branding():
    return _read_branding()


@router.post("/branding")
async def set_branding(body: dict, request: Request):
    """Update brand fields (allowlisted, colors validated). Empty body or
    {"reset": true} restores deepotus defaults (and removes the custom logo)."""
    _require_localhost(request)
    bdir = _branding_dir()
    if not body or body.get("reset"):
        (bdir / "branding.json").unlink(missing_ok=True)
        (bdir / "logo.png").unlink(missing_ok=True)
        logger.info("branding reset to deepotus defaults")
        return _read_branding()
    clean = {}
    for k in BRAND_DEFAULTS:
        v = body.get(k)
        if not isinstance(v, str) or not v.strip():
            continue
        v = v.strip()
        if k.endswith("_color") and not re.match(_BRAND_COLOR_RE, v):
            raise HTTPException(400, f"{k} must be #RRGGBB (got: {v})")
        clean[k] = v[:60]
    existing = {}
    f = bdir / "branding.json"
    if f.is_file():
        try:
            existing = json.loads(f.read_text(encoding="utf-8"))
        except (ValueError, OSError):
            existing = {}
    existing.update(clean)
    f.write_text(json.dumps(existing, indent=2), encoding="utf-8")
    logger.info(f"branding updated: {list(clean.keys())}")
    return _read_branding()


@router.get("/branding/logo")
async def get_branding_logo():
    """The brand logo: custom upload if present, else the bundled deepotus
    mark. Cache disabled so a rebrand shows immediately."""
    custom = _branding_dir() / "logo.png"
    if custom.is_file():
        return FileResponse(str(custom), media_type="image/png",
                            headers={"Cache-Control": "no-cache"})
    bundled = (Path(__file__).resolve().parents[2].parent
               / "frontend" / "public" / "deepotus-logo.png")
    if bundled.is_file():
        return FileResponse(str(bundled), media_type="image/png",
                            headers={"Cache-Control": "no-cache"})
    raise HTTPException(404, "No logo available")


@router.post("/branding/logo")
async def upload_branding_logo(request: Request, file: UploadFile = File(...)):
    _require_localhost(request)
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in (".png", ".jpg", ".jpeg", ".webp"):
        raise HTTPException(400, "Logo must be .png, .jpg or .webp")
    data = await file.read()
    if len(data) > 5 * 1024 * 1024:
        raise HTTPException(400, "Logo too large (max 5 MB)")
    # Normalize to PNG via Pillow (also validates it's a real image).
    try:
        import io
        from PIL import Image as PILImg
        img = PILImg.open(io.BytesIO(data)).convert("RGBA")
        img.save(_branding_dir() / "logo.png", format="PNG")
    except Exception as e:
        raise HTTPException(400, f"Not a valid image: {e}")
    logger.info("custom brand logo uploaded")
    return _read_branding()


# ============ v1.14: editable caption pack (Telegram Premium tags) ============
# The Scheduler caption editor offers one-tap branded tags. The default set is
# deepotus-flavoured; this endpoint lets the user (or a white-label reseller)
# edit the entries and upload a custom icon per entry. Stored next to branding.

_DEFAULT_CAPTION_PACK = [
    {"id": "deepotus-protocol", "emoji": "\U0001F419", "label": "Deepotus Protocol", "icon": "/pack/deepotus-protocol.png"},
    {"id": "rippled-signal",    "emoji": "\U0001F30A", "label": "Rippled Signal",    "icon": "/pack/rippled-signal.png"},
    {"id": "prophet",           "emoji": "\U0001F441", "label": "Prophet",           "icon": ""},
    {"id": "chapter-drop",      "emoji": "\U0001F4D6", "label": "Chapter Drop",      "icon": ""},
    {"id": "board-game",        "emoji": "\U0001F3B4", "label": "Board Game",        "icon": ""},
    {"id": "dnd",               "emoji": "\U0001F3B2", "label": "D&D",               "icon": ""},
    {"id": "mobile-devlog",     "emoji": "\U0001F4F1", "label": "Mobile Devlog",     "icon": ""},
    {"id": "deep",              "emoji": "\U0001FA99", "label": "$DEEP",             "icon": ""},
    {"id": "gencoin",           "emoji": "\U0001F9EC", "label": "Gencoin",           "icon": ""},
    {"id": "video-engine",      "emoji": "\U0001F3AC", "label": "Video Engine",      "icon": ""},
]


def _slug(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", str(s).lower()).strip("-")[:40] or "tag"


def _caption_pack_file() -> Path:
    return _branding_dir() / "caption-pack.json"


def _pack_icons_dir() -> Path:
    p = _branding_dir() / "pack-icons"
    p.mkdir(parents=True, exist_ok=True)
    return p


def _clean_pack_entry(e: dict) -> dict | None:
    if not isinstance(e, dict):
        return None
    label = str(e.get("label") or "").strip()
    if not label:
        return None
    return {
        "id": (str(e.get("id") or "").strip() or _slug(label))[:40],
        "emoji": str(e.get("emoji") or "").strip()[:8],
        "label": label[:40],
        "icon": str(e.get("icon") or "").strip()[:300],
    }


def _read_caption_pack() -> list:
    f = _caption_pack_file()
    if f.is_file():
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
            if isinstance(data, list):
                out = [c for c in (_clean_pack_entry(e) for e in data) if c]
                if out:
                    return out
        except (ValueError, OSError) as ex:
            logger.warning(f"caption-pack.json unreadable: {ex}")
    return [dict(e) for e in _DEFAULT_CAPTION_PACK]


@router.get("/caption-pack")
async def get_caption_pack():
    return {"pack": _read_caption_pack(), "is_default": not _caption_pack_file().is_file()}


@router.post("/caption-pack")
async def set_caption_pack(body: dict, request: Request):
    """Save the caption pack. {"reset": true} restores deepotus defaults and
    clears uploaded icons."""
    _require_localhost(request)
    if not body or body.get("reset"):
        _caption_pack_file().unlink(missing_ok=True)
        import shutil
        shutil.rmtree(_pack_icons_dir(), ignore_errors=True)
        logger.info("caption pack reset to defaults")
        return {"pack": _read_caption_pack(), "is_default": True}
    items = body.get("pack")
    if not isinstance(items, list):
        raise HTTPException(400, "pack must be a list")
    clean = [c for c in (_clean_pack_entry(e) for e in items[:40]) if c]
    _caption_pack_file().write_text(
        json.dumps(clean, indent=2, ensure_ascii=False), encoding="utf-8")
    logger.info(f"caption pack saved: {len(clean)} entries")
    return {"pack": clean, "is_default": False}


@router.get("/caption-pack/icon/{slot}")
async def get_pack_icon(slot: str):
    safe = re.sub(r"[^a-zA-Z0-9_-]", "", slot)[:40]
    p = _pack_icons_dir() / f"{safe}.png"
    if not p.is_file():
        raise HTTPException(404, "icon not found")
    return FileResponse(str(p), media_type="image/png", headers={"Cache-Control": "no-cache"})


@router.post("/caption-pack/icon/{slot}")
async def upload_pack_icon(slot: str, request: Request, file: UploadFile = File(...)):
    _require_localhost(request)
    safe = re.sub(r"[^a-zA-Z0-9_-]", "", slot)[:40] or "tag"
    data = await file.read()
    if len(data) > 5 * 1024 * 1024:
        raise HTTPException(400, "Icon too large (max 5 MB)")
    try:
        import io
        from PIL import Image as PILImg
        img = PILImg.open(io.BytesIO(data)).convert("RGBA")
        img.thumbnail((128, 128), PILImg.LANCZOS)
        img.save(_pack_icons_dir() / f"{safe}.png", format="PNG")
    except Exception as e:
        raise HTTPException(400, f"Not a valid image: {e}")
    logger.info(f"caption pack icon uploaded: {safe}")
    return {"ok": True, "icon": f"/api/caption-pack/icon/{safe}"}


# ============ v1.9: MARKETING PLAN + SCHEDULER + IMAGE GEN ============

from datetime import datetime as _dt, timedelta as _td
from sqlalchemy import select as _select, delete as _delete, \
    or_ as _or, and_ as _and, func as _func
from app.services.storage import ScheduledPost, JobRecord, async_session_factory
from app.services import marketing


def _post_to_dict(p: ScheduledPost) -> dict:
    brief = None
    if getattr(p, "brief", None):
        try:
            brief = json.loads(p.brief)
        except (ValueError, TypeError):
            brief = None
    return {
        "brief": brief,
        "id": p.id,
        "title": p.title,
        "caption": p.caption,
        "channels": [c for c in (p.channels or "").split(",") if c],
        "run_at": (p.run_at.isoformat() + "Z") if p.run_at else None,
        "status": p.status,
        "mode": p.mode,
        "job_id": p.job_id,
        "format": p.format,
        "hook": p.hook,
        "script_idea": p.script_idea,
        "image_idea": p.image_idea,
        "plan_id": p.plan_id,
        "error": p.error,
        "created_at": (p.created_at.isoformat() + "Z") if p.created_at else None,
        "posted_at": (p.posted_at.isoformat() + "Z") if p.posted_at else None,
        "x_post_id": p.x_post_id,
        "metrics": p.metrics,
        "source_image": p.source_image,
    }


@router.get("/schedule")
async def list_schedule(days_back: int = 365, days_forward: int = 365):
    """Scheduled posts. Les posts encore actionnables (draft/scheduled/ready)
    sont TOUJOURS renvoyés quel que soit leur run_at : l'ancienne fenêtre
    [−30 j, +90 j] les faisait « disparaître » silencieusement du Scheduler
    dès qu'elle glissait au-delà (incident 20/07/2026 — rien n'était perdu
    en base). La fenêtre ne borne plus que l'historique posted/failed."""
    lo = _dt.utcnow() - _td(days=days_back)
    hi = _dt.utcnow() + _td(days=days_forward)
    async with async_session_factory() as session:
        res = await session.execute(
            _select(ScheduledPost)
            .where(_or(
                ScheduledPost.status.in_(("draft", "scheduled", "ready")),
                _and(ScheduledPost.run_at >= lo, ScheduledPost.run_at <= hi)))
            .order_by(ScheduledPost.run_at.asc()))
        return [_post_to_dict(p) for p in res.scalars().all()]


@router.post("/schedule")
async def create_scheduled_post(body: dict):
    """Create one post. Body: {title, caption, channels[], run_at (UTC ISO),
    status?, mode?, job_id?, format?}."""
    run_at_raw = (body.get("run_at") or "").replace("Z", "")
    try:
        run_at = _dt.fromisoformat(run_at_raw)
    except ValueError:
        raise HTTPException(400, f"Invalid run_at: {body.get('run_at')}")
    p = ScheduledPost(
        id=str(uuid4()),
        title=(body.get("title") or "Untitled post")[:200],
        caption=body.get("caption") or "",
        channels=",".join(body.get("channels") or ["x"]),
        run_at=run_at,
        status=body.get("status") if body.get("status") in
               ("draft", "scheduled", "ready") else "draft",
        mode=body.get("mode") if body.get("mode") in
             ("auto", "assisted") else "assisted",
        job_id=body.get("job_id"),
        format=body.get("format"),
        hook=body.get("hook"),
        script_idea=body.get("script_idea"),
        image_idea=body.get("image_idea"),
        source_image=body.get("source_image") or None,
        brief=(json.dumps(body["brief"], ensure_ascii=False)
               if isinstance(body.get("brief"), dict) and body["brief"]
               else None),
    )
    async with async_session_factory() as session:
        session.add(p)
        await session.commit()
    return _post_to_dict(p)


@router.patch("/schedule/{post_id}")
async def update_scheduled_post(post_id: str, body: dict):
    """Patch any editable field of a post."""
    async with async_session_factory() as session:
        res = await session.execute(
            _select(ScheduledPost).where(ScheduledPost.id == post_id))
        p = res.scalar_one_or_none()
        if not p:
            raise HTTPException(404, "Post not found")
        if "title" in body:
            p.title = (body["title"] or "")[:200]
        if "caption" in body:
            p.caption = body["caption"] or ""
        if "channels" in body:
            p.channels = ",".join(body["channels"] or [])
        if "run_at" in body and body["run_at"]:
            try:
                p.run_at = _dt.fromisoformat(
                    str(body["run_at"]).replace("Z", ""))
            except ValueError:
                raise HTTPException(400, f"Invalid run_at: {body['run_at']}")
        if "status" in body and body["status"] in (
                "draft", "scheduled", "ready", "posted", "failed"):
            p.status = body["status"]
        if "mode" in body and body["mode"] in ("auto", "assisted"):
            p.mode = body["mode"]
        if "job_id" in body:
            p.job_id = body["job_id"] or None
        if "format" in body:
            p.format = body["format"] or None
        if "source_image" in body:
            p.source_image = body["source_image"] or None
        if "brief" in body:
            p.brief = (json.dumps(body["brief"], ensure_ascii=False)
                       if isinstance(body["brief"], dict) and body["brief"]
                       else None)
        await session.commit()
        await session.refresh(p)
        return _post_to_dict(p)


@router.delete("/schedule/{post_id}")
async def delete_scheduled_post(post_id: str):
    async with async_session_factory() as session:
        res = await session.execute(
            _select(ScheduledPost).where(ScheduledPost.id == post_id))
        if not res.scalar_one_or_none():
            raise HTTPException(404, "Post not found")
        await session.execute(
            _delete(ScheduledPost).where(ScheduledPost.id == post_id))
        await session.commit()
    return {"deleted": post_id}


@router.post("/schedule/{post_id}/fire")
async def fire_scheduled_post(post_id: str):
    """Publish NOW on auto-capable channels (Telegram). Channels without an
    auto adapter stay listed in `pending` — the post flips to `ready` so the
    user can publish manually with the caption + downloaded render."""
    result = await marketing.fire_post(post_id)
    if not result.get("ok") and result.get("error") == "post not found":
        raise HTTPException(404, "Post not found")
    return result


def _render_poster_frame(jobrec) -> str | None:
    """Cached poster frame (~1s in) extracted from a render's video, used as the
    post-preview hero when the render has no still image (e.g. HeyGen avatars).
    Returns the PNG path or None."""
    import subprocess
    vp = getattr(jobrec, "final_video_path", None)
    if not vp:
        return None
    src = Path(vp)
    if not src.is_file():
        return None
    cache = settings.outputs_path / "_cache"
    cache.mkdir(parents=True, exist_ok=True)
    out = cache / f"poster_{jobrec.id}.png"
    if out.is_file() and out.stat().st_mtime >= src.stat().st_mtime:
        return str(out)
    try:
        subprocess.run(
            ["ffmpeg", "-y", "-ss", "1", "-i", str(src),
             "-frames:v", "1", "-q:v", "3", str(out)],
            check=True, capture_output=True, timeout=30)
        return str(out) if out.is_file() else None
    except Exception as e:
        logger.warning(f"poster extraction failed for {getattr(jobrec,'id','?')}: {e}")
        return None


@router.get("/schedule/{post_id}/preview.png")
async def scheduled_post_preview(post_id: str, channel: str = "x",
                                 caption: str | None = None,
                                 img: str | None = None,
                                 job: str | None = None):
    """Compose the final post (hero image + caption) as a platform-styled PNG
    so the user can visualize it before publishing. channel = x | telegram.

    Hero image resolves to the post's source_image, else the attached render's
    still (job.image_filename), else a placeholder. The caption is rendered in
    the platform's usage format (X: handle + 280-char card; Telegram: channel
    bubble with image-then-caption).

    Optional query overrides (caption, img, job) let the inspector preview
    live, still-unsaved edits without a DB round-trip."""
    from app.services import post_preview as _pp
    async with async_session_factory() as session:
        res = await session.execute(
            _select(ScheduledPost).where(ScheduledPost.id == post_id))
        p = res.scalar_one_or_none()
        if not p:
            raise HTTPException(404, "Post not found")
        eff_src = img if img is not None else p.source_image
        eff_job = job if job is not None else p.job_id
        hero = None
        if eff_src:
            cand = settings.images_path / Path(eff_src).name
            if cand.is_file():
                hero = str(cand)
        if not hero and eff_job:
            jr = await session.execute(
                _select(JobRecord).where(JobRecord.id == eff_job))
            jobrec = jr.scalar_one_or_none()
            if jobrec:
                if jobrec.image_filename:
                    cand = settings.images_path / Path(jobrec.image_filename).name
                    if cand.is_file():
                        hero = str(cand)
                if not hero:
                    hero = await asyncio.to_thread(_render_poster_frame, jobrec)
        caption = caption if caption is not None else (p.caption or p.title or "")
    brand = _read_branding()
    name = (brand.get("app_name") or "Deepotus").strip().title() or "Deepotus"
    handle = re.sub(r"[^a-z0-9_]", "", name.lower())[:15] or "deepotus"
    try:
        png = await asyncio.to_thread(
            _pp.render_preview, channel=channel, caption=caption,
            hero_path=hero, display_name=name, handle=handle)
    except Exception as e:
        logger.exception("post preview render failed")
        raise HTTPException(500, f"Preview failed: {e}")
    return Response(content=png, media_type="image/png",
                    headers={"Cache-Control": "no-store"})


# ============ COST WIDGET (v1.15.1) ============

@router.post("/cost/estimate")
async def cost_estimate(body: dict, request: Request):
    """Preview budget for a planned op (single edit / plan / campaign)."""
    _require_localhost(request)
    from app.services import pricing as _pricing
    return _pricing.estimate(body or {})


# Providers dont un job terminé N'ENGAGE AUCUNE DÉPENSE, et POURQUOI — le
# libellé part dans le `breakdown`, donc le zéro se lit au lieu de se deviner.
_JOBS_SANS_DEPENSE = {
    "montage":       "Montage (assemblage ffmpeg local)",
    "montage_proxy": "Proxy de montage (transcodage ffmpeg local)",
    "animation":     "Animation (ffmpeg + PIL locaux)",
    "news":          "News reel (ffmpeg local)",
    "ugc":           "Fichier téléversé par l'utilisateur",
    "card3d":        "Publication d'une carte 3D (rien n'est fabriqué)",
    "template":      "Rendu template (job parent — les sous-jobs paient)",
    "composition":   "Composition (job parent — les sous-jobs paient)",
}
# ... et ceux dont la « campagne » (1 image + N s de vidéo) EST la vérité.
# `provider IS NULL` arrive ici par le `or "seedance"` ci-dessous, et c'est
# JUSTE : les 13 jobs `done` de la base réelle qui portent NULL datent d'avant
# la colonne et sont bien des rendus Seedance (13/13 portent une image de
# départ ET une vidéo produite).
_JOBS_CAMPAGNE = {"seedance"}


def _job_to_cost(job, p):
    """Le devis d'UN job terminé — par LISTE BLANCHE, jamais par défaut.

    LA DÉCISION (P7, tâche 8b), et la mesure qui l'a tranchée. Cette fonction
    nommait `heygen`, `episode` et `sprite2d`, puis retombait sur une branche
    « campaign » PAR DÉFAUT qui facturait `duration_s or 10` secondes de
    Seedance PLUS une image FLUX — 0,403 USD par job aux tarifs par défaut de
    `pricing.load()` — à TOUT provider qu'elle ne nommait pas.

    PROTOCOLE DE MESURE, ET IL A DÛ ÊTRE REFAIT. La base tourne en mode WAL :
    au 05/09/2026 `deepotus.db` pesait 4,60 Mo et son `-wal` 4,54 Mo. Une
    COPIE D'OCTETS du seul `.db` perd donc tout ce que le WAL porte — elle
    rendait 105 jobs `done` là où la base en compte 116, et un total de
    43,41 USD là où l'application en affichait 51,90. Les chiffres ci-dessous
    sont pris sur un instantané COHÉRENT (`sqlite3.Connection.backup()`, qui
    fusionne le WAL), jamais sur la base vivante. L'ANCIENNE branche y est
    REJOUÉE, et sa transcription est VÉRIFIÉE plutôt que supposée : son total
    reproduit au cent près le `total_usd` que l'API du backend installé rend
    sur la vraie base (51,90). Sans cette égalité, la colonne « avant » ne
    vaudrait rien.

    MESURÉ (116 jobs `done` ; `scratchpad/mesure_avant.py`, `mesure_defaut.py`)
    — 98,9 % du total affiché passait par la branche par défaut :

                    n      avant     après
        seedance   35    18,885    18,885   1 image + N s de vidéo : JUSTE
        template   33    13,299     0,000   RIEN — job PARENT, les sous-jobs
                                            portent déjà leur propre dépense
        montage     4     6,332     0,000   RIEN — assemblage ffmpeg local
        <NULL>     13     5,239     5,239   Seedance d'avant la colonne : JUSTE
        ugc         9     4,307     0,000   RIEN — fichier TÉLÉVERSÉ (`await
                                            file.read()`, aucune API appelée)
        asset3d     7     2,821     2,480   un maillage, désormais au tarif du
                                            maillage et non d'une vidéo
        news        1     0,403     0,000   RIEN — reel ffmpeg local
        animation   1     0,323     0,000   RIEN — ffmpeg + PIL locaux
        heygen      5     0,240     0,240   branche nommée, juste
        sprite2d    8     0,048     0,048   branche nommée, juste
        TOTAL     116    51,900    26,890

    DÉPENSE FABRIQUÉE (template + montage + ugc + news + animation) :
    24,664 USD sur 51,900 AFFICHÉS, soit 47,5 % du chiffre montré à
    l'utilisateur.

    POURQUOI UNE LISTE BLANCHE plutôt qu'une entrée de tarif par provider :
    le dépôt n'écrit que TREIZE valeurs de `provider` (relevé exhaustif des
    `provider=` / `.provider =` sur `JobRecord`), dont HUIT ne dépensent RIEN
    du tout — leur « tarif » serait 0. Un seul, `asset3d`, dépense pour de
    vrai sans être tarifé, et `pricing.py` savait DÉJÀ le chiffrer (`kind`
    `asset3d` / `asset3d_texture`) : il n'était pas branché. Une table de
    tarifs aurait donc surtout été une table de zéros. Ce qui manquait
    n'était pas un prix, c'était le refus d'en inventer un.

    ET CE QU'UN PROVIDER INCONNU DEVIENT : zéro, mais un zéro qui SE NOMME —
    `by_provider` reçoit la clé `non-tarifé:<provider>`. Un blanc avoué se
    répare ; un chiffre inventé se croit. C'est la même décision que le
    `where` de `cost_usage` (65afc16), poussée jusqu'à sa conclusion.

    CE QUE CETTE FONCTION N'AFFIRME TOUJOURS PAS : que l'image FLUX d'une
    campagne ait été payée (un job parti d'une image FOURNIE n'a rien payé à
    FLUX — 0,003 USD par job, 0,3 % du total, et la ligne de base ne permet
    pas de trancher) ; ni que le texturage Meshy soit au bon palier
    (`cost_meta` n'enregistre pas la résolution, donc le devis prend le
    défaut 2k de `credits_retexture`). Deux approximations DÉCLARÉES sur des
    dépenses RÉELLES — pas des dépenses inventées.

    Tenu par `tests/test_cost_usage.py`.
    """
    from app.services import pricing as _pricing
    prov = (job.provider or "seedance").lower()
    dur = job.duration_s or 10
    if prov in _JOBS_SANS_DEPENSE:
        return _pricing.no_spend(_JOBS_SANS_DEPENSE[prov])
    if prov == "asset3d":
        import json as _json
        try:
            meta = _json.loads(job.cost_meta or "{}")
        except Exception:
            meta = {}
        if meta.get("texturier") == "meshy":
            # texturage d'un maillage DÉJÀ généré : facturé en crédits Meshy,
            # jamais chez fal (chaîne Tripo → Meshy).
            return _pricing.estimate({"kind": "asset3d_texture"}, p)
        return _pricing.estimate({"kind": "asset3d",
                                  "engine": meta.get("engine") or "tripo"}, p)
    if prov == "heygen":
        return _pricing.estimate({"kind": "heygen", "minutes": max(0.2, dur / 60.0)}, p)
    if prov == "episode":
        import json as _json
        try:
            meta = _json.loads(job.cost_meta or "{}")
        except Exception:
            meta = {}
        return _pricing.estimate({"kind": "episode",
                                  "images": int(meta.get("images", 1) or 1),
                                  "chars": float(meta.get("chars", 0) or 0)}, p)
    if prov == "sprite2d":
        import json as _json
        try:
            meta = _json.loads(job.cost_meta or "{}")
        except Exception:
            meta = {}
        return _pricing.estimate({"kind": "sprite2d",
                                  "frames": int(meta.get("frames", 0) or 0),
                                  "remove_bg": meta.get("remove_bg", "none")}, p)
    if prov in _JOBS_CAMPAGNE:
        return _pricing.estimate({"kind": "campaign", "ops": [
            {"kind": "image"},
            {"kind": "seedance", "duration_s": dur,
             "model": getattr(job, "video_model", None) or ""}]}, p)
    # LE BLANC AVOUÉ — voir la docstring. Le provider part dans la CLÉ pour
    # qu'un coup d'œil à `by_provider` dise LEQUEL n'est pas tarifé.
    return _pricing.no_spend(f"Non tarifé — provider « {prov} »",
                             f"non-tarifé:{prov}")


@router.get("/cost/usage")
async def cost_usage():
    """Cumulative ESTIMATED spend, computed from finished job records.

    LES PRÉCALCULS DU MONTAGE SONT ÉCARTÉS, et il le faut : `montage_proxy`
    est un transcodage ffmpeg LOCAL et GRATUIT (P7, l'aperçu 480p du
    balayage). Il ne porte ni `duration_s` ni `video_model`, donc
    `_job_to_cost` retombait sur sa branche « campaign » par défaut et
    facturait `duration_s or 10` secondes de Seedance PLUS une image FLUX.
    MESURÉ (tarifs par défaut de `pricing.load()`) : 0,403 USD PAR JOB,
    imputés à `fal` (0,400 « Seedance video » + 0,003 « FLUX image x1 »),
    soit 4,84 USD AFFICHÉS pour une timeline de douze clips — de l'argent
    montré à l'utilisateur pour une dépense qui n'existe pas. C'est la même
    décision que la fenêtre de `GET /api/jobs` (`Pipeline.list_jobs`) : un
    cache n'est ni un plan, ni une dépense.

    `coalesce(provider, '')` pour la même raison qu'ailleurs : `NULL != …`
    vaut NULL en SQL et écarterait les 13 jobs `done` sans provider.

    LA BRANCHE PAR DÉFAUT DE `_job_to_cost` EST FERMÉE DEPUIS (tâche 8b) :
    elle facture désormais par LISTE BLANCHE, et un provider inconnu rend 0
    sous la clé `non-tarifé:<provider>`. Ce `where` reste, et il n'est pas
    redondant : sans lui la carte `by_provider` porterait une entrée `local`
    à 0 pour des précalculs qui ne sont pas des opérations de l'utilisateur.

    CE QUE CETTE ROUTE N'AFFIRME TOUJOURS PAS : que le total soit juste au
    centime. Il reste DIRECTIONNEL par construction — voir les deux
    approximations déclarées dans la docstring de `_job_to_cost`.
    """
    from app.services import pricing as _pricing
    from app.services.montage_service import _PROXY_PROVIDER
    p = _pricing.load()
    per = {}
    total = 0.0
    async with async_session_factory() as session:
        res = await session.execute(
            _select(JobRecord).where(
                JobRecord.status == "done",
                _func.coalesce(JobRecord.provider, "") != _PROXY_PROVIDER))
        for job in res.scalars().all():
            e = _job_to_cost(job, p)
            total += e["total_usd"]
            for ln in e["breakdown"]:
                per[ln["provider"]] = round(per.get(ln["provider"], 0) + ln["usd"], 4)
    return {"total_usd": round(total, 2), "by_provider": per}


@router.get("/cost/balances")
async def cost_balances():
    """Live remaining balances where a provider exposes them (HeyGen credits,
    ElevenLabs characters); pay-as-you-go otherwise."""
    from app.services import pricing as _pricing
    p = _pricing.load()
    out = {}
    if settings.has_heygen:
        try:
            q = await HeyGenClient().remaining_quota()
            rem = q.get("remaining_quota") if isinstance(q, dict) else None
            out["heygen"] = {"available": True, "credits": rem,
                             "usd": round((rem or 0) * p["heygen_credit_usd"], 2)}
        except Exception:
            out["heygen"] = {"available": False}
    from app.services import voice_providers as _VP
    if await asyncio.get_running_loop().run_in_executor(
            None, _VP.voicebox_reachable):
        # spec voicebox 2026-07-11 : statut abonnement = « local : gratuit »
        out["voicebox"] = {"available": True, "local": True, "free": True}
    if settings.has_voiceover:
        try:
            async with httpx.AsyncClient(timeout=15.0, verify=SSL_VERIFY) as c:
                r = await c.get("https://api.elevenlabs.io/v1/user/subscription",
                                headers={"xi-api-key": settings.ELEVENLABS_API_KEY})
                if r.status_code == 200:
                    d = r.json()
                    used = d.get("character_count")
                    lim = d.get("character_limit")
                    out["elevenlabs"] = {
                        "available": True, "used": used, "limit": lim,
                        "remaining": (lim - used) if (lim is not None and used is not None) else None}
                else:
                    out["elevenlabs"] = {"available": False}
        except Exception:
            out["elevenlabs"] = {"available": False}
    for prov, on in (("fal", bool(settings.FAL_KEY)),
                     ("anthropic", settings.has_summarizer),
                     ("openai", settings.has_openai),
                     ("gemini", settings.has_gemini)):
        if on:
            out[prov] = {"available": False, "mode": "pay-as-you-go"}
    return out


@router.get("/cost/pricing")
async def get_pricing():
    from app.services import pricing as _pricing
    return _pricing.load()


@router.post("/cost/pricing")
async def set_pricing(body: dict, request: Request):
    _require_localhost(request)
    from app.services import pricing as _pricing
    return _pricing.save(body or {})



@router.post("/marketing/plan")
async def marketing_plan(body: dict):
    """Prompt → structured posting plan.
    Body: {prompt, days=7, posts_per_day=1, channels=["x"], language="EN",
           persona?{name,tone,audience}, auto_materialize=false,
           start_date "YYYY-MM-DD", tz_offset_minutes (JS getTimezoneOffset),
           mode "auto"|"assisted"}.
    Uses Anthropic when ANTHROPIC_API_KEY is set; deterministic otherwise."""
    prompt = (body.get("prompt") or "").strip()
    if not prompt:
        raise HTTPException(400, "prompt is required")
    days = max(1, min(31, int(body.get("days") or 7)))
    ppd = max(1, min(6, int(body.get("posts_per_day") or 1)))
    plan = await marketing.generate_plan(
        prompt,
        days=days,
        posts_per_day=ppd,
        channels=body.get("channels") or ["x"],
        language=body.get("language") or "EN",
        persona=body.get("persona"),
    )
    materialized: list[str] = []
    if body.get("auto_materialize"):
        start = body.get("start_date") or _dt.now().strftime("%Y-%m-%d")
        materialized = await marketing.materialize_plan(
            plan["posts"],
            start_date=start,
            tz_offset_minutes=int(body.get("tz_offset_minutes") or 0),
            mode=body.get("mode") or "assisted",
        )
    return {**plan, "materialized_ids": materialized}


@router.post("/marketing/plan/import")
async def import_marketing_plan(
    file: UploadFile = File(...),
    days: int = 30,
    channels: str = "x",
    language: str = "EN",
):
    """Upload an EXISTING strategy document (.md / .txt / .docx / .pdf) and
    transcribe it into posting-plan slices. Human-in-the-loop: returns the
    structured plan only — the UI previews it and the user materializes
    explicitly via POST /marketing/plan with the returned posts, or via the
    modal's Add-to-calendar which re-sends with auto_materialize."""
    data = await file.read()
    if len(data) > 15 * 1024 * 1024:
        raise HTTPException(400, "File too large (max 15 MB)")
    try:
        text = await asyncio.to_thread(
            marketing.extract_document_text, file.filename or "", data)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        logger.error(f"document extraction failed: {e}")
        raise HTTPException(400, f"Could not read the document: {e}")
    try:
        plan = await marketing.plan_from_document(
            text,
            days=max(1, min(60, days)),
            channels=[c for c in channels.split(",") if c] or ["x"],
            language=language,
        )
    except ValueError as e:
        raise HTTPException(400, str(e))
    return {**plan, "chars_read": len(text), "filename": file.filename}


@router.post("/marketing/plan/materialize")
async def materialize_plan_route(body: dict):
    """Materialize an already-generated/imported plan (the human-approved
    posts array) into scheduled_posts. Used by the import flow where the
    posts were reviewed in the preview before being committed."""
    posts = body.get("posts")
    if not isinstance(posts, list) or not posts:
        raise HTTPException(400, "posts array is required")
    start = body.get("start_date") or _dt.now().strftime("%Y-%m-%d")
    ids = await marketing.materialize_plan(
        posts,
        start_date=start,
        tz_offset_minutes=int(body.get("tz_offset_minutes") or 0),
        mode=body.get("mode") or "assisted",
    )
    return {"materialized_ids": ids}


@router.post("/channels/test")
async def test_channel(body: dict):
    """Send a test message on a channel to validate the keys. Telegram sends
    a text message; X posts a real (deletable) tweet only when confirm=true,
    otherwise it just validates credentials by fetching the authed user."""
    ch = (body or {}).get("channel")
    if ch == "telegram":
        if not settings.has_telegram:
            raise HTTPException(400, "Telegram keys not set")
        ok, detail = await marketing.publish_telegram(
            "Deepotus Video Gen — test message. The deep hears you. 🐙")
        return {"ok": ok, "detail": detail}
    if ch == "x":
        if not settings.has_x:
            raise HTTPException(400, "X keys not set")
        import asyncio as _aio
        def _verify():
            try:
                me = marketing._x_client().get_me()
                return True, f"authenticated as @{me.data.username}"
            except Exception as e:
                return False, str(e)
        ok, detail = await _aio.to_thread(_verify)
        return {"ok": ok, "detail": detail}
    raise HTTPException(400, f"No test available for channel: {ch}")


@router.post("/images/import-url")
async def import_image_url(body: dict):
    """Download a remote image (e.g. the picture attached to a news item)
    into the images folder so it can be used as a Seedance start frame or a
    post still. Returns {filename}. Used by the plan's Sources step."""
    url = (body or {}).get("url", "").strip()
    if not url or not url.lower().startswith(("http://", "https://")):
        raise HTTPException(400, "A valid image URL is required")
    if _is_private_host(httpx.URL(url).host):
        raise HTTPException(400, "Refusing to fetch a private/loopback address")
    try:
        async with _remote_image_client() as client:
            async with client.stream("GET", url, headers={"User-Agent": "Mozilla/5.0"}) as r:
                r.raise_for_status()
                ctype = r.headers.get("content-type", "")
                if "image" not in ctype and not url.lower().split("?")[0].endswith(
                        (".png", ".jpg", ".jpeg", ".webp")):
                    raise HTTPException(400, f"URL is not an image ({ctype})")
                buf = bytearray()
                async for chunk in r.aiter_bytes():
                    buf += chunk
                    if len(buf) > 25 * 1024 * 1024:
                        raise HTTPException(400, "Image too large (max 25 MB)")
            import io
            from PIL import Image as _PILImage
            img = _PILImage.open(io.BytesIO(bytes(buf))).convert("RGB")
            fname = f"news_{uuid4().hex[:8]}.png"
            img.save(settings.images_path / fname, format="PNG")
    except HTTPException:
        raise
    except Exception as e:
        logger.warning(f"news image import failed: {e}")
        raise HTTPException(502, f"Could not import the image: {e}")
    logger.info(f"imported news image -> {fname}")
    await LI.noter([fname], "news")
    return {"filename": fname}


@router.post("/images/generate")
async def generate_image(body: dict, background_tasks: BackgroundTasks):
    """Provenance (28/08) : l'enveloppe indexe les images rendues par les
    trois chemins modèles d'un coup — source = hint `source` du body si
    slug connu (appelants ciblés : épisodes, planificateur…), sinon
    `generation`."""
    out = await _generate_image_core(body, background_tasks)
    try:
        src = str(body.get("source") or "").strip().lower()
        if src not in LI.SOURCES:
            src = "generation"
        if isinstance(out, dict) and out.get("images"):
            await LI.noter(out["images"], src)
    except Exception as e:  # noqa: BLE001 — l'index ne casse pas le tir
        logger.warning(f"index generation ignoré: {e}")
    return out


async def _generate_image_core(body: dict, background_tasks: BackgroundTasks):
    """Text-to-image via fal.ai FLUX (same FAL_KEY as Seedance). Saves the
    PNG(s) into the images folder so they're immediately usable as Seedance
    start frames. Body: {prompt, n=1, size="portrait_16_9"}. Synchronous —
    FLUX schnell returns in ~2-4s per image."""
    prompt = (body.get("prompt") or "").strip()
    if not prompt:
        raise HTTPException(400, "prompt is required")
    n = max(1, min(4, int(body.get("n") or 1)))
    size = body.get("size") or "portrait_16_9"
    # option de style (chantier vitrail 27/08): style="vitrail" (ou toute
    # famille de la fiche épinglée) → le prompt de l'appelant garde son sujet
    # et gagne le bloc de la famille + les garde-fous; un nom d'artiste tapé
    # par l'utilisateur est épuré avant l'envoi (doctrine du skill).
    style = (body.get("style") or "").strip().lower()
    if style:
        from app.services import style_vitrail as SV
        try:
            prompt = SV.appliquer(SV.epurer_noms(prompt), style)
        except KeyError as e:
            raise HTTPException(400, f"Style inconnu: {e}")
        except (ValueError, FileNotFoundError) as e:
            raise HTTPException(400, f"Style '{style}': {e}")
    model = (body.get("model") or "").strip().lower()
    if not model:
        # Callers that drive the API directly (plan agents, scripts) send no
        # model — honour the saved "Image generator" default instead of
        # silently falling back to FLUX.
        async with async_session_factory() as _s:
            model = (await _atelier_setting(
                _s, "image_model_default")).strip().lower()
        logger.info("images/generate: no model in request, saved default -> "
                    f"{model or 'flux (fallback)'}")
    import httpx as _httpx

    # --- OpenAI gpt-image / dall-e path (per the selected model) -----------
    if model.startswith("gpt-image") or model.startswith("dall-e"):
        if not settings.OPENAI_API_KEY:
            raise HTTPException(400, "OPENAI_API_KEY not configured. Add it in Settings.")
        osize = ("1024x1536" if "portrait" in size
                 else "1536x1024" if "landscape" in size else "1024x1024")
        payload = {"model": model, "prompt": prompt, "n": n, "size": osize}
        try:
            async with _httpx.AsyncClient(verify=SSL_VERIFY, timeout=180.0) as client:
                resp = await client.post(
                    "https://api.openai.com/v1/images/generations",
                    headers={"Authorization": f"Bearer {settings.OPENAI_API_KEY}"},
                    json=payload)
        except Exception as e:
            logger.error(f"OpenAI image gen failed: {e}")
            raise HTTPException(502, f"OpenAI: image generation failed: {e}")
        if resp.status_code != 200:
            logger.error(f"OpenAI image HTTP {resp.status_code}: {resp.text[:300]}")
            raise HTTPException(502, f"OpenAI image error: {resp.text[:200]}")
        data = (resp.json() or {}).get("data", [])
        import base64 as _b64
        saved: list[str] = []
        for it in data:
            fname = f"gen_{uuid4().hex[:8]}.png"
            dest = settings.images_path / fname
            if it.get("b64_json"):
                dest.write_bytes(_b64.b64decode(it["b64_json"]))
                saved.append(fname)
            elif it.get("url"):
                async with _httpx.AsyncClient(verify=SSL_VERIFY, timeout=60.0) as c2:
                    rr = await c2.get(it["url"])
                    rr.raise_for_status()
                    dest.write_bytes(rr.content)
                saved.append(fname)
        if not saved:
            raise HTTPException(502, "OpenAI returned no images")
        logger.info(f"OpenAI {model}: saved {len(saved)} image(s): {saved}")
        return {"images": saved, "prompt": prompt, "model": model}

    # --- Nano Banana (Gemini via fal) --------------------------------------
    if model == "nano-banana":
        from app.services import image_providers as IP
        try:
            out = await IP.generate("nano-banana", prompt, size, n)
        except RuntimeError as e:
            raise HTTPException(502, str(e))
        return {"images": out["images"], "prompt": prompt,
                "model": "nano-banana"}

    # --- fal.ai FLUX path (default) ---------------------------------------
    seed = body.get("seed")
    seed = int(seed) if isinstance(seed, (int, float)) else None
    out = await _flux_generate(prompt, size, n, seed=seed)
    return {"images": out["images"], "prompt": prompt, "model": "flux",
            "seed": out.get("seed")}


_CROP_RATIOS = {"9:16": (9, 16), "16:9": (16, 9), "1:1": (1, 1),
                "4:5": (4, 5), "3:4": (3, 4)}


@router.post("/images/process")
async def process_image(body: dict):
    """Provenance (28/08) : l'enveloppe indexe `retouche` d'un coup les
    fichiers rendus par les 8 ops du cœur ({images:[…]}) — un seul site."""
    out = await _process_image_core(body)
    try:
        if isinstance(out, dict) and out.get("images"):
            await LI.noter(out["images"], "retouche")
    except Exception as e:  # noqa: BLE001 — l'index ne casse pas la retouche
        logger.warning(f"index retouche ignoré: {e}")
    return out


async def _process_image_core(body: dict):
    """Post-traitements des nœuds image du Studio. Body: {op, filename, ...}.
    Ops:
      - crop        {ratio "9:16"}            — recadrage centré, local (PIL)
      - upscale     {mode "ai"|"simple", scale 2} — fal esrgan / PIL Lanczos
      - remove-bg   {method "api"|"local"}    — fal rembg / lib rembg locale
      - edit        {prompt, model, n=1}      — édition par prompt (gpt-image,
                    nano-banana, ou FLUX Kontext par défaut)
      - variations  {n=3, model, prompt?}     — N variantes proches
      - pixel       {target_px, colors|palette, dither, scale} — pixel-art
                    local (PIL, chantier 9b), palettes pico8/gameboy/nes/
                    sweetie16/onebit
      - tile-preview {grid 2|3}               — composite de raccord + score
                    seam_score 0-100 (0 = tuile parfaite, base du 9e)
      - seamless    {method offset|mirror, blend 5-45, target_px 0|64-1024,
                    square} — tuile raccordable locale (PIL, chantier 9e) +
                    seam_before/seam_after dans la réponse
    Retour {images:[filenames]} — sauvées dans la Library comme gen_*.png."""
    op = (body.get("op") or "").strip().lower()
    fname = (body.get("filename") or "").strip()
    src = settings.images_path / fname
    if not fname or not src.is_file():
        raise HTTPException(400, f"Source image not found: {fname or '(none)'}")
    from PIL import Image as _PILImage

    def _save_png(img) -> str:
        out_name = f"gen_{uuid4().hex[:8]}.png"
        img.save(settings.images_path / out_name, format="PNG")
        return out_name

    if op == "crop":
        ratio = body.get("ratio") or "9:16"
        if ratio not in _CROP_RATIOS:
            raise HTTPException(400, f"Unknown ratio: {ratio}")
        rw, rh = _CROP_RATIOS[ratio]
        img = _PILImage.open(src)
        w, h = img.size
        target = rw / rh
        if w / h > target:   # trop large -> rogner les côtés
            nw = int(h * target)
            box = ((w - nw) // 2, 0, (w + nw) // 2, h)
        else:                # trop haut -> rogner haut/bas
            nh = int(w / target)
            box = (0, (h - nh) // 2, w, (h + nh) // 2)
        out_name = _save_png(img.crop(box))
        logger.info(f"images/process crop {ratio}: {fname} -> {out_name}")
        return {"images": [out_name], "op": op}

    if op == "upscale":
        mode = (body.get("mode") or "ai").lower()
        scale = max(2, min(4, int(body.get("scale") or 2)))
        if mode == "simple":
            img = _PILImage.open(src)
            up = img.resize((img.width * scale, img.height * scale),
                            _PILImage.LANCZOS)
            out_name = _save_png(up)
            logger.info(f"images/process upscale simple x{scale}: "
                        f"{fname} -> {out_name}")
            return {"images": [out_name], "op": op}
        if not settings.FAL_KEY:
            raise HTTPException(400, "FAL_KEY not configured (Settings) — "
                                     "use the 'simple' mode instead.")
        import fal_client
        from app.services.fal_service import FalSeedanceClient
        url = await FalSeedanceClient.upload_image(src)
        try:
            result = await fal_client.subscribe_async(
                "fal-ai/esrgan", arguments={"image_url": url, "scale": scale})
        except Exception as e:
            logger.error(f"esrgan upscale failed: {e}")
            raise HTTPException(502, f"fal.ai esrgan: {e}")
        out_url = ((result or {}).get("image") or {}).get("url") or \
            next((im.get("url") for im in (result or {}).get("images", [])
                  if im.get("url")), None)
        if not out_url:
            raise HTTPException(502, "esrgan returned no image")
        from app.services.image_providers import _download
        saved = await _download([out_url])
        logger.info(f"images/process upscale ai x{scale}: "
                    f"{fname} -> {saved[0]}")
        return {"images": saved, "op": op}

    if op == "remove-bg":
        method = (body.get("method") or "api").lower()
        if method == "local":
            try:
                from rembg import remove as _rembg_remove
            except ImportError:
                raise HTTPException(
                    400, "rembg is not installed in this runtime — use the "
                         "'API cloud (fal)' method, or install it with: "
                         "pip install rembg")
            loop = asyncio.get_running_loop()
            data = await loop.run_in_executor(
                None, _rembg_remove, src.read_bytes())
            out_name = f"gen_{uuid4().hex[:8]}.png"
            (settings.images_path / out_name).write_bytes(data)
            logger.info(f"images/process remove-bg local: "
                        f"{fname} -> {out_name}")
            return {"images": [out_name], "op": op}
        if not settings.FAL_KEY:
            raise HTTPException(400, "FAL_KEY not configured (Settings) — "
                                     "use the 'local (rembg)' method.")
        import fal_client
        from app.services.fal_service import FalSeedanceClient
        url = await FalSeedanceClient.upload_image(src)
        try:
            result = await fal_client.subscribe_async(
                "fal-ai/imageutils/rembg", arguments={"image_url": url})
        except Exception as e:
            logger.error(f"fal rembg failed: {e}")
            raise HTTPException(502, f"fal.ai rembg: {e}")
        out_url = ((result or {}).get("image") or {}).get("url") or \
            next((im.get("url") for im in (result or {}).get("images", [])
                  if im.get("url")), None)
        if not out_url:
            raise HTTPException(502, "rembg returned no image")
        from app.services.image_providers import _download
        saved = await _download([out_url])
        logger.info(f"images/process remove-bg api: {fname} -> {saved[0]}")
        return {"images": saved, "op": op}

    if op == "pixel":
        from app.services.pixel_ops import normalize_pixel_opts, pixelate
        try:
            popts = normalize_pixel_opts(body)
        except ValueError as e:
            raise HTTPException(400, str(e))
        with _PILImage.open(src) as img:
            out = pixelate(img, popts)
        out_name = _save_png(out)
        logger.info(f"images/process pixel {popts['palette'] or popts['colors']}"
                    f"@{popts['target_px']}px: {fname} -> {out_name}")
        return {"images": [out_name], "op": op, "pixel": popts,
                "size": list(out.size)}

    if op == "seamless":
        from app.services.pixel_ops import (make_seamless,
                                            normalize_seamless_opts,
                                            seam_score)
        try:
            sopts = normalize_seamless_opts(body)
        except ValueError as e:
            raise HTTPException(400, str(e))
        with _PILImage.open(src) as img:
            before = seam_score(img)
            out = make_seamless(img, sopts)
        after = seam_score(out)
        out_name = _save_png(out)
        logger.info(f"images/process seamless {sopts['method']} "
                    f"blend={sopts['blend']} seam {before}->{after}: "
                    f"{fname} -> {out_name}")
        return {"images": [out_name], "op": op, "method": sopts["method"],
                "seam_before": before, "seam_after": after,
                "size": list(out.size)}

    if op == "tile-preview":
        from app.services.pixel_ops import tile_preview
        try:
            grid = int(body.get("grid") or 2)
        except (TypeError, ValueError):
            grid = 0
        if grid not in (2, 3):
            raise HTTPException(400, "grid must be 2 or 3")
        with _PILImage.open(src) as img:
            comp, score = tile_preview(img, grid)
        out_name = _save_png(comp)
        logger.info(f"images/process tile-preview {grid}x{grid} "
                    f"score={score}: {fname} -> {out_name}")
        return {"images": [out_name], "op": op, "grid": grid,
                "seam_score": score}

    if op in ("edit", "variations"):
        n = max(1, min(4, int(body.get("n") or (3 if op == "variations"
                                                else 1))))
        prompt = (body.get("prompt") or "").strip()
        if op == "variations" and not prompt:
            prompt = ("a close variation of the same subject, same style, "
                      "same framing, slightly different details")
        if not prompt:
            raise HTTPException(400, "prompt is required for edit")
        model = (body.get("model") or "").strip().lower()
        if not model:
            async with async_session_factory() as _s:
                model = (await _atelier_setting(
                    _s, "image_model_default")).strip().lower()
        size = body.get("size") or "portrait_16_9"
        if model.startswith("gpt-image") or model.startswith("dall-e") \
                or model == "nano-banana":
            from app.services import image_providers as IP
            try:
                out = await IP.generate(model, prompt, size, n,
                                        image_path=src)
            except RuntimeError as e:
                raise HTTPException(502, str(e))
            logger.info(f"images/process {op} via {model}: "
                        f"{fname} -> {out['images']}")
            return {"images": out["images"], "op": op, "model": model}
        # défaut: FLUX Kontext (génération conditionnée par l'image)
        ratio = _EDIT_RATIO.get(size)
        out = await _flux_generate(prompt, size, n,
                                   model="fal-ai/flux-kontext/dev",
                                   image_path=src, ratio=ratio)
        logger.info(f"images/process {op} via kontext: "
                    f"{fname} -> {out['images']}")
        return {"images": out["images"], "op": op, "model": "flux-kontext"}

    raise HTTPException(400, f"Unknown op: {op}")


# cadre → ratio des modèles edit (resolution_mode Kontext / aspect Banana)
_EDIT_RATIO = {"portrait_16_9": "9:16", "portrait_4_3": "3:4",
               "square_hd": "1:1", "square": "1:1",
               "landscape_4_3": "4:3", "landscape_16_9": "16:9"}


async def _flux_generate(prompt: str, size: str, n: int,
                         seed: int | None = None,
                         model: str = "fal-ai/flux/schnell",
                         image_path: Path | None = None,
                         ratio: str | None = None,
                         guidance: float | None = None) -> dict:
    """Génération d'image fal → Library (v1.20: multi-modèles).

    - schnell (défaut, /images/generate): rapide et pas cher.
    - fal-ai/flux/dev: meilleure adhérence aux consignes de mise en page —
      utilisé pour les planches de référence.
    - fal-ai/flux-kontext/dev + image_path: génération CONDITIONNÉE par une
      image de référence (préserve l'identité du sujet) — utilisé quand
      l'entité a une image d'inspiration. Kontext cadre par défaut sur
      l'image d'entrée (match_input): `ratio` (ex. "9:16") force un autre
      cadre via resolution_mode — indispensable pour les panneaux corps en
      pied chaînés sur un headshot (sinon figure tassée ou coupée).
    Retourne {"images": [filenames], "seed": <seed utilisé>}."""
    import httpx as _httpx
    if not settings.FAL_KEY:
        raise HTTPException(400, "FAL_KEY not configured. Add it in Settings.")
    if size not in ("square_hd", "square", "portrait_4_3", "portrait_16_9",
                    "landscape_4_3", "landscape_16_9"):
        size = "portrait_16_9"
    import fal_client
    arguments: dict = {"prompt": prompt, "num_images": n}
    if image_path is not None:
        from app.services.fal_service import FalSeedanceClient
        arguments["image_url"] = await FalSeedanceClient.upload_image(image_path)
        if ratio:
            arguments["resolution_mode"] = ratio
    else:
        arguments["image_size"] = size
        arguments["enable_safety_checker"] = True
    if guidance is not None:
        # panneaux corps Kontext: adhérence renforcée aux consignes de
        # proportions (défaut 2.5 — trop lâche, la tête de la réf gagne)
        arguments["guidance_scale"] = guidance
    if seed is not None:
        arguments["seed"] = seed
    try:
        result = await fal_client.subscribe_async(model, arguments=arguments)
    except Exception as e:
        logger.error(f"image generation failed ({model}): {e}")
        raise HTTPException(502, f"fal.ai ({model}): image generation failed: {e}")
    urls = [im.get("url") for im in (result or {}).get("images", [])
            if im.get("url")]
    if not urls:
        raise HTTPException(502, "FLUX returned no images")
    used_seed = result.get("seed")
    used_seed = int(used_seed) if isinstance(used_seed, (int, float)) else seed
    saved = []
    async with _httpx.AsyncClient(verify=SSL_VERIFY, timeout=60.0) as client:
        for u in urls:
            fname = f"gen_{uuid4().hex[:8]}.png"
            dest = settings.images_path / fname
            r = await client.get(u)
            r.raise_for_status()
            dest.write_bytes(r.content)
            saved.append(fname)
    logger.info(f"FLUX: saved {len(saved)} image(s), seed={used_seed}: {saved}")
    return {"images": saved, "seed": used_seed}


@router.post("/images/fetch")
async def fetch_image(body: dict):
    """Download a remote image URL into the images folder so it's usable as a
    Studio slot (e.g. a news headline's own image). Body: {url}. -> {filename}."""
    url = (body.get("url") or "").strip()
    if not url.startswith(("http://", "https://")):
        raise HTTPException(400, "A valid image URL is required")
    if _is_private_host(httpx.URL(url).host):
        raise HTTPException(400, "Refusing to fetch a private/loopback address")
    try:
        async with _remote_image_client() as client:
            r = await client.get(url, headers={"User-Agent": "Mozilla/5.0"})
            r.raise_for_status()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(502, f"Image fetch failed: {e}")
    if len(r.content) > 25 * 1024 * 1024:
        raise HTTPException(400, "Image too large (max 25 MB)")
    ct = (r.headers.get("content-type") or "").lower()
    low = url.lower().split("?")[0]
    if "image" not in ct and not low.endswith((".jpg", ".jpeg", ".png", ".webp", ".gif")):
        raise HTTPException(415, "That URL is not an image")
    ext = ".jpg" if ("jpeg" in ct or "jpg" in ct or low.endswith((".jpg", ".jpeg"))) \
        else ".webp" if ("webp" in ct or low.endswith(".webp")) else ".png"
    fname = f"gen_{uuid4().hex[:8]}{ext}"
    (settings.images_path / fname).write_bytes(r.content)
    logger.info(f"Fetched image -> {fname} ({len(r.content) // 1024} KB)")
    await LI.noter([fname], "import_url")
    return {"filename": fname}


@router.get("/video-models")
async def list_video_models():
    """W-a — video-generation models for the Studio Generator node + Quick.

    Every registry entry is returned with an `available` flag (fal entries
    need FAL_KEY, Google-native ones need GEMINI_API_KEY) so the select can
    show-but-disable what a missing key would unlock. `usd_per_s` reflects
    pricing.json overrides; the card badge and topbar ≈$ derive from it."""
    from app.services.fal_service import VIDEO_MODELS, DEFAULT_VIDEO_MODEL
    from app.services import pricing as _pricing
    p = _pricing.load()
    rates = p.get("video_usd_per_s") or {}
    out = []
    for mid, m in VIDEO_MODELS.items():
        available = (settings.has_gemini if m["provider"] == "google"
                     else bool(settings.FAL_KEY))
        out.append({
            "id": mid,
            "label": m["label"],
            "provider": m["provider"],
            "available": available,
            "durations": m["durations"],
            "ratios": m["ratios"],
            "resolutions": m["resolutions"],
            "end_image": m["end_image"],
            "seed": m["seed"],
            "audio_included": (m["provider"] == "google"),
            "usd_per_s": rates.get(mid) or {},
        })
    return {"models": out, "default": DEFAULT_VIDEO_MODEL}


@router.get("/image-models")
async def list_image_models():
    """Image-generation models available given the registered API keys. The
    Studio image picker uses this to show only usable models; the chosen id is
    sent back as `model` to POST /images/generate. Persisted client-side."""
    out = []
    if settings.FAL_KEY:
        out.append({"id": "flux", "label": "FLUX schnell",
                    "provider": "fal", "note": "fast, low cost"})
        out.append({"id": "nano-banana", "label": "Nano Banana (Gemini)",
                    "provider": "fal", "note": "strong edits"})
        # Nano Banana Pro était au catalogue (image_providers, pricing, série
        # Cardforge) mais ABSENT de cette liste — donc d'aucun sélecteur UI,
        # qui la lisent tous. Corrigé 28/08/2026.
        out.append({"id": "nano-banana-pro", "label": "Nano Banana Pro (Gemini 3)",
                    "provider": "fal", "note": "2K/4K, 14 refs"})
    if settings.OPENAI_API_KEY:
        out.append({"id": "gpt-image-2", "label": "GPT Image 2",
                    "provider": "openai", "note": "best quality"})
        out.append({"id": "gpt-image-1", "label": "GPT Image 1",
                    "provider": "openai", "note": "balanced"})
        out.append({"id": "gpt-image-1-mini", "label": "GPT Image 1 mini",
                    "provider": "openai", "note": "cheapest OpenAI"})
    async with async_session_factory() as _s:
        configured = (await _atelier_setting(
            _s, "image_model_default")).strip().lower()
    if configured and configured not in {m["id"] for m in out}:
        configured = ""  # stale default (key removed) — ignore it
    return {"models": out, "configured": configured,
            "default": configured or ("flux" if settings.FAL_KEY
                                      else (out[0]["id"] if out else ""))}


# ═════════════════════ Atelier Chapitre (v1.17, P1) ═════════════════════
# Persistent story bible (characters/places/objects with a seeded reference
# image) + chapters (script text + annotated entity spans). Consumed by the
# /atelier page. Spec: docs/superpowers/specs/2026-07-05-atelier-chapitre-design.md

_ENTITY_KINDS = ("character", "place", "object", "date", "ambiance", "decor")
# v1.20 — chaque kind génère une PLANCHE DE RÉFÉRENCE multi-vues en UNE seule
# image (un seul seed = identité cohérente sous tous les angles), à la manière
# des model sheets de studio. La planche est l'ancre de cohérence; la recette
# exacte (prompt+seed) est stockée et rejouable à l'identique.
_KIND_SIZE = {"character": "landscape_16_9", "place": "landscape_16_9",
              "object": "landscape_16_9", "date": "landscape_16_9",
              "ambiance": "landscape_16_9", "decor": "landscape_16_9"}
_KIND_PREFIX = {
    # Personnage v3 = DEUX passes chaînées (une rangée par image, jamais deux):
    # 1) turnaround plein pied, 2) gros plans visage via Kontext conditionné
    # sur la passe 1 (même visage garanti). Prompts ci-dessous.
    "character": ("character model sheet (turnaround), wide landscape "
                  "composition: exactly FOUR full-body views of the SAME "
                  "character in ONE single row, all standing on one shared "
                  "ground line, evenly spaced, same scale — (1) front view, "
                  "(2) left profile, (3) right profile, (4) back view — "
                  "identical outfit, hairstyle and colors in all four views; "
                  "accurate realistic human proportions, figures about seven "
                  "and a half heads tall, natural standing posture; flat "
                  "light studio background; absolutely no text, no titles, "
                  "no lettering, no logos"),
    "place": ("location reference board of the SAME location, consistent "
              "architecture and palette: one wide establishing shot, one "
              "alternate angle, one key detail close-up, no characters, "
              "no text"),
    "object": ("prop reference sheet, the SAME object with identical design "
               "in 3 angles side by side: front, three-quarter, back, plus "
               "one detail close-up; flat neutral background, no text"),
    "date": ("era/period reference board: three evocative frames of the "
             "same time period side by side (architecture, costume, "
             "technology), consistent palette, no text"),
    "ambiance": ("lighting and atmosphere mood board: three frames of the "
                 "SAME mood side by side (light, weather, tone) plus a "
                 "color palette strip, no text"),
    "decor": ("set-dressing reference board: furniture, materials and "
              "textures of the SAME set in 3 framed views plus one texture "
              "close-up, consistent palette, no text"),
}

# Passe 2 des personnages : gros plans visage, Kontext conditionné sur le
# turnaround (identité de visage garantie par le chaînage).
_FACES_PROMPT = ("using the EXACT same character as in the reference sheet — "
                 "same face, same hairstyle, same outfit: one single row of "
                 "exactly THREE head-and-shoulders close-up portraits, evenly "
                 "spaced, same scale — (1) front view, (2) left profile, "
                 "(3) right profile; flat light studio background; absolutely "
                 "no text, no titles, no lettering")


async def _atelier_setting(session, key: str) -> str:
    from app.services.storage import AtelierSetting
    row = await session.get(AtelierSetting, key)
    return (row.value or "").strip() if row else ""


@router.get("/atelier/settings")
async def get_atelier_settings():
    from app.services.storage import AtelierSetting, async_session_factory
    from sqlalchemy import select
    async with async_session_factory() as session:
        rows = (await session.execute(select(AtelierSetting))).scalars().all()
        return {"settings": {r.key: r.value or "" for r in rows}}


@router.get("/atelier/providers")
async def list_image_providers():
    """Générateurs d'images disponibles (selon les clés configurées), avec
    l'indicateur seeds (déterminisme des recettes)."""
    from app.services import image_providers as IP
    return {"providers": IP.available(), "default": "flux"}


@router.get("/voice/providers")
async def list_voice_providers():
    """Fournisseurs de voix (spec voicebox 2026-07-11) : disponibilité de
    chacun (clé ElevenLabs / Voicebox local joignable), réglage atelier
    voice_provider et provider effectivement résolu."""
    from app.services import voice_providers as VP
    from app.services.storage import async_session_factory
    async with async_session_factory() as session:
        configured = await _atelier_setting(session, "voice_provider")
    loop = asyncio.get_running_loop()
    providers = await loop.run_in_executor(None, VP.available)
    resolved = await loop.run_in_executor(
        None, lambda: VP.resolve_provider(configured))
    return {"providers": providers, "configured": configured,
            "resolved": resolved}


@router.post("/atelier/style/propose")
async def propose_art_direction(body: dict):
    """v1.23 (DA) — l'agent lit un extrait représentatif du manuscrit (ton,
    époque, genre, indices visuels rédigés) et propose 4 directions
    artistiques motivées. Persistées dans atelier_settings.style_proposals.
    Body: {chapter_id?} — sinon: tous les chapitres, concaténés."""
    from app.services import manuscript_agent as MA
    from app.services.storage import (Chapter, BibleEntity, AtelierSetting,
                                      async_session_factory)
    from app.services.summarizer import available
    from sqlalchemy import select
    if not available():
        raise HTTPException(400, "Aucun LLM configuré (Réglages → clés API).")
    async with async_session_factory() as session:
        if body.get("chapter_id"):
            ch = await session.get(Chapter, body["chapter_id"])
            if not ch:
                raise HTTPException(404, "Chapter not found")
            texts = [ch.script_text or ""]
        else:
            rows = (await session.execute(select(Chapter))).scalars().all()
            texts = [c.script_text or "" for c in rows]
        excerpt = "\n\n".join(t[:3000] for t in texts if t.strip())[:9000]
        if len(excerpt) < 200:
            raise HTTPException(400, "Pas assez de texte — importe le manuscrit d'abord.")
        ents = (await session.execute(select(BibleEntity))).scalars().all()
        names = [e.name for e in ents]
    loop = asyncio.get_running_loop()
    props = await loop.run_in_executor(
        None, lambda: MA.propose_styles(excerpt, names))
    if not props:
        raise HTTPException(502, "La proposition de DA a échoué — réessaie.")
    async with async_session_factory() as session:
        row = await session.get(AtelierSetting, "style_proposals")
        val = json.dumps(props, ensure_ascii=False)
        if row:
            row.value = val
        else:
            session.add(AtelierSetting(key="style_proposals", value=val))
        await session.commit()
    return {"proposals": props, "presets": MA.STYLE_PRESETS}


@router.put("/atelier/settings")
async def put_atelier_settings(body: dict):
    """Upsert de réglages {key: value}. Clés: global_style, …"""
    from app.services.storage import AtelierSetting, async_session_factory
    async with async_session_factory() as session:
        for k, v in (body or {}).items():
            if not isinstance(k, str) or len(k) > 60:
                continue
            row = await session.get(AtelierSetting, k)
            if row:
                row.value = str(v or "")
            else:
                session.add(AtelierSetting(key=k, value=str(v or "")))
        await session.commit()
    return await get_atelier_settings()


def _entity_dict(e) -> dict:
    import json as _json

    def _jload(v):
        try:
            return _json.loads(v) if v else []
        except Exception:
            return []
    return {"id": e.id, "kind": e.kind, "name": e.name,
            "description": e.description or "",
            "ref_image": e.ref_image, "seed": e.seed,
            "style_notes": e.style_notes or "",
            "inspiration_images": _jload(e.inspiration_images),
            "aliases": _jload(getattr(e, "aliases", None)),
            "evidence": _jload(getattr(e, "evidence", None)),
            "has_recipe": bool(getattr(e, "prompt_recipe", None)),
            "face_image": getattr(e, "face_image", None),
            "voice_id": getattr(e, "voice_id", None),
            "voice_name": getattr(e, "voice_name", None),
            "voice_prev": getattr(e, "voice_prev", None),
            "model3d_job": getattr(e, "model3d_job", None),
            "model3d_file": getattr(e, "model3d_file", None),
            "created_at": e.created_at.isoformat() if e.created_at else None,
            "updated_at": e.updated_at.isoformat() if e.updated_at else None}


@router.get("/bible/entities")
async def list_bible_entities(kind: str | None = None):
    """List bible entities, optionally filtered by kind."""
    from app.services.storage import BibleEntity, async_session_factory
    from sqlalchemy import select
    async with async_session_factory() as session:
        q = select(BibleEntity).order_by(BibleEntity.created_at.asc())
        if kind in _ENTITY_KINDS:
            q = q.where(BibleEntity.kind == kind)
        rows = (await session.execute(q)).scalars().all()
    return {"entities": [_entity_dict(e) for e in rows]}


@router.post("/bible/entities")
async def create_bible_entity(body: dict):
    """Create an entity. Body: {kind, name, description?, style_notes?,
    inspiration_images?}."""
    from app.services.storage import BibleEntity, async_session_factory
    import json as _json
    kind = (body.get("kind") or "").strip()
    name = (body.get("name") or "").strip()
    if kind not in _ENTITY_KINDS:
        raise HTTPException(400, f"kind must be one of {_ENTITY_KINDS}")
    if not name:
        raise HTTPException(400, "name is required")
    e = None
    async with async_session_factory() as session:
        from app.services.storage import BibleEntity as BE
        eid = str(uuid4())
        e = BE(id=eid, kind=kind, name=name[:120],
               description=body.get("description") or "",
               style_notes=body.get("style_notes") or "",
               inspiration_images=_json.dumps(
                   body.get("inspiration_images") or []),
               created_at=datetime.utcnow(), updated_at=datetime.utcnow())
        session.add(e)
        await session.commit()
        await session.refresh(e)
    return _entity_dict(e)


@router.put("/bible/entities/{entity_id}")
async def update_bible_entity(entity_id: str, body: dict):
    """Update name/description/style_notes/inspiration_images (partial)."""
    from app.services.storage import BibleEntity, async_session_factory
    import json as _json
    async with async_session_factory() as session:
        e = await session.get(BibleEntity, entity_id)
        if not e:
            raise HTTPException(404, "Entity not found")
        if "name" in body and (body["name"] or "").strip():
            e.name = body["name"].strip()[:120]
        if "description" in body:
            e.description = body["description"] or ""
        if "style_notes" in body:
            e.style_notes = body["style_notes"] or ""
        if "inspiration_images" in body:
            e.inspiration_images = _json.dumps(body["inspiration_images"] or [])
        if "aliases" in body:
            e.aliases = _json.dumps(body["aliases"] or [])
        # v1.21 — casting voix (choix manuel ou application d'une suggestion)
        for vk in ("voice_id", "voice_name", "voice_prev",
                   "model3d_job", "model3d_file"):
            if vk in body:
                setattr(e, vk, body[vk] or None)
        # v1.17.1 — allow re-linking a reference / pinning a seed directly
        # (used by recovery tooling and future "use this Library image as ref").
        if "ref_image" in body:
            e.ref_image = (Path(body["ref_image"]).name
                           if body["ref_image"] else None)
        if "seed" in body:
            sd = body["seed"]
            e.seed = int(sd) if isinstance(sd, (int, float)) else None
        e.updated_at = datetime.utcnow()
        await session.commit()
        await session.refresh(e)
        return _entity_dict(e)


@router.post("/bible/entities/{entity_id}/model3d")
async def generate_bible_model3d(entity_id: str, background_tasks: BackgroundTasks,
                                 body: dict = None):
    """Verrouille une entité de la bible EN 3D — spec Magnific §9.1.

    « Employer le flux image → 3D lorsque l'application a besoin de
    verrouiller un produit, accessoire, véhicule, élément de décor ou
    personnage stylisé » : jusqu'ici la bible verrouillait en 2D seulement
    (planche + seed). Le maillage devient l'ancrage géométrique, réutilisable
    par tous les chapitres — et exportable en GLB vers un moteur.

    Body: {image_filename?, besoin?, engine?, multiview?, views?, textures?,
    quality?, formats?}. `besoin` (hero|prop|decor|rig|realtime|brouillon)
    choisit le moteur ET ses options, avec la justification rendue.

    Piège traité : la référence d'un personnage est une PLANCHE composite
    (board_*.png, plusieurs vues côte à côte). La donner à un moteur
    image→3D produirait un monstre — la route refuse et le dit, au lieu de
    payer une génération absurde.
    """
    from datetime import datetime as _dtu
    import json as _json
    from app.services import asset3d_service as A3
    from app.services.asset3d_service import generate_asset3d
    from app.services.storage import BibleEntity, JobRecord, async_session_factory

    body = body or {}
    if not settings.FAL_KEY:
        raise HTTPException(400, "FAL_KEY not configured. Add it in Settings.")

    async with async_session_factory() as session:
        e = await session.get(BibleEntity, entity_id)
        if e is None:
            raise HTTPException(404, "Entity not found")
        nom, kind = e.name, e.kind
        ref = e.ref_image

    # 1. quelle image nourrit le moteur ?
    demande = Path(str(body.get("image_filename") or "")).name
    src = demande or (ref or "")
    if not src:
        raise HTTPException(
            400, f"« {nom} » n'a pas de référence : génère d'abord sa planche "
                 "(🎨 Planche), puis choisis une vue unique.")
    # Le refus porte sur le nom EFFECTIVEMENT retenu, d'où qu'il vienne : le
    # tester seulement sur le chemin de repli laissait passer la planche dès
    # que le client la nommait — c'est-à-dire toujours, puisque le sélecteur
    # de la Bibliothèque affiche aussi les planches.
    if src.startswith("board_") and not body.get("force_planche"):
        raise HTTPException(
            400, f"« {src} » est une PLANCHE composite : plusieurs vues sur "
                 "une même image. Un moteur image→3D a besoin d'UNE vue — il "
                 "reconstruirait les quatre côte à côte. Choisis une vue "
                 "seule dans la Bibliothèque, ou passe `force_planche: true` "
                 "si tu sais ce que tu fais.")
    if not (settings.images_path / src).is_file():
        raise HTTPException(400, f"Image introuvable dans la Bibliothèque : {src}")

    # 2. moteur : besoin motivé, ou choix explicite
    besoin = str(body.get("besoin") or "").strip().lower()
    reco = None
    opts: dict = {}
    if besoin:
        try:
            reco = A3.recommend_engine(besoin)
        except ValueError as ve:
            raise HTTPException(400, str(ve))
        opts = dict(reco["opts"])
        engine = reco["engine"]
    else:
        engine = str(body.get("engine") or "tripo").lower()
    if body.get("engine"):
        engine = str(body["engine"]).lower()
    try:
        caps = A3.engine_caps(engine)
    except ValueError as ve:
        raise HTTPException(400, str(ve))

    for k in ("multiview", "views", "textures", "quality", "tpose", "formats"):
        if k in body:
            opts[k] = body[k]
    # Mêmes validations d'entrée que la route sœur POST /assets/3d : sans
    # elles, `views` non entier ferait tomber le job en arrière-plan, et
    # CHAQUE format supplémentaire est une génération facturée en plus
    # (asset3d_service ré-appelle le moteur par format manquant).
    if "views" in opts:
        try:
            int(opts["views"])
        except (TypeError, ValueError):
            raise HTTPException(400, "views doit être un entier (1-4).")
    if "formats" in opts:
        fmts = opts["formats"]
        if not isinstance(fmts, list) or not all(isinstance(f, str) for f in fmts):
            raise HTTPException(400, "formats doit être une liste de chaînes.")
        inconnus = sorted({f.lower() for f in fmts} - set(caps["formats"]))
        if inconnus:
            raise HTTPException(
                400, f"{caps['label']} n'exporte pas {', '.join(inconnus)} "
                     f"(formats : {', '.join(caps['formats'])}).")
        extra = len({f.lower() for f in fmts} - {"glb"})
        if extra and not body.get("confirm_formats"):
            raise HTTPException(
                400, f"{extra} format(s) en plus du GLB = {extra} "
                     "génération(s) facturée(s) de plus. Renvoie avec "
                     "`confirm_formats: true` pour accepter.")
    if opts.get("multiview") and not caps["multiview"]:
        raise HTTPException(
            400, f"{caps['label']} ne prend qu'une vue (max_images "
                 f"{caps['max_images']}) — désactive multiview ou change de moteur.")
    if opts.get("tpose") and not caps["tpose"]:
        raise HTTPException(
            400, f"{caps['label']} ne sait pas demander une T-pose — "
                 "utilise le besoin « rig » (Rodin).")

    payload = {**opts, "engine": engine, "image_filename": src,
               "subject": f"{nom} — the same {kind}, consistent design",
               "title": f"3D · {nom}"}
    job_id = str(uuid4())
    short = job_id[:8]
    async with async_session_factory() as s:
        s.add(JobRecord(
            id=job_id, status=JobStatus.GENERATING_VIDEO.value, progress=10,
            title=f"3D · {nom} ({caps['label']})",
            image_filename=f"asset3d_{short}",
            provider="asset3d", current_step="Generating 3D"))
        await s.commit()

    async def on_step(label, pct):
        async with async_session_factory() as s2:
            jr2 = await s2.get(JobRecord, job_id)
            if jr2 is not None:
                jr2.current_step, jr2.progress = label, int(pct)
                await s2.commit()

    async def _run():
        try:
            r = await generate_asset3d(payload, short, on_step=on_step)
            async with async_session_factory() as s:
                jr = await s.get(JobRecord, job_id)
                if jr is not None:
                    jr.status, jr.progress = JobStatus.DONE.value, 100
                    jr.final_video_path = r.get("glb")
                    if r.get("preview"):
                        jr.image_filename = "preview.png"
                    jr.current_step = "Complete"
                    jr.completed_at = _dtu.utcnow()
                    jr.cost_meta = _json.dumps(
                        {"engine": r["engine"], "files": r["files"],
                         "shots": r["shots"], "job": short,
                         "entity_id": entity_id,
                         "skipped_formats": r.get("skipped_formats") or []})
                # l'ancrage : l'entité porte désormais son maillage
                ent = await s.get(BibleEntity, entity_id)
                if ent is not None:
                    ent.model3d_job = short
                    ent.model3d_file = "model.glb"
                    ent.updated_at = _dtu.utcnow()
                await s.commit()
        except Exception as ex:
            logger.exception(f"bible model3d {job_id} failed: {ex}")
            async with async_session_factory() as s:
                jr = await s.get(JobRecord, job_id)
                if jr is not None:
                    jr.status, jr.error = JobStatus.FAILED.value, str(ex)
                    jr.current_step = "Failed"
                    await s.commit()

    background_tasks.add_task(_run)
    return {"job_id": job_id, "status": "queued", "asset3d_job": short,
            "engine": engine, "image_filename": src,
            "besoin": reco["besoin"] if reco else None,
            "why": reco["why"] if reco else None,
            # étape suivante conseillée du besoin (ex. budget de triangles
            # après un asset « temps réel ») — rendue pour être exécutable,
            # jamais lancée d'office
            "apres_generation": (reco or {}).get("apres_generation")}


@router.delete("/bible/entities/{entity_id}")
async def delete_bible_entity(entity_id: str):
    from app.services.storage import BibleEntity, async_session_factory
    async with async_session_factory() as session:
        e = await session.get(BibleEntity, entity_id)
        if not e:
            raise HTTPException(404, "Entity not found")
        await session.delete(e)
        await session.commit()
    return {"ok": True}


@router.post("/bible/entities/{entity_id}/generate")
async def generate_bible_reference(entity_id: str, body: dict):
    """Generate the entity's canonical reference BOARD (multi-view sheet,
    FLUX + seed).

    Body: {seed?: int, use_recipe?: bool}. use_recipe=true rejoue la recette
    stockée À L'IDENTIQUE (même prompt + même seed → même image, FLUX est
    déterministe) — l'ancre de cohérence. Sans seed → aléatoire; le seed et
    la recette exacte sont stockés sur l'entité."""
    from app.services.storage import BibleEntity, async_session_factory
    import json as _json
    async with async_session_factory() as session:
        e = await session.get(BibleEntity, entity_id)
        if not e:
            raise HTTPException(404, "Entity not found")
        recipe = None
        if body.get("use_recipe") and e.prompt_recipe:
            try:
                recipe = _json.loads(e.prompt_recipe)
            except Exception:
                recipe = None
        # v1.20.2 — PLANCHE COMPOSITE: chaque panneau est généré séparément
        # (net, bien proportionné) puis la planche est assemblée PAR CODE —
        # layout garanti. Identité: panneau 1 conditionné sur l'inspiration
        # de l'utilisateur (Kontext) si présente, panneaux suivants chaînés
        # sur le panneau 1.
        from app.services import board_service as BS
        plan = BS.PANEL_PLANS.get(e.kind) or BS.PANEL_PLANS["object"]
        insp_file = None
        try:
            for f in (_json.loads(e.inspiration_images)
                      if e.inspiration_images else []):
                if (settings.images_path / Path(f).name).is_file():
                    insp_file = Path(f).name
                    break
        except Exception:
            insp_file = None
        seeds_by_key: dict = {}
        # provider du projet (réglage DA) — la recette fige le sien
        provider = await _atelier_setting(session, "image_provider") or "flux"
        style_ref = await _atelier_setting(session, "style_ref_image")
        if style_ref and not (settings.images_path / Path(style_ref).name).is_file():
            style_ref = ""
        if recipe and recipe.get("v") == 2:
            insp_file = recipe.get("ref_file") or insp_file
            provider = recipe.get("provider") or provider
            style_ref = recipe.get("style_ref") or style_ref
            seeds_by_key = {p["key"]: p.get("seed")
                            for p in (recipe.get("panels") or [])}
        elif not recipe:
            pass
        desc = (e.description or "").strip()
        if not desc and not (recipe and recipe.get("v") == 2):
            raise HTTPException(400, "Add a description before generating")
        subj = f" Subject: {e.name}. {desc}"
        # style: l'override de l'entité prime, sinon le STYLE GLOBAL du projet
        # (réglage atelier_settings.global_style) — cohérence de réalisation
        # sur toutes les planches, override ponctuel par entité possible.
        style_src = (e.style_notes or "").strip() or \
            await _atelier_setting(session, "global_style")
        style = f". Style: {style_src}" if style_src else ""
        # canon de proportions (DA2): la recette fige le sien, sinon le
        # réglage explicite style_canon, sinon auto-détection par mots-clés
        # du style — défaut: canon académique De Vinci (7.5-8 têtes).
        from app.services import manuscript_agent as MA
        canon_pref = await _atelier_setting(session, "style_canon")
        if recipe and recipe.get("v") == 2 and recipe.get("canon"):
            canon_pref = recipe["canon"]
        canon_key = MA.resolve_canon(
            style_src, canon_pref if canon_pref != "auto" else None)
        canon = MA.PROPORTION_CANONS[canon_key]
        req_seed = body.get("seed")
        req_seed = int(req_seed) if isinstance(req_seed, (int, float)) else None
        panels: dict[str, str] = {}
        recipe_panels = []
        from app.services import image_providers as IP
        from app.services import proportion_qc as PQC
        # leçons apprises (QC proportions passés): consigne corrective
        # persistée par canon — appliquée d'office aux prompts corps.
        lessons = PQC.load_lessons(await _atelier_setting(session,
                                                          "canon_lessons"))
        lesson_hint = PQC.lesson_hint(lessons, canon_key)
        for key, ptxt, chain_on, p1size in plan["panels"]:
            # injection du canon: proportions du corps ({PROPORTIONS}) et
            # traits du visage ({FACE}) selon le style de la DA.
            is_body = p1size == "CANON"      # panneau corps en pied (v7)
            # v1.25.1: le canon anatomique ("X heads tall") est doublé de sa
            # contrainte en coordonnées IMAGE ("framing": la tête n'occupe
            # qu'un N-ième de la hauteur du cadre) — la diffusion respecte
            # mieux les fractions du cadre que les têtes anatomiques.
            prop = canon["char"] + ("; " + canon["framing"]
                                    if canon.get("framing") else "")
            ptxt = ptxt.replace("{PROPORTIONS}", prop) \
                       .replace("{FACE}", canon["face"])
            ratio = None
            if chain_on:
                prompt = ptxt + style
                if is_body and lesson_hint:
                    prompt += ". " + lesson_hint
                img = settings.images_path / panels[chain_on]
                if is_body:
                    # cadre vertical du canon (leçon tests A/B: sans lui le
                    # modèle edit garde le cadre du headshot → corps tassé)
                    size = canon["frame"]
                    ratio = _EDIT_RATIO.get(size)
                else:
                    size = "landscape_16_9"  # les modèles edit cadrent la réf
            else:
                prompt = ptxt + "." + subj + style
                if e.kind in ("place", "decor", "ambiance", "date"):
                    # lieux/décors: perspective et échelle du même canon
                    prompt += ". " + canon["decor"]
                if insp_file:
                    prompt = ("Using the exact same subject, face and design "
                              "as the reference image, keep its identity and "
                              "art style, but remove any text or lettering: "
                              + prompt)
                    img = settings.images_path / insp_file
                elif style_ref:
                    # référence de STYLE du projet (pas d'identité propre):
                    # conditionne le panneau maître sur son rendu.
                    prompt = ("Reproduce the exact ART STYLE of the "
                              "reference image (medium, line, palette, "
                              "rendering) applied to a NEW subject: " + prompt)
                    img = settings.images_path / Path(style_ref).name
                else:
                    img = None
                size = p1size or "landscape_16_9"
            seed = seeds_by_key.get(key)
            if seed is None and not chain_on:
                seed = req_seed

            async def _gen(p: str):
                if provider == "flux":
                    mdl = ("fal-ai/flux-kontext/dev" if img is not None
                           else "fal-ai/flux/dev")
                    return mdl, await _flux_generate(
                        p, size, 1, seed=seed, model=mdl,
                        image_path=img, ratio=ratio,
                        guidance=(3.5 if is_body and img is not None
                                  else None))
                try:
                    return provider, await IP.generate(
                        provider, p, size, 1, seed=seed,
                        image_path=img, ratio=ratio)
                except RuntimeError as pe:
                    raise HTTPException(502, str(pe))

            model, out = await _gen(prompt)
            # QC proportions (panneau corps maître, hors rejeu de recette):
            # un contrôle vision mesure le nombre de têtes; hors canon →
            # jusqu'à DEUX régénérations correctives (seed libre: chaque
            # retry explore une composition différente), on garde la
            # meilleure mesure, et la leçon est persistée pour que l'agent
            # ne reproduise plus l'erreur sur ce canon.
            if (is_body and key == "front" and e.kind == "character"
                    and not seeds_by_key):
                import asyncio as _aio
                m = await _aio.to_thread(
                    PQC.measure, settings.images_path / out["images"][0])
                verdict = PQC.judge(m, canon)
                had_fail, tries = False, 0
                while verdict and not verdict["ok"] and tries < 2:
                    tries += 1
                    had_fail = True
                    fix = PQC.corrective_clause(verdict, canon)
                    logger.info(f"proportion QC {canon_key}: "
                                f"{verdict['note']} → retry {tries}/2")
                    model2, out2 = await _gen(prompt + ". " + fix)
                    m2 = await _aio.to_thread(
                        PQC.measure, settings.images_path / out2["images"][0])
                    v2 = PQC.judge(m2, canon)
                    if PQC.better(v2, verdict):
                        model, out, prompt = model2, out2, prompt + ". " + fix
                        verdict = v2
                    lessons = PQC.record_lesson(lessons, canon_key,
                                                verdict, fix)
                if verdict:
                    if verdict["ok"] and not had_fail:
                        lessons = PQC.record_success(lessons, canon_key)
                    await put_atelier_settings(
                        {"canon_lessons": PQC.dump_lessons(lessons)})
            panels[key] = out["images"][0]
            recipe_panels.append({"key": key, "prompt": prompt,
                                  "seed": out.get("seed"), "model": model})
        # profils droits = miroir logiciel du profil gauche (direction
        # opposée garantie — la diffusion confond gauche/droite)
        for tgt, src in (plan.get("mirrors") or {}).items():
            panels[tgt] = BS.mirror_panel(settings.images_path, panels[src])
        if plan.get("compose") == "character":
            board = BS.compose_character_board(settings.images_path, panels)
        else:
            rows = [[panels[k] for k in row] for row in plan["rows"]]
            board = BS.compose_board(
                settings.images_path, rows, plan["row_heights"],
                palette_from=(list(panels.values()) if plan.get("palette") else None))
        # provenance : panneaux (gen_*), miroirs et planche (board_*)
        await LI.noter(sorted(set(panels.values())) + [board], "atelier")
        e.ref_image = board
        e.face_image = None            # tout est dans la planche composite
        e.seed = recipe_panels[0].get("seed")
        e.prompt_recipe = _json.dumps(
            {"v": 2, "kind": e.kind, "ref_file": insp_file,
             "provider": provider, "style_ref": style_ref or None,
             "canon": canon_key, "panels": recipe_panels}, ensure_ascii=False)
        e.updated_at = datetime.utcnow()
        await session.commit()
        await session.refresh(e)
        return _entity_dict(e)


async def _fetch_11l_voices() -> list[dict]:
    """Voix ElevenLabs du compte, avec labels (genre/âge/accent/description)
    et preview_url — la matière du casting voix."""
    if not settings.has_voiceover:
        raise HTTPException(400, "Clé ElevenLabs non configurée (Réglages).")
    async with httpx.AsyncClient(timeout=20, verify=SSL_VERIFY) as c:
        r = await c.get("https://api.elevenlabs.io/v1/voices",
                        headers={"xi-api-key": settings.ELEVENLABS_API_KEY})
        r.raise_for_status()
        data = r.json()
    out = []
    for v in (data.get("voices") or []):
        lbl = v.get("labels") or {}
        out.append({"voice_id": v.get("voice_id"), "name": v.get("name"),
                    "category": v.get("category"), "labels": lbl,
                    "preview_url": v.get("preview_url")})
    return [v for v in out if v["voice_id"]]


async def _fetch_casting_voices() -> tuple[str, list[dict]]:
    """Catalogue de voix du provider actif (spec voicebox, étape 3) :
    ElevenLabs (labels riches) ou Voicebox (/profiles mappés au même format).
    Retourne (provider, voices) ; 400 si aucun provider utilisable."""
    from app.services import voice_providers as VP
    from app.services.storage import async_session_factory
    async with async_session_factory() as session:
        configured = await _atelier_setting(session, "voice_provider")
    loop = asyncio.get_running_loop()
    provider = await loop.run_in_executor(
        None, lambda: VP.resolve_provider(configured))
    if provider == "voicebox":
        return "voicebox", await loop.run_in_executor(
            None, VP.list_voicebox_voices)
    if provider == "elevenlabs":
        return "elevenlabs", await _fetch_11l_voices()
    raise HTTPException(400, "Aucun fournisseur de voix disponible — "
                             "configure la clé ElevenLabs ou lance Voicebox.")


@router.post("/bible/entities/{entity_id}/suggest-voice")
async def suggest_entity_voice(entity_id: str, body: dict):
    """v1.21 (B) — casting voix: l'agent croise la fiche du personnage
    (genre, âge, ton déduits de la description) avec le catalogue du provider
    actif (ElevenLabs ou Voicebox, v1.26 étape 3) et propose LA voix + des
    alternatives du même profil. La suggestion est appliquée à l'entité
    (modifiable ensuite)."""
    from app.services.storage import BibleEntity, async_session_factory
    from app.services.summarizer import available, _chat_dispatch
    if not available():
        raise HTTPException(400, "Aucun LLM configuré (Réglages → clés API).")
    provider, voices = await _fetch_casting_voices()
    if not voices:
        raise HTTPException(502, f"Aucune voix disponible ({provider}) — "
                                 "crée des profils dans Voicebox (Voices)."
                            if provider == "voicebox" else
                            "Aucune voix ElevenLabs disponible sur le compte.")
    async with async_session_factory() as session:
        e = await session.get(BibleEntity, entity_id)
        if not e:
            raise HTTPException(404, "Entity not found")
        if e.kind != "character":
            raise HTTPException(400, "Le casting voix ne s'applique qu'aux personnages.")
        roster = [{"voice_id": v["voice_id"], "name": v["name"],
                   "labels": v["labels"],
                   } for v in voices][:120]
        system = ("You are a casting director assigning narration/dialogue "
                  "voices to characters of a narrated animation. Return ONLY "
                  "valid JSON.")
        prompt = (
            f"Character sheet:\nName: {e.name}\nDescription: "
            f"{(e.description or '')[:600]}\n\n"
            f"Available voices from provider '{provider}' (with labels):\n"
            f"{json.dumps(roster, ensure_ascii=False)}\n\n"
            f"Pick the voice that best matches the character's gender, age "
            f"and personality, plus up to 4 ALTERNATES of the same profile "
            f"(same gender / similar age & tone). Some voices may have sparse "
            f"labels (local/cloned voices): infer gender, age and tone from "
            f"the voice name, description and personality fields; prefer a "
            f"voice whose language matches the character sheet's language. "
            f"Return ONLY JSON: "
            f"{{\"best\": \"<voice_id>\", \"alternates\": [\"<voice_id>\", …], "
            f"\"why\": \"<one short sentence in French>\"}}")
        loop = asyncio.get_running_loop()
        out, _prov = await loop.run_in_executor(
            None, lambda: _chat_dispatch(prompt, system, 1200))
        txt = (out or "").strip()
        if txt.startswith("```"):
            txt = re.sub(r"^```[a-zA-Z]*\n?", "", txt)
            txt = re.sub(r"\n?```$", "", txt).strip()
        i, j = txt.find("{"), txt.rfind("}")
        try:
            data = json.loads(txt[i:j + 1]) if i >= 0 and j > i else {}
        except Exception:
            data = {}
        by_id = {v["voice_id"]: v for v in voices}
        best = by_id.get(str(data.get("best") or ""))
        if not best:
            raise HTTPException(502, "La suggestion de voix a échoué — réessaie.")
        alternates = [by_id[a] for a in (data.get("alternates") or [])
                      if a in by_id and a != best["voice_id"]][:4]
        e.voice_id = best["voice_id"]
        e.voice_name = best["name"]
        e.voice_prev = best.get("preview_url")
        e.updated_at = datetime.utcnow()
        await session.commit()
        await session.refresh(e)
        return {"entity": _entity_dict(e), "suggested": best,
                "alternates": alternates,
                "why": str(data.get("why") or "")[:300]}


def _chapter_dict(ch) -> dict:
    import json as _json
    try:
        spans = _json.loads(ch.spans) if ch.spans else []
    except Exception:
        spans = []
    return {"id": ch.id, "title": ch.title, "series": ch.series,
            "script_text": ch.script_text or "", "spans": spans,
            "created_at": ch.created_at.isoformat() if ch.created_at else None,
            "updated_at": ch.updated_at.isoformat() if ch.updated_at else None}


@router.get("/chapters")
async def list_chapters():
    from app.services.storage import Chapter, async_session_factory
    from sqlalchemy import select
    async with async_session_factory() as session:
        rows = (await session.execute(
            select(Chapter).order_by(Chapter.created_at.asc()))).scalars().all()
    return {"chapters": [{"id": c.id, "title": c.title, "series": c.series,
                          "updated_at": c.updated_at.isoformat()
                          if c.updated_at else None} for c in rows]}


@router.post("/chapters")
async def create_chapter(body: dict):
    from app.services.storage import Chapter, async_session_factory
    import json as _json
    async with async_session_factory() as session:
        ch = Chapter(id=str(uuid4()),
                     title=(body.get("title") or "Sans titre")[:200],
                     series=(body.get("series") or None),
                     script_text=body.get("script_text") or "",
                     spans=_json.dumps(body.get("spans") or []),
                     created_at=datetime.utcnow(), updated_at=datetime.utcnow())
        session.add(ch)
        await session.commit()
        await session.refresh(ch)
        return _chapter_dict(ch)


@router.get("/chapters/{chapter_id}")
async def get_chapter(chapter_id: str):
    from app.services.storage import Chapter, async_session_factory
    async with async_session_factory() as session:
        ch = await session.get(Chapter, chapter_id)
        if not ch:
            raise HTTPException(404, "Chapter not found")
        return _chapter_dict(ch)


@router.put("/chapters/{chapter_id}")
async def update_chapter(chapter_id: str, body: dict):
    from app.services.storage import Chapter, async_session_factory
    import json as _json
    async with async_session_factory() as session:
        ch = await session.get(Chapter, chapter_id)
        if not ch:
            raise HTTPException(404, "Chapter not found")
        if "title" in body and (body["title"] or "").strip():
            ch.title = body["title"].strip()[:200]
        if "series" in body:
            ch.series = body["series"] or None
        if "script_text" in body:
            ch.script_text = body["script_text"] or ""
        if "spans" in body:
            ch.spans = _json.dumps(body["spans"] or [])
        ch.updated_at = datetime.utcnow()
        await session.commit()
        await session.refresh(ch)
        return _chapter_dict(ch)


@router.delete("/chapters/{chapter_id}")
async def delete_chapter(chapter_id: str):
    from sqlalchemy import delete as _delete
    from app.services.storage import (Chapter, VectorDocLink,
                                      async_session_factory)
    async with async_session_factory() as session:
        ch = await session.get(Chapter, chapter_id)
        if not ch:
            raise HTTPException(404, "Chapter not found")
        # phase 6 Vectorlab : les liaisons d'instanciation partent avec le
        # chapitre (les documents, eux, ne bougent pas)
        await session.execute(_delete(VectorDocLink)
                              .where(VectorDocLink.chapter_id == chapter_id))
        await session.delete(ch)
        await session.commit()
    return {"ok": True}


# ───────── Atelier P2 (v1.18): storyboard shots ─────────
# Découpage du chapitre en plans (IA entity-aware ou paragraphes), croquis
# low-cost par plan (validation cadrage/rythme AVANT toute prod payante).

_SHOT_TYPES = ("establishing", "wide", "medium", "close-up",
               "extreme close-up", "over-shoulder", "POV", "insert")
_CAMERA_MOVES = ("slow push-in", "slow pull-out", "360-degree orbit",
                 "tracking shot", "handheld with subtle shake",
                 "static, locked-off", "low angle dramatic",
                 "rack focus reveal", "dolly zoom (vertigo effect)",
                 "whip pan transition", "crane shot descending")
_SKETCH_STYLE = ("rough storyboard sketch, loose pencil strokes, monochrome "
                 "gray, simple composition lines, no color, no text")


def _shot_dict(s) -> dict:
    import json as _json
    try:
        ents = _json.loads(s.entities) if s.entities else []
    except Exception:
        ents = []
    return {"id": s.id, "chapter_id": s.chapter_id, "idx": s.idx,
            "source_text": s.source_text or "", "action": s.action or "",
            "entities": ents, "shot_type": s.shot_type,
            "camera_move": s.camera_move, "duration_s": s.duration_s,
            "sketch_image": s.sketch_image, "sketch_seed": s.sketch_seed,
            "prompt": s.prompt or "",
            "motion_recipe": s.motion_recipe, "energy": s.energy}


async def _list_shots(session, chapter_id: str):
    from app.services.storage import Shot
    from sqlalchemy import select
    return (await session.execute(
        select(Shot).where(Shot.chapter_id == chapter_id)
        .order_by(Shot.idx.asc()))).scalars().all()


async def _reindex(session, chapter_id: str):
    for i, s in enumerate(await _list_shots(session, chapter_id)):
        s.idx = i


def _paragraph_shots(script: str) -> list[dict]:
    """Fallback sans LLM : un plan par paragraphe, durée estimée à la lecture
    (~150 mots/min, bornée 3–12 s)."""
    parts = [p.strip() for p in re.split(r"\n\s*\n", script) if p.strip()]
    out = []
    for p in parts:
        words = len(p.split())
        dur = max(3.0, min(12.0, round(words / 2.5, 1)))
        out.append({"source_text": p, "action": p[:200], "entities": [],
                    "shot_type": "medium", "camera_move": "static, locked-off",
                    "duration_s": dur, "prompt": "",
                    "motion_recipe": None, "energy": None})
    return out


def _ai_shots(script: str, bible: list[dict], lang: str) -> list[dict]:
    """Découpage réalisateur : plans avec action, entités présentes (mappées
    sur la bible), type de plan, caméra, durée, prompt d'illustration, et
    depuis v1.22 (W-d) la grammaire video-shotcraft (recette motion validée
    contre le catalogue du skill installé + courbe d'énergie 1-5)."""
    from app.services import shotcraft_service
    from app.services.summarizer import _chat_dispatch
    langname = "French" if lang.startswith("fr") else "English"
    n = max(4, min(40, len(script.split()) // 60 + 2))
    roster = "\n".join(f"- {e['name']} ({e['kind']}): {(e['description'] or '')[:120]}"
                       for e in bible) or "(none)"
    sc_block = shotcraft_service.prompt_block()
    system = ("You are a film director doing the storyboard breakdown of a "
              "narrated animation chapter"
              + (", applying the video-shotcraft doctrine and motion recipe "
                 "catalog provided" if sc_block else "")
              + ". Return ONLY valid JSON.")
    prompt = (
        f"Break this chapter into about {n} sequential storyboard shots.\n"
        + (sc_block + "\n\n" if sc_block else "")
        + f"Known entities (use these EXACT names when present in a shot):\n{roster}\n\n"
        f"For each shot return an object with:\n"
        f"\"source_excerpt\": the chapter text covered, copied verbatim, in order;\n"
        f"\"action\": what we SEE, one vivid sentence in {langname};\n"
        f"\"entities\": array of entity names present (from the list, may be empty);\n"
        f"\"shot_type\": one of {list(_SHOT_TYPES)};\n"
        f"\"camera_move\": one of {list(_CAMERA_MOVES)};\n"
        f"\"duration_s\": number 3-12;\n"
        + (f"\"motion_recipe\": the catalog slug whose motion best serves "
           f"this beat, or null;\n"
           f"\"energy\": integer 1 (calm) to 5 (peak), designing a "
           f"deliberate arc across the sequence;\n" if sc_block else "")
        + f"\"prompt\": a cinematic illustration prompt in {langname}.\n"
        f"Return ONLY a JSON array.\n\nChapter:\n{script[:12000]}")
    out, _prov = _chat_dispatch(prompt, system, 6000)
    if not out:
        return []
    txt = out.strip()
    if txt.startswith("```"):
        txt = re.sub(r"^```[a-zA-Z]*\n?", "", txt)
        txt = re.sub(r"\n?```$", "", txt).strip()
    i, j = txt.find("["), txt.rfind("]")
    if i >= 0 and j > i:
        txt = txt[i:j + 1]
    try:
        data = json.loads(txt)
    except Exception:
        return []
    name2id = {e["name"].strip().lower(): e["id"] for e in bible}
    sc_slugs = shotcraft_service.valid_slugs()
    shots = []
    for it in (data if isinstance(data, list) else []):
        if not isinstance(it, dict):
            continue
        st = str(it.get("shot_type") or "medium")
        cm = str(it.get("camera_move") or "static, locked-off")
        try:
            dur = max(3.0, min(12.0, float(it.get("duration_s") or 4)))
        except (TypeError, ValueError):
            dur = 4.0
        ents = [name2id[str(x).strip().lower()]
                for x in (it.get("entities") or [])
                if str(x).strip().lower() in name2id]
        mr = str(it.get("motion_recipe") or "").strip().lower()
        try:
            en = max(1, min(5, int(it.get("energy"))))
        except (TypeError, ValueError):
            en = None
        shots.append({
            "source_text": str(it.get("source_excerpt") or "").strip(),
            "action": str(it.get("action") or "").strip(),
            "entities": ents,
            "shot_type": st if st in _SHOT_TYPES else "medium",
            "camera_move": cm if cm in _CAMERA_MOVES else "static, locked-off",
            "duration_s": dur,
            "prompt": str(it.get("prompt") or "").strip(),
            "motion_recipe": mr if mr in sc_slugs else None,
            "energy": en,
        })
    return [s for s in shots if s["source_text"] or s["action"]]


@router.get("/chapters/{chapter_id}/shots")
async def list_chapter_shots(chapter_id: str):
    from app.services.storage import async_session_factory
    async with async_session_factory() as session:
        return {"shots": [_shot_dict(s)
                          for s in await _list_shots(session, chapter_id)]}


@router.post("/chapters/{chapter_id}/storyboard/decoupe")
async def storyboard_decoupe(chapter_id: str, body: dict):
    """Découpe le chapitre en plans et REMPLACE le storyboard existant.
    Body: {method: "ai"|"paragraph", language?}."""
    from app.services.storage import Chapter, Shot, async_session_factory
    import json as _json
    method = (body.get("method") or "paragraph").lower()
    lang = str(body.get("language") or "fr").lower()
    async with async_session_factory() as session:
        ch = await session.get(Chapter, chapter_id)
        if not ch:
            raise HTTPException(404, "Chapter not found")
        script = (ch.script_text or "").strip()
        if not script:
            raise HTTPException(400, "Le chapitre est vide")
        drafts = []
        if method == "ai":
            from app.services.summarizer import available
            if not available():
                return {"shots": [], "method": "ai",
                        "error": "Aucun LLM configuré (Réglages → clés API). "
                                 "Utilise le découpage par paragraphe."}
            ents_resp = await list_bible_entities(None)
            loop = asyncio.get_running_loop()
            drafts = await loop.run_in_executor(
                None, lambda: _ai_shots(script, ents_resp["entities"], lang))
            if not drafts:
                return {"shots": [], "method": "ai",
                        "error": "Le découpage IA a échoué — réessaie ou "
                                 "utilise les paragraphes."}
        else:
            drafts = _paragraph_shots(script)
        for s in await _list_shots(session, chapter_id):
            await session.delete(s)
        rows = []
        for i, d in enumerate(drafts):
            s = Shot(id=str(uuid4()), chapter_id=chapter_id, idx=i,
                     source_text=d["source_text"], action=d["action"],
                     entities=_json.dumps(d["entities"]),
                     shot_type=d["shot_type"], camera_move=d["camera_move"],
                     duration_s=d["duration_s"], prompt=d["prompt"],
                     motion_recipe=d.get("motion_recipe"),
                     energy=d.get("energy"),
                     created_at=datetime.utcnow(), updated_at=datetime.utcnow())
            session.add(s)
            rows.append(s)
        await session.commit()
        return {"shots": [_shot_dict(s) for s in rows], "method": method}


@router.post("/chapters/{chapter_id}/shots")
async def insert_shot(chapter_id: str, body: dict):
    """Insère un plan vide après `after_id` (ou en fin)."""
    from app.services.storage import Chapter, Shot, async_session_factory
    async with async_session_factory() as session:
        if not await session.get(Chapter, chapter_id):
            raise HTTPException(404, "Chapter not found")
        shots = await _list_shots(session, chapter_id)
        pos = len(shots)
        after = body.get("after_id")
        if after:
            for i, s in enumerate(shots):
                if s.id == after:
                    pos = i + 1
                    break
        s = Shot(id=str(uuid4()), chapter_id=chapter_id, idx=pos,
                 action=body.get("action") or "", entities="[]",
                 created_at=datetime.utcnow(), updated_at=datetime.utcnow())
        for later in shots[pos:]:
            later.idx += 1
        session.add(s)
        await session.commit()
        await _reindex(session, chapter_id)
        await session.commit()
        await session.refresh(s)
        return _shot_dict(s)


@router.put("/shots/{shot_id}")
async def update_shot(shot_id: str, body: dict):
    from app.services.storage import Shot, async_session_factory
    import json as _json
    async with async_session_factory() as session:
        s = await session.get(Shot, shot_id)
        if not s:
            raise HTTPException(404, "Shot not found")
        for k in ("source_text", "action", "prompt"):
            if k in body:
                setattr(s, k, body[k] or "")
        if "shot_type" in body and body["shot_type"] in _SHOT_TYPES:
            s.shot_type = body["shot_type"]
        if "camera_move" in body and body["camera_move"] in _CAMERA_MOVES:
            s.camera_move = body["camera_move"]
        if "duration_s" in body:
            try:
                s.duration_s = max(0.5, min(60.0, float(body["duration_s"])))
            except (TypeError, ValueError):
                pass
        if "motion_recipe" in body:
            from app.services import shotcraft_service
            mr = str(body["motion_recipe"] or "").strip().lower()
            s.motion_recipe = mr if mr in shotcraft_service.valid_slugs() else None
        if "energy" in body:
            try:
                s.energy = max(1, min(5, int(body["energy"])))
            except (TypeError, ValueError):
                s.energy = None
        if "entities" in body:
            s.entities = _json.dumps(body["entities"] or [])
        s.updated_at = datetime.utcnow()
        await session.commit()
        await session.refresh(s)
        return _shot_dict(s)


@router.delete("/shots/{shot_id}")
async def delete_shot(shot_id: str):
    from app.services.storage import Shot, async_session_factory
    async with async_session_factory() as session:
        s = await session.get(Shot, shot_id)
        if not s:
            raise HTTPException(404, "Shot not found")
        cid = s.chapter_id
        await session.delete(s)
        await session.commit()
        await _reindex(session, cid)
        await session.commit()
    return {"ok": True}


@router.post("/shots/{shot_id}/sketch")
async def generate_shot_sketch(shot_id: str, body: dict):
    """Croquis low-cost du plan (FLUX, style storyboard). Body: {seed?}."""
    from app.services.storage import Shot, BibleEntity, async_session_factory
    import json as _json
    async with async_session_factory() as session:
        s = await session.get(Shot, shot_id)
        if not s:
            raise HTTPException(404, "Shot not found")
        action = (s.action or s.source_text or "").strip()
        if not action:
            raise HTTPException(400, "Décris l'action du plan avant le croquis")
        try:
            eids = _json.loads(s.entities) if s.entities else []
        except Exception:
            eids = []
        descs = []
        for eid in eids[:4]:
            e = await session.get(BibleEntity, eid)
            if e:
                descs.append(f"{e.name}: {(e.description or '')[:100]}")
        prompt = (f"{_SKETCH_STYLE}. Shot: {s.shot_type}, camera: "
                  f"{s.camera_move}. {action}")
        if s.motion_recipe:
            # v1.22 (W-d) — la recette shotcraft colore la composition du
            # croquis (intention de mouvement + niveau d'énergie).
            from app.services import shotcraft_service
            prompt += f". Motion intent: {shotcraft_service.gloss(s.motion_recipe)}"
        tone = {1: "calm, still composition", 2: "quiet composition",
                4: "dynamic composition",
                5: "explosive, high-energy composition"}.get(s.energy or 0)
        if tone:
            prompt += f", {tone}"
        if descs:
            prompt += ". Characters/places: " + "; ".join(descs)
        seed = body.get("seed")
        seed = int(seed) if isinstance(seed, (int, float)) else None
        out = await _flux_generate(prompt, "portrait_16_9", 1, seed=seed)
        await LI.noter([out["images"][0]], "atelier")
        s.sketch_image = out["images"][0]
        s.sketch_seed = out.get("seed")
        s.updated_at = datetime.utcnow()
        await session.commit()
        await session.refresh(s)
        return _shot_dict(s)


# ── Vectorlab (phase 0) : documents vectoriels versionnés ────────────────────
# Plan : docs/superpowers/plans/2026-08-27-editeur-vectoriel-vitrail.md.
# Contenu en fichiers (services/vector_store), index et ancrage en SQLite
# (storage.VectorDoc). La suppression ARCHIVE — l'historique reste sur disque.

_VECTOR_ROLES = ("decor", "lumiere", "personnage", "libre")


def _vector_meta(row) -> dict:
    from app.services import vector_store as VS
    return {"id": row.id, "name": row.name, "chapter_id": row.chapter_id,
            "entity_id": row.entity_id, "deck_id": row.deck_id,
            "role": row.role,
            "version": row.version,
            "vignette": VS.a_vignette(row.id),
            "liaison": False,          # la liste par chapitre marque les liés
            "updated_at": (row.updated_at.isoformat()
                           if row.updated_at else None)}


@router.post("/vector/docs")
async def create_vector_doc(body: dict):
    """Body: {name, role, doc, chapter_id?, entity_id?} → {id, version:1}."""
    from app.services import vector_store as VS
    from app.services.storage import VectorDoc, async_session_factory
    name = str(body.get("name") or "").strip()[:120] or "Sans titre"
    role = str(body.get("role") or "libre").strip().lower()
    if role not in _VECTOR_ROLES:
        raise HTTPException(400, f"role invalide: {role} "
                                 f"(valides: {', '.join(_VECTOR_ROLES)})")
    try:
        did = VS.creer(body.get("doc") or {})
    except ValueError as e:
        raise HTTPException(400, str(e))
    async with async_session_factory() as session:
        session.add(VectorDoc(id=did, name=name,
                              chapter_id=(body.get("chapter_id") or None),
                              entity_id=(body.get("entity_id") or None),
                              deck_id=(body.get("deck_id") or None),
                              role=role, version=1))
        await session.commit()
    return {"id": did, "version": 1}


@router.get("/vector/docs")
async def list_vector_docs(chapter_id: str = "", role: str = "", q: str = "",
                           deck_id: str = ""):
    """`chapter_id` FUSIONNE les docs propres et les docs LIÉS au chapitre
    (méta `liaison: true`) ; `deck_id` (pont cartes) ne rend que les docs
    du jeu ; `role` et `q` (sous-chaîne insensible à la casse sur le nom)
    se cumulent ; tri updated_at desc."""
    from app.services.storage import (VectorDoc, VectorDocLink,
                                      async_session_factory)
    from sqlalchemy import select
    async with async_session_factory() as session:
        def _affiner(query):
            if role:
                query = query.where(VectorDoc.role == role)
            if q:
                query = query.where(VectorDoc.name.ilike(f"%{q}%"))
            return query
        propres = select(VectorDoc)
        if chapter_id:
            propres = propres.where(VectorDoc.chapter_id == chapter_id)
        if deck_id:
            propres = propres.where(VectorDoc.deck_id == deck_id)
        paires = [(r, False) for r in
                  (await session.execute(_affiner(propres))).scalars().all()]
        if chapter_id:
            lies = _affiner(
                select(VectorDoc)
                .join(VectorDocLink, VectorDocLink.doc_id == VectorDoc.id)
                .where(VectorDocLink.chapter_id == chapter_id))
            paires += [(r, True) for r in
                       (await session.execute(lies)).scalars().all()]
        paires.sort(key=lambda p: p[0].updated_at or datetime.min,
                    reverse=True)
        return {"docs": [{**_vector_meta(r), "liaison": lie}
                         for r, lie in paires]}


@router.get("/vector/docs/{doc_id}")
async def get_vector_doc(doc_id: str):
    from app.services import vector_store as VS
    from app.services.storage import VectorDoc, async_session_factory
    async with async_session_factory() as session:
        row = await session.get(VectorDoc, doc_id)
        if not row:
            raise HTTPException(404, "Document introuvable")
        try:
            doc = VS.lire(doc_id)
        except FileNotFoundError:
            raise HTTPException(404, "Contenu du document introuvable")
        return {"meta": _vector_meta(row), "doc": doc}


@router.put("/vector/docs/{doc_id}")
async def update_vector_doc(doc_id: str, body: dict):
    """Body: {doc, name?} → {id, version} — bump + historique disque."""
    from app.services import vector_store as VS
    from app.services.storage import VectorDoc, async_session_factory
    async with async_session_factory() as session:
        row = await session.get(VectorDoc, doc_id)
        if not row:
            raise HTTPException(404, "Document introuvable")
        try:
            v = VS.ecrire(doc_id, body.get("doc") or {})
        except FileNotFoundError:
            raise HTTPException(404, "Contenu du document introuvable")
        except ValueError as e:
            raise HTTPException(400, str(e))
        row.version = v
        if body.get("name"):
            row.name = str(body["name"]).strip()[:120]
        row.updated_at = datetime.utcnow()
        await session.commit()
        return {"id": doc_id, "version": v}


@router.delete("/vector/docs/{doc_id}")
async def delete_vector_doc(doc_id: str):
    """Archive le contenu (dernière version sur disque) et retire l'index —
    ses liaisons d'instanciation partent avec lui (jamais d'orphelines)."""
    from sqlalchemy import delete as _delete
    from app.services import vector_store as VS
    from app.services.storage import (VectorDoc, VectorDocLink,
                                      async_session_factory)
    async with async_session_factory() as session:
        row = await session.get(VectorDoc, doc_id)
        if not row:
            raise HTTPException(404, "Document introuvable")
        try:
            VS.supprimer(doc_id)
        except FileNotFoundError:
            pass                      # index orphelin : on nettoie quand même
        await session.execute(_delete(VectorDocLink)
                              .where(VectorDocLink.doc_id == doc_id))
        await session.delete(row)
        await session.commit()
    return {"ok": True}


@router.post("/vector/docs/{doc_id}/duplicate")
async def duplicate_vector_doc(doc_id: str, body: dict):
    """Body: {chapter_id?, name?} — la copie INDÉPENDANTE qui fait diverger :
    contenu COURANT relu du disque, id neuf, version 1, rôle copié,
    entity_id non copié ; ancrée au chapitre demandé — ou au JEU demandé
    (`deck_id`, pont cartes) — sinon bibliothèque. Si la liaison
    (chapter_id, source) existe elle est RETIRÉE — la copie remplace la
    référence. La vignette du source est copiée."""
    from app.services import vector_store as VS
    from app.services.storage import (VectorDoc, VectorDocLink,
                                      async_session_factory)
    chapter_id = (body.get("chapter_id") or None)
    deck_id = (body.get("deck_id") or None)
    async with async_session_factory() as session:
        src = await session.get(VectorDoc, doc_id)
        if not src:
            raise HTTPException(404, "Document introuvable")
        try:
            nid = VS.creer(VS.lire(doc_id))
        except FileNotFoundError:
            raise HTTPException(404, "Contenu du document introuvable")
        VS.copier_vignette(doc_id, nid)
        name = (str(body.get("name") or "").strip()
                or f"{src.name} (copie)")[:120]
        session.add(VectorDoc(id=nid, name=name, chapter_id=chapter_id,
                              deck_id=deck_id, role=src.role, version=1))
        if chapter_id:
            lien = await session.get(VectorDocLink, (chapter_id, doc_id))
            if lien:
                await session.delete(lien)
        await session.commit()
    return {"id": nid, "version": 1}


@router.post("/vector/docs/{doc_id}/export")
async def export_vector_doc(doc_id: str, body: dict):
    """Body: {svg}. Le CLIENT compile (compilateur unique du Vectorlab,
    verrouillé au snapshot qa) ; le serveur stocke `<id>.svg` à côté du
    JSON et le sert au GET export.svg. Le ré-export remplace."""
    from app.services import vector_store as VS
    from app.services.storage import VectorDoc, async_session_factory
    svg = (body.get("svg") or "").strip()
    if not svg.startswith("<svg"):
        raise HTTPException(400, "svg requis (le document compilé du client)")
    async with async_session_factory() as session:
        if not await session.get(VectorDoc, doc_id):
            raise HTTPException(404, "Document introuvable")
    try:
        return {"filename": VS.ecrire_svg(doc_id, svg)}
    except FileNotFoundError:
        raise HTTPException(404, "Contenu du document introuvable")


@router.get("/vector/docs/{doc_id}/export.svg")
async def get_vector_export_svg(doc_id: str):
    from fastapi.responses import Response
    from app.services import vector_store as VS
    svg = VS.lire_svg(doc_id)
    if svg is None:
        raise HTTPException(404, "Aucun export encore : exporte d'abord "
                                 "depuis l'éditeur (bouton Exporter → SVG)")
    return Response(content=svg, media_type="image/svg+xml")


# Les liaisons d'instanciation (phase 6) : un chapitre RÉFÉRENCE un doc de
# la bibliothèque globale ou d'un autre chapitre, sans copie — un seul
# document, l'édition se voit partout. Jamais d'orphelines : DELETE du doc
# et DELETE du chapitre emportent leurs liaisons.

def _vector_link_dict(l) -> dict:
    return {"chapter_id": l.chapter_id, "doc_id": l.doc_id,
            "created_at": l.created_at.isoformat() if l.created_at else None}


@router.post("/vector/links")
async def create_vector_link(body: dict):
    """Body: {chapter_id, doc_id} — instancie le doc dans le chapitre par
    référence. 409 déjà lié ou déjà propre au chapitre, 404 doc inconnu."""
    from app.services.storage import (VectorDoc, VectorDocLink,
                                      async_session_factory)
    chapter_id = str(body.get("chapter_id") or "").strip()
    doc_id = str(body.get("doc_id") or "").strip()
    if not chapter_id or not doc_id:
        raise HTTPException(400, "chapter_id et doc_id requis")
    async with async_session_factory() as session:
        doc = await session.get(VectorDoc, doc_id)
        if not doc:
            raise HTTPException(404, "Document introuvable")
        if doc.chapter_id == chapter_id:
            raise HTTPException(409, "Ce document est déjà propre à ce "
                                     "chapitre — rien à instancier")
        if await session.get(VectorDocLink, (chapter_id, doc_id)):
            raise HTTPException(409, "Déjà instancié dans ce chapitre")
        session.add(VectorDocLink(chapter_id=chapter_id, doc_id=doc_id))
        await session.commit()
    return {"ok": True, "chapter_id": chapter_id, "doc_id": doc_id}


@router.get("/vector/links")
async def list_vector_links(chapter_id: str = "", doc_id: str = ""):
    from app.services.storage import VectorDocLink, async_session_factory
    from sqlalchemy import select
    async with async_session_factory() as session:
        q = select(VectorDocLink)
        if chapter_id:
            q = q.where(VectorDocLink.chapter_id == chapter_id)
        if doc_id:
            q = q.where(VectorDocLink.doc_id == doc_id)
        rows = (await session.execute(
            q.order_by(VectorDocLink.created_at.desc()))).scalars().all()
        return {"links": [_vector_link_dict(l) for l in rows]}


@router.delete("/vector/links")
async def delete_vector_link(chapter_id: str = "", doc_id: str = ""):
    """Retire la liaison — le document, lui, ne bouge pas."""
    from app.services.storage import VectorDocLink, async_session_factory
    if not chapter_id or not doc_id:
        raise HTTPException(400, "chapter_id et doc_id requis")
    async with async_session_factory() as session:
        l = await session.get(VectorDocLink, (chapter_id, doc_id))
        if not l:
            raise HTTPException(404, "Liaison introuvable")
        await session.delete(l)
        await session.commit()
    return {"ok": True}


_PNG_MAGIC = b"\x89PNG\r\n\x1a\n"


@router.post("/vector/docs/{doc_id}/vignette")
async def set_vector_vignette(doc_id: str, request: Request):
    """Corps binaire image/png : la MINI-vignette rasterisée par le client
    au save. Stockée `<id>.png` à côté du JSON — jamais par /images/upload,
    la Library réelle reste propre."""
    from app.services import vector_store as VS
    from app.services.storage import VectorDoc, async_session_factory
    octets = await request.body()
    if not octets.startswith(_PNG_MAGIC):
        raise HTTPException(400, "vignette: un PNG est attendu")
    async with async_session_factory() as session:
        if not await session.get(VectorDoc, doc_id):
            raise HTTPException(404, "Document introuvable")
    try:
        return {"filename": VS.ecrire_vignette(doc_id, octets)}
    except FileNotFoundError:
        raise HTTPException(404, "Contenu du document introuvable")


@router.get("/vector/docs/{doc_id}/vignette.png")
async def get_vector_vignette(doc_id: str):
    from app.services import vector_store as VS
    octets = VS.lire_vignette(doc_id)
    if octets is None:
        raise HTTPException(404, "Aucune vignette encore : elle naît au "
                                 "premier Sauver dans l'éditeur")
    return Response(content=octets, media_type="image/png")


@router.get("/vector/vitrail")
async def vector_vitrail():
    """Le mode vitrail du Vectorlab lit la FICHE ÉPINGLÉE
    (services/style_vitrail.json, copie du skill vitrail-mloda-polska) —
    l'unique source des ancres, bornes et motifs. Le banc compare la
    réponse à l'octet avec le fichier : toute divergence rougit."""
    import json as _json
    from pathlib import Path as _P
    fiche = (_P(__file__).resolve().parent.parent / "services"
             / "style_vitrail.json")
    try:
        data = _json.loads(fiche.read_text(encoding="utf-8"))
        return {"famille": data["familles"]["vitrail"],
                "source": "style_vitrail.json (copie épinglée du skill "
                          "vitrail-mloda-polska)"}
    except Exception as e:
        raise HTTPException(503, f"fiche style_vitrail.json indisponible: {e}")


@router.get("/atelier/shotcraft")
async def shotcraft_info():
    """v1.22 (W-d) — état du pont video-shotcraft + catalogue des recettes
    motion (fiches du skill installé si présent, sinon catalogue embarqué).
    Consommé par l'Atelier (badge + selects des plans)."""
    from app.services import shotcraft_service
    d = shotcraft_service.catalog()
    cards = [{"slug": c["slug"], "cat": c.get("cat"),
              "energy": c.get("energy"), "anim": bool(c.get("anim")),
              "gloss": c.get("gloss") or ""}
             for c in d["cards"].values()]
    cards.sort(key=lambda c: (not c["anim"], c["cat"] or "", c["slug"]))
    return {"status": shotcraft_service.status(), "cards": cards}


@router.delete("/chapters/{chapter_id}/shots")
async def reset_storyboard(chapter_id: str):
    """Réinitialise le storyboard du chapitre (supprime tous les plans)."""
    from app.services.storage import async_session_factory
    async with async_session_factory() as session:
        n = 0
        for s in await _list_shots(session, chapter_id):
            await session.delete(s)
            n += 1
        await session.commit()
    return {"ok": True, "deleted": n}


@router.post("/chapters/{chapter_id}/storyboard/reorder")
async def reorder_shots(chapter_id: str, body: dict):
    """Body: {ids: [shot ids dans le nouvel ordre]}."""
    from app.services.storage import async_session_factory
    ids = body.get("ids") or []
    async with async_session_factory() as session:
        shots = {s.id: s for s in await _list_shots(session, chapter_id)}
        if set(ids) != set(shots):
            raise HTTPException(400, "ids must be a permutation of the chapter's shots")
        for i, sid in enumerate(ids):
            shots[sid].idx = i
        await session.commit()
        return {"shots": [_shot_dict(shots[sid]) for sid in ids]}


# ───────── Atelier v1.19: agent d'ingestion de manuscrit ─────────
# Manuscrit complet → segmentation en chapitres (titres importés) →
# extraction LLM chapitre par chapitre (6 kinds) → relecture globale de
# consolidation (avec le fichier compagnon de l'auteur en autorité) →
# surlignage automatique. Job long: progression pollable.

_MS_JOBS: dict[str, dict] = {}


def _ms_register(jid: str, state: dict) -> dict:
    """Register a manuscript/VO/adapt job, evicting finished ones first.

    In-memory only (status is polled, never persisted), so without eviction a
    long session grows this dict forever.
    """
    done = [k for k, v in _MS_JOBS.items()
            if v.get("phase") in ("done", "error") and k != jid]
    for k in done[:-20]:  # keep the last 20 finished for late pollers
        _MS_JOBS.pop(k, None)
    _MS_JOBS[jid] = state
    return state


def _read_upload_text(name: str, data: bytes) -> str:
    """Texte brut d'un upload txt/docx/pdf. Pour docx, les paragraphes stylés
    Heading sont préfixés du marqueur \\x1f — le segmenteur les importe comme
    titres de chapitres même sans convention typographique."""
    import io as _io
    name = (name or "").lower()
    if name.endswith(".docx"):
        import docx
        doc = docx.Document(_io.BytesIO(data))
        parts = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if not t:
                continue
            style = (p.style.name or "").lower() if p.style is not None else ""
            if style.startswith("heading") or style.startswith("titre"):
                parts.append("\x1f" + t)
            else:
                parts.append(t)
        return "\n\n".join(parts)
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(_io.BytesIO(data))
        return "\n\n".join((pg.extract_text() or "").strip()
                           for pg in reader.pages if (pg.extract_text() or "").strip())
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(enc)
        except Exception:
            continue
    return ""


@router.post("/atelier/manuscript")
async def import_manuscript(background_tasks: BackgroundTasks,
                            manuscript: UploadFile = File(...),
                            companion: UploadFile | None = File(None),
                            series: str = Form("")):
    """Lance l'agent d'ingestion sur un manuscrit complet (+ fichier compagnon
    optionnel). Retourne {job_id} — suivre GET /atelier/manuscript/{job_id}."""
    from app.services.summarizer import available
    if not available():
        raise HTTPException(400, "Aucun LLM configuré (Réglages → clés API) — "
                                 "l'agent manuscrit a besoin d'un modèle.")
    raw = await manuscript.read()
    if len(raw) > 30 * 1024 * 1024:
        raise HTTPException(400, "Manuscrit trop lourd (max 30 MB)")
    try:
        text = (await asyncio.to_thread(
            _read_upload_text, manuscript.filename or "", raw)).strip()
    except Exception as e:
        raise HTTPException(422, f"Lecture du manuscrit impossible: {e}")
    if len(text) < 200:
        raise HTTPException(400, "Le manuscrit semble vide (moins de 200 caractères)")
    comp_text = ""
    if companion is not None:
        try:
            comp_text = (await asyncio.to_thread(
                _read_upload_text, companion.filename or "",
                await companion.read())).strip()
        except Exception as e:
            logger.warning(f"fichier compagnon illisible (ignoré): {e}")
    stem = Path(manuscript.filename or "Manuscrit").stem
    series = (series or "").strip() or stem
    jid = str(uuid4())
    _ms_register(jid, {"job_id": jid, "phase": "segmentation", "chapter_i": 0,
                       "chapter_n": 0, "message": "Segmentation en chapitres…",
                       "done": False, "error": None, "stats": {}, "series": series})
    background_tasks.add_task(_run_manuscript_job, jid, text, comp_text, series)
    return {"job_id": jid, "series": series, "chars": len(text),
            "companion_chars": len(comp_text)}


@router.get("/atelier/manuscript/{job_id}")
async def manuscript_job_status(job_id: str):
    st = _MS_JOBS.get(job_id)
    if not st:
        raise HTTPException(404, "Job inconnu")
    return st


async def _run_manuscript_job(jid: str, text: str, companion: str, series: str):
    from app.services import manuscript_agent as MA
    from app.services.storage import (BibleEntity, Chapter,
                                      async_session_factory)
    from sqlalchemy import select
    import json as _json

    def upd(**kw):
        _MS_JOBS[jid].update(kw)

    try:
        loop = asyncio.get_running_loop()
        # 1. segmentation
        segs = MA.segment_chapters(text)
        upd(phase="extraction", chapter_n=len(segs),
            message=f"{len(segs)} chapitres détectés")
        # roster initial = bible existante (stabilité des noms)
        async with async_session_factory() as session:
            existing = (await session.execute(select(BibleEntity))).scalars().all()
            roster = [{"name": e.name, "kind": e.kind} for e in existing]
        # 2. extraction chapitre par chapitre
        raw: list[dict] = []
        qmap: dict[tuple, list] = {}   # (kind, lower name/alias) -> evidence
        for i, seg in enumerate(segs):
            upd(chapter_i=i + 1, message=f"Extraction — {seg['title']}")
            found = await loop.run_in_executor(
                None, lambda s=seg: MA.extract_chapter(s["title"], s["text"], roster))
            for f in found:
                raw.append(f)
                keys = [(f["kind"], f["name"].strip().lower())] + \
                       [(f["kind"], a.strip().lower()) for a in f.get("aliases") or []]
                for k in keys:
                    qmap.setdefault(k, [])
                for q in f.get("quotes") or []:
                    qmap[keys[0]].append({"chapter": seg["title"], "quote": q})
                if not any(r["name"].lower() == f["name"].lower()
                           and r["kind"] == f["kind"] for r in roster):
                    roster.append({"name": f["name"], "kind": f["kind"]})
        if not raw:
            raise RuntimeError("Aucune entité extraite — vérifie le manuscrit / le LLM")
        # 3. relecture globale
        upd(phase="consolidation", chapter_i=len(segs),
            message=f"Relecture globale — consolidation de {len(raw)} mentions…")
        final = await loop.run_in_executor(
            None, lambda: MA.consolidate(raw, companion))
        # 4. écriture bible (upsert) + chapitres + surlignage
        upd(phase="liens", message="Écriture de la bible et surlignage…")
        created_e = updated_e = created_c = updated_c = total_spans = 0
        async with async_session_factory() as session:
            rows = (await session.execute(select(BibleEntity))).scalars().all()
            by_key = {(e.kind, e.name.strip().lower()): e for e in rows}
            ent_pairs = []
            for fe in final:
                key = (fe["kind"], fe["name"].strip().lower())
                ev = list(qmap.get(key, []))
                for a in fe.get("aliases") or []:
                    ev.extend(qmap.get((fe["kind"], a.strip().lower()), []))
                e = by_key.get(key)
                if e:
                    if fe.get("description") and \
                            len(fe["description"]) > len(e.description or ""):
                        e.description = fe["description"]
                    old_alias = set(_json.loads(e.aliases) if e.aliases else [])
                    e.aliases = _json.dumps(sorted(old_alias | set(fe.get("aliases") or [])))
                    e.evidence = _json.dumps(ev[:20])
                    e.updated_at = datetime.utcnow()
                    updated_e += 1
                else:
                    e = BibleEntity(id=str(uuid4()), kind=fe["kind"],
                                    name=fe["name"],
                                    description=fe.get("description") or "",
                                    aliases=_json.dumps(fe.get("aliases") or []),
                                    evidence=_json.dumps(ev[:20]),
                                    inspiration_images="[]",
                                    created_at=datetime.utcnow(),
                                    updated_at=datetime.utcnow())
                    session.add(e)
                    by_key[key] = e
                    created_e += 1
                ent_pairs.append((e, fe))
            await session.commit()
            for e, _fe in ent_pairs:
                await session.refresh(e)
            ents_for_spans = [{"id": e.id, "name": e.name,
                               "aliases": _json.loads(e.aliases) if e.aliases else [],
                               "quotes": fe.get("quotes") or []}
                              for e, fe in ent_pairs]
            ch_rows = (await session.execute(select(Chapter))).scalars().all()
            ch_by_key = {(c.series or "", c.title): c for c in ch_rows}
            for seg in segs:
                spans = MA.compute_spans(seg["text"], ents_for_spans)
                total_spans += len(spans)
                key = (series, seg["title"][:200])
                c = ch_by_key.get(key)
                if c:
                    c.script_text = seg["text"]
                    c.spans = _json.dumps(spans)
                    c.updated_at = datetime.utcnow()
                    updated_c += 1
                else:
                    c = Chapter(id=str(uuid4()), title=seg["title"][:200],
                                series=series, script_text=seg["text"],
                                spans=_json.dumps(spans),
                                created_at=datetime.utcnow(),
                                updated_at=datetime.utcnow())
                    session.add(c)
                    ch_by_key[key] = c
                    created_c += 1
            await session.commit()
        # 5. direction artistique — l'agent propose 4 styles motivés par le
        # manuscrit (best-effort: un échec n'invalide pas l'ingestion).
        n_da = 0
        try:
            upd(phase="direction artistique",
                message="Propositions de direction artistique…")
            excerpt = "\n\n".join(s["text"][:3000] for s in segs[:4])[:9000]
            names = [fe["name"] for fe in final]
            props = await loop.run_in_executor(
                None, lambda: MA.propose_styles(excerpt, names))
            if props:
                from app.services.storage import AtelierSetting
                async with async_session_factory() as session:
                    row = await session.get(AtelierSetting, "style_proposals")
                    val = _json.dumps(props, ensure_ascii=False)
                    if row:
                        row.value = val
                    else:
                        session.add(AtelierSetting(key="style_proposals",
                                                   value=val))
                    await session.commit()
                n_da = len(props)
        except Exception as de:
            logger.warning(f"DA proposals skipped: {de}")
        upd(phase="terminé", done=True,
            message="Ingestion terminée — bible consolidée, chapitres surlignés"
                    + (f", {n_da} directions artistiques proposées (🎨)." if n_da
                       else "."),
            stats={"chapitres_crees": created_c, "chapitres_mis_a_jour": updated_c,
                   "entites_creees": created_e, "entites_enrichies": updated_e,
                   "zones_surlignees": total_spans,
                   "directions_proposees": n_da})
        logger.success(f"manuscrit {jid}: {created_c}+{updated_c} chapitres, "
                       f"{created_e}+{updated_e} entités, {total_spans} spans")
    except Exception as e:
        logger.exception(f"manuscrit {jid} échec: {e}")
        upd(phase="échec", done=True, error=str(e))


# ───────── Atelier v1.20 (phase A): passe Scénario — adaptation ─────────
# Roman → scènes de scénario (Fountain) SANS toucher le manuscrit : sluglines
# INT/EXT + lieu de la bible + moment, éclairage/caméra/mood motivés par le
# narratif, entités (perso + décor) liées et créées au besoin pour la
# réutilisation inter-chapitres.

def _scene_dict(s) -> dict:
    import json as _json
    try:
        ents = _json.loads(s.entities) if s.entities else []
    except Exception:
        ents = []
    return {"id": s.id, "chapter_id": s.chapter_id, "idx": s.idx,
            "slugline": s.slugline, "int_ext": s.int_ext,
            "location_entity_id": s.location_entity_id,
            "time_of_day": s.time_of_day,
            "fountain_text": s.fountain_text or "",
            "lighting": s.lighting or "", "camera_notes": s.camera_notes or "",
            "mood": s.mood or "", "entities": ents,
            "source_text": s.source_text or "",
            "duration_s": s.duration_s, "vo_audio": s.vo_audio}


async def _list_scenes(session, chapter_id: str):
    from app.services.storage import Scene
    from sqlalchemy import select
    return (await session.execute(
        select(Scene).where(Scene.chapter_id == chapter_id)
        .order_by(Scene.idx.asc()))).scalars().all()


@router.get("/chapters/{chapter_id}/scenes")
async def list_chapter_scenes(chapter_id: str):
    from app.services.storage import async_session_factory
    async with async_session_factory() as session:
        return {"scenes": [_scene_dict(s)
                           for s in await _list_scenes(session, chapter_id)]}


@router.get("/chapters/{chapter_id}/screenplay")
async def get_chapter_screenplay(chapter_id: str, format: str = "json"):
    """Scénario Fountain assemblé du chapitre. ?format=fountain → texte brut
    téléchargeable."""
    from app.services import manuscript_agent as MA
    from app.services.storage import Chapter, async_session_factory
    async with async_session_factory() as session:
        ch = await session.get(Chapter, chapter_id)
        if not ch:
            raise HTTPException(404, "Chapter not found")
        scenes = [_scene_dict(s) for s in await _list_scenes(session, chapter_id)]
    text = MA.assemble_fountain(ch.title, scenes)
    if format == "fountain":
        return Response(content=text, media_type="text/plain; charset=utf-8",
                        headers={"Content-Disposition":
                                 f'attachment; filename="{ch.title[:60]}.fountain"'})
    return {"title": ch.title, "fountain": text, "scene_count": len(scenes)}


@router.put("/scenes/{scene_id}")
async def update_scene(scene_id: str, body: dict):
    from app.services.storage import Scene, async_session_factory
    async with async_session_factory() as session:
        s = await session.get(Scene, scene_id)
        if not s:
            raise HTTPException(404, "Scene not found")
        for k in ("fountain_text", "camera_notes"):
            if k in body:
                setattr(s, k, body[k] or "")
        for k, lim in (("lighting", 120), ("mood", 120)):
            if k in body:
                setattr(s, k, (body[k] or "")[:lim])
        if "int_ext" in body and str(body["int_ext"]).upper() in ("INT", "EXT", "INT/EXT"):
            s.int_ext = str(body["int_ext"]).upper()
        if "time_of_day" in body and str(body["time_of_day"]).upper() in \
                ("JOUR", "NUIT", "AUBE", "CRÉPUSCULE", "MATIN", "SOIR"):
            s.time_of_day = str(body["time_of_day"]).upper()
        if "location" in body and (body["location"] or "").strip():
            loc = body["location"].strip().upper()[:150]
            s.slugline = f"{s.int_ext}. {loc} - {s.time_of_day}"
        else:
            # slugline recomposée si int_ext / time_of_day ont changé
            m = re.match(r"^(?:INT\.|EXT\.|INT\./EXT\.)\s*(.+?)\s*-\s*[A-ZÉÈ]+$",
                         s.slugline or "")
            loc = m.group(1) if m else (s.slugline or "LIEU")
            s.slugline = f"{s.int_ext}. {loc} - {s.time_of_day}"
        s.updated_at = datetime.utcnow()
        await session.commit()
        await session.refresh(s)
        return _scene_dict(s)


# ───────── Atelier v1.22 (C): voice-over minuté par scène ─────────
# Mode audiobook hybride: la narration est lue par le personnage
# « Narrateur » (voix castée en B), chaque réplique par la voix castée de
# son personnage. La durée réelle de l'audio devient la durée de la scène —
# le storyboard et la production héritent de durées EXACTES.

def _vo_audio_dir() -> Path:
    d = settings.images_path.parent / "audio"
    d.mkdir(parents=True, exist_ok=True)
    return d


def _concat_audio(paths: list, dest: Path) -> None:
    """Concatène des mp3 en un seul (ré-encodage: robuste). Patchable en test."""
    if len(paths) == 1:
        import shutil as _sh
        _sh.copy2(paths[0], dest)
        return
    from app.services.composition_service import _run_ffmpeg
    cmd = ["ffmpeg", "-y"]
    for p in paths:
        cmd += ["-i", str(p)]
    labels = "".join(f"[{i}:a]" for i in range(len(paths)))
    cmd += ["-filter_complex", f"{labels}concat=n={len(paths)}:v=0:a=1[outa]",
            "-map", "[outa]", "-c:a", "libmp3lame", "-b:a", "128k", str(dest)]
    _run_ffmpeg(cmd, dest)


def _audio_duration(path: Path) -> float:
    from app.services.template_service import _probe_duration
    return _probe_duration(path)


def _fold_name(s: str) -> str:
    import unicodedata as _ud
    return "".join(c for c in _ud.normalize("NFD", (s or "").lower())
                   if _ud.category(c) != "Mn").strip()


async def _voice_cast(session) -> tuple:
    """(narrateur {voice_id,name} | None, map nom/alias replié → voix perso)."""
    from app.services.storage import BibleEntity
    from sqlalchemy import select
    import json as _json
    rows = (await session.execute(
        select(BibleEntity).where(BibleEntity.kind == "character"))).scalars().all()
    narrator, cues = None, {}
    for e in rows:
        v = {"voice_id": e.voice_id, "name": e.name}
        if _fold_name(e.name) in ("narrateur", "narrator"):
            narrator = v if e.voice_id else None
            continue
        if not e.voice_id:
            continue
        cues[_fold_name(e.name)] = v
        try:
            for a in (_json.loads(e.aliases) if e.aliases else []):
                cues.setdefault(_fold_name(a), v)
        except Exception:
            pass
    return narrator, cues


async def _generate_scene_vo(session, scene, lang: str) -> dict:
    """Génère l'audio d'UNE scène (segments → TTS par voix → concat → durée)."""
    from app.services import manuscript_agent as MA
    from app.services.elevenlabs_service import VoiceoverService
    import tempfile as _tf
    segments = MA.parse_fountain_segments(scene.fountain_text or "")
    segments = [s for s in segments if len(s["text"].strip()) >= 2]
    if not segments:
        raise HTTPException(400, "La scène n'a pas de texte à lire.")
    narrator, cues = await _voice_cast(session)
    if any(s["kind"] == "narration" for s in segments) and not narrator:
        raise HTTPException(
            400, "Crée un personnage « Narrateur » dans la bible et caste sa "
                 "voix (🎙 Suggérer) — c'est lui qui lit la narration.")
    voice = VoiceoverService()
    loop = asyncio.get_running_loop()
    if not await loop.run_in_executor(None, VoiceoverService.is_enabled):
        raise HTTPException(400, "Aucune voix disponible : configure la clé "
                                 "ElevenLabs ou lance Voicebox (Réglages).")
    tmp = Path(_tf.mkdtemp(prefix="dz_vo_"))
    parts, plan = [], []
    l11 = "FR" if lang.startswith("fr") else "EN"
    for i, seg in enumerate(segments):
        if seg["kind"] == "dialogue":
            v = cues.get(_fold_name(seg["character"] or ""))
            vid = (v or narrator or {}).get("voice_id")
            speaker = (v or narrator or {}).get("name")
        else:
            vid = narrator["voice_id"]
            speaker = narrator["name"]
        dest = tmp / f"part_{i:03d}.mp3"
        await loop.run_in_executor(
            None, lambda s=seg, d=dest, vv=vid: voice.generate_long(
                text=s["text"], output_path=d, language=l11, voice_id=vv))
        parts.append(dest)
        plan.append({"kind": seg["kind"], "speaker": speaker,
                     "chars": len(seg["text"])})
    fname = f"vo_{scene.id[:8]}_{uuid4().hex[:6]}.mp3"
    out = _vo_audio_dir() / fname
    _concat_audio(parts, out)
    dur = round(float(_audio_duration(out)), 2)
    scene.vo_audio = fname
    scene.duration_s = dur
    scene.updated_at = datetime.utcnow()
    await session.commit()
    await session.refresh(scene)
    return {"scene": _scene_dict(scene), "segments": plan, "duration_s": dur}


@router.post("/scenes/{scene_id}/voiceover")
async def scene_voiceover(scene_id: str, body: dict):
    """Génère le voice-over d'une scène. Body: {language?}."""
    from app.services.storage import Scene, async_session_factory
    lang = str(body.get("language") or "fr").lower()
    async with async_session_factory() as session:
        s = await session.get(Scene, scene_id)
        if not s:
            raise HTTPException(404, "Scene not found")
        return await _generate_scene_vo(session, s, lang)


@router.post("/chapters/{chapter_id}/voiceover")
async def chapter_voiceover(chapter_id: str, body: dict,
                            background_tasks: BackgroundTasks):
    """Voice-over de TOUTES les scènes du chapitre (job). Body: {language?,
    force?} — force=true régénère aussi les scènes déjà minutées. Suivre
    GET /atelier/manuscript/{job_id} (job store commun)."""
    from app.services.storage import Chapter, async_session_factory
    from app.services.elevenlabs_service import VoiceoverService
    if not await asyncio.get_running_loop().run_in_executor(
            None, VoiceoverService.is_enabled):
        raise HTTPException(400, "Aucune voix disponible : configure la clé "
                                 "ElevenLabs ou lance Voicebox (Réglages).")
    async with async_session_factory() as session:
        if not await session.get(Chapter, chapter_id):
            raise HTTPException(404, "Chapter not found")
        scenes = await _list_scenes(session, chapter_id)
    if not scenes:
        raise HTTPException(400, "Pas de scénario — lance 🎭 Adapter d'abord.")
    lang = str(body.get("language") or "fr").lower()
    force = bool(body.get("force"))
    jid = str(uuid4())
    _ms_register(jid, {"job_id": jid, "phase": "voice-over", "chapter_i": 0,
                       "chapter_n": len(scenes), "message": "Voix en cours…",
                       "done": False, "error": None, "stats": {}})
    background_tasks.add_task(_run_vo_job, jid, chapter_id, lang, force)
    return {"job_id": jid, "scenes": len(scenes)}


async def _run_vo_job(jid: str, chapter_id: str, lang: str, force: bool):
    from app.services.storage import async_session_factory

    def upd(**kw):
        _MS_JOBS[jid].update(kw)

    done = skipped = 0
    total = 0.0
    try:
        async with async_session_factory() as session:
            scenes = await _list_scenes(session, chapter_id)
            for i, s in enumerate(scenes):
                upd(chapter_i=i + 1,
                    message=f"Scène {i + 1}/{len(scenes)} — {(s.slugline or '')[:50]}")
                if s.vo_audio and s.duration_s and not force:
                    skipped += 1
                    total += float(s.duration_s or 0)
                    continue
                r = await _generate_scene_vo(session, s, lang)
                total += r["duration_s"]
                done += 1
        upd(phase="terminé", done=True,
            message="Voice-over terminé — les scènes sont minutées.",
            stats={"scenes_generees": done, "scenes_conservees": skipped,
                   "duree_totale_s": round(total, 1)})
    except HTTPException as e:
        upd(phase="échec", done=True, error=str(e.detail))
    except Exception as e:
        logger.exception(f"vo job {jid}: {e}")
        upd(phase="échec", done=True, error=str(e))


@router.delete("/chapters/{chapter_id}/scenes")
async def reset_screenplay(chapter_id: str):
    """Réinitialise le scénario du chapitre (supprime toutes les scènes).
    Le manuscrit n'est évidemment pas touché."""
    from app.services.storage import async_session_factory
    async with async_session_factory() as session:
        n = 0
        for s in await _list_scenes(session, chapter_id):
            await session.delete(s)
            n += 1
        await session.commit()
    return {"ok": True, "deleted": n}


@router.post("/chapters/{chapter_id}/screenplay/adapt")
async def adapt_chapter_endpoint(chapter_id: str, body: dict,
                                 background_tasks: BackgroundTasks):
    """Lance la passe d'adaptation (roman → scénario) sur un chapitre.
    REMPLACE les scènes existantes. Retourne {job_id} — suivre
    GET /atelier/manuscript/{job_id} (job store commun)."""
    from app.services.summarizer import available
    from app.services.storage import Chapter, async_session_factory
    if not available():
        raise HTTPException(400, "Aucun LLM configuré (Réglages → clés API).")
    async with async_session_factory() as session:
        ch = await session.get(Chapter, chapter_id)
        if not ch:
            raise HTTPException(404, "Chapter not found")
        if not (ch.script_text or "").strip():
            raise HTTPException(400, "Le chapitre est vide")
    lang = str(body.get("language") or "fr").lower()
    jid = str(uuid4())
    _ms_register(jid, {"job_id": jid, "phase": "adaptation", "chapter_i": 1,
                       "chapter_n": 1, "message": "Adaptation en scénario…",
                       "done": False, "error": None, "stats": {}})
    background_tasks.add_task(_run_adapt_job, jid, chapter_id, lang)
    return {"job_id": jid}


async def _run_adapt_job(jid: str, chapter_id: str, lang: str):
    from app.services import manuscript_agent as MA
    from app.services.storage import (BibleEntity, Chapter, Scene,
                                      async_session_factory)
    from sqlalchemy import select
    import json as _json

    def upd(**kw):
        _MS_JOBS[jid].update(kw)

    try:
        loop = asyncio.get_running_loop()
        async with async_session_factory() as session:
            ch = await session.get(Chapter, chapter_id)
            title, text = ch.title, ch.script_text or ""
            rows = (await session.execute(select(BibleEntity))).scalars().all()
            bible = [_entity_dict(e) for e in rows]
        upd(message=f"Adaptation de « {title} »…")
        drafts = await loop.run_in_executor(
            None, lambda: MA.adapt_chapter(title, text, bible, lang))
        if not drafts:
            raise RuntimeError("L'adaptation n'a produit aucune scène — réessaie.")
        upd(phase="liens", message="Scènes, lieux et décors…")
        created_e = 0
        async with async_session_factory() as session:
            rows = (await session.execute(select(BibleEntity))).scalars().all()
            by_key = {}
            for e in rows:
                by_key[(e.kind, e.name.strip().lower())] = e
                for a in (_json.loads(e.aliases) if e.aliases else []):
                    by_key.setdefault((e.kind, a.strip().lower()), e)

            def find_or_create(kind, name, desc=""):
                nonlocal created_e
                key = (kind, name.strip().lower())
                # les lieux sont souvent cités en CAPS dans la slugline
                e = by_key.get(key) or by_key.get((kind, name.strip().title().lower()))
                if e:
                    return e
                e = BibleEntity(id=str(uuid4()), kind=kind, name=name[:120],
                                description=desc, aliases="[]", evidence="[]",
                                inspiration_images="[]",
                                created_at=datetime.utcnow(),
                                updated_at=datetime.utcnow())
                session.add(e)
                by_key[key] = e
                created_e += 1
                return e

            for s in await _list_scenes(session, chapter_id):
                await session.delete(s)
            n_scenes = 0
            for i, d in enumerate(drafts):
                loc_ent = find_or_create(
                    "place", d["slugline_location"].title(),
                    f"Lieu établi par l'adaptation de « {title} ».")
                ent_ids = [loc_ent.id]
                for cn in d["characters"]:
                    e = by_key.get(("character", cn.strip().lower()))
                    if e:
                        ent_ids.append(e.id)
                for dn in d["decor"]:
                    e = find_or_create("decor", dn,
                                       f"Élément de décor — {d['slugline_location'].title()}.")
                    ent_ids.append(e.id)
                slug = f"{d['int_ext']}. {d['slugline_location'].upper()} - {d['time_of_day']}"
                session.add(Scene(
                    id=str(uuid4()), chapter_id=chapter_id, idx=i,
                    slugline=slug, int_ext=d["int_ext"],
                    location_entity_id=loc_ent.id,
                    time_of_day=d["time_of_day"],
                    fountain_text=d["fountain"], lighting=d["lighting"],
                    camera_notes=d["camera_notes"], mood=d["mood"],
                    entities=_json.dumps(list(dict.fromkeys(ent_ids))),
                    source_text=d["source_excerpt"],
                    created_at=datetime.utcnow(), updated_at=datetime.utcnow()))
                n_scenes += 1
            await session.commit()
        upd(phase="terminé", done=True,
            message=f"Scénario prêt — {n_scenes} scènes.",
            stats={"scenes": n_scenes, "entites_creees": created_e})
        logger.success(f"adaptation {jid}: {n_scenes} scènes, "
                       f"{created_e} entités créées")
    except Exception as e:
        logger.exception(f"adaptation {jid} échec: {e}")
        upd(phase="échec", done=True, error=str(e))


# ═════════════════════ Material Forge (matières PBR) ═════════════════════
# SPEC Material Forge section 3. Une matière vit dans
# outputs/materials/mat_xxxxxxxx/ (meta.json + 8 maps PNG) — cf.
# app/services/material_store.py, seul endroit qui touche à ces chemins.
#
# Génération : prompt -> image (routage FLUX / OpenAI déjà en place) ou
# fichier de la Library pris tel quel, puis raccord seamless (pixel_ops),
# score de raccord avant/après, dérivation des maps (pbr_service), écriture.
# Job asynchrone pollable, même patron que les jobs manuscrit (_MS_JOBS) :
# état en mémoire, jamais persisté, éviction des terminés.

_MAT_JOBS: dict[str, dict] = {}


def _mat_job_register(jid: str, state: dict) -> dict:
    done = [k for k, v in _MAT_JOBS.items()
            if v.get("status") in ("done", "failed") and k != jid]
    for k in done[:-20]:            # on garde les 20 derniers pour les retards
        _MAT_JOBS.pop(k, None)
    _MAT_JOBS[jid] = state
    return state


def _mat_or_404(mid: str) -> dict:
    """Matière existante, ou l'erreur qui va bien. `mid` hors
    ^mat_[0-9a-f]{8}$ = 400 (aucune traversée ne peut aller plus loin)."""
    from app.services import material_store as MS
    if not MS.is_valid_mid(mid):
        raise HTTPException(400, "Identifiant de matière invalide")
    mat = MS.read_material(mid)
    if mat is None:
        raise HTTPException(404, "Matière introuvable")
    return mat


def _mat_library_path(filename) -> Path:
    """Fichier de la Library désigné par le client : basename uniquement +
    confinement dans images_path (doctrine de l'audit sécurité)."""
    fn = Path(str(filename or "")).name
    if not fn:
        raise HTTPException(400, "Nom de fichier manquant")
    src = (settings.images_path / fn).resolve()
    if not str(src).startswith(str(settings.images_path.resolve())) \
            or not src.is_file():
        raise HTTPException(400, f"Image introuvable dans la Librairie: {fn!r}")
    return src


def _mat_square(img: "PILImage.Image", res: int) -> "PILImage.Image":
    """Recadrage centré carré puis mise à la résolution demandée."""
    rgb = img.convert("RGB")
    w, h = rgb.size
    if w != h:
        s = min(w, h)
        x0, y0 = (w - s) // 2, (h - s) // 2
        rgb = rgb.crop((x0, y0, x0 + s, y0 + s))
    if rgb.size != (res, res):
        rgb = rgb.resize((res, res), PILImage.LANCZOS)
    return rgb


def _mat_build_maps(base: "PILImage.Image", derive: dict) -> dict:
    """basecolor + les 7 maps dérivées (pbr_service). Synchrone : appelé
    dans un thread par les routes."""
    from app.services import material_store as MS
    from app.services import pbr_service as PBR
    maps = PBR.derive_maps(base, derive, list(MS.SECONDARY_MAPS))
    if not isinstance(maps, dict):
        raise RuntimeError("pbr_service.derive_maps n'a pas renvoyé de maps")
    out = {k: v for k, v in maps.items() if v is not None}
    out["basecolor"] = base
    return out


@router.get("/materials")
async def list_materials():
    """Toutes les matières, plus récente d'abord."""
    from app.services import material_store as MS
    return {"materials": await asyncio.to_thread(MS.list_materials)}


@router.get("/materials/presets")
async def list_material_presets():
    """Préréglages de matière (« Appliquer un préréglage »)."""
    from app.services import material_store as MS
    return {"presets": MS.PRESETS}


@router.get("/materials/namings")
async def list_material_namings():
    """Les conventions d'export, avec l'emplacement de destination de chaque
    fichier dans le moteur visé.

    C'est ici que vit la vérité moteur, pas dans une chaîne écrite en dur à
    l'écran. La version précédente annonçait « Slots URP / HDRP : BaseMap,
    MaskMap, Occlusion » pour une cible « unity » unique — or la documentation
    Unity dit autre chose :

    - URP Lit n'a AUCUNE propriété Mask Map. Ses emplacements sont Base Map,
      Metallic Map (R=métal, A=smoothness), Normal Map, Height Map,
      Occlusion Map, Emission Map. Une texture packée est admise, à condition
      d'être assignée aux DEUX emplacements Metallic et Occlusion.
    - HDRP Lit lit un Mask Map (R=métal, V=occlusion, B=masque de détail,
      A=smoothness) ; l'occlusion y est déjà, donc pas d'emplacement Occlusion
      séparé.

    Les deux cibles sont donc séparées, et chacune livre exactement ce que son
    moteur branche."""
    from app.services import material_store as MS
    return {"namings": MS.naming_catalog(),
            "aliases": dict(MS.NAMING_ALIASES),
            "render_note": MS.RENDER_NOTE}


@router.get("/materials/seam-scale")
async def material_seam_scale():
    """L'échelle du score de raccord — pour que l'écran n'invente ni seuil ni
    couleur.

    `pixel_ops.seam_score` (le couple `seam.before` / `seam.after`) compare la
    colonne 0 à la colonne w-1 ; la passe seamless termine en les rendant
    IDENTIQUES, donc l'« après » vaut 0.00 pour toute matière corrigée : c'est
    une tautologie, pas une réussite. Le chiffre qui décide est
    `seam.ratio` — la marche à la jonction rapportée à la marche interne
    médiane du même motif, à trois échelles, le pire des trois. 1.00 = la
    jonction ne dépasse pas le grain normal de la matière.

    Les paliers viennent d'un test à deux alternatives en aveugle sur neuf
    tuiles réelles (fenêtre de 256 px à cheval sur la jonction contre fenêtre
    sans jonction) : 0 détection sur 3 à 0.96 / 0.98 / 1.36, 6 sur 6 de 1.27 à
    6.23. La détection est donc certaine à partir de 2.0."""
    from app.services import pbr_service as PBR
    return {
        "metric": "ratio",
        "scales_px": list(PBR.SEAM_SCALES),
        "grades": [{"max": limit, "grade": label,
                    "label": {"invisible": "Invisible",
                              "discret": "Discret",
                              "visible": "Visible"}.get(label, label)}
                   for limit, label in PBR.SEAM_GRADES]
        + [{"max": None, "grade": "cassé", "label": "Cassé"}],
        "visible_from": 2.0,
        "note": ("1.00 = la jonction ne dépasse pas la variation interne du "
                 "motif. Mesuré visible en aveugle à partir de 2.0."),
        "edge_note": ("seam.before / seam.after sont l'ancien score de bord "
                      "(0-100) ; l'après vaut 0.00 par construction."),
    }


@router.get("/materials/envs")
async def list_material_envs():
    """Les 7 ambiances du viewport."""
    from app.services import material_store as MS
    return {"envs": MS.env_list()}


@router.get("/materials/envs/{name}.jpg")
async def get_material_env(name: str):
    """Équirectangulaire d'une ambiance (liste blanche, cache disque)."""
    from app.services import material_store as MS
    try:
        data = await asyncio.to_thread(MS.env_jpeg, name)
    except ValueError as e:
        raise HTTPException(404, str(e))
    return Response(content=data, media_type="image/jpeg",
                    headers={"Cache-Control": "public, max-age=86400"})


@router.post("/materials/generate")
async def generate_material(body: dict, background_tasks: BackgroundTasks):
    """Lance une génération de matière. Corps :
    {prompt?, filename?, model?, res, seamless, seam_method, enhance}.
    `filename` (image de la Library) l'emporte sur `prompt`. Retour
    {job_id} — suivre GET /materials/jobs/{job_id}."""
    from app.services import material_store as MS

    body = body or {}
    filename = str(body.get("filename") or "").strip()
    prompt = str(body.get("prompt") or "").strip()
    res = MS.clean_res(body.get("res"), 2048)
    seamless = bool(body.get("seamless", True))
    method = str(body.get("seam_method") or "offset").strip().lower()
    if method not in MS.SEAM_METHODS:
        raise HTTPException(400, "seam_method doit valoir 'offset' ou 'mirror'")
    enhance = bool(body.get("enhance"))

    spec = {"res": res, "seamless": seamless, "seam_method": method,
            "enhance": enhance, "name": MS.clean_name(
                body.get("name") or prompt or Path(filename).stem or "Matière")}

    if filename:
        src = _mat_library_path(filename)      # fail fast (400 immédiat)
        spec.update({"kind": "library", "filename": src.name,
                     "prompt": prompt, "full_prompt": "", "model": None})
    else:
        if not prompt:
            raise HTTPException(400, "prompt ou filename est requis")
        model = MS.clean_model(body.get("model"))
        if model.startswith("gpt-image") or model.startswith("dall-e"):
            if not settings.OPENAI_API_KEY:
                raise HTTPException(400, "OPENAI_API_KEY non configurée "
                                         "(Réglages).")
        elif not settings.FAL_KEY:
            raise HTTPException(400, "FAL_KEY non configurée (Réglages).")
        spec.update({"kind": "prompt", "filename": None, "prompt": prompt,
                     "full_prompt": MS.build_full_prompt(prompt, enhance),
                     "model": model})

    jid = str(uuid4())
    _mat_job_register(jid, {"job_id": jid, "status": "pending",
                            "step": "En attente", "pct": 0,
                            "material": None, "error": None})
    background_tasks.add_task(_run_material_job, jid, spec)
    return {"job_id": jid}


@router.get("/materials/jobs/{job_id}")
async def material_job_status(job_id: str):
    st = _MAT_JOBS.get(job_id)
    if not st:
        raise HTTPException(404, "Job inconnu")
    return st


async def _run_material_job(jid: str, spec: dict):
    """prompt|Library -> image carrée -> raccord seamless (score avant/après)
    -> dérivation des 8 maps -> écriture disque."""
    from app.services import material_store as MS
    from app.services import pixel_ops as PX

    def upd(**kw):
        st = _MAT_JOBS.get(jid)
        if st is not None:
            st.update(kw)

    mid = None
    try:
        res = spec["res"]
        upd(status="running", step="Génération de l'image", pct=5)

        # 1. image de base
        if spec["kind"] == "library":
            src = _mat_library_path(spec["filename"])
        else:
            model = spec["model"]
            if model.startswith("gpt-image") or model.startswith("dall-e") \
                    or model == "nano-banana":
                from app.services import image_providers as IP
                out = await IP.generate(model, spec["full_prompt"],
                                        "square_hd", 1)
            else:
                out = await _flux_generate(spec["full_prompt"], "square_hd", 1)
            names = out.get("images") or []
            if not names:
                raise RuntimeError("Le générateur n'a renvoyé aucune image")
            src = settings.images_path / Path(names[0]).name
            spec["filename"] = src.name
            await LI.noter([src.name], "matieres")

        upd(step="Préparation", pct=40)
        with PILImage.open(src) as im:
            base = await asyncio.to_thread(_mat_square, im.copy(), res)

        # 2. raccord + scores (un chiffre, pas une promesse)
        upd(step="Raccord de tuile", pct=55)
        seam_before = await asyncio.to_thread(PX.seam_score, base)
        source_img = base
        if spec["seamless"]:
            opts = {"method": spec["seam_method"], "blend": 20,
                    "target_px": 0, "square": True}
            tiled = await asyncio.to_thread(PX.make_seamless, base, opts)
            base = tiled.convert("RGB")
            if base.size != (res, res):
                base = base.resize((res, res), PILImage.LANCZOS)
        seam_after = await asyncio.to_thread(PX.seam_score, base)

        # 3. enregistrement (le dossier existe avant la dérivation, pour que
        #    la source reste sur disque même si la dérivation échoue)
        mat = await asyncio.to_thread(
            MS.create_material, name=spec["name"], prompt=spec["prompt"],
            full_prompt=spec["full_prompt"], res=res,
            seamless=spec["seamless"],
            seam={"before": seam_before, "after": seam_after},
            source={"kind": spec["kind"], "model": spec.get("model"),
                    "filename": spec.get("filename")})
        mid = mat["id"]
        await asyncio.to_thread(MS.write_source, mid, source_img)

        # 4. dérivation des maps secondaires (PIL local, gratuit, hors ligne)
        upd(step="Dérivation des maps", pct=70, material=mat)
        maps = await asyncio.to_thread(_mat_build_maps, base, mat["derive"])

        upd(step="Écriture des maps", pct=90)
        await asyncio.to_thread(MS.save_maps, mid, maps)
        # Niveaux de départ MESURÉS sur les maps produites, au lieu des 1.00 /
        # 0.00 de principe : le curseur affiche dès l'ouverture la rugosité que
        # la texture porte réellement, et la cuire ne change alors rien.
        mat = await asyncio.to_thread(MS.read_material, mid)
        mat["props"] = MS.merge_props(mat["props"],
                                      MS.natural_levels(maps))
        # Ce que chaque map contient VRAIMENT + le rapport de couture après
        # correction : mesuré ici une fois, relu ensuite (l'écran n'invente
        # aucun des deux).
        mat = await asyncio.to_thread(MS.refresh_report, mat, maps)
        await asyncio.to_thread(MS.write_material, mat)
        mat = await asyncio.to_thread(MS.read_material, mid)

        upd(status="done", step="Terminé", pct=100, material=mat)
        logger.success(f"matière {mid} générée — raccord "
                       f"{seam_before} -> {seam_after}")
    except HTTPException as e:
        logger.warning(f"matière {jid} refusée: {e.detail}")
        upd(status="failed", step="Échec", error=str(e.detail))
    except Exception as e:
        logger.exception(f"matière {jid} échec: {e}")
        upd(status="failed", step="Échec", error=str(e),
            material=(MS.read_material(mid) if mid else None))


@router.get("/materials/{mid}")
async def get_material(mid: str):
    return {"material": _mat_or_404(mid)}


@router.patch("/materials/{mid}")
async def patch_material(mid: str, body: dict):
    """Fusion PARTIELLE de {name?, props?, derive?}. Toute valeur absente ou
    invalide garde/reprend son défaut — jamais d'erreur 500."""
    from app.services import material_store as MS
    mat = _mat_or_404(mid)
    body = body if isinstance(body, dict) else {}
    if "name" in body:
        mat["name"] = MS.clean_name(body.get("name"), fallback=mat["name"])
    before = mat["props"]
    if "props" in body:
        mat["props"] = MS.merge_props(mat["props"], body.get("props"))
    if "derive" in body:
        mat["derive"] = MS.merge_derive(mat["derive"], body.get("derive"))
    # metallic et roughness sont CUITS dans les maps livrées : changer le
    # réglage change ce que valent metallic.png, roughness.png et l'ORM. Les
    # statistiques annoncées doivent suivre, sinon l'écran reparlerait de
    # l'ancien niveau. Recalcul analytique, aucune image relue.
    if any(abs(mat["props"][k] - before[k]) > 1e-9
           for k in ("metallic", "roughness")):
        from app.services import pbr_service as PBR
        stats = dict(mat.get("map_stats") or {})
        eff = {}
        for kind, key in (("metallic", "metallic"), ("roughness", "roughness"),
                          ("orm", "roughness")):
            st = stats.get(kind)
            if not isinstance(st, dict) or not isinstance(st.get("pattern"), dict):
                continue
            new = PBR.level_stats(st["pattern"], mat["props"][key])
            # Moyenne EXACTE des octets cuits, tiree de l'histogramme du motif.
            # C'est elle qui alimente `render.effective` : sans ce recalcul, le
            # curseur bougeait, le GLB suivait, et le bloc `render` continuait
            # d'annoncer la valeur de la derniere derivation.
            hist = st.get("pattern_hist")
            if isinstance(hist, list) and len(hist) == 256:
                m = PBR.level_mean(hist, mat["props"][key])
                new["mean"] = round(m * 255.0, 2)
                if kind in ("metallic", "roughness"):
                    eff[key] = round(m, 3)
            note = ""
            if not new["informative"]:
                note = (f"uniforme — {'la matière est ' if key == 'metallic' else 'rugosité '}"
                        f"{mat['props'][key]:.2f}"
                        f"{' métallique partout' if key == 'metallic' else ' partout'}")
            stats[kind] = dict(st, **new, note=note)
        mat["map_stats"] = stats
        mat["maps_informative"] = sum(1 for v in stats.values()
                                      if v.get("informative"))
        # Le contrat de composition suit le curseur, ou se declare non mesure.
        mat["render"] = MS.render_block(mat["props"])
        if eff:
            mat["render"]["effective"].update(eff)
            mat["render"]["measured"] = True
    await asyncio.to_thread(MS.write_material, mat)
    return {"material": mat}


@router.post("/materials/{mid}/derive")
async def rederive_material(mid: str, body: dict = None):
    """Re-dérive les maps secondaires depuis la basecolor. Corps
    {derive?, res?}. Local et gratuit : la barre facture chaque ajustement."""
    from app.services import material_store as MS
    mat = _mat_or_404(mid)
    body = body if isinstance(body, dict) else {}
    if "derive" in body:
        mat["derive"] = MS.merge_derive(mat["derive"], body.get("derive"))
    base_p = MS.map_path(mid, "basecolor")
    if not base_p.is_file():
        raise HTTPException(409, "Cette matière n'a pas de basecolor sur "
                                 "disque — relancer une génération")
    res = MS.clean_res(body.get("res"), mat["res"]) if body.get("res") else mat["res"]
    mat["res"] = res
    try:
        with PILImage.open(base_p) as im:
            base = await asyncio.to_thread(_mat_square, im.copy(), res)
        maps = await asyncio.to_thread(_mat_build_maps, base, mat["derive"])
    except ImportError as e:
        raise HTTPException(503, f"Module de dérivation indisponible: {e}")
    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"re-dérivation {mid} échec: {e}")
        raise HTTPException(500, f"Dérivation impossible: {e}")
    await asyncio.to_thread(MS.save_maps, mid, maps)
    mat = await asyncio.to_thread(MS.refresh_report, mat, maps)
    await asyncio.to_thread(MS.write_material, mat)
    return {"material": MS.read_material(mid)}


@router.post("/materials/{mid}/duplicate")
async def duplicate_material(mid: str):
    from app.services import material_store as MS
    _mat_or_404(mid)
    mat = await asyncio.to_thread(MS.duplicate_material, mid)
    if mat is None:
        raise HTTPException(404, "Matière introuvable")
    return {"material": mat}


@router.put("/materials/{mid}/thumb")
async def put_material_thumb(mid: str, request: Request):
    """Vignette 512x512 poussée par le client (rendu du viewport), PNG brut."""
    from app.services import material_store as MS
    _mat_or_404(mid)
    data = await request.body()
    if not data:
        raise HTTPException(400, "Corps vide (PNG attendu)")
    if len(data) > 8 * 1024 * 1024:
        raise HTTPException(400, "Vignette trop lourde (max 8 MB)")
    ok, msg = await asyncio.to_thread(MS.write_thumb, mid, data)
    if not ok:
        raise HTTPException(400, f"Vignette invalide: {msg}")
    mat = MS.read_material(mid)
    if mat:
        await asyncio.to_thread(MS.write_material, mat)
    return {"ok": True}


@router.get("/materials/{mid}/thumb.png")
async def get_material_thumb(mid: str):
    from app.services import material_store as MS
    _mat_or_404(mid)
    d = MS.material_dir(mid)
    p = d / "thumb.png"
    # Une vignette est un rendu 3D figé : perimee par un changement de
    # geometrie/UV, elle remettrait a l'ecran la matiere d'avant (en miroir).
    # On la tient pour absente ; la carte retombe sur la couleur de base.
    if not MS.thumb_is_current(d):
        raise HTTPException(404, "Pas de vignette")
    return FileResponse(p, media_type="image/png")


@router.delete("/materials/{mid}")
async def delete_material(mid: str):
    from app.services import material_store as MS
    _mat_or_404(mid)
    if not await asyncio.to_thread(MS.delete_material, mid):
        raise HTTPException(500, "Suppression impossible")
    return {"ok": True}


@router.get("/materials/{mid}/map/{kind}.png")
async def get_material_map(mid: str, kind: str, res: int = 0):
    """Une map. `kind` est une liste blanche ; `res` (optionnel) redimensionne
    à la volée."""
    from app.services import material_store as MS
    _mat_or_404(mid)
    try:
        p = MS.map_path(mid, kind)
    except ValueError as e:
        raise HTTPException(400, str(e))
    if not p.is_file():
        raise HTTPException(404, "Map absente")
    kind = kind.lower()
    mat = MS.read_material(mid) or {}
    target = MS.clean_preview_res(res, 1024) if res else 0

    # La vignette de map servie ici est celle de l'EXPORT : mêmes niveaux cuits
    # (bake_levels), donc ce que l'inspecteur montre sous « METALLIC » est le
    # PNG que le moteur recevra. Le fichier brut sur disque ne porte que le
    # motif — il n'est jamais servi tel quel pour ces deux maps.
    def _served() -> bytes:
        with PILImage.open(p) as im:
            maps = {kind: im.copy()}
            if kind == "orm":
                maps[kind] = maps[kind].convert("RGB")
            if target:
                maps = MS.resize_maps(maps, target)
            maps = MS.bake_levels(maps, mat.get("props"))
        return MS.png_bytes(maps[kind], kind, 8)

    if not res and kind not in ("metallic", "roughness", "orm"):
        return FileResponse(p, media_type="image/png")
    return Response(content=await asyncio.to_thread(_served),
                    media_type="image/png")


@router.get("/materials/{mid}/preview.glb")
async def material_preview_glb(request: Request, mid: str,
                               mesh: str = "sphere", res: int = 1024,
                               stage: int = 0, scale: int = 1):
    """GLB d'aperçu (maillage + matériau + textures embarquées).

    La galerie demande un GLB par carte : sans cache, chaque scroll relance une
    reconstruction PIL de ~0.9 s par carte. On cache sur disque (clé = maillage
    + résolution + propriétés + mtime des maps) et on répond 304 quand le
    navigateur a déjà le bon ETag."""
    from app.services import material_store as MS
    mat = _mat_or_404(mid)
    mesh = str(mesh or "sphere").strip().lower()
    if mesh not in MS.MESHES:
        raise HTTPException(400, f"mesh doit être l'un de: "
                                 f"{', '.join(MS.MESHES)}")
    res = MS.clean_preview_res(res, 1024)
    stage = 1 if str(stage) not in ("0", "false", "") else 0
    scale = 0 if str(scale) in ("0", "false") else 1
    key = await asyncio.to_thread(MS.preview_cache_key, mat, mesh, res)
    # décor et échelle changent le GLB ; la version du décor aussi, sinon un
    # cache disque servirait l'ancienne grille après retouche.
    try:
        from app.services.stage_service import STAGE_VERSION as _SV
    except Exception:
        _SV = 0
    # L'échelle de matière fait partie de la géométrie servie : si MESH_UV
    # change (la correction du pavage fractionnaire, par exemple), le cache
    # disque doit se périmer tout seul — sinon on continuerait de servir
    # l'ancien maillage, artefact compris.
    try:
        from app.services.gltf_builder import MESH_UV as _MUV
        _uv = "x".join(str(v) for v in (_MUV.get(mesh) or (1, 1)))
    except Exception:
        _uv = "0"
    # La clé composée doit rester HEXADÉCIMALE : material_store.preview_cache_get
    # et preview_cache_put refusent (silencieusement) toute clé hors
    # [0-9a-f]{1,40}. Concaténée telle quelle — « ...-s1v2u1-4.0x2.0 » — elle
    # était rejetée aux deux bouts : le cache disque n'a jamais servi et chaque
    # carte de galerie reconstruisait son GLB. On la ré-empreinte donc.
    # `MESH_VERSION` entre dans l'empreinte : une géométrie qui change (la
    # densité polaire de la sphère, par exemple) doit périmer le cache toute
    # seule, sinon on continuerait de servir l'ancien maillage.
    try:
        from app.services.gltf_builder import MESH_VERSION as _MV
    except Exception:
        _MV = 0
    key = hashlib.sha1(
        f"{key}-s{stage}v{_SV}u{scale}-{_uv}-m{_MV}".encode("utf-8")
    ).hexdigest()[:24]
    etag = f'W/"{key}"'
    head = {"Content-Disposition": f'inline; filename="{mat["id"]}.glb"',
            "ETag": etag, "Cache-Control": "private, max-age=900"}
    if request.headers.get("if-none-match", "").find(key) >= 0:
        return Response(status_code=304, headers=head)
    data = await asyncio.to_thread(MS.preview_cache_get, mat["id"], key)
    if data is None:
        data = await asyncio.to_thread(_mat_glb, mat, mesh, res, None,
                                       bool(stage), bool(scale))
        await asyncio.to_thread(MS.preview_cache_put, mat["id"], key, data)
    return Response(content=data, media_type="model/gltf-binary", headers=head)


def _mat_glb(mat: dict, mesh: str, res: int, kinds, stage: bool = False,
             scale: bool = False) -> bytes:
    """Construit le GLB via gltf_builder (textures PNG 8 bits embarquées).

    `stage` ajoute le sol d'aperçu (grille + flaque de contact) et `scale`
    applique l'échelle de matière par maillage : les deux ne servent QUE
    l'aperçu du viewport. L'export reste une géométrie nue, UV 0..1, sans
    décor — c'est ce qu'un moteur attend."""
    from app.services import material_store as MS
    try:
        from app.services import gltf_builder as GB
    except ImportError as e:
        raise HTTPException(503, f"Module GLB indisponible: {e}")
    # L'ORM est le seul canal par lequel glTF sait lire rugosité, métal et
    # occlusion : on la garde toujours, même si la sélection d'export ne l'a
    # pas cochée — sans elle le GLB retomberait sur des facteurs et perdrait
    # la variation des maps.
    if kinds:
        kinds = list(dict.fromkeys(list(kinds) + ["orm"]))
    maps = MS.load_maps(mat["id"], kinds)
    if not maps:
        raise HTTPException(409, "Cette matière n'a aucune map sur disque")
    maps = MS.resize_maps(maps, res)
    # niveaux cuits AVANT encodage : le GLB sort avec metallicFactor =
    # roughnessFactor = 1.0 (gltf_builder), donc la map fait foi.
    maps = MS.bake_levels(maps, mat["props"])
    payload = {k: MS.png_bytes(v, k, 8) for k, v in maps.items()}
    ground = None
    if stage:
        try:
            from app.services import stage_service as ST
            ground = ST.stage_png()
        except Exception as e:                # un décor absent n'empêche rien
            logger.warning(f"sol d'aperçu indisponible: {e}")
    uvr = GB.MESH_UV.get(mesh) if scale else None
    try:
        return GB.build_glb(payload, mat["props"], mesh, stage_png=ground,
                            uv_repeat=uvr)
    except HTTPException:
        raise
    except Exception as e:
        logger.exception(f"GLB {mat['id']} échec: {e}")
        raise HTTPException(500, f"Construction GLB impossible: {e}")


@router.get("/materials/{mid}/export/manifest")
async def export_material_manifest(mid: str, format: str = "zip",
                                   naming: str = "standard", res: int = 0,
                                   bits: int = 8, maps: str = "",
                                   mesh: str = "sphere"):
    """Le bordereau de l'export AVANT téléchargement : nom de l'archive, liste
    des fichiers avec leur nom moteur, canaux, profondeur et poids, plus le
    total. Aucune image n'est ré-encodée : c'est une lecture de tailles sur
    disque (exactes) plus une mise à l'échelle quand la résolution change."""
    from app.services import material_store as MS
    mat = _mat_or_404(mid)
    fmt = str(format or "zip").strip().lower()
    if fmt not in MS.EXPORT_FORMATS:
        raise HTTPException(400, f"format doit être l'un de: "
                                 f"{', '.join(MS.EXPORT_FORMATS)}")
    nm = str(naming or "standard").strip().lower()
    # `unity` reste accepte (alias -> unity_urp) : une convention enregistree
    # cote client ne doit pas se mettre a repondre 400 apres la separation
    # URP / HDRP.
    if nm not in MS.NAMINGS and nm not in MS.NAMING_ALIASES:
        raise HTTPException(400, f"naming doit être l'un de: "
                                 f"{', '.join(MS.NAMINGS)}")
    nm = MS.clean_naming(nm)
    ms = str(mesh or "sphere").strip().lower()
    if ms not in MS.MESHES:
        ms = "sphere"
    wanted = [k.strip().lower() for k in str(maps or "").split(",") if k.strip()]
    allowed = tuple(MS.MAP_KINDS) + MS.EXPORT_EXTRA_KINDS
    bad = [k for k in wanted if k not in allowed]
    if bad:
        raise HTTPException(400, f"maps inconnues: {', '.join(bad)}")
    return await asyncio.to_thread(MS.export_manifest, mat, fmt, nm,
                                   res, bits, wanted or None, ms)


@router.get("/materials/{mid}/export")
async def export_material(mid: str, format: str = "zip",
                          naming: str = "standard", res: int = 0,
                          bits: int = 8, maps: str = "",
                          mesh: str = "sphere"):
    """Export : ZIP complet (maps + material.json + LISEZMOI), GLB, ou glTF
    autonome. `naming` ∈ standard|unity_urp|unity_hdrp|unreal|godot, `bits` ∈ 8|16 (honoré
    pour height et normal), `maps` = liste blanche séparée par des virgules."""
    from app.services import material_store as MS
    mat = _mat_or_404(mid)
    fmt = str(format or "zip").strip().lower()
    if fmt not in MS.EXPORT_FORMATS:
        raise HTTPException(400, f"format doit être l'un de: "
                                 f"{', '.join(MS.EXPORT_FORMATS)}")
    nm = str(naming or "standard").strip().lower()
    # `unity` reste accepte (alias -> unity_urp) : une convention enregistree
    # cote client ne doit pas se mettre a repondre 400 apres la separation
    # URP / HDRP.
    if nm not in MS.NAMINGS and nm not in MS.NAMING_ALIASES:
        raise HTTPException(400, f"naming doit être l'un de: "
                                 f"{', '.join(MS.NAMINGS)}")
    nm = MS.clean_naming(nm)
    try:
        bits = int(bits or 8)
    except (TypeError, ValueError):
        bits = 8
    if bits not in (8, 16):
        raise HTTPException(400, "bits doit valoir 8 ou 16")
    wanted = [k.strip().lower() for k in str(maps or "").split(",") if k.strip()]
    allowed = tuple(MS.MAP_KINDS) + MS.EXPORT_EXTRA_KINDS
    bad = [k for k in wanted if k not in allowed]
    if bad:
        raise HTTPException(400, f"maps inconnues: {', '.join(bad)}")
    kinds = wanted or None
    target = MS.clean_res(res, mat["res"]) if res else mat["res"]
    fname = MS.export_filename(mat, fmt, nm)

    # Le maillage EXPORTÉ est celui qui a été jugé à l'écran. Il était figé sur
    # "sphere" pendant que le bordereau annonçait « géométrie torus, 96 000 o » :
    # l'écran promettait une géométrie que l'archive ne contenait pas.
    ms = str(mesh or "sphere").strip().lower()
    if ms not in MS.MESHES:
        raise HTTPException(400, f"mesh doit être l'un de: "
                                 f"{', '.join(MS.MESHES)}")

    if fmt in ("glb", "gltf"):
        # `scale=True` : la MÊME échelle de matière que l'aperçu. Sans elle, la
        # sphère exportée portait UNE tuile étirée sur 0..1 là où le lab en
        # montrait 4x2 — le fichier livré ne ressemblait pas à ce qui avait été
        # validé. Le décor de scène (`stage`), lui, reste à l'aperçu : c'est du
        # mobilier, pas de la matière.
        glb = await asyncio.to_thread(_mat_glb, mat, ms, target, kinds,
                                      False, True)
        if fmt == "glb":
            return Response(content=glb, media_type="model/gltf-binary",
                            headers={"Content-Disposition":
                                     f'attachment; filename="{fname}"'})
        data = await asyncio.to_thread(MS.glb_to_gltf, glb)
        return Response(content=data, media_type="model/gltf+json",
                        headers={"Content-Disposition":
                                 f'attachment; filename="{fname}"'})

    def _zip() -> bytes:
        want = kinds
        if want is None:
            want = [k for k in MS.default_export_maps(nm) if k in MS.MAP_KINDS]
        need_mask = MS.MASKMAP in (kinds or MS.default_export_maps(nm))
        # le MaskMap se fabrique à partir de metallic / ao / roughness : il
        # faut les charger même si l'archive ne les emporte pas.
        load = list(want) + (["metallic", "ao", "roughness"] if need_mask else [])
        loaded = MS.load_maps(mid, list(dict.fromkeys(load)))
        if not loaded:
            raise HTTPException(409, "Cette matière n'a aucune map sur disque")
        loaded = MS.resize_maps(loaded, target)
        loaded = MS.bake_levels(loaded, mat["props"])
        if need_mask:
            mask = MS.build_maskmap(loaded)
            loaded = {k: v for k, v in loaded.items() if k in want}
            if mask is not None:
                loaded[MS.MASKMAP] = mask
        thumb_dir = MS.material_dir(mid)
        thumb_p = thumb_dir / "thumb.png"
        thumb = (thumb_p.read_bytes()
                 if MS.thumb_is_current(thumb_dir) else None)
        return MS.export_zip(mat, loaded, nm, bits, thumb)

    return Response(content=await asyncio.to_thread(_zip),
                    media_type="application/zip",
                    headers={"Content-Disposition":
                             f'attachment; filename="{fname}"'})


# =============================================================================
# Rack VFX — catalogue et APERÇUS d'effets
#
# Jusqu'ici un effet se choisissait sur son nom, appliqué à l'aveugle puis
# jugé après un rendu complet. /effects/preview rend UNE image fixe passée par
# la même chaîne ffmpeg que le rendu final : ce que la vignette montre est ce
# que le rendu produira. Le résultat est mis en cache sur disque, donc la même
# combinaison ne relance jamais ffmpeg.
# =============================================================================

@router.get("/effects/catalog")
async def effects_catalog():
    """Catalogue complet : catégories, libellés FR, paramètres et bornes."""
    from app.services import effects_preview as FXP
    return FXP.catalog_payload()


@router.get("/effects/preview")
async def effects_preview(request: Request):
    """Vignette JPEG d'un effet.

    Paramètres : `type` (obligatoire), `source` (`mire` par défaut,
    `image:<nom>` pour une image de la Bibliothèque, `job:<id>` pour une frame
    d'un rendu), `t` (instant, pour les effets animés), `w` (largeur), plus
    les paramètres propres à l'effet (intensity, speed, angle, c0, ...).
    """
    from app.services import effects_preview as FXP

    q = dict(request.query_params)
    etype = q.pop("type", "")
    source = q.pop("source", "") or ""
    t = q.pop("t", FXP.T_DEFAULT)
    width = q.pop("w", FXP.W_DEFAULT)

    job_video = None
    if str(source).startswith("job:"):
        from app.services.storage import JobRecord, async_session_factory
        jid = str(source)[4:]
        async with async_session_factory() as session:
            jr = await session.get(JobRecord, jid)
        fp = jr and (jr.final_video_path or jr.video_path)
        if not fp or not Path(fp).is_file():
            raise HTTPException(404, "Rendu introuvable pour cet aperçu")
        job_video = Path(fp)

    try:
        p = await asyncio.to_thread(
            FXP.render_preview, etype, q, source=source, t=t, width=width,
            job_video=job_video)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except RuntimeError as e:
        raise HTTPException(500, str(e))
    return FileResponse(p, media_type="image/jpeg",
                        headers={"Cache-Control": "public, max-age=86400"})


# =============================================================================
# Sous-titres — piste S1 du Montage
#
# Trois couches, trois responsabilités :
#   * `subtitle_service`  — le format (ASS/SRT/VTT, styles, karaoké, qualité) ;
#   * `transcribe_service`— le CALAGE (texte connu → mots datés, gratuit ;
#                           ou transcription payante quand le texte est inconnu) ;
#   * `subtitle_ui`       — la traduction entre le vocabulaire du panneau
#                           (`frontend/patches/subs.js`) et celui du moteur.
#
# Ces routes ne font que les brancher. Le vrai point d'arrivée est la GRAVURE :
# `POST /api/montage/render` lit la clé `subtitles` du payload, écrit l'ASS et
# le passe au filtre ffmpeg `subtitles=` — c'est là que le style vu dans
# l'aperçu devient le style de la vidéo exportée.
#
# Toutes les entrées venant du client sont réduites au basename et vérifiées
# contenues dans leur dossier (même garde-fou que `_lut_path`) — ici via
# `transcribe_service.resolve_media` et la table close `FONT_FILES`.
# =============================================================================

#: Travaux de transcription en cours : id → {status, pct, step, segments, error}.
#: En mémoire volontairement — un travail dure quelques secondes et le panneau
#: interroge tout de suite ; rien à faire survivre à un redémarrage.
_SUBS_JOBS: dict[str, dict] = {}
_SUBS_JOBS_MAX = 40


def _subs_track_path() -> Path:
    """Piste de sous-titres persistée, à côté de `montage_saved.json`."""
    return settings.images_path.parent / "subtitles_track.json"


def _subs_canvas(ratio=None, preview: bool = False) -> tuple[int, int]:
    from app.services import subtitle_ui as SU
    return SU.canvas_for_ratio(ratio, preview)


def _subs_body_style(body: dict, canvas: tuple[int, int]) -> dict:
    """Style posté par le panneau → style du moteur, résolu."""
    from app.services import subtitle_service as S
    from app.services import subtitle_ui as SU
    st = body.get("style")
    if isinstance(st, str):                     # nom de préréglage moteur
        return S.resolve_style(st)
    return SU.ui_to_style(st if isinstance(st, dict) else {}, canvas)


def _subs_body_segments(body: dict) -> list:
    segs = body.get("segments")
    if not isinstance(segs, list):
        segs = body.get("clips") if isinstance(body.get("clips"), list) else []
    return [s for s in segs if isinstance(s, dict)]


def _subs_body_normes(body: dict) -> dict:
    """Seuils de lisibilite POSTES par le panneau -> kwargs de `check_quality`.

    Les normes de sous-titrage varient selon le diffuseur (EBU 20 c/s, Netflix
    17 en francais, reseaux sociaux 25) : le panneau les AFFICHE et les laisse
    regler, donc il doit pouvoir les faire suivre au moteur. Sans cela l'ecran
    montrerait « 17 c/s » pendant que le moteur continuerait de marquer a 20 —
    deux regles dans la meme image, c'est-a-dire aucune.

    Toute valeur absente, non numerique ou hors bornes retombe sur la constante
    du moteur : un corps hostile ne peut pas eteindre le controle qualite en
    postant `{"cps_warn": 0}` ou une chaine.
    """
    from app.services import subtitle_service as S
    src = body.get("normes")
    if not isinstance(src, dict):
        src = {}

    def _b(key, default, lo, hi):
        try:
            v = float(src[key])
        except (KeyError, TypeError, ValueError):
            return default
        if v != v or v in (float("inf"), float("-inf")) or not (lo <= v <= hi):
            return default
        return v

    cps_warn = _b("cps_warn", S.CPS_WARN, 8.0, 40.0)
    out = {
        "cps_warn": cps_warn,
        "cps_error": _b("cps_error", round(cps_warn * 1.35, 2),
                        cps_warn, 120.0),
        "min_duration": _b("min_duration", S.MIN_DURATION, 0.2, 4.0),
        "max_duration": _b("max_duration", S.MAX_DURATION, 2.0, 20.0),
        "min_gap": _b("min_gap", S.MIN_GAP, 0.0, 0.5),
    }
    # un minimum au-dessus du maximum marquerait TOUTES les repliques des deux
    # cotes a la fois : la borne basse cede, comme dans le panneau.
    if out["min_duration"] > out["max_duration"] - 0.1:
        out["min_duration"] = max(0.2, round(out["max_duration"] - 0.1, 2))
    return out


#: Code du moteur → genre affiché par le panneau (sert au tri et au libellé
#: de la puce, pas au geste : le geste vient du PLAN, qui porte lui-même son
#: libellé, ses conséquences et ses `ops`).
_SUBS_W_UI = {
    "texte_vide": ("vide", None),
    "duree_nulle": ("court", None),
    "trop_court": ("court", None),
    "trop_long": ("long", None),
    "debit_eleve": ("vitesse", None),
    "debit_illisible": ("vitesse", None),
    "chevauchement": ("chevauche", None),
    "intervalle_court": ("intervalle", None),
    "ligne_trop_large": ("large", None),
    "trop_de_lignes": ("lignes", None),
    "mots_incoherents": ("mots", None),
    "fond_translucide_karaoke": ("style", None),
}
_SUBS_SEV_UI = {"erreur": "err", "avertissement": "warn"}

#: LA MESURE qui a motivé la pastille, en trois caractères ou presque.
#: Une ligne repliée du panneau ne montrait qu'un temps de début : les chiffres
#: qu'elle demandait de croire ne pouvaient être rattachés à aucun calage
#: concret sans déplier. Le moteur mesure (`value`, `limit`) — il envoie donc
#: SON chiffre, y compris là où le panneau ne sait pas le refaire : la largeur
#: de ligne en PIXELS avec la vraie fonte (`ligne_trop_large`).
def _subs_fr(v: float, dec: int = 2) -> str:
    """Nombre à la française, zéros de queue retirés (« 1,35 », « 0,8 »)."""
    s = ("%.*f" % (dec, float(v)))
    if "." in s:
        s = s.rstrip("0").rstrip(".")
    return s.replace(".", ",") or "0"


_SUBS_MES_UI = {
    "texte_vide": lambda v, l: "0 car.",
    "duree_nulle": lambda v, l: "0 s",
    "trop_court": lambda v, l: _subs_fr(v) + " s",
    "trop_long": lambda v, l: _subs_fr(v, 1) + " s",
    "debit_eleve": lambda v, l: "%d c/s" % round(v),
    "debit_illisible": lambda v, l: "%d c/s" % round(v),
    "chevauchement": lambda v, l: "−%d ms" % round(abs(v) * 1000),
    "intervalle_court": lambda v, l: "%d ms" % round(v * 1000),
    "ligne_trop_large": lambda v, l: ("%d/%d px" % (round(v), round(l))
                                      if l else "%d px" % round(v)),
    "trop_de_lignes": lambda v, l: "%d lignes" % round(v),
    "mots_incoherents": lambda v, l: "%d mots" % round(v),
}


def _subs_mes(code: str, w: dict) -> str:
    """Mesure courte d'un avertissement, ou chaîne vide si le moteur n'en a
    pas fourni (le panneau retombe alors sur son propre calcul local)."""
    fn = _SUBS_MES_UI.get(str(code))
    if fn is None:
        return ""
    try:
        v = w.get("value")
        if v is None and code not in ("texte_vide", "duree_nulle"):
            return ""
        return fn(float(v or 0.0), float(w.get("limit") or 0.0))
    except (TypeError, ValueError):                    # noqa: BLE001
        return ""

#: Correctif de STYLE du moteur → réglage du panneau, avec ce que le bouton
#: fait. Un avertissement de style n'a pas d'index de segment : il se répare
#: dans l'onglet Style, c'est là qu'on le pose.
_SUBS_STYLE_FIX_UI = {
    "back_opacity": lambda v: {
        "champ": "bgOpacity", "valeur": int(round(float(v) * 100)),
        "label": "Passer le fond en opaque",
        "effect": "L'opacité du fond passe à 100 %. Les coutures entre les "
                  "mots disparaissent au rendu ; le fond ne laissera plus "
                  "voir l'image derrière."},
}


def _subs_plan_ui(plan: dict | None) -> dict | None:
    """Plan négocié du moteur → contrat du panneau (mêmes mots, mêmes `ops`).

    Les `ops` sont des DONNÉES, pas du code : `{op, index, id, start, end,
    text}`. Le panneau les applique tel quel — il n'a rien à recalculer, donc
    rien à faire dériver du moteur.
    """
    if not isinstance(plan, dict):
        return None
    out = {"action": plan.get("action"), "ok": bool(plan.get("ok")),
           "label": str(plan.get("label") or ""),
           "effect": str(plan.get("effect") or ""),
           "ops": plan.get("ops") or [], "touches": plan.get("touches") or []}
    if plan.get("blocked"):
        out["blocked"] = str(plan["blocked"])
    for k in ("granted", "requested"):
        if plan.get(k) is not None:
            out[k] = plan[k]
    alt = plan.get("alt")
    if isinstance(alt, dict) and alt.get("ok"):
        out["alt"] = _subs_plan_ui(alt)
    return out


def _subs_warnings_ui(raw: list, segs: list) -> tuple[list, list]:
    """Avertissements du moteur → vocabulaire du panneau.

    **Aucun ré-ancrage.** L'ancienne version déplaçait `chevauchement` et
    `intervalle_court` sur `index - 1` pour que le bouton « Séparer » tombe
    sur le segment dont il raccourcissait la fin — mais le MESSAGE, lui,
    restait écrit du point de vue du segment suivant. Résultat constaté :
    « 60 ms depuis le segment précédent » s'affichait sur la carte 12 alors
    que le seul écart de 60 ms précède la carte 13. Le message et la carte
    doivent parler du même segment.

    Désormais le moteur ancre chaque avertissement sur le segment qu'il
    MESURE, nomme son partenaire par son rang dans le message, et le correctif
    est un PLAN qui porte lui-même la liste des segments qu'il modifie
    (`ops`) : plus personne n'a besoin de deviner qui bouge.

    Retourne (avertissements de SEGMENT, avertissements de STYLE) — les
    seconds n'ont pas d'index et se réparent dans l'onglet Style.
    """
    seg_w, style_w = [], []
    for w in raw or []:
        code = str(w.get("code") or "")
        kind, _legacy = _SUBS_W_UI.get(code, (code or "regle", None))
        sev = _SUBS_SEV_UI.get(str(w.get("severity")), "warn")
        i = w.get("index")
        msg = str(w.get("message") or "")
        about = [int(x) for x in (w.get("about") or [])]
        if i is None:
            fx = w.get("fix") or {}
            maker = _SUBS_STYLE_FIX_UI.get(str(fx.get("champ")))
            d = {"kind": kind, "sev": sev, "msg": msg, "code": code,
                 "style": w.get("style"), "about": about}
            if maker and fx.get("valeur") is not None:
                d["fix"] = maker(fx["valeur"])
            style_w.append(d)
            continue
        i = int(i)
        if i >= len(segs):
            continue
        if code == "trop_long":
            sev = "info"
        d = {"i": i, "id": segs[i].get("id"), "kind": kind, "sev": sev,
             "msg": msg, "code": code, "about": about or [i],
             # la MESURE, à afficher à côté de la pastille de la ligne repliée
             "mes": _subs_mes(code, w)}
        p = _subs_plan_ui(w.get("plan"))
        if p:
            d["plan"] = p
        seg_w.append(d)
    return seg_w, style_w


# ------------------------------------------------------- styles et fontes ---

@router.get("/subtitles/presets")
async def subtitles_presets(ratio: str = "9:16"):
    """Préréglages de style, dans le vocabulaire du panneau.

    Une seule source de vérité (`subtitle_service.STYLES`, neuf préréglages
    sur fontes EMBARQUÉES) : le panneau propose donc exactement les styles qui
    savent se graver, au lieu de son repli local sur des fontes système que
    libass ne trouvera pas forcément.
    """
    from app.services import subtitle_service as S
    from app.services import subtitle_ui as SU
    canvas = _subs_canvas(ratio)
    fonts = S.check_fonts()
    return {"ok": True, "ratio": ratio, "canvas": list(canvas),
            "default": S.DEFAULT_STYLE,
            "presets": SU.ui_presets(canvas),
            "fonts_missing": fonts["missing"]}


@router.get("/subtitles/fonts")
async def subtitles_fonts():
    """Fontes EMBARQUÉES gravables, avec l'URL du fichier.

    L'URL sert à l'aperçu : le panneau déclare une `@font-face` par famille,
    et le texte à l'écran est dessiné avec la fonte QUI SERA GRAVÉE. Sans
    cela l'aperçu montrerait une fonte système et le rendu une autre.
    """
    from app.services import subtitle_service as S
    st = S.check_fonts()
    # `lh` : hauteur de ligne REELLE de libass, en multiples du corps demande.
    # L'aperçu s'en sert comme `line-height` ; sans lui, deux lignes gravees
    # seraient plus espacees a l'ecran qu'a l'image (mesure : libass avance
    # d'exactement un Fontsize, or Fontsize = em x lh depuis la correction
    # d'echelle de `to_ass`).
    fonts = [{"id": fam, "label": fam, "file": S.FONT_FILES[fam],
              "lh": round(S.font_line_height(fam), 4),
              "url": f"/api/subtitles/fonts/{fam}"}
             for fam in st["ok"]]
    return {"ok": True, "fonts": fonts, "dir": st["dir"],
            "missing": st["missing"]}


@router.get("/subtitles/fonts/{family}")
async def subtitles_font_file(family: str):
    """Fichier .ttf d'une famille embarquée.

    Garde-fou : `family` doit être une clé de la table CLOSE `FONT_FILES` —
    aucun chemin ne vient du client, donc rien d'autre que les fontes livrées
    ne peut sortir.
    """
    from app.services import subtitle_service as S
    fam = str(family or "").strip()
    if fam.lower().endswith(".ttf"):
        fam = fam[:-4]
    p = S.font_path(fam)
    if p is None:
        raise HTTPException(404, f"Fonte inconnue ou non livrée : {family}")
    return FileResponse(p, media_type="font/ttf",
                        headers={"Cache-Control": "public, max-age=604800"})


@router.post("/subtitles/style")
async def subtitles_style(request: Request):
    """Style du panneau → style du MOTEUR (diagnostic).

    Renvoie le style ASS résolu, la fonte réellement retenue, et surtout
    `unsupported` : ce que l'aperçu montre et que la gravure ne peut PAS
    reproduire (coins arrondis, flou d'ombre, interligne, animations…).
    """
    from app.services import subtitle_ui as SU
    try:
        body = await request.json()
    except Exception:
        body = {}
    if not isinstance(body, dict):
        body = {}
    canvas = _subs_canvas(body.get("ratio"), bool(body.get("preview")))
    st = _subs_body_style(body, canvas)
    return {"ok": True, "canvas": list(canvas), "style": st,
            "karaoke": SU.ui_karaoke(body.get("style")),
            "karaoke_mode": SU.ui_karaoke_mode(body.get("style")),
            "unsupported": SU.ui_unsupported(
                body.get("style") if isinstance(body.get("style"), dict) else {},
                canvas)}


# --------------------------------------------------------------- qualité ---

@router.post("/subtitles/check")
async def subtitles_check(request: Request):
    """Avertissements de qualité, traduits pour le panneau.

    Body : `{segments:[{start,end,text,words?}], style?, ratio?, dur?}`.
    Le moteur mesure la largeur de ligne en PIXELS avec la vraie fonte (ce que
    le calcul local du panneau ne sait pas faire) ; on lui rend donc la main
    dès qu'il répond, plans de réparation compris.

    `dur` (durée du montage) sert au dernier segment : sans plafond, un
    étirement croirait disposer d'un silence infini après la fin de la vidéo.
    """
    from app.services import subtitle_service as S
    from app.services import subtitle_ui as SU
    try:
        body = await request.json()
    except Exception:
        body = {}
    if not isinstance(body, dict):
        body = {}
    canvas = _subs_canvas(body.get("ratio"))
    segs_in = _subs_body_segments(body)
    st = _subs_body_style(body, canvas)
    kar = SU.ui_karaoke(body.get("style"))
    try:
        dur = float(body.get("dur"))
        media_dur = dur if dur > 0 else None
    except (TypeError, ValueError):
        media_dur = None
    # keep_empty=True : la MÊME indexation que `check_quality`, donc que les
    # cartes du panneau. Écarter les segments vides ici décalerait tous les
    # avertissements suivants d'un cran.
    segs = S.normalize_segments(segs_in, sort=False, clamp_words=False,
                                keep_empty=True)
    # LES SEUILS VIENNENT DU PANNEAU : ce sont ceux qu'il ECRIT a l'ecran.
    normes = _subs_body_normes(body)
    raw = S.check_quality(segs_in, st, canvas, karaoke=kar, media_dur=media_dur,
                          **normes)
    seg_w, style_w = _subs_warnings_ui(raw, segs)
    return {"ok": True, "warnings": seg_w, "style_warnings": style_w,
            "normes": normes,
            "unsupported": SU.ui_unsupported(
                body.get("style") if isinstance(body.get("style"), dict) else {},
                canvas),
            "segments": len(segs)}


@router.post("/subtitles/autofix")
async def subtitles_autofix(request: Request):
    """Correctifs TEMPORELS appliqués (chevauchements, segments trop courts,
    ordre) — le texte n'est jamais touché sauf `split_long:true`."""
    from app.services import subtitle_service as S
    try:
        body = await request.json()
    except Exception:
        body = {}
    if not isinstance(body, dict):
        body = {}
    canvas = _subs_canvas(body.get("ratio"))
    st = _subs_body_style(body, canvas)
    segs = S.autofix(_subs_body_segments(body), st,
                     split_long=bool(body.get("split_long")))
    return {"ok": True, "segments": segs, "count": len(segs)}


@router.post("/subtitles/emoji-hints")
async def subtitles_emoji_hints(request: Request):
    """Emoji suggérés par MOT-CLÉ dans une piste de sous-titres.

    Body : `{segments:[{start,end,text,words?}]}`.
    Réponse : `{ok, hints:[{t, word, emoji, file, png, url}], count, manifest}`
    — `t` est le début du MOT, `png` un chemin ABSOLU existant (le Montage le
    pose tel quel en `src:{file_path:…}` d'un clip d'overlay) et `url` la
    forme servie au navigateur (`/emoji/<f>.png`, le même dessin que le
    sélecteur).

    `manifest` est le NOMBRE d'emoji indexés. Sans lui, un manifeste illisible
    (`emoji_manifest()` rend `{}` sans bruit, et c'est le bon choix : une
    suggestion ne doit pas empêcher un rendu) sortait `count: 0` — exactement
    comme un texte sans mot-clé, et l'appelant accusait le texte de
    l'utilisateur. `manifest: 0` sépare les deux cas.

    Cette route ne MODIFIE rien : elle propose. C'est le Montage qui décide
    d'en faire des clips, et ces clips s'annulent comme les autres.
    """
    from app.services import subtitle_service as S
    try:
        body = await request.json()
    except Exception:
        body = {}
    if not isinstance(body, dict):
        body = {}
    hints = S.emoji_hints(_subs_body_segments(body))
    return {"ok": True, "hints": hints, "count": len(hints),
            "manifest": len(S.emoji_manifest())}


@router.post("/subtitles/fillers")
async def subtitles_fillers(request: Request):
    """Plages de MOTS DE REMPLISSAGE (« euh », « hum », « um »…) dans un texte
    horodaté.

    Body : `{words:[{i?, w, start, end}]}` — la liste PLATE des mots — ou
    `{segments:[{words:[…]}]}`, dont les mots sont alors aplatis ET
    renumérotés sur la liste plate (sans quoi deux répliques auraient chacune
    un mot d'indice 0, et l'écran surlignerait le mauvais bouton).
    `lang` : `fr` par défaut, `en` reconnu ; toute autre langue retombe sur le
    sac français plutôt que de ne rien proposer.

    `kind` : `all` (défaut) rend les deux natures, `hesitation` ne rend que
    les NON-MOTS (euh, hum, um, uh). C'est la distinction qui décide de ce
    qu'un bouton a le droit d'emporter sans qu'on relise : mesuré, une
    narration française sans une seule hésitation rend quand même cinq
    plages de `tic` — « Voilà », « Enfin », « genre », « quoi » — dont
    quatre portent la phrase.

    Réponse : `{ok, lang, kind, spans:[{start, end, kind, words:[i]}], count}`.

    Cette route ne MODIFIE rien : elle propose des plages. C'est le Montage
    qui décide d'en faire des coupes (`DzTracks.rippleCut`), et ces coupes
    passent par l'historique — « annuler » ramène les clips.
    """
    from app.services import transcribe_service as T
    try:
        body = await request.json()
    except Exception:
        body = {}
    if not isinstance(body, dict):
        body = {}
    lang = str(body.get("lang") or "fr")[:5]
    kind = str(body.get("kind") or "all")[:16]
    words = body.get("words")
    if not isinstance(words, list):
        words = []
        for seg in _subs_body_segments(body):
            for w in seg.get("words") or []:
                # l'indice est la POSITION dans la liste plate — y compris
                # pour une entrée illisible, sinon les suivantes glissent.
                if isinstance(w, dict):
                    words.append({"i": len(words),
                                  "w": w.get("w", w.get("word", "")),
                                  "start": w.get("start"), "end": w.get("end")})
                else:
                    words.append({"i": len(words), "w": ""})
    spans = T.find_fillers(words, lang, kind=kind)
    return {"ok": True, "lang": lang, "kind": kind, "spans": spans,
            "count": len(spans)}


# ---------------------------------------------------------------- export ---

_SUBS_MIME = {"srt": "application/x-subrip", "vtt": "text/vtt",
              "ass": "text/x-ssa", "txt": "text/plain"}


@router.post("/subtitles/export")
async def subtitles_export(request: Request):
    """Export d'une piste en `srt`, `vtt`, `ass` (ou `txt`).

    L'ASS est le seul des trois qui porte le STYLE et le KARAOKÉ — c'est
    exactement le fichier que ffmpeg grave au rendu, donc l'export permet de
    vérifier hors de l'app ce qui sortira dans la vidéo.
    Body : `{format, segments, style?, ratio?, karaoke?, name?}`.
    """
    from app.services import subtitle_service as S
    from app.services import subtitle_ui as SU
    try:
        body = await request.json()
    except Exception:
        body = {}
    if not isinstance(body, dict):
        body = {}
    fmt = str(body.get("format") or "srt").strip().lower().lstrip(".")
    if fmt not in _SUBS_MIME:
        raise HTTPException(400, f"Format inconnu : {fmt} — srt, vtt, ass ou txt.")
    segs = _subs_body_segments(body)
    if not segs:
        raise HTTPException(400, "Aucun segment à exporter.")
    canvas = _subs_canvas(body.get("ratio"))
    st = _subs_body_style(body, canvas)
    kar = (bool(body["karaoke"]) if "karaoke" in body
           else SU.ui_karaoke(body.get("style")))
    if fmt == "srt":
        text = S.to_srt(segs)
    elif fmt == "vtt":
        text = S.to_vtt(segs, word_timings=kar)
    elif fmt == "txt":
        text = "\n".join(s["text"].replace("\n", " ")
                         for s in S.normalize_segments(segs, with_words=False))
    else:
        text = S.to_ass(segs, st, canvas=canvas, karaoke=kar,
                        karaoke_mode=SU.ui_karaoke_mode(body.get("style")))
    base = re.sub(r"[^\w\-. ]+", "_", str(body.get("name") or "sous-titres"))[:60]
    return Response(content=text.encode("utf-8"),
                    media_type=f"{_SUBS_MIME[fmt]}; charset=utf-8",
                    headers={"Content-Disposition":
                             f'attachment; filename="{base}.{fmt}"'})


@router.post("/subtitles/import")
async def subtitles_import(request: Request):
    """Lecture d'un .srt / .vtt collé → segments (mots relus si le VTT les
    porte). Body : `{text}` ou `{content}`."""
    from app.services import subtitle_service as S
    try:
        body = await request.json()
    except Exception:
        body = {}
    text = str((body or {}).get("text") or (body or {}).get("content") or "")
    if not text.strip():
        raise HTTPException(400, "Rien à lire — collez un .srt ou un .vtt.")
    fmt = S.sniff_format(text)
    segs = S.parse_subtitles(text)
    if not segs:
        raise HTTPException(400, f"Aucun sous-titre lisible (format vu : {fmt}).")
    return {"ok": True, "format": fmt, "segments": segs, "count": len(segs)}


# ------------------------------------------------- lecture / écriture piste ---

@router.get("/subtitles/track")
async def subtitles_track_get():
    """Piste enregistrée (segments + style), ou piste vide."""
    p = _subs_track_path()
    try:
        if p.is_file():
            data = json.loads(p.read_text(encoding="utf-8"))
            if isinstance(data, dict) and isinstance(data.get("segments"), list):
                return {"ok": True, "saved": True, **data}
    except (OSError, ValueError):
        logger.warning("subtitles: subtitles_track.json illisible — piste vide")
    return {"ok": True, "saved": False, "segments": [], "style": None}


@router.put("/subtitles/track")
async def subtitles_track_put(request: Request):
    """Écrit la piste (segments normalisés + style du panneau tel quel).

    Écriture atomique (tmp voisin + replace), comme `montage_service`.
    """
    from app.services import subtitle_service as S
    try:
        body = await request.json()
    except Exception:
        body = {}
    if not isinstance(body, dict):
        raise HTTPException(400, "Objet {segments:[…], style?} attendu.")
    segs = S.normalize_segments(_subs_body_segments(body))
    if len(segs) > 2000:
        raise HTTPException(400, f"{len(segs)} segments — 2000 au maximum.")
    data = {"segments": segs,
            "style": body.get("style") if isinstance(body.get("style"), dict) else None,
            "ratio": str(body.get("ratio") or "9:16")[:12],
            "saved_at": datetime.utcnow().replace(microsecond=0).isoformat() + "Z"}
    p = _subs_track_path()
    tmp = p.with_name(f"{p.name}.{uuid4().hex[:8]}.tmp")
    try:
        await asyncio.to_thread(
            tmp.write_text, json.dumps(data, ensure_ascii=False), "utf-8")
        await asyncio.to_thread(tmp.replace, p)
    except OSError as e:
        try:
            tmp.unlink()
        except OSError:
            pass
        raise HTTPException(500, f"Écriture de la piste impossible : {e}")
    return {"ok": True, "count": len(segs), "saved_at": data["saved_at"]}


# ------------------------------------------------------ calage / transcription ---

def _subs_audio_dir() -> Path:
    return settings.images_path.parent / "audio"


async def _subs_resolve(src) -> Path | None:
    """Source du panneau → fichier réel, avec le garde-fou de chemin.

    Accepte le même vocabulaire que le Montage (`{job_id}`, `{audio}`,
    `{image}`, `{file_path}`) et, en plus, un simple nom de fichier — réduit
    au basename et vérifié contenu dans le dossier audio ou dans
    `outputs/uploads` (`transcribe_service.resolve_media`).
    """
    from app.services import transcribe_service as T
    from app.services.montage_service import _resolve_src as _mres
    if isinstance(src, dict):
        p = await _mres(src)
        if p is not None:
            return p
        src = src.get("name") or src.get("audio") or src.get("filename")
    if isinstance(src, str) and src.strip():
        for folder in (_subs_audio_dir(),
                       settings.outputs_path / "uploads",
                       settings.outputs_path / "videos"):
            q = T.resolve_media(src.strip(), folder)
            if q is not None:
                return q
    return None


def _subs_cues_to_segments(cues: list) -> list:
    from app.services import subtitle_service as S
    return S.normalize_segments(
        [{"start": c["start"], "end": c["end"], "text": c["text"],
          "words": c.get("words")} for c in cues or []])


# ── P13 (06/09/2026) — LA TRANSCRIPTION VISE LA PISTE DE DIALOGUE ET DÉCALE ──
# Mesuré sur le journal du 06/09 : `transcribe: s1_drift-746849.mp3 (11.8s)`
# — le vieux MP3 de A1, jamais la vidéo ; et `_subs_cues_to_segments(cues)`
# posait les mots AU TEMPS DU FICHIER, sans le `start` ni le `srcIn` du clip.
# Ça « marchait avant » parce que le vestige A1 était à t = 0. SIX aides
# suivent ; QUATRE sont pures (`_subs_src_key`, `_subs_bornes`,
# `_subs_carrier`, `_subs_shift_words`), DEUX résolvent (`_subs_dialogue_ids`
# lit `_tracks_meta`, `_subs_dialogue_sources` résout les chemins). Toutes
# testées par backend/tests/test_subs_transcribe_cible.py.

def _subs_src_key(src) -> str:
    """Clé de comparaison d'une source : JSON aux clés triées — la même loi
    que la couche (`dzmSrcKey`, JSON.stringify des clés triées), pour que
    `{job_id:"x"}` et `{"job_id": "x"}` soient la même source."""
    try:
        return json.dumps(src, sort_keys=True, ensure_ascii=False, default=str)
    except (TypeError, ValueError):
        return repr(src)


def _subs_bornes(c: dict) -> tuple[float, float, float]:
    """`(start, end, srcIn)` d'un clip du payload, en secondes, bornés :
    jamais négatifs, `end >= start`. Le tiroir envoie `start`/`end` arrondis
    au millième et — depuis P13 (section M24i du patcher montage) — `srcIn` ;
    un payload d'avant, ou une sauvegarde, peut ne pas le porter : 0."""
    def _f(v, d=0.0):
        try:
            x = float(v)
        except (TypeError, ValueError):
            return d
        return x if x == x and x >= 0 else d      # NaN et négatifs → défaut
    start = _f(c.get("start"))
    end = max(start, _f(c.get("end"), start))
    return start, end, _f(c.get("srcIn"))


def _subs_dialogue_ids(tracks) -> set[str]:
    """Les pistes de dialogue du payload : MÊME LOI que le rendu
    (`montage_service._tracks_meta` — `tracks` absent ⇒ la table historique,
    donc `a1`) : de genre audio et de bus « dialogue », jamais une piste
    BOUCLÉE (le rendu joue une piste `loop` d'un bout à l'autre du film en
    ignorant les bornes de son premier clip : son `start` ne décale rien).
    Reste connu, dit ici : la couche (`DzTracks.dialogueTrack`) retombe sur
    `a1` PAR IDENTIFIANT même quand un payload la re-buse ailleurs ; la
    route suit `_tracks_meta` (le bus fait foi). Les deux ne diffèrent que
    sur une a1 explicitement sortie du bus dialogue."""
    from app.services.montage_service import _tracks_meta
    meta = _tracks_meta(tracks)
    return {tid for tid, m in meta.items()
            if m["kind"] == "audio" and m["bus"] == "dialogue" and not m["loop"]}


def _subs_carrier(clips: list, src, dial: set[str]) -> dict | None:
    """Le clip du payload qui PORTE la source explicite `src` — c'est lui
    qui décale les répliques. Préférence : piste de dialogue, puis v1, puis
    le reste ; à rang égal, le plus TÔT (`start`). Le geste PAR PLAN du tiroir
    envoie déjà les seuls clips qui chevauchent le plan (mesuré, bundle
    `function transcribe(plan){` : filtre `subsN(c.end,0)>plan.start+.05…`,
    aucun décalage côté client — le décalage vit ICI, une seule fois).
    `None` quand aucun clip ne la porte : décalage 0, comme avant P13."""
    key = _subs_src_key(src)
    cands = [c for c in clips
             if c.get("src") is not None and _subs_src_key(c.get("src")) == key]
    if not cands:
        return None

    def _rang(c):
        tr = str(c.get("tr") or "")
        return (0 if tr in dial else 1 if tr == "v1" else 2, _subs_bornes(c)[0])
    return min(cands, key=_rang)


async def _subs_dialogue_sources(clips: list, dial: set[str]) -> list:
    """Sans `src` explicite : TOUS les clips des pistes de dialogue porteurs
    d'une source résoluble, triés par `start` — chacun sera transcrit et
    décalé. Sans aucun : la PREMIÈRE v1 (au plus tôt) qui résout, seule,
    décalée de même. Rend une liste de `(chemin, clip)`."""
    out = []
    for c in sorted((c for c in clips if str(c.get("tr") or "") in dial and c.get("src")),
                    key=lambda c: _subs_bornes(c)[0]):
        p = await _subs_resolve(c.get("src"))
        if p is not None:
            out.append((p, c))
    if out:
        return out
    for c in sorted((c for c in clips if c.get("tr") == "v1" and c.get("src")),
                    key=lambda c: _subs_bornes(c)[0]):
        p = await _subs_resolve(c.get("src"))
        if p is not None:
            return [(p, c)]
    return []


def _subs_shift_words(words: list, clip: dict | None) -> list:
    """Mots au temps du FICHIER → temps de la TIMELINE : `+ (start − srcIn)`,
    COUPÉS à `[start, end]` (un mot qui déborde est rogné, un mot hors du
    clip est jeté). Chaque mot est marqué du clip qui le porte (`clip`) :
    `group_words` coupe une réplique à chaque changement de `clip`, donc deux
    sources ne fusionnent jamais dans une même réplique. Sans clip porteur :
    copie telle quelle (décalage 0 — le comportement d'avant P13).
    DETTE DITE (revue du 06/09) : la vitesse C4 d'un clip V1 (`speed`,
    montage_service `_v1_speed`) n'est PAS appliquée — un plan V1 à ×2
    transcrit par `src` explicite SANS jumeau A1 verrait ses mots à
    `start + (t − srcIn)` et non `/speed`. Mesuré : le tiroir n'envoie pas
    `speed` (`subsSrcClips` écrit id/tr/src/name/srcIn/start/end, 0
    occurrence de `speed`), et le jumeau A1 — le porteur préféré — garde sa
    vitesse (l'audio d'un V1 n'entre jamais dans le graphe du rendu)."""
    if clip is None:
        return [dict(w) for w in words or []]
    start, end, src_in = _subs_bornes(clip)
    off = start - src_in
    out = []
    for w in words or []:
        try:
            s = float(w["start"]) + off
            e = float(w["end"]) + off
        except (KeyError, TypeError, ValueError):
            continue
        if e <= start or s >= end:
            continue
        w2 = dict(w)
        w2["start"] = round(max(s, start), 3)
        w2["end"] = round(min(e, end), 3)
        try:
            se = float(w.get("speech_end", w["end"])) + off
        except (TypeError, ValueError):
            se = e
        w2["speech_end"] = round(min(max(se, w2["start"]), end), 3)
        w2["clip"] = str(clip.get("id") or f"{clip.get('tr')}@{start}")
        out.append(w2)
    return out


@router.post("/subtitles/from-narration")
async def subtitles_from_narration(request: Request):
    """Sous-titres CRÉÉS depuis la narration — gratuit, hors ligne, exact.

    Le texte de la voix off est déjà connu (c'est nous qui l'avons écrit et
    fait dire) : le caler vaut mieux que le faire re-deviner par une
    transcription payante, qui écrit « Dipotus » là où le script dit
    « Deepotus ». Les silences RÉELS du fichier audio sont mesurés
    (`silencedetect`) et les mots répartis entre les travées de parole au
    prorata de leur poids syllabique.

    Body : `{clips:[…]}` (modèle client du Montage) ou `{text, start, end |
    src}` pour un bloc isolé ; `cps?` = caractères par réplique.
    Sans `clips`, la sauvegarde du Montage (`montage_saved.json`) fait foi.
    """
    from app.services import transcribe_service as T
    try:
        body = await request.json()
    except Exception:
        body = {}
    if not isinstance(body, dict):
        body = {}
    lang = str(body.get("lang") or "fr")[:5]
    cps = int(body.get("cps") or T.CHARS_PER_SUBTITLE_DEFAULT)

    text = str(body.get("text") or "").strip()
    if text:
        p = await _subs_resolve(body.get("src"))
        start = max(0.0, float(body.get("start") or 0))
        end = body.get("end")
        if p is not None and end is None:
            res = await asyncio.to_thread(T.align_to_audio, text, p,
                                          start=start, lang=lang)
        else:
            if end is None:
                raise HTTPException(400, "Sans audio résoluble, `end` (ou "
                                         "`duration_s`) est nécessaire.")
            res = await asyncio.to_thread(
                T.align_known_text, text, start=start, end=float(end), lang=lang)
        cues = T.group_words(res["words"], max_chars=cps)
        # `aligned` dans LES DEUX branches : un contrat asymetrique se
        # decouvre au pire moment. Ici les mots ne portent pas de `clip` —
        # il n'y a pas de clip, c'est un bloc isole ; l'appelant qui veut
        # repartir un texte doit poster des `clips`.
        return {"ok": True, "source": "align", "words": len(res["words"]),
                "aligned": res["words"],
                "segments": _subs_cues_to_segments(cues)}

    clips = body.get("clips")
    if not isinstance(clips, list):
        from app.services.montage_service import _load_saved
        saved = await asyncio.to_thread(_load_saved)
        clips = (saved or {}).get("clips") or []
    clips = [dict(c) for c in clips if isinstance(c, dict)]
    narr = [c for c in clips
            if c.get("tr") in ("a1", "a3") and str(c.get("text") or "").strip()]
    if not narr:
        raise HTTPException(
            400, "Aucun clip de narration porteur de texte sur A1/A3 — écrivez "
                 "la voix off dans le tiroir Narration, ou lancez une "
                 "transcription (POST /api/subtitles/transcribe).")
    for c in narr:
        c["_path"] = await _subs_resolve(c.get("src"))
    res = await asyncio.to_thread(T.align_narration_clips, narr,
                                  lambda c: c.get("_path"), lang=lang)
    cues = T.group_words(res["words"], max_chars=cps)
    # `aligned` : les mots PLATS, chacun sachant DE QUEL CLIP il vient. Les
    # `segments` ne le disent pas (normalize_segments ne garde que w/start/end),
    # et sans cette information le Montage ne peut pas répartir le texte d'un
    # clip de narration FENDU entre ses deux moitiés — les deux garderaient la
    # phrase entière, et un recalage la compterait deux fois. Additif : la clé
    # `words` reste le COMPTE qu'elle a toujours été.
    return {"ok": True, "source": "align", "words": len(res["words"]),
            "blocks": res["blocks"], "clips": len(narr),
            "aligned": res["words"],
            "segments": _subs_cues_to_segments(cues)}


@router.get("/subtitles/estimate")
async def subtitles_estimate(duration_s: float = 0.0, provider: str = "",
                             src: str = ""):
    """COÛT et DURÉE d'une transcription, AVANT de la lancer (convention de
    l'app). `ok:false` + `reason` si aucune clé n'est configurée — l'UI
    propose alors le chemin gratuit (calage de la narration)."""
    from app.services import transcribe_service as T
    dur = float(duration_s or 0)
    if src and dur <= 0:
        p = await _subs_resolve(src)
        if p is not None:
            dur = await asyncio.to_thread(T.probe_duration, p)
    try:
        return T.estimate_transcription(dur, provider or None)
    except ValueError as e:
        raise HTTPException(400, str(e))


def _subs_job_set(jid: str, **kw) -> None:
    j = _SUBS_JOBS.setdefault(jid, {"status": "pending", "pct": 0, "step": ""})
    j.update(kw)


@router.post("/subtitles/transcribe")
async def subtitles_transcribe(request: Request, background_tasks: BackgroundTasks):
    """Sous-titres automatiques → `{job_id}`, puis GET /api/subtitles/jobs/{id}.

    DEUX chemins, et le gratuit passe d'abord :

    1. **Calage** (`align`) — des clips de narration portent leur texte : on
       le cale sur les silences réels. Gratuit, hors ligne, orthographe exacte.
    2. **Transcription** (`stt`) — texte inconnu : appel payant ElevenLabs
       Scribe ou OpenAI Whisper, horodatage AU MOT.

    Body : `{src?, clips?, tracks?, lang?, cps?, provider?, mode?}` —
    `mode:"stt"` force la transcription même si un texte est disponible.

    P13 (06/09/2026) — LA CIBLE, LE DÉCALAGE, LA LANGUE :
    · sans `src`, en chemin STT, les SOURCES sont tous les clips des pistes
      de dialogue (`tracks` du payload, même loi que le rendu, sinon `a1`)
      porteurs d'une source résoluble, triés par `start` ; chacun est
      transcrit, ses mots décalés de `start − srcIn` et coupés à
      `[start, end]`, les répliques concaténées ; `step` nomme chaque fichier
      (« transcription de kapwing_sample.mp4 (1/2) »), `usd` cumule,
      `provider` est celui du dernier, `sources` liste les fichiers. UN MÊME
      FICHIER porté par plusieurs clips (la lame coupe un clip en deux de
      même `src`) part UNE fois : `transcribe` envoie et facture le fichier
      ENTIER à chaque appel (`probe_duration`, mesuré) — un cache par chemin
      résolu, puis un décalage par clip porteur. Sans aucune : la première
      v1 qui résout, décalée de même ;
    · avec `src` (le geste PAR PLAN) : UNE source, décalée par le clip du
      payload qui la porte (dialogue avant v1, au plus tôt) — sans porteur,
      décalage 0 comme avant. Le client ne décale pas (mesuré) ;
    · `lang` vide ou « auto » ⇒ `None` chez le moteur (sa détection, la
      branche `if language:` de `transcribe` redevient vivante) ; le calage
      gratuit reçoit « fr » dans ce cas ;
    · un `provider` inconnu ⇒ 400 (la `ValueError` de `resolve_provider`
      sortait en 500, mesuré).
    CE QUI N'EST PAS FAIT ICI : quand une narration écrite existe (`text` sur
    a1/a3), le calage gratuit l'emporte et la STT ne tourne pas
    (`use_align`) — inchangé.
    """
    from app.services import transcribe_service as T
    try:
        body = await request.json()
    except Exception:
        body = {}
    if not isinstance(body, dict):
        body = {}
    # P13 — « auto » (ou rien) : le MOTEUR détecte ; le calage gratuit,
    # lui, a besoin d'un code (« fr » — `syllables` n'en fait rien de plus).
    lang_raw = str(body.get("lang") or "").strip().lower()[:5]
    lang_stt = None if lang_raw in ("", "auto") else lang_raw
    lang = lang_stt or "fr"
    cps = int(body.get("cps") or T.CHARS_PER_SUBTITLE_DEFAULT)
    mode = str(body.get("mode") or "auto").lower()
    provider = str(body.get("provider") or "") or None

    clips = body.get("clips")
    if not isinstance(clips, list):
        from app.services.montage_service import _load_saved
        saved = await asyncio.to_thread(_load_saved)
        clips = (saved or {}).get("clips") or []
    clips = [dict(c) for c in clips if isinstance(c, dict)]
    narr = [c for c in clips
            if c.get("tr") in ("a1", "a3") and str(c.get("text") or "").strip()]

    use_align = bool(narr) and mode != "stt"
    # P13 — `(chemin, clip porteur)` : le clip décale, le chemin se transcrit.
    sources: list = []
    if not use_align:
        dial = _subs_dialogue_ids(body.get("tracks"))
        src_path = await _subs_resolve(body.get("src"))
        if src_path is not None:
            sources = [(src_path, _subs_carrier(clips, body.get("src"), dial))]
        else:
            sources = await _subs_dialogue_sources(clips, dial)
        if not sources:
            raise HTTPException(
                400, "Aucun média à transcrire : posez un plan ou une voix off "
                     "sur la timeline, ou écrivez la narration (le calage d'un "
                     "texte connu est gratuit et plus exact).")
        try:
            pid = T.resolve_provider(provider)
        except ValueError as e:
            raise HTTPException(400, str(e))
        if pid is None:
            raise HTTPException(
                400, "Aucune clé de transcription configurée (Réglages : "
                     "ElevenLabs ou OpenAI). Le calage de la narration reste "
                     "disponible, gratuit et hors ligne.")

    for c in narr:
        c["_path"] = await _subs_resolve(c.get("src"))

    # P13 — les FICHIERS distincts, dans l'ordre de leur premier clip : un
    # même chemin porté par deux clips n'est transcrit (et facturé) qu'une
    # fois ; `noms` et `sources` le nomment une fois.
    fichiers: list = []
    for p, _c in sources:
        if p not in fichiers:
            fichiers.append(p)
    noms = [p.name for p in fichiers]
    jid = uuid4().hex[:12]
    if len(_SUBS_JOBS) > _SUBS_JOBS_MAX:
        for k in list(_SUBS_JOBS)[:len(_SUBS_JOBS) - _SUBS_JOBS_MAX]:
            _SUBS_JOBS.pop(k, None)
    _subs_job_set(jid, status="pending", pct=5,
                  step="calage de la narration" if use_align
                       else "transcription", segments=None, error=None,
                  source="align" if use_align else "stt")

    async def _run():
        try:
            _subs_job_set(jid, status="running", pct=20)
            if use_align:
                res = await asyncio.to_thread(
                    T.align_narration_clips, narr, lambda c: c.get("_path"),
                    lang=lang)
                _subs_job_set(jid, pct=75, step="découpe en répliques")
                cues = T.group_words(res["words"], max_chars=cps)
                src_kind = "align"
                extra = {"words": len(res["words"]), "blocks": res["blocks"]}
            else:
                words, usd, prov, texts = [], 0.0, None, []
                # UN appel par FICHIER distinct (`fichiers`), puis UN décalage
                # par clip porteur (`sources`) : deux clips de même source —
                # la lame — ne paient qu'une transcription.
                n = len(fichiers)
                cache: dict = {}
                for k, p in enumerate(fichiers):
                    _subs_job_set(jid, pct=25 + int(50 * k / n),
                                  step=f"transcription de {p.name} ({k + 1}/{n})")
                    res = await asyncio.to_thread(T.transcribe, p,
                                                  provider=provider,
                                                  language=lang_stt)
                    cache[p] = res.get("words") or []
                    usd += float(res.get("usd_estimated") or 0.0)
                    # `transcribe` nomme le fournisseur `source` (pas `provider`)
                    prov = res.get("source") or prov
                    if res.get("text"):
                        texts.append(str(res["text"]))
                for p, c in sources:
                    words.extend(_subs_shift_words(cache[p], c))
                if not words:
                    raise RuntimeError(
                        "Le moteur n'a rendu aucun mot horodaté dans la "
                        "fenêtre des clips (" + ", ".join(noms) + ") — média "
                        "muet, format refusé, ou clip hors du fichier.")
                for i, w in enumerate(words):
                    w["i"] = i
                _subs_job_set(jid, pct=75, step="découpe en répliques")
                cues = T.group_words(words, max_chars=cps)
                src_kind = "stt"
                extra = {"words": len(words), "usd": round(usd, 4),
                         "provider": prov, "text": "\n".join(texts),
                         "sources": noms}
            segs = _subs_cues_to_segments(cues)
            _subs_job_set(jid, status="done", pct=100, step="terminé",
                          segments=segs, source=src_kind, **extra)
        except Exception as e:                          # noqa: BLE001
            logger.exception(f"subtitles job {jid} failed: {e}")
            _subs_job_set(jid, status="failed", pct=100, step="échec",
                          error=str(e))

    background_tasks.add_task(_run)
    return {"ok": True, "job_id": jid,
            "source": "align" if use_align else "stt",
            "sources": noms,
            "message": ("Calage de la narration lancé (gratuit)."
                        if use_align else
                        "Transcription lancée sur " + ", ".join(noms) + ".")}


@router.get("/subtitles/jobs/{jid}")
async def subtitles_job(jid: str):
    """État d'un travail de sous-titrage : `{status, pct, step, segments?,
    error?}` — `status` ∈ pending | running | done | failed."""
    j = _SUBS_JOBS.get(str(jid))
    if j is None:
        raise HTTPException(404, "Travail de sous-titrage inconnu (ou expiré "
                                 "avec le redémarrage du backend).")
    return {"ok": True, "job_id": jid, **j}


# ── Impression 3D : exports vers le slicer (plan 2026-08-27) ─────────────────
# Service print3d : 100 % local, python pur (lecteur GLB minimal, STL/3MF
# stdlib). Un dossier par export sous assets/print3d ; « Ouvrir dans le
# slicer » = association Windows du .3mf, repli SLICER_PATH du .env.

def _print3d_base() -> Path:
    base = settings.outputs_path.parent / "print3d"
    base.mkdir(parents=True, exist_ok=True)
    return base


@router.post("/print3d/from-assets3d/{job}")
async def print3d_from_assets3d(job: str, body: dict):
    """Body: {cible_mm?, nom?} — convertit le maillage d'un job Game
    Assets 3D en dossier d'impression (STL + 3MF aux mm). Préfère le
    `model.stl` du moteur quand il existe (zéro conversion = zéro risque),
    sinon lit `model.glb` ; `model.opt.glb` SEUL → 409 parlant (meshopt)."""
    from app.services import print3d as P3
    d = settings.outputs_path / "assets3d" / Path(job).name
    if not d.is_dir():
        raise HTTPException(404, f"Job 3D introuvable: {job}")
    cible = body.get("cible_mm")
    cible = float(cible) if cible not in (None, "") else None
    nom = str(body.get("nom") or Path(job).name).strip()[:80]
    try:
        if (d / "model.stl").is_file():
            tris = P3.lire_stl((d / "model.stl").read_bytes())
        elif (d / "model.glb").is_file():
            tris = P3.lire_glb_triangles((d / "model.glb").read_bytes())
        elif (d / "model.opt.glb").is_file():
            raise HTTPException(409,
                "seul model.opt.glb existe — il est compressé meshopt (pour "
                "les moteurs de jeu) ; il faut model.glb, la source non "
                "compressée")
        else:
            raise HTTPException(404,
                "aucun maillage lisible (model.stl / model.glb) dans ce job")
        export = await asyncio.to_thread(
            P3.creer_export, _print3d_base(), nom, tris, cible,
            f"assets3d:{Path(job).name}", "inconnue")
    except ValueError as e:
        raise HTTPException(409, str(e))
    return export


@router.post("/print3d/from-stl")
async def print3d_from_stl(request: Request, nom: str = "objet",
                           cible_mm: float | None = None,
                           source: str = "stl", etanche: str = "inconnue"):
    """Corps binaire = STL BINAIRE (la voie de la Forge 3D cartes et du
    Vectorlab). `cible_mm` absent = « tel quel » (les producteurs mm) ;
    `etanche=garantie` seulement quand le producteur le PROUVE (gate
    forge3d)."""
    from app.services import print3d as P3
    octets = await request.body()
    try:
        tris = P3.lire_stl(octets)
    except ValueError as e:
        raise HTTPException(400, str(e))
    export = await asyncio.to_thread(
        P3.creer_export, _print3d_base(), str(nom)[:80], tris, cible_mm,
        str(source)[:40],
        "garantie" if etanche == "garantie" else "inconnue")
    return export


@router.get("/print3d/exports")
async def print3d_exports():
    from app.services import print3d as P3
    return {"exports": P3.lister_exports(_print3d_base())}


@router.post("/print3d/open")
async def print3d_open(body: dict):
    """Ouvre le `.3mf` d'un export dans le slicer. Le nom de dossier est
    CONTENU dans assets/print3d — jamais de startfile arbitraire."""
    from app.services import print3d as P3
    nom = str(body.get("dossier") or "")
    safe = Path(nom).name
    if not safe or safe != nom or safe in (".", ".."):
        raise HTTPException(400, "nom de dossier invalide")
    dossier = _print3d_base() / safe
    if not dossier.is_dir():
        raise HTTPException(404, "export introuvable")
    mf3 = sorted(dossier.glob("*.3mf"))
    if not mf3:
        raise HTTPException(404, "aucun .3mf dans cet export")
    try:
        mode = P3.ouvrir_dans_slicer(mf3[0])
    except RuntimeError as e:
        raise HTTPException(409, str(e))
    return {"ok": True, "mode": mode, "fichier": mf3[0].name}


# ── Bibliothèque unifiée (28/08) : import Figma → Bibliothèque ──────────────

@router.post("/images/import-figma")
async def import_figma(body: dict):
    """Body: {url} — le lien d'un CALQUE Figma (node-id présent) devient un
    PNG de la Bibliothèque (`figma_<clé>_<node>.png`, réécrit en place au
    ré-import). Jeton = FIGMA_TOKEN du .env des données ; absent → 409
    parlant. Le service ne sort que vers api.figma.com."""
    from app.services import figma_import as FI
    jeton = str(getattr(settings, "FIGMA_TOKEN", "") or "").strip()
    if not jeton:
        raise HTTPException(409,
            "FIGMA_TOKEN absent — crée un Personal Access Token Figma "
            "(figma.com → Settings → Security), pose FIGMA_TOKEN=... dans "
            "le .env des données (DeepotusVideoGenData\.env) puis relance "
            "l'application.")
    try:
        nom = await FI.importer(str(body.get("url") or ""), jeton,
                                settings.images_path)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except RuntimeError as e:
        raise HTTPException(502, str(e))
    await LI.noter([nom], "figma")
    return {"filename": nom}


# ═══ Finition — spec Magnific §13 phase D, les deux derniers points ═════════
#
# « Tester upscale seulement après verrouillage du plan ; mesurer gain visuel
#   vs coût et dérive. »
# « Générer des exports de montage avec audio séparé, puis comparer au son
#   natif du clip. »

def _finition_dir() -> Path:
    d = settings.outputs_path / "finition"
    d.mkdir(parents=True, exist_ok=True)
    return d


async def _resoudre_video(body: dict) -> Path:
    """`job_id` (rendu fini) ou `filename` (outputs/videos) → chemin réel."""
    from app.services.storage import JobRecord, async_session_factory
    jid = str(body.get("job_id") or "").strip()
    if jid:
        async with async_session_factory() as s:
            jr = await s.get(JobRecord, jid)
        if jr is None:
            raise HTTPException(404, f"Job inconnu : {jid}")
        p = jr.final_video_path or jr.video_path
        if not p or not Path(p).is_file():
            raise HTTPException(400, f"Le job {jid} n'a pas de vidéo sur le "
                                     f"disque (statut « {jr.status} »).")
        return Path(p)
    fn = Path(str(body.get("filename") or "")).name
    if not fn:
        raise HTTPException(400, "Passe job_id (rendu) ou filename (vidéo).")
    p = settings.outputs_path / "videos" / fn
    if not p.is_file():
        raise HTTPException(404, f"Vidéo introuvable : {fn}")
    return p


@router.post("/finition/stems")
async def finition_stems(body: dict):
    """Un montage → vidéo MUETTE + stem audio WAV, et la comparaison au son
    natif d'un rush.

    Body: {job_id? | filename?, natif?}. `natif` = nom d'un rush
    (outputs/videos) ou d'un fichier audio : sa loudness est mesurée avec le
    MÊME filtre ebur128 que POST /api/montage/measure, donc les chiffres se
    comparent.

    Honnêteté déclarée dans la réponse (`fidelite`) : le WAV est le DÉCODAGE
    de la piste livrée (AAC du mp4) — c'est ce que le spectateur entend, la
    bonne référence pour cette comparaison, mais PAS un master pré-encodage.
    Un vrai master exigerait de rejouer le graphe audio du montage.
    Local et gratuit : 0 $."""
    from app.services import finition as F
    src = await _resoudre_video(body)
    loop = asyncio.get_running_loop()
    try:
        out = await loop.run_in_executor(
            None, lambda: F.separer(src, _finition_dir()))
    except RuntimeError as e:
        raise HTTPException(502, str(e))

    natif = Path(str(body.get("natif") or "")).name
    if natif:
        cand = [settings.outputs_path / "videos" / natif,
                settings.images_path.parent / "audio" / natif,
                settings.outputs_path / "audio" / natif,
                settings.images_path / natif]
        pn = next((c for c in cand if c.is_file()), None)
        if pn is None:
            raise HTTPException(404, f"Source native introuvable : {natif}")
        try:
            out["comparaison"] = await loop.run_in_executor(
                None, lambda: F.comparer_audio(src, pn))
        except RuntimeError as e:
            raise HTTPException(502, str(e))
    return {"ok": True, "source": src.name, "dir": "outputs/finition", **out}


@router.post("/finition/upscale-measure")
async def finition_upscale_measure(body: dict):
    """Agrandit et MESURE : gain de netteté, dérive au retour, coût.

    Body: {filename, scale?=2, ai?=false, confirm?=false, shot_id?}.

    Porte avant dépense (doctrine §2.3) : la variante LOCALE (PIL Lanczos,
    0 $) tourne toujours ; la variante payante (fal-ai/esrgan) exige
    `ai: true` ET `confirm: true` — sans confirmation, la route rend la
    mesure gratuite et dit ce que coûterait l'autre. `shot_id`, s'il est
    fourni, doit désigner un plan existant : le verrou par plan
    (shot.status = keyframe_ok) arrivera avec le lot « image-clé » du plan
    d'ensemble et se branchera ici.

    Aucun gagnant n'est déclaré : un upscale n'a pas de vérité terrain. On
    rend deux mesures orthogonales — le détail produit et la fidélité à la
    source — et l'humain tranche."""
    from app.services import finition as F
    fn = Path(str(body.get("filename") or "")).name
    if not fn:
        raise HTTPException(400, "filename requis.")
    src = settings.images_path / fn
    if not src.is_file():
        raise HTTPException(404, f"Image introuvable : {fn}")
    try:
        scale = max(2, min(4, int(body.get("scale") or 2)))
    except (TypeError, ValueError):
        raise HTTPException(400, "scale doit être un entier (2-4).")

    sid = str(body.get("shot_id") or "").strip()
    if sid:
        from app.services.storage import Shot, async_session_factory
        async with async_session_factory() as s:
            if await s.get(Shot, sid) is None:
                raise HTTPException(404, f"Plan inconnu : {sid}")

    loop = asyncio.get_running_loop()
    variantes = []

    r = await _process_image_core({"op": "upscale", "filename": fn,
                                   "scale": scale, "mode": "simple"})
    nom = (r.get("images") or [None])[0]
    if not nom:
        raise HTTPException(502, "L'agrandissement local n'a rien produit.")
    try:
        await LI.noter([nom], "retouche")
    except Exception:
        pass
    variantes.append(await loop.run_in_executor(
        None, lambda: F.mesurer_variante(src, settings.images_path / nom,
                                         mode="simple", scale=scale, usd=0.0)))

    veut_ai = bool(body.get("ai"))
    if veut_ai and not body.get("confirm"):
        return {"ok": True, "source": F.fiche_image(src), "scale": scale,
                "variantes": variantes,
                "en_attente": {
                    "mode": "ai", "provider": "fal-ai/esrgan", "usd": None,
                    "usd_note": "tarif fal-ai/esrgan non renseigné dans "
                                "pricing.json — la clé reste vide tant que le "
                                "catalogue fal n'a pas été relu (même règle "
                                "que le 1080p de Seedance 2.5).",
                    "message": "Renvoie la requête avec confirm:true pour "
                               "lancer l'agrandissement payant."}}
    if veut_ai:
        if not settings.FAL_KEY:
            raise HTTPException(400, "FAL_KEY not configured (Settings) — "
                                     "la variante locale est déjà mesurée.")
        r2 = await _process_image_core({"op": "upscale", "filename": fn,
                                        "scale": scale, "mode": "ai"})
        nom2 = (r2.get("images") or [None])[0]
        if not nom2:
            raise HTTPException(502, "esrgan n'a rien produit.")
        try:
            await LI.noter([nom2], "retouche")
        except Exception:
            pass
        variantes.append(await loop.run_in_executor(
            None, lambda: F.mesurer_variante(
                src, settings.images_path / nom2, mode="ai", scale=scale,
                usd=None)))

    return {"ok": True, "source": F.fiche_image(src), "scale": scale,
            "variantes": variantes,
            "lecture": "Compare `nettete` ENTRE variantes (même taille cible, "
                       "même source) : c'est elle qui départage. `derive` dit "
                       "l'écart à la source au retour. Forte netteté + forte "
                       "dérive = détail INVENTÉ ; forte netteté + faible "
                       "dérive = détail restitué. `gain_nettete` compare à une "
                       "image de taille différente : indicatif seulement, et "
                       "négatif pour un Lanczos, qui lisse sans rien inventer."}


# ── l'Établi : inspection et chirurgie de maillage ───────────────────────────
# Spec docs/superpowers/specs/2026-08-29-etabli-inspecteur-3d-design.md.
# Le navigateur envoie des PARAMÈTRES ; l'écriture du GLB vit dans mesh_edit.

def _etabli_glb(job: str, version) -> bytes:
    """Les octets d'une version d'un job, ou un 404 parlant.

    LE PORTEUR DE LECTURE — `GET /etabli/rig`, dont FastAPI type déjà `job` et
    `version`. Les cinq routes d'ÉCRITURE passent par `_etabli_glb_cible`, plus
    bas, qui juge un corps JSON venu du réseau (entier ≥ 1, chaîne, deux gardes
    de chemin) et rend `depuis`. Deux porteurs, donc, et c'est dit ici."""
    from app.services import mesh_report
    d = mesh_report.job_dir(Path(str(job)).name)
    v = int(version or 1)
    nom = "model.glb" if v <= 1 else f"model.v{v}.glb"
    p = d / nom
    if not p.is_file():
        raise HTTPException(404, f"{Path(str(job)).name}/{nom} introuvable")
    return p.read_bytes()


def _etabli_ecrire(job: str, sortie: bytes, operation: str, detail: dict):
    """Dépose la sortie comme nouvelle version. Les ValueError du socle
    deviennent des 400 : trois tâches les ont durcies pour cette marche-ci."""
    from app.services import mesh_edit
    return mesh_edit.ecrire_version(job, sortie, operation=operation,
                                    detail=detail)


@router.get("/etabli/sources")
async def etabli_sources(limit: int = 60):
    """La chronologie unifiée : jobs assets3d et tâches Meshy rapatriées.

    `mesh_sources.lister()` fait de l'E/S disque SYNCHRONE — sans
    `asyncio.to_thread`, elle gèlerait la boucle d'événements pendant tout le
    parcours des dossiers, donc toutes les requêtes du serveur.
    """
    from app.services import mesh_sources
    jobs = await asyncio.to_thread(mesh_sources.lister)
    return {"jobs": jobs, "meshy": await mesh_sources.lister_meshy(limit)}


# ── la catégorie « Établi » de la Bibliothèque ───────────────────────────────
# « Rajouter les dossiers générés dans une catégorie spécifique de la
# librairie pour pouvoir facilement la retrouver. » Moitié SERVEUR de cette
# demande ; l'onglet de la Bibliothèque est écrit ailleurs et ne doit rien
# avoir à réinventer — d'où la forme de sortie, ci-dessous.


def _etabli_vignette(d: Path, job: str, v: int) -> str | None:
    """L'image à montrer pour une version, ou `None` s'il n'y en a aucune.

    ORDRE : vignette du canevas → `preview.png` → `shot_0.png` → silhouette →
    `None`.

    LA VIGNETTE DU CANEVAS D'ABORD, PARCE QU'ELLE MONTRE CETTE VERSION-CI.
    C'est la seule image de la liste qui soit une photo du maillage LISTÉ :
    l'Établi la capture au moment où il écrit la version (voir
    `etabli_vignette_ecrire`). Elle a été ajoutée parce que la correction
    d'ordre ci-dessous, juste, était impuissante — sur les trois jobs mesurés
    de l'utilisateur, un seul portait un vrai rendu de moteur ; un autre avait
    un `preview.png` ET un `shot_0.png` qui sont le MÊME aplat ambré (14
    couleurs distinctes) ; le troisième n'avait que des masques. Aucun ordre
    de préférence n'invente une image absente du disque.

    ELLE NAÎT À L'ÉCRITURE SEULEMENT, décision de l'utilisateur : ni
    rattrapage à l'ouverture, ni régénération par lots. Les productions
    antérieures retombent donc sur la suite, inchangée.

    LA SUITE — `preview.png` → `shot_0.png` → silhouette — est celle d'hier,
    et elle ne bouge pas d'un cran. Elle avait déjà été RETOURNÉE : la
    silhouette passait d'abord, au motif qu'elle est celle de
    la VERSION demandée là où le rendu ne montre que le brouillon —
    raisonnement juste, résultat faux, et c'est lui qui a produit le
    défaut : « les nouvelles versions apparaissent bien dans la librairie,
    mais les illustrations ne se montrent pas », six vignettes blanches
    sur huit.

    LE RENDU DU MOTEUR AVANT LE MASQUE, PARCE QU'IL ILLUSTRE.
    `sil_v<n>/silhouette_face.png` n'est pas une image de l'objet mais un
    MASQUE de contrôle : `mesh_report.silhouettes()` l'écrit pour la
    comparaison de silhouette du QC — une forme blanche pleine sur fond
    noir. Mesuré sur les données de l'utilisateur : 60 % de pixels clairs
    pour la silhouette, 0 % pour `preview.png`. Elle s'affichait
    parfaitement ; elle ne montrait rien.

    LE PRIX, ASSUMÉ. `preview.png` et `shot_0.png` sont déposés par le
    moteur au premier tir : ils montrent la VERSION 1, pas la version
    listée. C'est acceptable parce que la carte porte déjà le numéro en
    toutes lettres — « 6e0a8a5f · v5 · transformer » — et qu'une image du
    bon objet renseigne mieux qu'un masque de la bonne version.

    LA SILHOUETTE RESTE, EN QUEUE. Quand elle sort, c'est qu'il n'y a NI
    preview NI shot dans le dossier : elle est alors la seule image qui
    existe, et un masque vaut mieux qu'une case vide. Pis-aller, pas
    choix — l'état du job `6e0a8a5f`, quatre silhouettes et rien d'autre.
    L'absence totale est DITE (`None`) plutôt que servie en lien mort.

    POUR L'ONGLET : `null` doit être traité EXPLICITEMENT. La carte 3D du
    bundle porte un repli `onError` vers `shot/0`, mais React omet
    l'attribut `src` quand la valeur est `null` — aucune requête, donc aucun
    évènement `error`, donc le repli ne joue jamais.
    """
    if _etabli_vignette_path(job, v).is_file():
        return f"/api/assets/3d/{job}/vignette?v={int(v)}"
    if (d / "preview.png").is_file():
        return f"/api/assets/3d/{job}/preview"
    if (d / "shot_0.png").is_file():
        return f"/api/assets/3d/{job}/shot/0"
    if (d / f"sil_v{int(v)}" / "silhouette_face.png").is_file():
        return f"/api/assets/3d/{job}/silhouette/face?v={int(v)}"
    return None


def _etabli_entree(ligne: dict, etape: dict, operation: str, d: Path) -> dict:
    """Une production, à la FORME de la carte 3D de la Bibliothèque.

    Le bundle construit son onglet 3D ainsi — `{name, kind, size, date,
    provider, jobId, short, url, thumb}` — et la carte qui les affiche existe
    déjà. Épouser cette forme, c'est offrir l'onglet « Établi » sans écrire
    une carte de plus :

      * `kind: "asset3d"` donne la carte 3D — la vignette, le lien, la
        silhouette. Voir `imprimable` pour ce qu'il NE donne pas ;
      * `short` porte le NOM DE DOSSIER ENTIER, jamais un préfixe de huit.
        Le bundle coupe `job_id.slice(0, 8)` parce qu'un job `assets3d`
        normal a pour dossier les 8 premiers caractères de son UUID ; un job
        adopté s'appelle `meshy_<task_id>` et les routes `/api/assets/3d/…`
        prennent le segment LITTÉRALEMENT (`Path(job).name`, aucune
        résolution de préfixe). Couper ici donnerait des vignettes et des
        téléchargements morts ;
      * `url` est reprise TELLE QUELLE de `mesh_sources`, seule à savoir
        composer le lien d'une version — la recomposer ici dupliquerait sa
        convention de nommage ;
      * `size` reste vide comme dans le bundle (qui a `go(bytes)` pour ça) ;
        `date` porte l'ISO BRUT, là où le bundle envoie à sa carte une date
        déjà passée par `mo()` — l'onglet devra appliquer `mo(created_at)`.

    Le reste est ce que la carte 3D n'a pas : de quoi retrouver le maillage
    et distinguer une version ÉCRITE d'une ADOPTION, sans repasser par une
    autre route. `job` y DOUBLE `jobId` à dessein : la carte du bundle parle
    camelCase, ces champs-ci parlent la langue des routes `/etabli/*` (`job`,
    `version`, `operation`), et les mélanger dans un même bloc se relit mal.
    Alias de lecture, donc, assumé comme tel.

    `imprimable` DÉMENT une capacité que `kind: "asset3d"` semble offrir. Le
    menu « Envoyer vers » propose « Impression 3D » à toute carte `asset3d`
    portant un `short`, et appelle `POST /api/print3d/from-assets3d/<short>`
    — or cette route lit `model.stl` sinon `model.glb`, JAMAIS
    `model.v<n>.glb`, et n'accepte aucun numéro de version. Une entrée
    « v2 · reparer » y ferait donc imprimer le BROUILLON NON RÉPARÉ, en
    silence. L'onglet doit masquer l'entrée d'impression quand ce champ est
    faux ; le rendre vrai pour toutes demanderait d'abord d'apprendre une
    version à `print3d_from_assets3d`, ce qui n'est pas de cette tâche-ci.
    """
    job = ligne["id"]
    v = etape["version"]
    return {
        "name": f"{ligne['nom']} · v{v} · {operation}",
        "kind": "asset3d",
        "size": "",
        "date": etape["created_at"] or "",
        "provider": "Établi",
        "jobId": job,
        "short": job,
        "url": etape["url"],
        "thumb": _etabli_vignette(d, job, v),
        "job": job,
        "version": v,
        "operation": operation,
        "origine": "adoption" if operation == "adoption" else "version",
        "imprimable": v == 1,
        "fichier": etape["file"],
        "bytes": etape["bytes"],
        "sha256": etape["sha256"],
        "triangles": etape["triangles"],
        "created_at": etape["created_at"],
        "moteur": ligne["moteur"],
    }


def _etabli_du_job(ligne: dict) -> list[dict]:
    """Ce que l'Établi a produit dans UN job. Lève si le job est abîmé :
    c'est l'appelant qui tient le filet."""
    from app.services import mesh_report

    job = ligne["id"]
    d = mesh_report.job_dir(job)
    try:
        registre = mesh_report.read_registry(job)
    except FileNotFoundError:
        registre = {}          # aucune fiche : l'Établi n'a rien écrit ici
    except ValueError as e:
        # ALIGNÉ SUR `mesh_sources._versions_du_job`, qui fait exactement ce
        # `except (FileNotFoundError, ValueError)` : un registre illisible
        # vaut SANS FICHE, pas dossier perdu. Sans cet alignement, le job
        # s'affichait dans la chronologie et dans /etabli/sources tout en
        # disparaissant du seul onglet « Établi » — et une adoption dont le
        # `asset.json` est intact était perdue avec lui. On le DIT quand même.
        logger.warning(f"etabli/productions: report.json de {job} illisible "
                       f"({e}) — lu comme sans fiche")
        registre = {}
    # ASYMÉTRIE DÉLIBÉRÉE : `[1, 2, 3]` est du JSON VALIDE — `read_registry`
    # ne lève pas, et le `.get` ci-dessous part en AttributeError. On ne le
    # rattrape pas ici : un flux d'octets tronqué est une écriture à
    # moitié faite, tandis qu'un registre d'un AUTRE TYPE dit que ce dossier
    # n'est plus ce qu'on croit. Le filet de l'appelant le traite comme abîmé.
    fiches = {str(f.get("file")): f
              for f in (registre.get("entries") or [])
              if isinstance(f, dict)}

    out: list[dict] = []
    vues: set[int] = set()
    for etape in ligne["etapes"]:
        v = etape["version"]
        if v is None:
            # « décimée » (model.opt.glb) : pas une production de l'Établi.
            # Ce qui l'exclut aujourd'hui est l'absence de fiche
            # `outil == "etabli"` sur ce fichier ; cette ligne-ci tient le
            # jour où une fiche y atterrirait — `int(None)` lèverait dans la
            # vignette, le filet du job rattraperait, et le job ENTIER
            # disparaîtrait de la catégorie, sa v2 légitime comprise.
            continue
        src = (fiches.get(etape["file"]) or {}).get("source")
        if not isinstance(src, dict) or src.get("outil") != "etabli":
            continue
        # `src` porte tout le détail de l'opération : pour « couper », le compte
        # rendu du couteau — nœuds produits (`_avant` / `_apres`), capuchons
        # posés ou non et pourquoi (format en tête de la section couteau de
        # mesh_cut.py) — et, pour toute écriture, `depuis` : la version dont
        # elle part. L'entrée n'en remonte que le nom ; qui voudra afficher les
        # capuchons ou la lignée le lira ici.
        out.append(_etabli_entree(
            ligne, etape, str(src.get("operation") or "?"), d))
        vues.add(v)

    if ligne["phase"] == "adopte" and 1 not in vues:
        # Repêchage : une adoption dont la fiche manque ou a perdu son
        # `source` (écriture interrompue — `adopter_meshy` garde ses trois
        # écritures séparées —, ou registre illisible lu comme vide
        # ci-dessus). Le `asset.json` suffit alors à la reconnaître.
        prem = next((e for e in ligne["etapes"] if e["version"] == 1), None)
        if prem is not None:
            out.append(_etabli_entree(ligne, prem, "adoption", d))
    return out


def _etabli_productions() -> list[dict]:
    """Les maillages que l'Établi a produits. SYNCHRONE : lecture de disque.

    Deux marqueurs, posés par P1, et il faut les DEUX :

      1. une version écrite — `mesh_edit.ecrire_version` fait poser à
         `mesh_report.write_report` un `source: {"outil": "etabli", …}` sur
         l'entrée de registre de cette version ;
      2. une tâche Meshy adoptée — `mesh_edit.adopter_meshy` écrit un
         `asset.json` `stage == "adopte"`.

    L'adoption pose en fait les deux (sa fiche porte `operation: "adoption"`),
    mais ses trois écritures sont gardées SÉPARÉES exprès : une adoption
    interrompue laisse l'un sans l'autre. On les additionne donc en
    dédoublonnant par version, sans quoi la même v1 sortirait deux fois.

    Le brouillon `model.glb` d'un job vient du MOTEUR : il n'entre pas ici,
    même quand le job contient par ailleurs une production.
    """
    from app.services import mesh_sources

    out: list[dict] = []
    for ligne in mesh_sources.lister():
        try:
            out.extend(_etabli_du_job(ligne))
        except Exception as e:      # noqa: BLE001 — un job abîmé n'éteint pas la liste
            # Même filet que `mesh_sources.lister()`, pour la même raison :
            # ce dossier est ouvert aux mains de l'utilisateur. Ce job-ci est
            # perdu, jamais la catégorie entière — et on le DIT.
            logger.warning(f"etabli/productions: job {ligne.get('id')} "
                           f"illisible ({e}) — ignoré")

    # `mesh_sources.lister()` trie par NOM de dossier — un préfixe d'UUID,
    # donc sans rapport avec le temps, et sa docstring demande à l'appelant de
    # retrier. Ce que la personne cherche, c'est son DERNIER dossier.
    # Comparer les ISO en TEXTE est sûr ici : ces dates ont une seule plume,
    # `mesh_report.report()`, qui écrit toujours
    # `datetime.now(timezone.utc).isoformat(timespec="seconds")` — même
    # fuseau `+00:00`, même largeur, donc ordre lexical = ordre du temps.
    out.sort(key=lambda e: (e["created_at"] or "", e["job"], e["version"]),
             reverse=True)
    return out


@router.get("/etabli/productions")
async def etabli_productions():
    """Ce que l'Établi a produit — la catégorie « Établi » de la Bibliothèque.

    `_etabli_productions` fait de l'E/S disque SYNCHRONE (le parcours des
    dossiers, plus un `report.json` par job) : sans `asyncio.to_thread`, elle
    gèlerait la boucle d'événements pendant tout le parcours, donc toutes les
    requêtes du serveur — pas seulement la sienne.
    """
    return {"items": await asyncio.to_thread(_etabli_productions)}


# ── la vignette du canevas : le navigateur voit, PYTHON écrit ────────────────
# 2 Mio. La page réduit à 512 px de plus grand côté avant d'envoyer — un PNG
# de cet ordre pèse quelques dizaines de kilo-octets — mais le serveur ne peut
# pas croire l'appelant sur parole : rien ne garantit que c'est notre page, et
# un canevas 2000×1500 non réduit ferait plusieurs mégaoctets. La borne laisse
# une marge large et refuse PARLANT au lieu d'avaler.
_ETABLI_VIGNETTE_MAX = 2 * 1024 * 1024


@router.post("/etabli/vignette")
async def etabli_vignette_ecrire(request: Request, job: str, version: int):
    """Dépose la vignette d'une version : le canevas de l'Établi, capturé.

    LA RÈGLE STRUCTURANTE TIENT, ET IL FAUT LE DIRE. « Le navigateur voit et
    manipule, PYTHON ÉCRIT » (spec §2.1) porte sur l'AUTORITÉ DU MAILLAGE :
    aucun GLB ne naît côté client, la page ne contient toujours pas l'ombre
    d'un GLTFExporter, et `test_la_page_ne_fabrique_jamais_un_glb` reste vert.
    Une vignette PNG n'est pas un maillage — c'est une PHOTO de ce que la page
    affiche déjà — et le fichier sur le disque est écrit ICI, en Python. Le
    navigateur voit ; Python écrit. Sans ce paragraphe, le prochain lecteur
    croira à une entorse.

    À L'ÉCRITURE SEULEMENT, décision de l'utilisateur : la page n'appelle
    cette route qu'après une écriture réussie, jamais à l'ouverture, jamais
    par lots. Rien ici ne l'impose — c'est le client qui la tient, et un banc
    qui compte ses sites d'appel — mais la route n'offre AUCUN moyen de
    rattraper une lignée entière, et c'est délibéré.

    LES GARDES, dans l'ordre où elles mordent. Une route d'écriture non gardée
    est une porte ouverte, et celle-ci reçoit des octets venus du navigateur :

      1. `version` doit être un entier — FastAPI le refuse en 422 sinon — et
         un entier POSITIF : `0` n'est le numéro de rien ;
      2. `job` passe par `_etabli_vignette_cible`, qui porte DEUX gardes de
         natures différentes — le refus des noms dégénérés (`..` survit à
         `Path(...).name`, mesuré) et le confinement du chemin RÉSOLU sous le
         dossier des jobs. Le détail, et la mesure, sont chez elle ;
      3. la VERSION DOIT EXISTER sur le disque. Une vignette sans maillage est
         un mensonge en attente : la carte de la Bibliothèque montrerait une
         image pour une version que le disque ignore. C'est aussi ce qui
         empêche de fabriquer un dossier de job à volonté. ATTENTION : ce
         n'est PAS une garde de traversée, et l'avoir crue telle a coûté un
         défaut — voir `_etabli_vignette_cible` ;
      4. la TAILLE est bornée AVANT tout examen du contenu — c'est la garde la
         moins chère, elle passe donc en premier des deux ;
      5. la SIGNATURE PNG est vérifiée. L'en-tête `Content-Type` ne prouve
         rien : seul le magic le fait. Même geste que
         `POST /vector/docs/{doc_id}/vignette`, qui a le même problème.

    FAIBLESSE CONNUE, ASSUMÉE, ET ÉCRITE ICI POUR QU'ON NE LA « CORRIGE » PAS
    DE TRAVERS. `await request.body()` bufferise le corps ENTIER avant que la
    borne de 2 Mio ne morde : un client qui ment peut donc faire allouer plus
    que la borne. Ce n'est pas corrigé, et c'est délibéré — `routes.py` compte
    cinq `await request.body()` et celui-ci est le SEUL à borner quoi que ce
    soit ; passer en `request.stream()` ferait de cette route la seule
    cérémonieuse au milieu de quatre voisines grandes ouvertes, sur une API
    sans authentification, atteignable seulement par un processus local (qui a
    des moyens plus directs) ou par une page du navigateur. Gain nul,
    incohérence réelle. Un pré-contrôle de `Content-Length` serait une
    COURTOISIE — un 413 immédiat pour un client honnête — jamais une garde,
    puisqu'un attaquant ment sur l'en-tête.

    ÉCRITURE ATOMIQUE (`.tmp` puis `Path.replace`, qui est `os.replace`) :
    une écriture interrompue laisserait sinon une vignette tronquée SERVIE —
    pire qu'une absence, qui est dite proprement. Un `.tmp` ORPHELIN survit à
    une interruption et rien ne le nettoie : sans conséquence, la préférence
    de `_etabli_vignette` nommant des fichiers exacts, et dit plutôt que tu.
    """
    if version < 1:
        raise HTTPException(400, f"vignette : version {version} — les "
                                 "versions sont numérotées à partir de 1")
    # UNE SEULE PORTE POUR LE NOM, et le dossier se déduit du chemin plutôt
    # que d'être recomposé. Mesuré : avec un `Path(job).name` ICI *en plus* de
    # celui du composeur, retirer l'un OU l'autre laissait le banc de
    # traversée entièrement VERT — deux gardes IDENTIQUES se couvrent l'une
    # l'autre, et plus rien ne dit laquelle tient. `mesh_report.job_dir()` est
    # évité pour la même raison : il aplatit lui aussi, et l'aurait masqué une
    # troisième fois. (Les deux gardes de `_etabli_vignette_cible`, elles, ne
    # sont pas identiques : elles composent, et deux bancs les séparent.)
    p = _etabli_vignette_cible(job, version)
    d = p.parent
    glb = d / ("model.glb" if version <= 1 else f"model.v{version}.glb")
    if not glb.is_file():
        raise HTTPException(404, f"vignette : {d.name}/{glb.name} introuvable "
                                 "— une vignette sans maillage ne dit rien")
    octets = await request.body()
    if len(octets) > _ETABLI_VIGNETTE_MAX:
        raise HTTPException(413, f"vignette : {len(octets)} octets, la borne "
                                 f"est à {_ETABLI_VIGNETTE_MAX // 1024 // 1024}"
                                 " Mio — réduisez avant d'envoyer")
    if not octets.startswith(_PNG_MAGIC):
        raise HTTPException(400, "vignette : un PNG est attendu (signature "
                                 "absente, l'en-tête ne prouve rien)")
    tmp = p.parent / f"{p.name}.tmp"
    tmp.write_bytes(octets)
    # `Path.replace` EST `os.replace` — même appel système, même atomicité,
    # sans l'import local qui a déjà masqué un NameError dans ce dépôt.
    tmp.replace(p)
    return {"ok": True, "fichier": p.name, "octets": len(octets)}


@router.get("/etabli/rig")
async def etabli_rig(job: str, version: int = 1):
    from app.services import mesh_edit
    return mesh_edit.rig_inventory(_etabli_glb(job, version))


@router.post("/etabli/adopter")
async def etabli_adopter(body: dict):
    """Fait entrer une tâche Meshy dans le monde des jobs (spec §6.2), pour
    qu'une correction ait où être versionnée."""
    from app.services import mesh_edit
    try:
        job = mesh_edit.adopter_meshy(str(body.get("task_id") or ""),
                                      str(body.get("fichier") or "model.glb"))
    except FileNotFoundError as e:
        raise HTTPException(404, str(e))
    return {"job": job, "version": 1,
            "url": f"/api/assets/3d/{job}/version/1"}


# LES CINQ ROUTES D'ÉCRITURE PASSENT LA MÊME PORTE, `_etabli_glb_cible` (revue
# du lot B) : `version` un entier ≥ 1, `job` une chaîne, les deux gardes de
# chemin, et `depuis` — la version dont l'écriture part — dans la fiche. Avant,
# les trois routes de P1 prenaient `int(version or 1)` : « 1 », 0, 1,5 ou True
# écrivaient une version, [1] faisait un 500, et `job = 12` un 404 là où les
# routes du lot B disaient 400.

@router.post("/etabli/extraire")
async def etabli_extraire(body: dict):
    from app.services import mesh_edit
    job, data, depuis = _etabli_glb_cible(body.get("job"), body.get("version"),
                                          "extraction")
    noeuds = body.get("noeuds")
    try:
        sortie = mesh_edit.extraire(data, noeuds or [])
    except ValueError as e:
        raise HTTPException(400, str(e))
    return _etabli_ecrire(job, sortie, "extraire",
                          {"depuis": depuis, "noeuds": list(noeuds or [])})


@router.post("/etabli/transformer")
async def etabli_transformer(body: dict):
    from app.services import mesh_edit
    job, data, depuis = _etabli_glb_cible(body.get("job"), body.get("version"),
                                          "transformation")
    try:
        sortie = mesh_edit.transformer(data, body.get("transforms"))
    except ValueError as e:
        raise HTTPException(400, str(e))
    return _etabli_ecrire(job, sortie, "transformer",
                          {"depuis": depuis,
                           "transforms": body.get("transforms") or {}})


@router.post("/etabli/reparer")
async def etabli_reparer(body: dict):
    from app.services import mesh_edit
    job, data, depuis = _etabli_glb_cible(body.get("job"), body.get("version"),
                                          "réparation")
    try:
        sortie = mesh_edit.reparer(
            data, axe_haut=body.get("axe_haut"),
            echelle=body.get("echelle"),
            recentrer=bool(body.get("recentrer")))
    except ValueError as e:
        # un GLB compressé refuse le RECENTRAGE seul : le message le dit
        raise HTTPException(400, str(e))
    return _etabli_ecrire(job, sortie, "reparer",
                          {"depuis": depuis,
                           "axe_haut": body.get("axe_haut"),
                           "echelle": body.get("echelle"),
                           "recentrer": bool(body.get("recentrer"))})


# ── le lot B de la plaque façon slicer : deux écritures de plus ──────────────
# « Poser sur une face » et le couteau. Les deux passent la porte de
# `_etabli_cible_sous_jobs` (nom dégénéré refusé, chemin résolu confiné), et
# leurs corps sont jugés ICI avant toute lecture — `mesh_edit` refuse à son
# tour en ValueError, traduites en 400.

def _etabli_glb_cible(job, version, quoi: str) -> tuple[str, bytes, dict]:
    """Les octets d'une version pour un `job` venu du RÉSEAU : les deux gardes
    de `_etabli_cible_sous_jobs`, puis le 404 franc d'un fichier absent. Rend
    aussi le nom de dossier APLATI, celui sous lequel la version s'écrira, et
    `depuis` — {version, fichier} — que chaque route dépose dans la fiche :
    sans lui, aucune fiche ne disait de quelle version une écriture partait, et
    le couteau est la première opération qui renomme des pièces."""
    from app.services import mesh_report
    if not isinstance(job, str):
        raise HTTPException(400, f"{quoi} : job « {job} » — le nom du dossier "
                                 "de job est attendu")
    if not _etabli_entier(version) or version < 1:
        raise HTTPException(400, f"{quoi} : version « {version} » — un entier "
                                 "à partir de 1")
    nom = "model.glb" if version <= 1 else f"model.v{version}.glb"
    p = _etabli_cible_sous_jobs(
        job, lambda j: mesh_report.job_dir(Path(str(j)).name) / nom, quoi)
    if not p.is_file():
        raise HTTPException(404, f"{quoi} : {p.parent.name}/{nom} introuvable")
    return p.parent.name, p.read_bytes(), {"version": version, "fichier": nom}


def _etabli_vecteur(v, quoi: str, *, direction: bool = False) -> list[float]:
    """Trois nombres finis — et pour une direction, pas tous nuls."""
    if not isinstance(v, list) or len(v) != 3 or not all(_etabli_nombre(c) for c in v):
        raise HTTPException(400, f"{quoi} — trois nombres finis [x, y, z] "
                                 "sont attendus")
    if direction and all(c == 0 for c in v):
        raise HTTPException(400, f"{quoi} — une direction ne peut pas être "
                                 "(0, 0, 0)")
    return [float(c) for c in v]


@router.post("/etabli/assise")
async def etabli_assise(body: dict):
    """Pose le modèle sur la face désignée : `normale` (monde, unitaire ou
    non) et `point` (le pivot, facultatif). Écrit une version de plus, comme
    `/etabli/reparer` dont c'est le geste en un clic — voir
    `mesh_edit.assise` pour l'invariant du nœud de correction NEUF."""
    from app.services import mesh_edit
    job, data, depuis = _etabli_glb_cible(body.get("job"), body.get("version"),
                                          "assise")
    normale = _etabli_vecteur(body.get("normale"), "assise : normale",
                              direction=True)
    point = body.get("point")
    if point is not None:
        point = _etabli_vecteur(point, "assise : point")
    try:
        sortie = mesh_edit.assise(data, normale=normale, point=point)
    except ValueError as e:
        # un GLB compressé refuse : la translation de contact lit la géométrie
        raise HTTPException(400, str(e))
    return _etabli_ecrire(job, sortie, "assise",
                          {"depuis": depuis, "normale": normale, "point": point})


_ETABLI_GARDER = ("deux", "a", "b")


@router.post("/etabli/couper")
async def etabli_couper(body: dict):
    """Le couteau : coupe les nœuds `noeuds` par le plan (`point`, `normale`)
    et écrit une version de plus. `garder` ∈ deux | a | b (a : le côté vers
    lequel pointe la normale). Le compte rendu de `mesh_cut.couper` — nœuds
    produits (`_avant` / `_apres`), capuchons posés ou non et pourquoi — et
    `depuis` deviennent le `source` de la fiche : c'est de là que
    `/etabli/productions` et la Bibliothèque peuvent le lire (l'entrée n'en
    remonte que le nom), et le format est décrit en tête de la section couteau
    de mesh_cut.py.

    Aucun plan de plaque n'est reporté sur la version coupée, et c'est voulu
    (voir le format du plan) : ses index de nœud ne sont plus ceux d'avant.
    """
    from app.services import mesh_cut
    job, data, depuis = _etabli_glb_cible(body.get("job"), body.get("version"),
                                          "couteau")
    noeuds = body.get("noeuds")
    if not isinstance(noeuds, list) or not noeuds:
        raise HTTPException(400, "couteau : `noeuds` doit être une liste non "
                                 "vide d'index de nœud — le couteau ne tranche "
                                 "jamais tout le modèle par défaut")
    for n in noeuds:
        if not _etabli_entier(n) or n < 0:
            raise HTTPException(400, f"couteau : nœud « {n} » — un entier ≥ 0 "
                                     "(index de nœud glTF)")
    point = _etabli_vecteur(body.get("point"), "couteau : point")
    normale = _etabli_vecteur(body.get("normale"), "couteau : normale",
                              direction=True)
    garder = body.get("garder", "deux")
    if garder not in _ETABLI_GARDER:
        raise HTTPException(400, f"couteau : garder « {garder} » — deux, a ou b")
    try:
        sortie, rapport = mesh_cut.couper(data, noeuds, point, normale, garder)
    except ValueError as e:
        raise HTTPException(400, str(e))
    return _etabli_ecrire(job, sortie, "couper", {"depuis": depuis, **rapport})
